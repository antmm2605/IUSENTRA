from __future__ import annotations

import base64
import hashlib
import io
import json
from datetime import datetime
from pathlib import Path
from types import SimpleNamespace
from zoneinfo import ZoneInfo

import pytest

from pct.notifiche_legali import (
    LEGAL_NOTIFICATION_SUBJECT,
    RECIPIENT_NOTIFICATION_DIRECTIVES,
    UNEP_REQUEST_TYPES,
    available_template_fields,
    build_attestazione_conformita_payload,
    build_client_communication,
    build_public_register_confirmation_evidence,
    generate_attestazione_conformita_docx,
    generate_attestazione_conformita_pdf_bytes,
    generate_relata_pdf_bytes,
    build_notification_attachment_manifest,
    build_notification_normative_checks,
    build_notification_send_plan,
    build_notification_signature_plan,
    build_notification_timing_plan,
    client_communication_templates_version,
    legal_notification_automation_payload,
    list_client_communication_templates,
    list_notification_templates,
    normalise_public_register,
    public_register_capability,
    get_notification_template,
    is_plausible_pec_address,
    notification_directive_matrix,
    office_notification_evidence_from_pec,
    preview_legal_relata,
    prepare_pst_failed_notification_workflow,
    released_office_documents_from_pec,
    template_catalog_version,
    validate_legal_notification,
    validate_non_pec_notification_tracking,
    validate_unep_notification_request,
    _pec_verification_matches,
)
from pct.prova_deposito_notifica import validate_deposit_notification_proof
from pct.clienti import TipoCliente
from pct.fascicoli import TipoDocumento, TipoFascicolo
from pct.soggetti import Recapiti, RuoloSoggetto, TipoSoggetto
from tests.test_web_bootstrap import _cfg_web, _write_studio_config
from web.app import create_app
from web.blueprints import api_v1_react as react_api
from web.helpers import get_clienti, get_fascicoli, get_soggetti
from web.services import react_notifiche_legali_bridge
from web.services.react_notifiche_legali_bridge import build_react_notifiche_legali_payload


def _app(tmp_path: Path):
    _write_studio_config(tmp_path / "config" / "studio.json")
    app = create_app(_cfg_web(tmp_path))
    app.config["API_KEY"] = "react-test-key"
    return app


def _pec_evidence(source: str, pec: str, tax_code: str, checked_at: str) -> dict[str, object]:
    confirmation = source not in {"reginde", "registro_ppaa"}
    capability = public_register_capability(source)
    document = {
        "source": source,
        "source_label": capability["label"],
        "official_url": capability["official_url"],
        "subject": "Soggetto verificato",
        "codice_fiscale": tax_code,
        "pec": pec,
        "consulted_at": checked_at,
        "confirmed_at": checked_at,
        "confirmed_by": "Avvocato di prova",
        "verification_method": "official_register_user_confirmation",
    }
    raw = json.dumps(document, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8") if confirmation else (
        f"<risposta><codiceFiscale>{tax_code}</codiceFiscale>"
        f"<postaElettronicaCertificata>{pec}</postaElettronicaCertificata>"
        "<stato>ATTIVO</stato></risposta>"
    ).encode("utf-8")
    evidence = {
        "source": source,
        "verified": True,
        "found": True,
        "pec_attesa": pec,
        "codice_fiscale": tax_code,
        "verified_at": checked_at,
        "checked_at": checked_at,
        "evidence_sha256": hashlib.sha256(raw).hexdigest(),
        "evidence_body_b64": base64.b64encode(raw).decode("ascii"),
    }
    if confirmation:
        evidence.update({
            "verification_method": "official_register_user_confirmation",
            "confirmed_by": "Avvocato di prova",
            "confirmed_at": checked_at,
            "consulted_at": checked_at,
            "official_url": capability["official_url"],
        })
    return evidence


def _legal_payload() -> dict[str, object]:
    payload: dict[str, object] = {
        "operazione": "notifica_pec_l53",
        "oggetto_pec": LEGAL_NOTIFICATION_SUBJECT,
        "avvocato_nome": "Mario Rossi",
        "avvocato_cf": "RSSMRA80A01H501U",
        "avvocato_foro": "Roma",
        "studio_indirizzo": "Via Roma 1",
        "studio_cap": "00100",
        "studio_citta": "Roma",
        "studio_provincia": "RM",
        "mittente_pec": "studio@example.pec.it",
        "fonte_pec_mittente": "ReGIndE",
        "mittente_pec_pubblico_elenco": True,
        "mittente_avvocato_abilitato": True,
        "mittente_pec_validata": True,
        "assistito_nome": "Cliente S.r.l.",
        "assistito_cf": "01234567890",
        "ruolo_destinatario": "controparte",
        "destinatario_nome": "Controparte S.p.A.",
        "destinatario_cf": "01234567890",
        "destinatario_pec": "controparte@example.pec.it",
        "fonte_pec_destinatario": "registro_imprese",
        "destinatario_pec_pubblico_elenco": True,
        "data_verifica_pec": "2026-05-12T10:30",
        "data_relata": "2026-05-12",
        "ora_relata": "14:25",
        "procedimento_pendente": True,
        "ufficio_giudiziario": "Tribunale di Roma",
        "sezione": "III",
        "numero_rg": "1234",
        "anno_rg": "2026",
        "ricevuta_completa": True,
        "relata_firmata": True,
        "relata_documento_separato": True,
        "approvazione_avvocato": True,
        "documenti": [
            {
                "nome_file": "ricorso.pdf",
                "descrizione": "Ricorso notificato",
                "origine": "copia_fascicolo",
            }
        ],
        "attestazione_conformita": "che il file ricorso.pdf è copia informatica conforme al fascicolo informatico.",
    }
    payload["verifica_pec_mittente"] = _pec_evidence(
        "reginde",
        "studio@example.pec.it",
        "RSSMRA80A01H501U",
        "2026-05-12T10:30:00+02:00",
    )
    payload["verifiche_pec_destinatari"] = [_pec_evidence(
        "registro_imprese",
        "controparte@example.pec.it",
        "01234567890",
        "2026-05-12T10:30:01+02:00",
    )]
    return payload


@pytest.fixture
def legal_payload_due_destinatari() -> dict[str, object]:
    payload = _legal_payload()
    avvocatura = {
        "id": "destinatario-avvocatura-reggio-calabria",
        "nome": "Avvocatura Distrettuale dello Stato di Reggio Calabria",
        "codice_fiscale_piva": "80224030587",
        "pec": "ads.rc@mailcert.avvocaturastato.it",
        "ruolo": "difensore",
        "fonte_pec": "reginde",
        "parte_rappresentata": "Ministero dell'Istruzione e del Merito",
    }
    ministero = {
        "id": "destinatario-ministero-istruzione",
        "nome": "Ministero dell'Istruzione e del Merito",
        "codice_fiscale_piva": "80185250588",
        "pec": "dgosv@postacert.istruzione.it",
        "ruolo": "pa",
        "fonte_pec": "registro_ppaa",
    }
    payload.update({
        "caso_notifica": "in_corso_di_causa",
        "ruolo_destinatario": avvocatura["ruolo"],
        "destinatario_nome": avvocatura["nome"],
        "destinatario_cf": avvocatura["codice_fiscale_piva"],
        "destinatario_pec": avvocatura["pec"],
        "fonte_pec_destinatario": avvocatura["fonte_pec"],
        "destinatario_parte_rappresentata": avvocatura["parte_rappresentata"],
        "destinatari": [avvocatura, ministero],
        "verifiche_pec_destinatari": [
            _pec_evidence(
                "reginde",
                avvocatura["pec"],
                avvocatura["codice_fiscale_piva"],
                "2026-07-21T09:12:13+02:00",
            ),
            _pec_evidence(
                "registro_ppaa",
                ministero["pec"],
                ministero["codice_fiscale_piva"],
                "2026-07-21T11:47:59+02:00",
            ),
        ],
    })
    return payload


def test_notifica_l53_genera_relata_solo_con_controlli_completi():
    result = validate_legal_notification(_legal_payload())

    assert result.ok is True
    assert result.subject == LEGAL_NOTIFICATION_SUBJECT
    assert "RELATA DI NOTIFICA EX ART. 3-BIS L. 53/1994 E SUCC. MOD." in result.relata_text
    assert "A) - Ricorso notificato (File: ricorso.pdf)" in result.relata_text
    assert "B) - Attestazione di conformit" in result.relata_text
    assert "C) - Relata di notifica." in result.relata_text
    assert "Registro Imprese" in result.relata_text
    assert "RG: 1234/2026" in result.relata_text
    assert "Via Roma 1, CAP 00100" in result.relata_text
    assert "Roma (RM)" in result.relata_text


def test_notifica_pec_non_richiede_conferma_manual_abilitazione_avvocato():
    payload = _legal_payload()
    payload.pop("mittente_avvocato_abilitato", None)

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert all("AVVOCATO_ABILITATO" not in item for item in result.blockers)


def test_payload_notifiche_legali_completa_cap_dal_comune_dello_studio():
    config = SimpleNamespace(
        studio=SimpleNamespace(
            nome="Studio Legale Montagnese",
            avvocato="Giuseppe Montagnese",
            indirizzo="Via NINO BIXIO 4",
            cap="",
            city="Taurianova",
            province="RC",
        ),
        pec=SimpleNamespace(indirizzo="studio@example.pec.it"),
    )

    payload = build_react_notifiche_legali_payload(config_studio=config)

    assert payload["defaults"]["studioIndirizzo"] == "Via NINO BIXIO 4"
    assert payload["defaults"]["studioCap"] == "89029"
    assert payload["defaults"]["studioCitta"] == "Taurianova"
    assert payload["defaults"]["studioProvincia"] == "RC"


def test_notifica_non_blocca_la_preparazione_per_prova_pec_storica_non_coerente():
    payload = _legal_payload()
    payload["destinatario_cf"] = "01234567890"
    payload["verifica_pec_mittente"] = _pec_evidence(
        "reginde", "studio@example.pec.it", "RSSMRA80A01H501U", "2026-07-13T10:30:00+02:00",
    )
    payload["verifiche_pec_destinatari"] = [_pec_evidence(
        "registro_imprese", "controparte@example.pec.it", "01234567890", "2026-07-13T10:30:01+02:00",
    )]

    assert validate_legal_notification(payload).ok is True

    payload["destinatario_pec"] = "pec-diversa@example.pec.it"
    result = validate_legal_notification(payload)

    assert result.ok is True
    assert not any("PEC_DESTINATARIO_PROVA_NON_VALIDA" in item for item in result.blockers)
    assert not any("PEC_DESTINATARIO_PROVA_NON_VALIDA" in item for item in result.warnings)


def test_notifica_relata_ignora_prova_pubblico_elenco_manomessa():
    payload = _legal_payload()
    sender = dict(payload["verifica_pec_mittente"])
    sender["evidence_sha256"] = "f" * 64
    payload["verifica_pec_mittente"] = sender

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert not any("PEC_MITTENTE_PROVA_INTEGRA_REQUIRED" in item for item in result.blockers)
    assert result.output_plan is not None
    assert "expectedReceiptSubjects" not in result.output_plan["deliveryPlan"]
    assert result.output_plan["deliveryPlan"]["presidioPecAutomation"]["phase"] == "post_invio_reale"
    assert "RELATA DI NOTIFICA EX ART. 3-BIS" in result.relata_text


def test_notifica_l53_studio_telematico_non_blocca_verifica_live_rg_o_parte_rappresentata():
    payload = _legal_payload()
    payload.update({
        "template_id": "relata_pec_a_difensore_costituito",
        "caso_notifica": "ordinaria",
        "ruolo_destinatario": "difensore",
        "destinatario_nome": "Avvocatura distrettuale di Stato di Milano",
        "destinatario_cf": "97021490152",
        "destinatario_pec": "ads.mi@mailcert.avvocaturastato.it",
        "fonte_pec_destinatario": "reginde",
        "destinatario_parte_rappresentata": "",
        "procedimento_pendente": True,
        "ufficio_giudiziario": "",
        "sezione": "",
        "numero_rg": "",
        "anno_rg": "",
        "documenti": [{
            "nome_file": "ricorso-opposizione-con-decreto-fissazione.pdf",
            "descrizione": "Ricorso in opposizione a decreto ingiuntivo con decreto di fissazione udienza",
            "origine": "copia_fascicolo_informatico",
        }],
        "attestazione_conformita": "",
        "verifiche_pec_destinatari": [],
    })
    payload.pop("verifica_pec_mittente", None)

    result = validate_legal_notification(payload)
    all_messages = "\n".join([*result.blockers, *result.warnings])

    assert result.ok is True
    assert not result.blockers
    assert "PEC_DESTINATARIO_VERIFICA_REQUIRED" not in all_messages
    assert "PEC_DESTINATARIO_PUBBLICO_ELENCO_REQUIRED" not in all_messages
    assert "Per una notifica in corso" not in all_messages
    assert "Parte rappresentata" not in all_messages
    assert "[dato mancante: Parte rappresentata]" not in result.relata_text
    assert "[dato mancante: destinatario parte rappresentata]" not in result.relata_text
    assert "- Ricorso, in opposizione a decreto ingiuntivo con decreto di fissazione udienza;" in result.relata_text
    assert "- Decreto, emesso" not in result.relata_text
    assert "Attestazione di conformità.pdf" in result.output_plan["files"]
    assert "Relata di notifica.pdf" in result.output_plan["files"]

    preview = preview_legal_relata(payload)
    missing_labels = "\n".join(preview["missingFields"])
    assert "parte rappresentata" not in missing_labels.lower()

    attestazione = build_attestazione_conformita_payload(payload)
    assert attestazione["document_rows"][0]["title"] == "Ricorso"
    assert "Ricorso, in opposizione a decreto ingiuntivo con decreto di fissazione udienza;" in attestazione["text"]
    assert "Decreto, emesso" not in attestazione["text"]


def test_notifica_non_blocca_modello_con_origine_documento_da_verificare():
    payload = _legal_payload()
    payload.update({
        "template_id": "relata_pec_con_attestazione_scansione_analogica",
        "caso_notifica": "in_corso_di_causa",
        "documenti": [
            {
                "nome_file": "SentenzaDefinitiva_33581101.pdf",
                "descrizione": "Sentenza",
                "origine": "copia_fascicolo_informatico",
                "hash_sha256": "a" * 64,
            },
            {
                "nome_file": "VerbaleUdienza_33393309.pdf",
                "descrizione": "Verbale di udienza",
                "origine": "copia_fascicolo_informatico",
                "hash_sha256": "b" * 64,
            },
        ],
    })

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert not any("MODELLO_DOCUMENTO_INCOERENTE" in item for item in result.blockers)
    assert not any("MODELLO_DOCUMENTO_DA_VERIFICARE" in item for item in result.warnings)
    assert "A) - Sentenza" in result.relata_text
    assert "B) - Verbale di udienza" in result.relata_text
    assert "in data ," not in result.relata_text
    assert "è conforme alla copia informatica presente nel fascicolo informatico del relativo procedimento" in result.relata_text


def test_notifica_l53_normalizza_alias_studio_telematico_pubblici_elenchi():
    aliases = {
        "INIPEC-professionisti": "ini_pec",
        "RegistroImprese": "registro_imprese",
        "RegInde": "reginde",
        "IPA": "ipa",
        "altro": "altro_pubblico_elenco",
    }
    for raw, expected in aliases.items():
        assert normalise_public_register(raw) == expected

    payload = _legal_payload()
    payload["fonte_pec_destinatario"] = "RegistroImprese"

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert "Registro Imprese" in result.relata_text


def test_pubblici_elenchi_distinguono_servizio_autenticato_consultazione_e_fonti_non_valide():
    expected = {
        "reginde": ("authenticated_service", True, True),
        "registro_ppaa": ("authenticated_service", True, True),
        "ini_pec": ("assisted_browser", False, True),
        "registro_imprese": ("assisted_browser", False, True),
        "inad": ("assisted_browser", False, True),
        "anpr": ("not_notification_register", False, False),
        "altro_pubblico_elenco": ("documented_manual", False, True),
    }

    for source, (mode, automatic, valid) in expected.items():
        capability = public_register_capability(source)
        assert capability["verification_mode"] == mode
        assert capability["automatic"] is automatic
        assert capability["valid_for_notification"] is valid

    assert normalise_public_register("IPA") == "ipa"
    assert "ipa" not in expected


def test_consultazione_pubblico_elenco_produce_prova_verificabile_per_la_relata():
    now = datetime.now(ZoneInfo("Europe/Rome")).replace(microsecond=0).isoformat()
    evidence = build_public_register_confirmation_evidence({
        "source": "ini_pec",
        "pec": "controparte@example.pec.it",
        "codice_fiscale": "01234567890",
        "soggetto": "Controparte S.p.A.",
        "consulted_at": now,
        "fascicolo_id": "FASCICOLO-TEST",
    }, confirmed_by="Avvocato di prova")
    payload = _legal_payload()
    payload["fonte_pec_destinatario"] = "ini_pec"
    payload["verifiche_pec_destinatari"] = [evidence]

    assert validate_legal_notification(payload).ok is True
    assert evidence["verification_method"] == "official_register_user_confirmation"
    assert evidence["confirmed_by"] == "Avvocato di prova"
    assert evidence["consulted_at"]
    assert len(str(evidence["evidence_sha256"])) == 64

    raw = json.loads(base64.b64decode(str(evidence["evidence_body_b64"])).decode("utf-8"))
    raw["official_url"] = "https://example.invalid/"
    tampered_body = json.dumps(raw, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    tampered = {
        **evidence,
        "evidence_sha256": hashlib.sha256(tampered_body).hexdigest(),
        "evidence_body_b64": base64.b64encode(tampered_body).decode("ascii"),
    }
    payload["verifiche_pec_destinatari"] = [tampered]
    assert not _pec_verification_matches(
        tampered,
        expected_pec="controparte@example.pec.it",
        expected_cf="01234567890",
        expected_source="ini_pec",
    )
    assert validate_legal_notification(payload).ok is True


def test_anpr_non_puo_essere_registrato_come_prova_pec_di_notifica():
    with pytest.raises(ValueError, match="non certifica indirizzi PEC"):
        build_public_register_confirmation_evidence({
            "source": "anpr",
            "pec": "persona@example.pec.it",
            "codice_fiscale": "RSSMRA80A01H501U",
            "soggetto": "Mario Rossi",
            "consulted_at": datetime.now(ZoneInfo("Europe/Rome")).isoformat(),
        }, confirmed_by="Avvocato di prova")


def test_indirizzo_pec_plausibile_validazione_lineare_senza_redos():
    assert is_plausible_pec_address("Controparte+atti@example.pec.it")
    assert not is_plausible_pec_address("utente@gmail.com")
    assert not is_plausible_pec_address("%" * 5000)
    assert not is_plausible_pec_address("utente@-example.pec.it")


def test_payload_react_espone_modalita_e_azione_di_verifica_per_ogni_fonte():
    payload = build_react_notifiche_legali_payload(config_studio=None)
    sources = {item["value"]: item for item in payload["registriPec"]}

    assert sources["reginde"]["automatic"] is True
    assert sources["registro_ppaa"]["verificationMode"] == "authenticated_service"
    assert sources["registro_ppaa"]["automatic"] is True
    assert sources["registro_ppaa"]["requiresUserConfirmation"] is False
    assert sources["registro_ppaa"]["officialUrl"].endswith("/PST/it/pst_2_8.wp")
    assert sources["ini_pec"]["requiresUserConfirmation"] is True
    assert sources["anpr"]["validForNotification"] is False
    assert "non certifica PEC" in sources["anpr"]["label"]
    assert payload["azioni"]["verificaPecConsultata"].endswith("/verifica-pec-consultata")


def test_notifica_l53_normalizza_avvocato_e_blocco_procedimento():
    payload = _legal_payload()
    payload["avvocato_nome"] = "Avv. Giuseppe Montagnese"
    payload["sezione"] = ""
    payload["numero_rg"] = "466"
    payload["anno_rg"] = "2023"

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert "Io sottoscritto Avv. Giuseppe Montagnese C.F:" in result.relata_text
    assert "Avv. Avv." not in result.relata_text
    assert "Sezione ," not in result.relata_text
    assert "RG: 466/2023" in result.relata_text


def test_notifica_l53_attestazione_automatica_cumulativa_per_documenti_multipli():
    payload = _legal_payload()
    payload["template_id"] = "relata_pec_con_attestazione_fascicolo"
    payload["attestazione_conformita"] = ""
    payload["attestazione_multipla"] = True
    payload["documenti"] = [
        {
            "nome_file": "provvedimento.pdf",
            "descrizione": "Provvedimento",
            "origine": "copia_fascicolo_informatico",
            "attestazione_conformita_presente": True,
        },
        {
            "nome_file": "ordinanza.pdf",
            "descrizione": "Ordinanza",
            "origine": "copia_fascicolo_informatico",
            "attestazione_conformita_presente": True,
        },
        {
            "nome_file": "verbale.pdf",
            "descrizione": "Verbale",
            "origine": "copia_fascicolo_informatico",
            "attestazione_conformita_presente": True,
        },
    ]

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert result.relata_text.count("Attesto, ai sensi della normativa vigente") == 1
    assert "Attesto che il file provvedimento.pdf" not in result.relata_text
    assert "- Provvedimento, emesso dal Tribunale di Roma Sez. III;" in result.relata_text
    assert "- Ordinanza, emessa dal Tribunale di Roma Sez. III;" in result.relata_text
    assert "- Verbale, documento allegato alla notificazione;" in result.relata_text
    assert result.relata_text.count("sono conformi alle copie informatiche presenti") == 1
    assert "copia informatica conforme al corrispondente" not in result.relata_text


def test_notifica_l53_accetta_eml_scelto_come_allegato_non_autoproposto():
    payload = _legal_payload()
    payload["documenti"] = [
        {"nome_file": "ricorso.pdf", "descrizione": "Ricorso notificato", "origine": "copia_fascicolo"},
        {"nome_file": "richiesta_pagamento.eml", "descrizione": "PEC richiesta pagamento allegata", "origine": "originale_informatico"},
    ]

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert "B) - PEC richiesta pagamento allegata (File: richiesta_pagamento.eml)" in result.relata_text

    blocked = _legal_payload()
    blocked["documenti"] = [
        {"nome_file": "ricorso.pdf", "descrizione": "Ricorso notificato", "origine": "copia_fascicolo"},
        {"nome_file": "archivio.zip", "descrizione": "Archivio non ammesso", "origine": "originale_informatico"},
    ]

    blocked_result = validate_legal_notification(blocked)

    assert blocked_result.ok is True
    assert not any("PDF/PDF-A, file firmato, EML o MSG" in item for item in blocked_result.blockers)
    assert any("PDF/PDF-A, file firmato, EML o MSG" in item for item in blocked_result.warnings)


def test_notifica_l53_segnala_cliente_senza_bloccare_relata():
    payload = _legal_payload()
    payload["ruolo_destinatario"] = "cliente"
    payload["documenti"] = [{"nome_file": "scansione.pdf", "descrizione": "Provvedimento", "origine": "scansione"}]
    payload["attestazione_conformita"] = ""

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert not any("Comunicazione al cliente" in item for item in result.blockers)
    assert any("Comunicazione al cliente" in item for item in result.warnings)
    assert "RELATA DI NOTIFICA" in result.relata_text


def test_notifica_l53_genera_attestazione_automatica_da_origine_documento():
    payload = _legal_payload()
    payload["documenti"] = [{"nome_file": "scansione.pdf", "descrizione": "Provvedimento", "origine": "scansione"}]
    payload["attestazione_conformita"] = ""

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert not any("ATTESTAZIONE_REQUIRED" in item for item in result.blockers)
    assert "Attestazione di conformità.pdf" in result.output_plan["files"]
    assert "Relata di notifica.pdf" in result.output_plan["files"]


def test_notifica_l53_documento_nativo_digitale_non_richiede_attestazione():
    payload = _legal_payload()
    payload["documenti"] = [{"nome_file": "atto.pdf", "descrizione": "Atto nativo", "origine": "nativo_digitale"}]
    payload["attestazione_conformita"] = ""

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert "ATTESTAZIONE DI CONFORMITA" not in result.relata_text


def test_notifica_l53_riporta_piu_documenti_nell_elenco_allegati():
    payload = _legal_payload()
    payload["documenti"] = [
        {"nome_file": "ricorso.pdf", "descrizione": "Ricorso", "origine": "nativo_digitale"},
        {"nome_file": "procura.pdf", "descrizione": "Procura alle liti", "origine": "firmato_digitalmente"},
        {"nome_file": "provvedimento.pdf", "descrizione": "Provvedimento", "origine": "copia_fascicolo_informatico"},
    ]
    payload["attestazione_conformita"] = "Attesto la conformità del provvedimento estratto dal fascicolo informatico."

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert "A) - Ricorso (File: ricorso.pdf)" in result.relata_text
    assert "B) - Procura alle liti (File: procura.pdf)" in result.relata_text
    assert "C) - Provvedimento (File: provvedimento.pdf)" in result.relata_text
    assert "D) - Attestazione di conformit" in result.relata_text
    assert "E) - Relata di notifica." in result.relata_text


def test_attestazione_conformita_autocompila_fascicolo_cliente_e_documenti():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "ricorso.pdf",
            "descrizione": "Ricorso per il recupero delle annualità Carta del docente",
            "origine": "copia_fascicolo_informatico",
            "hash_sha256": "a" * 64,
        },
        {
            "nome_file": "procura.pdf",
            "descrizione": "Procura alle liti",
            "origine": "firmato_digitalmente",
            "hash_sha256": "b" * 64,
        },
        {
            "nome_file": "decreto_fissazione_udienza.pdf",
            "descrizione": "Decreto fissazione udienza",
            "origine": "comunicazione_cancelleria",
            "data_comunicazione_cancelleria": "2026-05-18",
            "hash_sha256": "c" * 64,
        },
    ]

    model = build_attestazione_conformita_payload(payload)

    assert model["ok"] is True
    assert model["missing_fields"] == []
    assert "ATTESTAZIONE DI CONFORMITÀ" in model["text"]
    assert "Avv. Mario Rossi" in model["text"]
    assert "R.G. n. 1234/2026" in model["text"]
    assert "Ricorso, per il recupero delle annualità Carta del docente;" in model["text"]
    assert "Procura alle liti" not in model["text"]
    assert len(model["documenti"]) == 2
    assert "comunicazione_cancelleria" in {item["origine"] for item in model["documenti"]}
    assert model["text"].count("\nAttesta\n") == 1
    assert "Dettaglio attestazioni" not in model["text"]
    assert "art. 196-undecies" in " ".join(model["normativa"])


def test_attestazione_conformita_docx_rispetta_modello_e_rimuove_evidenziazioni(tmp_path):
    output = tmp_path / "attestazione.docx"

    result = generate_attestazione_conformita_docx(_legal_payload(), output)

    assert result["ok"] is True
    assert output.exists()

    from docx import Document
    from docx.enum.text import WD_UNDERLINE
    from docx.oxml.ns import qn
    from zipfile import ZipFile

    with ZipFile(result["template_path"]) as template_zip, ZipFile(output) as output_zip:
        assert template_zip.namelist() == output_zip.namelist()
        changed_parts = {
            name
            for name in template_zip.namelist()
            if template_zip.read(name) != output_zip.read(name)
        }
    assert changed_parts == {"word/document.xml"}

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "ATTESTAZIONE DI CONFORMITÀ" in text
    assert "Avv. Mario Rossi" in text
    assert "Ricorso, notificato;" in text
    assert "Roma, " not in text
    assert len(document.sections) == 1
    section = document.sections[0]
    assert section.page_width.twips == 11910
    assert section.page_height.twips == 16840
    assert section.top_margin.twips == 1340
    assert section.bottom_margin.twips == 280
    assert section.left_margin.twips == 1020
    assert section.right_margin.twips == 1020
    list_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.style.name == "List Paragraph"]
    assert len(list_paragraphs) == 1
    assert list_paragraphs[0].runs[0].text == "Ricorso"
    assert list_paragraphs[0].runs[0].bold is True
    assert list_paragraphs[0].runs[0].underline == WD_UNDERLINE.THICK
    assert document.paragraphs[-2].paragraph_format.left_indent.twips == 5778
    assert document.paragraphs[-2].runs[0].bold is True
    assert document.paragraphs[-2].runs[0].italic is True
    assert document.paragraphs[-1].runs[0].italic is True
    assert not document.element.body.xpath(".//w:highlight")
    assert "Giuseppe Montagnese" not in text
    assert "MNTGPP94L01G791A" not in text
    assert not any(
        run._r.rPr is not None and run._r.rPr.find(qn("w:highlight")) is not None
        for paragraph in document.paragraphs
        for run in paragraph.runs
    )


def test_attestazione_conformita_preserva_descrizione_con_virgole():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "ricorso.pdf",
            "descrizione": (
                "Ricorso per il recupero delle annualità della Carta del docente richiesto "
                "a favore dell'assistito e contro il Ministero, i.p.l.r.p.t."
            ),
            "origine": "copia_fascicolo_informatico",
            "data_documento": "2026-05-18",
        }
    ]

    model = build_attestazione_conformita_payload(payload)

    assert model["document_rows"] == [
        {
            "title": "Ricorso",
            "detail": (
                "per il recupero delle annualità della Carta del docente richiesto "
                "a favore dell'assistito e contro il Ministero, i.p.l.r.p.t., "
                "depositato in data 18/05/2026"
            ),
            "text": (
                "Ricorso, per il recupero delle annualità della Carta del docente richiesto "
                "a favore dell'assistito e contro il Ministero, i.p.l.r.p.t., "
                "depositato in data 18/05/2026"
            ),
        }
    ]


def test_attestazione_conformita_pdf_riproduce_modello_con_tutti_i_documenti():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "ricorso.pdf",
            "descrizione": "Ricorso",
            "origine": "copia_fascicolo_informatico",
            "data_documento": "2026-05-18",
        },
        {
            "nome_file": "procura.pdf",
            "descrizione": "Procura alle liti",
            "origine": "copia_fascicolo_informatico",
            "data_documento": "2026-05-18",
        },
        {
            "nome_file": "decreto.pdf",
            "descrizione": "Decreto fissazione udienza",
            "origine": "copia_fascicolo_informatico",
            "data_documento": "2026-05-19",
        },
    ]
    payload["attestazione_multipla"] = True

    pdf = generate_attestazione_conformita_pdf_bytes(payload)

    from pypdf import PdfReader

    extracted = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(pdf)).pages)
    assert pdf.startswith(b"%PDF")
    assert "ATTESTAZIONE DI CONFORMIT" in extracted
    assert "Ricorso" in extracted
    assert "Procura" in extracted
    assert "alle liti" in extracted
    assert "Decreto fissazione udienza" in extracted
    assert "R.G. n. 1234/2026" in extracted


def test_attestazione_pdf_non_tratta_ricorso_come_sentenza_per_caso_globale():
    payload = _legal_payload()
    payload.update({
        "caso_notifica": "sentenza_termine_breve",
        "provvedimento_tipo": "SentenzaDefinitiva",
        "ufficio_giudiziario": "",
        "sezione": "",
        "provvedimento_data": "",
        "provvedimento_data_rilascio": "",
        "documenti": [
            {
                "nome_file": "sentenza_originale.pdf.p7m",
                "descrizione": "Sentenza da notificare",
                "origine": "originale_informatico",
            },
            {
                "nome_file": "ricorso_opposizione_decreto_ingiuntivo.pdf",
                "descrizione": "Ricorso in opposizione a decreto ingiuntivo con decreto di fissazione udienza",
                "origine": "copia_fascicolo_informatico",
            },
        ],
    })

    model = build_attestazione_conformita_payload(payload)
    pdf = generate_attestazione_conformita_pdf_bytes(payload)

    assert model["ok"] is True
    assert "procedimento.ufficio" not in model["missing_fields"]
    assert "provvedimento.data_rilascio" not in model["missing_fields"]
    assert model["document_rows"][0]["title"] == "Ricorso"
    assert "Ricorso, in opposizione a decreto ingiuntivo con decreto di fissazione udienza;" in model["text"]
    assert "Sentenza, emessa" not in model["text"]
    assert pdf.startswith(b"%PDF")


def test_attestazione_modificata_dall_avvocato_alimenta_testo_e_pdf():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "verbale_udienza.pdf",
            "descrizione": "Verbale di udienza",
            "origine": "copia_fascicolo_informatico",
            "data_documento": "2026-06-05",
        }
    ]
    override = (
        "ATTESTAZIONE DI CONFORMITÀ\n\n"
        "Il sottoscritto Avv. Mario Rossi attesta il testo modificato dall'avvocato.\n\n"
        "Mario Rossi\n"
        "Firmato digitalmente"
    )
    payload["attestazione_override_text"] = override

    model = build_attestazione_conformita_payload(payload)
    pdf = generate_attestazione_conformita_pdf_bytes(payload)

    from pypdf import PdfReader

    extracted = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(pdf)).pages)
    assert model["text"] == override + "\n"
    assert "testo modificato dall'avvocato" in extracted
    assert pdf.startswith(b"%PDF")


def test_piano_pec_allega_automaticamente_attestazione_salvata():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "verbale_udienza.pdf",
            "descrizione": "Verbale di udienza",
            "origine": "copia_fascicolo_informatico",
            "hash_sha256": "a" * 64,
        }
    ]
    payload["attestazione_conformita_file"] = "Attestazione_di_conformita_1234_2026.pdf"
    payload["attestazione_conformita_sha256"] = "b" * 64

    plan = build_notification_send_plan(payload)
    attestation = next(item for item in plan["attachments"] if item["id"] == "attestazione_conformita")

    assert attestation["filename"] == "Attestazione_di_conformita_1234_2026.pdf"
    assert attestation["sha256"] == "b" * 64
    assert attestation["required"] is True


def test_relata_elenca_tutti_documenti_attestazione_e_relata_senza_firma_doppia():
    payload = _legal_payload()
    payload["caso_notifica"] = "in_corso_di_causa"
    payload["documenti"] = [
        {
            "nome_file": "ricorso.pdf",
            "descrizione": "Ricorso",
            "origine": "copia_fascicolo_informatico",
            "attestazione_conformita_presente": True,
        },
        {
            "nome_file": "documento_allegato.pdf",
            "descrizione": "documento allegato",
            "origine": "originale_informatico",
        },
        {
            "nome_file": "decreto_fissazione_udienza.pdf",
            "descrizione": "Decreto fissazione udienza",
            "origine": "copia_fascicolo_informatico",
            "attestazione_conformita_presente": True,
        },
    ]

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert "A) - Ricorso" in result.relata_text
    assert "B) - documento allegato" in result.relata_text
    assert "C) - Decreto fissazione udienza" in result.relata_text
    assert "D) - Attestazione di conformit" in result.relata_text
    assert "E) - Relata di notifica." in result.relata_text
    assert "Contenuto del documento: attestazione unica di conformit" in result.relata_text
    assert "F.to digitalmente da\nAvv. Mario Rossi\nFirmato digitalmente" not in result.relata_text
    assert result.relata_text.count("Firmato digitalmente") == 0
    assert result.output_plan["files"][:5] == [
        "ricorso.pdf",
        "documento_allegato.pdf",
        "decreto_fissazione_udienza.pdf",
        "Attestazione di conformità.pdf",
        "Relata di notifica.pdf",
    ]


def test_attestazione_sentenza_autocompila_modello_word_e_firma():
    payload = _legal_payload()
    payload.update({
        "caso_notifica": "sentenza_termine_breve",
        "avvocato_nome": "Giuseppe",
        "avvocato_cognome": "Montagnese",
        "avvocato_cf": "MNTGPP94L01G791A",
        "avvocato_foro": "Palmi",
        "ufficio_giudiziario": "Tribunale di Palmi",
        "sezione": "Lavoro",
        "numero_rg": "704",
        "anno_rg": "2026",
        "provvedimento_tipo": "SentenzaDefinitiva",
        "provvedimento_data_rilascio": "2026-06-05",
        "documenti": [
            {
                "nome_file": "sentenza.pdf",
                "descrizione": "Sentenza",
                "origine": "copia_fascicolo_informatico",
                "data_documento": "2026-06-05",
                "provvedimentoTipo": "SentenzaDefinitiva",
            }
        ],
        "attestazione_multipla": True,
    })
    payload["verifica_pec_mittente"] = _pec_evidence(
        "reginde", "studio@example.pec.it", "MNTGPP94L01G791A", "2026-05-12T10:30:00+02:00",
    )

    model = build_attestazione_conformita_payload(payload)
    result = validate_legal_notification(payload)

    assert model["ok"] is True
    assert "Il sottoscritto Avv. Giuseppe Montagnese C. F. MNTGPP94L01G791A, del Foro di Palmi," in model["text"]
    assert "Sentenza, emessa dal Tribunale di Palmi Sez. Lavoro in data 05/06/2026" in model["text"]
    assert "SentenzaDefinitiva" not in model["text"]
    assert "Sentenza, Definitiva" not in model["text"]
    assert "R.G. n. 704/2026 dal quale è estratta." in model["text"]
    assert "Avv. Giuseppe Montagnese" in model["text"]
    assert "Firmato digitalmente" in model["text"]
    assert model["campi_database"]["avvocato"]["firma_in_calce"] == "Avv. Giuseppe Montagnese"
    assert model["campi_database"]["avvocato"]["firma_digitale_dicitura"] == "Firmato digitalmente"
    assert result.ok is True
    assert result.template_id == "relata_sentenza_attestazione_conformita"
    assert "ATTESTO" in result.relata_text
    assert "Sentenza, emessa dal Tribunale di Palmi Sez. Lavoro in data 05/06/2026" in result.relata_text
    assert "SentenzaDefinitiva" not in result.relata_text
    assert "Sentenza, Definitiva" not in result.relata_text


def test_relata_non_trascrive_sentenza_su_verbale_udienza_nello_stesso_elenco():
    payload = _legal_payload()
    payload.update({
        "caso_notifica": "sentenza_termine_breve",
        "ufficio_giudiziario": "Tribunale di Palmi",
        "sezione": "Lavoro",
        "numero_rg": "704",
        "anno_rg": "2026",
        "provvedimento_tipo": "SentenzaDefinitiva",
        "provvedimento_data_rilascio": "2026-06-05",
        "documenti": [
            {
                "nome_file": "SentenzaDefinitiva_33581101.pdf",
                "descrizione": "Sentenza",
                "origine": "copia_fascicolo_informatico",
                "hash_sha256": "a" * 64,
                "provvedimentoTipo": "SentenzaDefinitiva",
            },
            {
                "nome_file": "VerbaleUdienza_33393309.pdf",
                "descrizione": "Sentenza",
                "origine": "copia_fascicolo_informatico",
                "hash_sha256": "b" * 64,
                "provvedimentoTipo": "SentenzaDefinitiva",
                "provvedimentoData": "2025-12-16",
            },
        ],
    })

    result = validate_legal_notification(payload)
    model = build_attestazione_conformita_payload(payload)

    assert result.ok is True
    assert "A) - Sentenza (File: SentenzaDefinitiva_33581101.pdf)" in result.relata_text
    assert "B) - Verbale di udienza (File: VerbaleUdienza_33393309.pdf)" in result.relata_text
    assert "B) - Sentenza (File: VerbaleUdienza_33393309.pdf)" not in result.relata_text
    assert "- Verbale di udienza, estratto dal fascicolo informatico del Tribunale di Palmi Sez. Lavoro in data 16/12/2025;" in model["text"]
    assert "- Sentenza, emessa dal Tribunale di Palmi Sez. Lavoro in data 16/12/2025;" not in model["text"]
    assert "- Sentenza, emessa dal Tribunale di Palmi" in model["text"]


def test_relata_deriva_decreto_fissazione_da_suggerimento_documento_non_dal_nome_file():
    payload = _legal_payload()
    payload.pop("caso_notifica", None)
    payload.update({
        "provvedimento_tipo": "",
        "provvedimento_data": "",
        "documenti": [
            {
                "nome_file": "decreto_ingiuntivo_nome_sbagliato.pdf",
                "descrizione": "Documento letto dal fascicolo",
                "origine": "copia_fascicolo_informatico",
                "data_documento": "2026-05-18",
                "casoNotificaSuggerito": "provvedimento_giudice",
                "modelloRelataSuggerito": "relata_provvedimento_giudice",
                "provvedimentoTipo": "Decreto fissazione udienza",
                "criterioTipoDocumento": "testo documento letto",
            }
        ],
        "attestazione_multipla": True,
    })

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert result.template_id == "relata_provvedimento_giudice"
    assert result.output_plan["notificationDirective"]["caseId"] == "provvedimento_giudice"
    assert "Decreto ingiuntivo" not in result.template_label
    assert "Decreto fissazione udienza" in result.relata_text


def test_notifica_l53_audit_automatico_include_normativa_e_piu_allegati():
    payload = _legal_payload()
    payload["documenti"] = [
        {"nome_file": "ricorso.pdf", "descrizione": "Ricorso", "origine": "nativo_digitale", "hash_sha256": "a" * 64},
        {"nome_file": "procura.pdf", "descrizione": "Procura alle liti", "origine": "firmato_digitalmente", "hash_sha256": "b" * 64},
        {"nome_file": "provvedimento.pdf", "descrizione": "Provvedimento", "origine": "copia_fascicolo_informatico", "hash_sha256": "c" * 64},
    ]
    payload["attestazione_multipla"] = True

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert {item["id"] for item in legal_notification_automation_payload()["notifica"]} >= {"precompilazione", "pubblici_elenchi", "allegati"}
    assert result.output_plan["auditTrail"]["documentsCount"] == 3
    assert result.output_plan["workflowSteps"][0]["source"].startswith("L. 53/1994")
    assert any(item["id"] == "allegati" and item["status"] == "superato" for item in result.output_plan["normativeChecks"])
    assert any(item["id"] == "attestazioni" and item["status"] == "superato" for item in result.output_plan["normativeChecks"])


def test_notifica_l53_segnala_documento_ufficio_rilasciato_non_acquisito():
    payload = _legal_payload()
    payload["documento_ufficio_rilasciato"] = True
    payload["acquisizione_portale_richiesta"] = True
    payload["documenti"] = []

    result = validate_legal_notification(payload)
    checks = build_notification_normative_checks(payload)

    assert result.ok is True
    assert not any("DOCUMENTO_UFFICIO_ACQUISIZIONE_REQUIRED" in item for item in result.blockers)
    assert any("DOCUMENTO_UFFICIO_ACQUISIZIONE_REQUIRED" in item for item in result.warnings)
    assert any(item["id"] == "documento_ufficio_acquisito" and item["status"] == "da completare" for item in checks)


def test_notifica_l53_documento_ufficio_acquisito_dal_portale_supera_controllo():
    payload = _legal_payload()
    payload["documento_ufficio_rilasciato"] = True
    payload["acquisizione_portale_richiesta"] = True
    payload["pec_ufficio_eml_file"] = "pec-cancelleria.eml"
    payload["pec_ufficio_eml_sha256"] = "d" * 64
    payload["documenti"] = [
        {
            "nome_file": "ordinanza_da_notificare.pdf",
            "descrizione": "Ordinanza rilasciata dall'ufficio",
            "origine": "copia_fascicolo_informatico",
            "fonte_documento": "PORTALE_TELEMATICO",
            "servizio_portale": "PST",
            "riferimento_portale": "PST-REL-2026-0001",
            "documento_ufficio": True,
            "acquisito_da_portale": True,
            "notifica_richiesta": True,
            "data_rilascio_portale": "2026-05-23",
        }
    ]
    payload["attestazione_multipla"] = True

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert result.output_plan["auditTrail"]["officeDocumentAcquisition"]["acquired"] is True
    assert any(item["id"] == "documento_ufficio_acquisito" and item["status"] == "superato" for item in result.output_plan["normativeChecks"])


def test_rilascio_documento_ufficio_parte_da_pec_cancelleria_e_non_da_metadati_portale():
    fascicolo = SimpleNamespace(
        id="fasc-portal",
        numero="FASC-1",
        titolo="Rossi c/ Bianchi",
        tribunale="Tribunale di Roma",
        numero_rg="1234",
        anno_rg=2026,
        documenti=[],
        depositi_pct=[
            SimpleNamespace(
                id="dep-1",
                id_deposito_esterno="PST-REL-2026-0001",
                fonte="PST",
                servizio_portale="ConsultazioneFascicolo",
                tipo_atto="Ordinanza da notificare",
                data_deposito="2026-05-23",
                mittente="Tribunale di Roma",
                documenti_portale=[
                    {
                        "id_documento": "doc-rel-1",
                        "nome": "ordinanza_da_notificare.pdf",
                        "tipo": "Ordinanza da notificare",
                        "data_deposito": "2026-05-23",
                        "disponibile": True,
                    }
                ],
            )
        ],
    )

    assert released_office_documents_from_pec(fascicolo, []) == []

    pec = SimpleNamespace(
        id="pec-cancelleria-1",
        mittente="cancelleria.tribunale.roma@giustiziacert.it",
        oggetto="Tribunale di Roma R.G. 1234/2026 - provvedimento da notificare",
        data="2026-05-23T10:15:00",
        corpo_testo="Si comunica il provvedimento ordinanza_da_notificare.pdf da notificare nel procedimento R.G. 1234/2026.",
        allegati=[{"nome": "ordinanza_da_notificare.pdf", "sha256": "a" * 64}],
        message_id="<pec-cancelleria-1@giustizia>",
    )
    releases = released_office_documents_from_pec(fascicolo, [pec])

    assert len(releases) == 1
    assert releases[0]["nome"] == "ordinanza_da_notificare.pdf"
    assert releases[0]["fonteControllo"] == "pec_cancelleria"
    assert releases[0]["notificaRichiesta"] is True
    assert "single_document=1" in releases[0]["acquisitionHref"]
    assert "documento=ordinanza_da_notificare.pdf" in releases[0]["acquisitionHref"]
    assert "non_duplicare_documenti=1" in releases[0]["acquisitionHref"]

    fascicolo.documenti = [SimpleNamespace(id="doc-local", id_documento_portale="doc-rel-1", nome="ordinanza_da_notificare.pdf", hash_sha256="a" * 64)]
    assert released_office_documents_from_pec(fascicolo, [pec]) == []
    evidence = office_notification_evidence_from_pec(fascicolo, [pec])
    assert evidence[0]["acquisito"] is True

    altro_fascicolo_stesso_anno = SimpleNamespace(
        id="fasc-altro",
        numero="FASC-ALTRO",
        titolo="Altro fascicolo",
        tribunale="Tribunale di Roma",
        numero_rg="2026",
        anno_rg=2026,
        documenti=[],
    )
    assert released_office_documents_from_pec(altro_fascicolo_stesso_anno, [pec]) == []


def test_monitor_documenti_ufficio_esclude_ricevute_pct_e_nome_fittizio_notificato():
    fascicolo = SimpleNamespace(
        id="fasc-ricevute",
        numero="FASC-RICEVUTE",
        titolo="Romeo Maria c. MIM",
        tribunale="Tribunale di Palmi",
        numero_rg="1428",
        anno_rg=2026,
        documenti=[],
    )
    ricevute = [
        SimpleNamespace(
            id="pec-accettazione",
            cartella="CESTINO",
            mittente="tribunale.palmi@giustiziacert.it",
            oggetto="ACCETTAZIONE DEPOSITO TELEMATICO R.G. 1428/2026",
            corpo_testo="Accettato il ricorso Romeo Maria (originale notificato).pdf.",
            data="2026-05-20T10:01:00",
            message_id="<accettazione-1428@giustizia>",
            allegati=[],
        ),
        SimpleNamespace(
            id="pec-esito-inbox",
            cartella="INBOX",
            mittente="tribunale.palmi@giustiziacert.it",
            oggetto="ESITO CONTROLLI AUTOMATICI DEPOSITO R.G. 1428/2026",
            corpo_testo="Esito positivo per il ricorso Romeo Maria (originale notificato).pdf.",
            data="2026-05-20T10:02:00",
            message_id="<esito-1428@giustizia>",
            allegati=[],
        ),
        SimpleNamespace(
            id="pec-esito-cestino",
            cartella="CESTINO",
            mittente="tribunale.palmi@giustiziacert.it",
            oggetto="ESITO CONTROLLI AUTOMATICI DEPOSITO R.G. 1428/2026",
            corpo_testo="Esito positivo per il ricorso Romeo Maria (originale notificato).pdf.",
            data="2026-05-20T10:02:00",
            message_id="<esito-1428@giustizia>",
            allegati=[],
        ),
    ]

    assert office_notification_evidence_from_pec(fascicolo, ricevute) == []
    assert released_office_documents_from_pec(fascicolo, ricevute) == []


def test_monitor_documenti_ufficio_deduplica_copie_pec_ma_non_documenti_distinti():
    fascicolo = SimpleNamespace(
        id="fasc-duplicati",
        numero="FASC-DUP",
        titolo="Rossi c. Ministero",
        tribunale="Tribunale di Roma",
        numero_rg="1234",
        anno_rg=2026,
        documenti=[],
    )
    base = {
        "mittente": "cancelleria.tribunale.roma@giustiziacert.it",
        "oggetto": "R.G. 1234/2026 - ordinanza da notificare",
        "corpo_testo": "Si trasmette l'ordinanza da notificare.",
        "data": "2026-05-23T10:15:00",
        "message_id": "<ordinanza-1234@giustizia>",
        "allegati": [{"nome": "ordinanza.pdf", "sha256": "a" * 64}],
    }
    inbox = SimpleNamespace(id="pec-inbox", cartella="INBOX", **base)
    cestino = SimpleNamespace(id="pec-cestino", cartella="CESTINO", **base)

    evidence = office_notification_evidence_from_pec(fascicolo, [cestino, inbox])

    assert len(evidence) == 1
    assert evidence[0]["pecId"] == "pec-inbox"
    assert evidence[0]["pecSourceIds"] == ["pec-inbox", "pec-cestino"]

    distinto = SimpleNamespace(
        id="pec-distinto",
        cartella="INBOX",
        mittente=base["mittente"],
        oggetto=base["oggetto"],
        corpo_testo=base["corpo_testo"],
        data="2026-05-24T10:15:00",
        message_id="<ordinanza-distinta-1234@giustizia>",
        allegati=[{"nome": "ordinanza.pdf", "sha256": "b" * 64}],
    )
    evidence = office_notification_evidence_from_pec(fascicolo, [inbox, distinto])
    assert len(evidence) == 2
    assert {item["hashSha256"] for item in evidence} == {"a" * 64, "b" * 64}


def test_matrice_notifica_segnala_registro_incoerente_senza_bloccare_invio():
    payload = _legal_payload()
    payload.update({
        "ruolo_destinatario": "difensore",
        "destinatario_parte_rappresentata": "Controparte S.p.A.",
        "fonte_pec_destinatario": "registro_imprese",
        "caso_notifica": "appello_impugnazione",
        "provvedimento_tipo": "sentenza",
        "provvedimento_data": "2026-05-20",
    })

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert not any("PEC_DESTINATARIO_REGISTRO_INCOERENTE" in item for item in result.blockers)
    assert any("PEC_DESTINATARIO_REGISTRO_INCOERENTE" in item for item in result.warnings)


def test_matrice_notifica_caso_e_destinatario_generano_output_governato():
    payload = _legal_payload()
    payload.update({
        "ruolo_destinatario": "difensore",
        "destinatario_parte_rappresentata": "Controparte S.p.A.",
        "fonte_pec_destinatario": "reginde",
        "caso_notifica": "appello_impugnazione",
        "provvedimento_tipo": "sentenza",
        "provvedimento_numero": "101",
        "provvedimento_anno": "2026",
        "provvedimento_ufficio_origine": "Tribunale di Roma",
        "provvedimento_data": "2026-05-20",
        "provvedimento_data_rilascio": "2026-05-20",
    })
    payload["verifiche_pec_destinatari"] = [_pec_evidence(
        "reginde", "controparte@example.pec.it", "01234567890", "2026-05-12T10:30:01+02:00",
    )]

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert result.template_id == "relata_appello_impugnazione"
    directive = result.output_plan["notificationDirective"]
    assert directive["role"] == "difensore"
    assert directive["caseId"] == "appello_impugnazione"
    assert "reginde" in directive["allowedRegisters"]
    assert "difensore" in directive["allowedRecipientRoles"]
    assert directive["recipientRule"]
    assert directive["caseLegalBasis"]
    delivery = result.output_plan["deliveryPlan"]
    signature = result.output_plan["signaturePlan"]
    assert delivery["ready"] is True
    assert delivery["legalSubject"] == LEGAL_NOTIFICATION_SUBJECT
    assert delivery["studioTelematicoSubject"].startswith("Notificazione ai sensi della legge n. 53/1994")
    assert "[Notifica_ID:" in delivery["studioTelematicoSubject"]
    assert "expectedReceiptSubjects" not in delivery
    assert "receiptCorrelation" not in delivery
    assert delivery["recipients"][0]["role"] == "difensore"
    assert any(item["id"] == "relata_firmata" for item in delivery["attachments"])
    assert signature["requiredBeforeSend"][0]["id"] == "relata_notifica"
    assert signature["requiredBeforeSend"][0]["sourceFile"] == "relata_notifica.pdf"
    assert signature["requiredBeforeSend"][0]["signedFile"] == "relata_notifica.pdf.p7m"
    assert delivery["signaturePlan"]["requiredBeforeSend"][0]["id"] == "relata_notifica"
    assert delivery["sendPhase"] == "preparazione"
    assert not any(item["id"] == "allegati_pec" for item in delivery["sendChecks"])
    assert delivery["postSendDocumentArchive"]
    assert delivery["postSendDocumentArchive"][0]["archiveFilename"].endswith("(originale notificato).pdf")
    assert any(item["archiveFilename"] == "relata_notifica.pdf.p7m" for item in delivery["postSendDocumentArchive"])
    assert delivery["presidioPecAutomation"]["archiveTargets"] == ["fascicolo", "presidi_notifiche"]
    assert delivery["presidioPecAutomation"]["enabled"] is False
    assert delivery["localSendOnly"] is True
    assert delivery["presidioPecAutomation"]["localSendOnly"] is True


def test_relata_non_qualifica_pa_come_difensore_di_parte_rappresentata():
    payload = _legal_payload()
    payload.update({
        "ruolo_destinatario": "pa",
        "destinatario_nome": "Ministero dell'Istruzione e del Merito",
        "destinatario_cf": "80185250588",
        "destinatario_pec": "dgosv@postacert.istruzione.it",
        "fonte_pec_destinatario": "registro_ppaa",
        "destinatario_parte_rappresentata": "BANCA BETA S.P.A.",
    })

    result = validate_legal_notification(payload)

    assert "Ministero dell'Istruzione e del Merito" in result.relata_text
    assert "difensore di BANCA BETA S.P.A." not in result.relata_text


def test_piano_invio_studio_telematico_prepara_unico_to_e_allegati_reali():
    payload = _legal_payload()
    payload.update({
        "destinatari": [
            {
                "nome": "Controparte S.p.A.",
                "pec": "controparte@example.pec.it",
                "ruolo": "controparte",
                "fonte_pec": "registro_imprese",
            },
            {
                "nome": "Avv. Laura Bianchi",
                "pec": "laura.bianchi@example.pec.it",
                "ruolo": "difensore",
                "fonte_pec": "reginde",
                "parte_rappresentata": "Controparte S.p.A.",
            },
        ],
        "documenti": [
            {"nome_file": "ordinanza.pdf", "descrizione": "Ordinanza da notificare", "origine": "nativo_digitale", "hash_sha256": "a" * 64}
        ],
    })

    plan = build_notification_send_plan(payload)

    assert plan["ready"] is True
    assert plan["separatePecRequired"] is False
    assert plan["singleMessageToAllRecipients"] is True
    assert plan["messagesCount"] == 1
    assert plan["localSendOnly"] is True
    assert {item["pec"] for item in plan["recipients"]} == {"controparte@example.pec.it", "laura.bianchi@example.pec.it"}
    assert [item["filename"] for item in plan["attachments"]] == ["ordinanza.pdf", "relata_notifica.pdf.p7m"]
    assert [item["studioTelematicoArchiveRole"] for item in plan["attachments"]] == ["originale_notificato", "relata_notifica"]
    assert len(plan["messages"]) == 1
    message = plan["messages"][0]
    assert "[Notifica_ID:" in message["subject"]
    assert message["localSendOnly"] is True
    assert message["cc"] == []
    assert message["bcc"] == []
    assert "controparte@example.pec.it" in message["to"]
    assert "laura.bianchi@example.pec.it" in message["to"]
    assert "codice fiscale:" in message["to"]
    assert "pubblico elenco:" in message["to"]
    assert plan["studioTelematicoTo"] == message["to"]
    assert "postSendEvidenceRequired" not in plan
    assert "expectedReceiptSubjects" not in plan
    assert [item["archiveFilename"] for item in plan["postSendDocumentArchive"]] == [
        "ordinanza (originale notificato).pdf",
        "relata_notifica.pdf.p7m",
    ]
    assert plan["presidioPecAutomation"]["phase"] == "post_invio_reale"


def test_notifica_l53_preparazione_non_blocca_ricevute_o_approvazione_finale():
    payload = _legal_payload()
    payload.update({
        "ricevuta_completa": False,
        "relata_firmata": False,
        "approvazione_avvocato": False,
    })

    result = validate_legal_notification(payload, require_signed_relata=False)

    assert result.ok is True
    assert not any("RICEVUTA_COMPLETA_REQUIRED" in item for item in result.blockers)
    assert "ricevuta_accettazione.eml" not in result.output_plan["files"]
    assert "ricevuta_consegna_completa.eml" not in result.output_plan["files"]
    assert "eventuali_avvisi_errore.eml" not in result.output_plan["files"]
    assert "postSendFiles" not in result.output_plan
    delivery = result.output_plan["deliveryPlan"]
    assert delivery["sendPhase"] == "preparazione"
    send_checks = {item["id"]: item for item in delivery["sendChecks"]}
    assert "allegati_pec" not in send_checks
    assert "approvazione_avvocato" not in send_checks
    normative = {item["id"]: item for item in result.output_plan["normativeChecks"]}
    assert normative["allegati"]["status"] == "superato"
    assert "Relata separata" not in normative["allegati"]["detail"]
    assert "RAC" not in normative["allegati"]["detail"]
    assert "RdAC" not in normative["allegati"]["detail"]
    assert "Procura alle liti: non richiesta" in normative["allegati"]["detail"]
    assert "Attestazione di conformità: presente" in normative["allegati"]["detail"]
    assert "ricevuta_completa" not in normative


def test_notifica_l53_senza_documenti_non_inventa_allegato_generico():
    payload = _legal_payload()
    payload.update({
        "documenti": [],
        "nome_file": "",
        "descrizione_documento": "",
        "origine_documento": "copia_fascicolo_informatico",
        "attestazione_conformita": "Attestazione residua da ignorare se nessun documento e selezionato.",
        "relata_firmata": False,
    })

    preview = preview_legal_relata(payload)
    result = validate_legal_notification(payload, require_signed_relata=False)
    normative = {item["id"]: item for item in build_notification_normative_checks(payload)}

    assert preview["ok"] is True
    assert "documento allegato" not in preview["previewText"]
    assert "Attestazione di conform" not in preview["previewText"]
    assert "A) - Relata di notifica." in preview["previewText"]
    assert result.ok is True
    assert result.blockers == []
    assert result.warnings == ["Seleziona almeno un documento da notificare."]
    assert normative["allegati"]["status"] == "da completare"
    assert "Atto, provvedimento o documento da notificare: mancante" in normative["allegati"]["detail"]
    assert "Procura alle liti: non richiesta" in normative["allegati"]["detail"]
    assert "Attestazione di conformità: non richiesta" in normative["allegati"]["detail"]


def test_invio_finale_notifica_prepara_piano_senza_bloccare_firma_relata_e_approvazione_avvocato():
    payload = _legal_payload()
    payload.update({
        "operazione": "invio_pec_l53",
        "conferma_invio_pec": True,
        "invio_finale": True,
        "ricevuta_completa": False,
        "relata_firmata": False,
        "approvazione_avvocato": False,
    })

    result = validate_legal_notification(payload, require_signed_relata=True)

    assert result.ok is True
    assert not any("RELATA_FIRMATA_REQUIRED" in item for item in result.blockers)
    assert not any("approvazione finale dell'avvocato" in item for item in result.blockers)
    assert any("RELATA_FIRMATA_DA_COMPLETARE" in item for item in result.warnings)
    assert not any("Approvazione finale dell'avvocato" in item for item in result.warnings)
    assert not any("RICEVUTA_COMPLETA_REQUIRED" in item for item in result.blockers)
    delivery = result.output_plan["deliveryPlan"]
    assert delivery["sendPhase"] == "invio_finale"
    send_checks = {item["id"]: item for item in delivery["sendChecks"]}
    assert "allegati_pec" not in send_checks
    assert send_checks["documenti_notifica"]["status"] == "superato"
    assert send_checks["documenti_notifica"]["blocking"] is False
    assert send_checks["orario_notifica"]["status"] == "superato"
    assert send_checks["orario_notifica"]["blocking"] is False
    assert "automaticamente" in send_checks["orario_notifica"]["detail"]
    assert "Data del procedimento da confermare" not in send_checks["orario_notifica"]["detail"]
    assert "RAC" not in send_checks["orario_notifica"]["detail"]
    assert "RdAC" not in send_checks["orario_notifica"]["detail"]
    assert "approvazione_avvocato" not in send_checks
    signature_checks = {item["id"]: item for item in delivery["signaturePlan"]["checks"]}
    assert signature_checks["relata_firmata"]["status"] == "da completare"
    assert signature_checks["relata_firmata"]["blocking"] is False
    assert "inviabile" not in signature_checks["relata_firmata"]["detail"]


def test_relata_due_destinatari_usa_anteprima_testo_e_timestamp_individuali(
    legal_payload_due_destinatari: dict[str, object],
):
    preview = preview_legal_relata(legal_payload_due_destinatari)
    result = validate_legal_notification(legal_payload_due_destinatari)

    assert preview["ok"] is True
    assert result.ok is True
    for expected in (
        "Avvocatura Distrettuale dello Stato di Reggio Calabria",
        "ads.rc@mailcert.avvocaturastato.it",
        "Ministero dell'Istruzione e del Merito",
        "dgosv@postacert.istruzione.it",
    ):
        assert expected in preview["previewText"]
        assert expected in result.relata_text

    for rendered in (preview["previewText"], result.relata_text):
        avvocatura_start = rendered.index("Avvocatura Distrettuale dello Stato di Reggio Calabria")
        ministero_start = rendered.index("\n2) - Ministero dell'Istruzione e del Merito", avvocatura_start + 1)
        avvocatura_block = rendered[avvocatura_start:ministero_start]
        ministero_block = rendered[ministero_start:]
        assert "21/07/2026" in avvocatura_block
        assert "09:12" in avvocatura_block
        assert "11:47" not in avvocatura_block
        assert "21/07/2026" in ministero_block
        assert "11:47" in ministero_block


@pytest.mark.parametrize("difetto", ["pec_mancante", "registro_incoerente"])
def test_relata_due_destinatari_blocca_secondo_incompleto_o_registro_incoerente(
    legal_payload_due_destinatari: dict[str, object],
    difetto: str,
):
    secondo = legal_payload_due_destinatari["destinatari"][1]
    if difetto == "pec_mancante":
        secondo["pec"] = ""
    else:
        secondo["fonte_pec"] = "reginde"
        legal_payload_due_destinatari["verifiche_pec_destinatari"][1] = _pec_evidence(
            "reginde",
            secondo["pec"],
            secondo["codice_fiscale_piva"],
            "2026-07-21T11:47:59+02:00",
        )

    result = validate_legal_notification(legal_payload_due_destinatari)
    blockers = "\n".join(result.blockers)
    warnings = "\n".join(result.warnings)

    assert result.ok is True
    assert blockers == ""
    assert "Destinatario 2" in warnings
    if difetto == "pec_mancante":
        assert "PEC" in warnings
    else:
        assert "PEC_DESTINATARIO_REGISTRO_INCOERENTE" in warnings
        assert "Ministero dell'Istruzione e del Merito" in warnings


def test_override_relata_non_puo_eliminare_il_secondo_destinatario(
    legal_payload_due_destinatari: dict[str, object],
):
    canonical = validate_legal_notification(legal_payload_due_destinatari)
    assert canonical.ok is True
    legal_payload_due_destinatari["relata_override_text"] = canonical.relata_text.replace(
        "dgosv@postacert.istruzione.it",
        "",
    )

    result = validate_legal_notification(legal_payload_due_destinatari)

    assert result.ok is False
    assert any(
        "RELAZIONE_CONTENUTO_OBBLIGATORIO_REQUIRED" in blocker
        and "dgosv@postacert.istruzione.it" in blocker
        for blocker in result.blockers
    )


def test_modello_personalizzato_due_destinatari_richiede_elenco_completo(
    legal_payload_due_destinatari: dict[str, object],
):
    legal_payload_due_destinatari["template_id"] = "relata_personalizzata_due_destinatari"
    legal_payload_due_destinatari["template_personalizzato"] = {
        "id": "relata_personalizzata_due_destinatari",
        "label": "Relata personalizzata con destinatari multipli",
        "custom_body": "\n".join([
            "RELAZIONE DI NOTIFICAZIONE A MEZZO POSTA ELETTRONICA CERTIFICATA",
            "Avv. {{ avvocato.full_name }}, C.F. {{ avvocato.codice_fiscale }}, PEC {{ avvocato.pec }}",
            "per {{ cliente.nome_denominazione }}, C.F./P. IVA {{ cliente.codice_fiscale_piva }}",
            "{{ destinatari_righe }}",
            "{{ documenti_righe }}",
            "{{ blocco_procedimento }}",
            "{{ notifica.luogo }}, {{ notifica.data }} alle ore {{ notifica.ora }}",
        ]),
        "requires_proceeding": True,
    }

    preview = preview_legal_relata(legal_payload_due_destinatari)
    result = validate_legal_notification(legal_payload_due_destinatari)

    assert preview["ok"] is True
    assert result.ok is True
    assert "ads.rc@mailcert.avvocaturastato.it" in preview["previewText"]
    assert "dgosv@postacert.istruzione.it" in preview["previewText"]
    assert "ads.rc@mailcert.avvocaturastato.it" in result.relata_text
    assert "dgosv@postacert.istruzione.it" in result.relata_text

    legal_payload_due_destinatari["template_personalizzato"]["custom_body"] = (
        legal_payload_due_destinatari["template_personalizzato"]["custom_body"]
        .replace("{{ destinatari_righe }}", "a {{ destinatario.nome_denominazione }} presso {{ destinatario.pec }}")
    )
    blocked_preview = preview_legal_relata(legal_payload_due_destinatari)
    blocked_result = validate_legal_notification(legal_payload_due_destinatari)

    assert blocked_preview["ok"] is False
    assert blocked_result.ok is False
    assert any("MODELLO_DESTINATARI_MULTIPLI_REQUIRED" in item for item in blocked_preview["blockers"])
    assert any("MODELLO_DESTINATARI_MULTIPLI_REQUIRED" in item for item in blocked_result.blockers)


def test_log_audit_e_piano_invio_conservano_tutti_i_destinatari(
    legal_payload_due_destinatari: dict[str, object],
):
    result = validate_legal_notification(legal_payload_due_destinatari)

    assert result.ok is True
    expected_pecs = {
        "ads.rc@mailcert.avvocaturastato.it",
        "dgosv@postacert.istruzione.it",
    }
    assert len(result.log_json["destinatari"]) == 2
    assert {item["pec"] for item in result.log_json["destinatari"]} == expected_pecs

    audit = result.output_plan["auditTrail"]
    assert audit["recipientsCount"] == 2
    assert {item["pec"] for item in audit["recipients"]} == expected_pecs

    delivery = result.output_plan["deliveryPlan"]
    assert delivery["messagesCount"] == 1
    assert delivery["separatePecRequired"] is False
    assert delivery["singleMessageToAllRecipients"] is True
    assert {item["pec"] for item in delivery["recipients"]} == expected_pecs
    assert all(pec in delivery["studioTelematicoTo"] for pec in expected_pecs)
    assert all(item["recipientIdentityKey"] for item in delivery["recipients"])
    assert all(item["verificationEvidenceSha256"] for item in delivery["recipients"])

    assert result.log_json["destinatario"] == result.log_json["destinatari"][0]["nome"]
    assert result.log_json["pec_destinatario"] == result.log_json["destinatari"][0]["pec"]
    assert audit["recipient"] == audit["recipients"][0]["name"]
    assert audit["recipientPec"] == audit["recipients"][0]["pec"]


def test_relata_destinatario_singolo_preserva_contratto_legacy():
    payload = _legal_payload()
    preview = preview_legal_relata(payload)
    result = validate_legal_notification(payload)

    assert preview["ok"] is True
    assert result.ok is True
    assert "Controparte S.p.A." in preview["previewText"]
    assert "controparte@example.pec.it" in result.relata_text
    assert len(result.log_json["destinatari"]) == 1
    assert result.log_json["destinatario"] == "Controparte S.p.A."
    assert result.log_json["pec_destinatario"] == "controparte@example.pec.it"
    assert result.output_plan["auditTrail"]["recipientsCount"] == 1
    assert result.output_plan["deliveryPlan"]["messagesCount"] == 1


def test_piano_firma_seleziona_relata_e_non_rifirma_provvedimento_portale():
    payload = _legal_payload()
    payload.update({
        "documento_ufficio_rilasciato": True,
        "acquisizione_portale_richiesta": True,
        "pec_ufficio_eml_file": "pec-cancelleria.eml",
        "pec_ufficio_eml_sha256": "d" * 64,
        "documenti": [
            {
                "nome_file": "ordinanza_da_notificare.pdf",
                "descrizione": "Ordinanza rilasciata dall'ufficio",
                "origine": "copia_fascicolo_informatico",
                "fonte_documento": "PORTALE_TELEMATICO",
                "servizio_portale": "PST",
                "riferimento_portale": "PST-REL-2026-0001",
                "documento_ufficio": True,
                "acquisito_da_portale": True,
                "notifica_richiesta": True,
                "hash_sha256": "a" * 64,
            }
        ],
        "attestazione_multipla": True,
    })

    plan = build_notification_signature_plan(payload)

    assert plan["localSignerRequired"] is True
    assert plan["requiredBeforeSend"] == [
        {
            "id": "relata_notifica",
            "label": "Relata di notificazione",
            "sourceFile": "relata_notifica.pdf",
            "signedFile": "relata_notifica.pdf.p7m",
            "required": True,
            "phase": "prima_invio_pec",
            "format": "CADES",
            "signer": "Mario Rossi",
            "source": "L. 53/1994, art. 3-bis, comma 5; art. 56-bis disp. att. c.p.p. per il flusso penale",
            "reason": "La normativa richiede la relazione di notificazione su documento informatico separato, sottoscritta dall'avvocato prima dell'invio PEC.",
            "automaticSelection": True,
        }
    ]
    assert plan["notToSign"][0]["filename"] == "ordinanza_da_notificare.pdf"
    assert "relata firmata" in plan["notToSign"][0]["reason"]
    assert any(item["id"] == "relata_firmata" and item["status"] == "superato" for item in plan["checks"])


def test_piano_firma_include_solo_documento_marcato_come_atto_da_sottoscrivere():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "atto_principale.pdf",
            "descrizione": "Atto principale da notificare",
            "origine": "nativo_digitale",
            "firma_digitale_richiesta": True,
            "hash_sha256": "a" * 64,
        },
        {
            "nome_file": "provvedimento.pdf.p7m",
            "descrizione": "Provvedimento già firmato",
            "origine": "firmato_digitalmente",
            "hash_sha256": "b" * 64,
        },
    ]

    plan = build_notification_signature_plan(payload)

    assert [item["id"] for item in plan["requiredBeforeSend"]] == ["relata_notifica", "documento_1"]
    assert plan["requiredBeforeSend"][1]["signedFile"] == "atto_principale.pdf.p7m"
    assert plan["alreadySigned"] == [{"filename": "provvedimento.pdf.p7m", "reason": "Documento già firmato digitalmente o già in formato firmato."}]


def test_piano_orario_distingue_regime_corrente_e_storico():
    current = {"data_inizio_procedimento": "2026-01-10"}
    ordinario = build_notification_timing_plan({**current, "data_ora_invio_pec": "2026-05-24T10:40:57"})
    assert ordinario["ready"] is True
    assert ordinario["status"] == "fascia_ordinaria"
    assert ordinario["plannedAt"] == "24/05/2026 10:40:57"

    serale = build_notification_timing_plan({**current, "data_ora_invio_pec": "2026-05-24T21:30"})

    assert serale["ready"] is True
    assert serale["status"] == "fascia_con_differimento_destinatario"
    assert "RAC" in serale["senderEffect"]
    assert "07:00" in serale["recipientEffect"]
    assert any(item["id"] == "cpc_art147" for item in serale["legalBasis"])

    notturno = build_notification_timing_plan({**current, "data_ora_invio_pec": "2026-05-24T06:30"})

    assert notturno["ready"] is True
    assert notturno["status"] == "fascia_con_differimento_destinatario"


def test_area_web_pst_mancata_notifica_salva_art_3ter_e_avviso_eml():
    payload = {
        "pec_non_consegnata": True,
        "causa_imputabile_destinatario": True,
        "valutazione_avvocato": "Casella PEC satura risultante da avviso di mancata consegna.",
        "avviso_mancata_consegna": "mancata-consegna.eml",
        "atto_notificato": "atto.pdf",
        "atto_notificato_sha256": "a" * 64,
        "relata_firmata": "relata_notifica.pdf.p7m",
        "relata_sha256": "b" * 64,
        "pec_inviata": "notifica-inviata.eml",
        "pec_inviata_sha256": "c" * 64,
        "rac_file": "rac.eml",
        "rac_sha256": "d" * 64,
        "rdac_file": "rdac-completa.eml",
        "rdac_sha256": "e" * 64,
    }

    result = prepare_pst_failed_notification_workflow(payload)

    assert result.ok is True
    assert result.output_plan["legalBasis"][0]["id"] == "l53_art3ter"
    assert "Carica avviso di mancata consegna in formato EML" in result.output_plan["portalSteps"]


def test_matrice_notifica_segnala_destinatario_incoerente_con_caso_senza_bloccare():
    payload = _legal_payload()
    payload.update({
        "ruolo_destinatario": "controparte",
        "fonte_pec_destinatario": "registro_imprese",
        "caso_notifica": "chiamata_terzo",
        "provvedimento_data": "2026-05-20",
    })

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert not any("DESTINATARIO_CASO_INCOERENTE" in item for item in result.blockers)
    assert any("DESTINATARIO_CASO_INCOERENTE" in item for item in result.warnings)

    payload.update({
        "ruolo_destinatario": "terzo",
        "destinatario_nome": "Terzo Chiamato S.r.l.",
        "destinatario_qualifica": "terzo destinatario",
    })

    valid = validate_legal_notification(payload)

    assert valid.ok is True
    assert valid.output_plan["notificationDirective"]["caseId"] == "chiamata_terzo"


def test_matrice_notifica_esposta_per_ui_e_script():
    matrix = notification_directive_matrix()

    assert any(item["value"] == "difensore" and "reginde" in item["allowedRegisters"] for item in matrix["roles"])
    assert any(item["value"] == "sentenza_termine_breve" and item["templateId"] == "relata_sentenza_attestazione_conformita" for item in matrix["cases"])
    assert all(item["legalBasis"] for item in matrix["roles"])
    assert all(item["legalBasis"] and item["recipientRule"] and item["allowedRecipientRoles"] for item in matrix["cases"])
    assert any(item["id"] == "l53_art3ter" for item in matrix["roles"][0]["legalBasis"])
    provvedimento = next(item for item in matrix["cases"] if item["value"] == "provvedimento_giudice")
    assert any(item["id"] == "dgsia_2024_art22" for item in provvedimento["legalBasis"])
    sentenza = next(item for item in matrix["cases"] if item["value"] == "sentenza_termine_breve")
    assert any(item["id"] == "cpc_326" for item in sentenza["legalBasis"])
    assert any(item["id"] == "disp_att_cpc_196octies" for item in sentenza["legalBasis"])


def test_modello_sentenza_attestazione_esposto_con_campi_autocompilanti():
    template = get_notification_template("relata_sentenza_attestazione_conformita")
    tokens = {item["token"] for item in available_template_fields()}

    assert template is not None
    assert template["label"] == "Sentenza con attestazione di conformità"
    assert "avvocato.full_name" in template["required_fields"]
    assert "provvedimento.data_rilascio" in template["required_fields"]
    assert "{{ avvocato.firma_in_calce }}" in tokens
    assert "{{ avvocato.firma_digitale_dicitura }}" in tokens
    assert "{{ provvedimento.data_rilascio }}" in tokens


def test_allegati_notifica_ed_eml_pec_ufficio_sono_controllati():
    payload = _legal_payload()
    payload.update({
        "caso_notifica": "provvedimento_giudice",
        "provvedimento_data": "2026-05-23",
        "documento_ufficio_rilasciato": True,
        "pec_ufficio_rilascio": True,
        "pec_ufficio_eml_file": "pec-cancelleria.eml",
        "pec_ufficio_eml_sha256": "a" * 64,
        "documenti": [
            {
                "nome_file": "ordinanza.pdf",
                "descrizione": "Ordinanza comunicata dalla cancelleria",
                "origine": "comunicazione_cancelleria",
                "data_comunicazione_cancelleria": "2026-05-23",
                "hash_sha256": "b" * 64,
                "documento_ufficio": True,
                "notifica_richiesta": True,
                "acquisito_da_portale": True,
                "riferimento_portale": "pst-doc-1",
            }
        ],
        "attestazione_multipla": True,
    })

    manifest = build_notification_attachment_manifest(payload)
    result = validate_legal_notification(payload)

    assert result.ok is True
    assert all(item["present"] for item in manifest if item["required"])
    assert any(item["id"] == "eml_ufficio" and item["present"] for item in manifest)
    assert any(item["id"] == "eml_pec_ufficio" and item["status"] == "superato" for item in build_notification_normative_checks(payload))

    payload.pop("pec_ufficio_eml_file")
    payload.pop("pec_ufficio_eml_sha256")
    blocked = validate_legal_notification(payload)

    assert blocked.ok is True
    assert not any("PEC_UFFICIO_EML_REQUIRED" in item for item in blocked.blockers)
    assert any("PEC_UFFICIO_EML_REQUIRED" in item for item in blocked.warnings)


def test_notifica_l53_controllo_attestazioni_non_basta_per_un_solo_allegato():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "provvedimento.pdf",
            "descrizione": "Provvedimento",
            "origine": "copia_fascicolo_informatico",
            "attestazione_conformita_presente": True,
        },
        {"nome_file": "scansione.pdf", "descrizione": "Scansione", "origine": "scansione_analogico"},
    ]
    payload["attestazione_conformita"] = ""

    checks = build_notification_normative_checks(payload)

    assert any(item["id"] == "attestazioni" and item["status"] == "da completare" for item in checks)


def test_notifica_l53_modello_personalizzato_usa_campi_iusentra_e_note_avvocato():
    payload = _legal_payload()
    payload["luogo"] = "TAURIANOVA RC"
    payload["data_relata"] = "2026-05-14"
    payload["ora_relata"] = "16:40"
    payload["template_id"] = "relata_personalizzata_prova"
    payload["template_personalizzato"] = {
        "id": "relata_personalizzata_prova",
        "label": "Relata su misura",
        "custom_body": "\n".join([
            "RELAZIONE PERSONALIZZATA",
            "Avv. {{ avvocato.full_name }} per {{ cliente.nome_denominazione }}",
            "Destinatario: {{ destinatario.nome_denominazione }} - {{ destinatario.pec }}",
            "{{ documenti_righe }}",
            "{{ blocco_procedimento }}",
            "{{ attestazioni_testo }}",
            "{{ notifica.luogo }}, {{ notifica.data }} alle ore {{ notifica.ora }}",
        ]),
        "requires_proceeding": True,
    }
    payload["note_integrative_relata"] = "Precisazione finale aggiunta dall'avvocato."

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert "RELAZIONE PERSONALIZZATA" in result.relata_text
    assert "Avv. Mario Rossi per Cliente S.r.l." in result.relata_text
    assert "A) - Ricorso notificato (File: ricorso.pdf)" in result.relata_text
    assert "RG: 1234/2026" in result.relata_text
    assert "INTEGRAZIONE DELL'AVVOCATO" in result.relata_text
    assert "Precisazione finale aggiunta dall'avvocato." in result.relata_text
    assert "TAURIANOVA RC, 14/05/2026 alle ore 16:40" in result.relata_text
    assert "TAURIANOVA RC, 2026-05-14" not in result.relata_text


def test_modello_personalizzato_blocca_token_sconosciuto():
    payload = _legal_payload()
    payload["template_id"] = "relata_personalizzata_non_valida"
    payload["template_personalizzato"] = {
        "id": "relata_personalizzata_non_valida",
        "label": "Relata non valida",
        "custom_body": "Avv. {{ avvocato.full_name }} - {{ segreto.interno }}",
    }

    result = validate_legal_notification(payload)

    assert result.ok is False
    assert any("Campo automatico non consentito" in item for item in result.blockers)
    assert result.relata_text == ""


def test_modello_personalizzato_blocca_blocchi_jinja():
    payload = _legal_payload()
    payload["template_id"] = "relata_personalizzata_if"
    payload["template_personalizzato"] = {
        "id": "relata_personalizzata_if",
        "label": "Relata con istruzioni",
        "custom_body": "{% if avvocato %}Avv. {{ avvocato.full_name }}{% endif %}",
    }

    result = validate_legal_notification(payload)

    assert result.ok is False
    assert any("istruzioni Jinja" in item for item in result.blockers)


def test_modello_personalizzato_blocca_accesso_pericoloso():
    payload = _legal_payload()
    payload["template_id"] = "relata_personalizzata_globals"
    payload["template_personalizzato"] = {
        "id": "relata_personalizzata_globals",
        "label": "Relata pericolosa",
        "custom_body": "Accesso {{ cycler.__init__.__globals__ }}",
    }

    result = validate_legal_notification(payload)

    assert result.ok is False
    assert any("accesso riservato" in item or "non consentito" in item for item in result.blockers)


def test_modelli_standard_restano_renderizzabili():
    payload = _legal_payload()
    payload["template_id"] = "relata_pec_base_l53"

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert "RELATA DI NOTIFICA" in result.relata_text


def test_anteprima_relata_compilata_con_placeholder():
    full = preview_legal_relata(_legal_payload())
    missing_payload = _legal_payload()
    missing_payload["destinatario_pec"] = ""
    missing = preview_legal_relata(missing_payload)

    assert full["ok"] is True
    assert "Cliente S.r.l." in full["previewText"]
    assert "Via Roma 1, CAP 00100" in full["previewText"]
    assert "Roma (RM)" in full["previewText"]
    assert missing["ok"] is True
    assert "[dato mancante: PEC destinatario]" in missing["previewText"]
    assert "PEC destinatario" in missing["missingFields"]


def test_anteprima_non_segnala_attestazioni_automatiche_quando_non_servono():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "ricorso.pdf.p7m",
            "descrizione": "Ricorso firmato digitalmente",
            "origine": "firmato_digitalmente",
        }
    ]

    preview = preview_legal_relata(payload)

    assert preview["ok"] is True
    assert "[dato mancante: Attestazioni automatiche]" not in preview["previewText"]
    assert "Attestazioni automatiche" not in preview["missingFields"]


def test_anteprima_genera_attestazione_dell_avvocato_per_copia_da_fascicolo():
    payload = _legal_payload()
    payload["documenti"] = [
        {
            "nome_file": "sentenza.pdf",
            "descrizione": "Sentenza estratta dal fascicolo informatico",
            "origine": "copia_fascicolo_informatico",
        }
    ]

    preview = preview_legal_relata(payload)

    assert preview["ok"] is True
    assert "Attesto, ai sensi della normativa vigente" in preview["previewText"]
    assert "Sentenza estratta dal fascicolo informatico" in preview["previewText"]
    assert "[dato mancante: Attestazioni automatiche]" not in preview["previewText"]


def test_anteprima_standard_ricalcola_token_operativi_dai_documenti_correnti():
    payload = _legal_payload()
    payload.update({
        "template_id": "relata_pec_a_difensore_costituito",
        "ruolo_destinatario": "difensore",
        "destinatario_parte_rappresentata": "",
        "attestazioni_testo": (
            "Attesto, ai sensi della normativa vigente, che la seguente copia informatica:\n"
            "- Decreto, emesso dall'ufficio giudiziario indicato nel fascicolo;\n"
            "è conforme alla copia informatica presente nel fascicolo informatico dal quale è estratta."
        ),
        "documenti": [
            {
                "nome_file": "ricorso-opposizione-con-decreto-fissazione.pdf",
                "descrizione": (
                    "Ricorso in opposizione a decreto ingiuntivo con decreto di fissazione "
                    "udienza estratto dal fascicolo informatico"
                ),
                "origine": "copia_fascicolo_informatico",
            },
            {
                "nome_file": "procura-alle-liti.pdf",
                "descrizione": "Procura alle liti rilasciata su documento informatico separato",
                "origine": "originale_informatico",
            },
        ],
    })

    preview = preview_legal_relata(payload)

    assert preview["ok"] is True
    assert "Ricorso in opposizione a decreto ingiuntivo con decreto di fissazione udienza" in preview["previewText"]
    assert "Procura alle liti rilasciata su documento informatico separato" in preview["previewText"]
    assert "- Decreto, emesso dall'ufficio giudiziario indicato nel fascicolo" not in preview["previewText"]
    assert "[dato mancante: Parte rappresentata]" not in preview["previewText"]


def test_anteprima_modelli_standard_catalogo_non_bloccata():
    for template in list_notification_templates(kind="relata"):
        payload = _legal_payload()
        payload["template_id"] = template["id"]
        preview = preview_legal_relata(payload)
        assert preview["ok"] is True, template["id"]


def test_comunicazione_cliente_non_usa_relata_o_oggetto_l53():
    blocked = build_client_communication({
        "operazione": "comunicazione_cliente_non_notifica",
        "cliente_nome": "Cliente",
        "oggetto": LEGAL_NOTIFICATION_SUBJECT,
        "genera_relata": True,
    })
    ok = build_client_communication({
        "operazione": "comunicazione_cliente_non_notifica",
        "cliente_nome": "Cliente",
        "ufficio_giudiziario": "Tribunale di Roma",
        "numero_rg": "1234",
        "anno_rg": "2026",
        "provvedimento_descrizione": "Sentenza depositata",
    })

    assert blocked.ok is False
    assert any("non deve usare l'oggetto" in item for item in blocked.blockers)
    assert any("non genera una relata" in item for item in blocked.blockers)
    assert ok.ok is True
    assert ok.relata_text == ""
    assert ok.subject == "Aggiornamento pratica"
    assert ok.template_version == client_communication_templates_version()


def test_comunicazione_cliente_usa_modelli_separati():
    templates = list_client_communication_templates()
    result = build_client_communication({
        "operazione": "comunicazione_cliente_non_notifica",
        "template_id": "richiesta_documenti",
        "cliente_nome": "Cliente",
        "pratica_codice": "2026/001",
        "provvedimento_descrizione": "Documenti reddituali",
    })

    assert {item["id"] for item in templates} >= {"aggiornamento_pratica", "esito_notifica", "richiesta_documenti"}
    assert result.ok is True
    assert result.relata_text == ""
    assert result.template_id == "richiesta_documenti"
    assert "Richiesta documenti" in result.subject
    assert result.template_version != template_catalog_version()


def test_comunicazione_cliente_blocca_catalogo_relata_l53():
    result = build_client_communication({
        "operazione": "comunicazione_cliente_non_notifica",
        "template_id": "relata_pec_base_l53",
        "cliente_nome": "Cliente",
        "provvedimento_descrizione": "Provvedimento",
    })

    assert result.ok is False
    assert any("modello comunicazione cliente" in item for item in result.blockers)


def test_prova_deposito_richiede_rac_rdac_originali():
    blocked = validate_deposit_notification_proof({
        "atto_notificato": "ricorso.pdf",
        "atto_sha256": "a" * 64,
        "relata_firmata": "relata.pdf.p7m",
        "relata_sha256": "b" * 64,
        "pec_inviata": "pec_inviata.eml",
        "pec_inviata_sha256": "c" * 64,
        "destinatario_nome": "Controparte",
        "destinatario_cf": "01234567890",
        "destinatario_pec": "controparte@example.pec.it",
        "fonte_pec_destinatario": "RegistroImprese",
        "rac_file": "accettazione.pdf",
        "rac_sha256": "d" * 64,
        "rdac_file": "consegna.eml",
        "rdac_sha256": "e" * 64,
        "ricevuta_completa": True,
    })
    ok = validate_deposit_notification_proof({
        "atto_notificato": "ricorso.pdf",
        "atto_sha256": "a" * 64,
        "relata_firmata": "relata.pdf.p7m",
        "relata_sha256": "b" * 64,
        "pec_inviata": "pec_inviata.eml",
        "pec_inviata_sha256": "c" * 64,
        "destinatario_nome": "Controparte",
        "destinatario_cf": "01234567890",
        "destinatario_pec": "controparte@example.pec.it",
        "fonte_pec_destinatario": "RegistroImprese",
        "rac_file": "accettazione.eml",
        "rac_sha256": "d" * 64,
        "rdac_file": "consegna.eml",
        "rdac_sha256": "e" * 64,
        "ricevuta_completa": True,
        "dati_atto_ricevute": "RAC e RdAC indicizzate",
    })

    assert blocked.ok is False
    assert any("originale digitale .eml o .msg" in item for item in blocked.blockers)
    assert ok.ok is True
    assert ok.output_plan["workflowSteps"][0]["id"] == "pacchetto_prova_deposito"
    assert any(item["id"] == "atti" for item in ok.output_plan["workflowSteps"])
    assert any(item["id"] == "rac_rdac" and item["status"] == "superato" for item in ok.output_plan["normativeChecks"])


def test_prova_deposito_accetta_piu_atti_notificati_con_hash():
    result = validate_deposit_notification_proof({
        "atti_notificati": [
            {"nome_file": "pst:JPW_SIGP:2182464 - ricorso.pdf", "hash_sha256": "a" * 64},
            {"nome_file": "procura.pdf", "hash_sha256": "f" * 64},
        ],
        "relata_firmata": "relata_notifica.pdf.p7m",
        "relata_sha256": "b" * 64,
        "pec_inviata": "pec_inviata.eml",
        "pec_inviata_sha256": "c" * 64,
        "destinatario_nome": "Controparte",
        "destinatario_cf": "01234567890",
        "destinatario_pec": "controparte@example.pec.it",
        "fonte_pec_destinatario": "RegistroImprese",
        "rac_file": "accettazione.eml",
        "rac_sha256": "d" * 64,
        "rdac_file": "consegna.eml",
        "rdac_sha256": "e" * 64,
        "ricevuta_completa": True,
        "dati_atto_ricevute": "RAC e RdAC indicizzate",
    })

    assert result.ok is True
    items = result.output_plan["evidencePack"]["items"]
    assert any(item["kind"] == "atto" and "pst:JPW_SIGP:2182464" in item["filename"] for item in items)
    assert any(item["kind"] == "allegato_2" and item["filename"] == "procura.pdf" for item in items)
    assert result.output_plan["auditTrail"]["documentsCount"] == 2


def test_prova_deposito_richiede_metadati_destinatario_studio_telematico():
    result = validate_deposit_notification_proof({
        "atto_notificato": "ricorso.pdf",
        "atto_sha256": "a" * 64,
        "relata_firmata": "relata_notifica.pdf.p7m",
        "relata_sha256": "b" * 64,
        "pec_inviata": "pec_inviata.eml",
        "pec_inviata_sha256": "c" * 64,
        "destinatario_nome": "Controparte",
        "rac_file": "accettazione.eml",
        "rac_sha256": "d" * 64,
        "rdac_file": "consegna.eml",
        "rdac_sha256": "e" * 64,
        "ricevuta_completa": True,
        "dati_atto_ricevute": "RAC e RdAC indicizzate",
    })

    assert result.ok is False
    assert any("DESTINATARIO_CF_REQUIRED" in item for item in result.blockers)
    assert any("DESTINATARIO_PEC_REQUIRED" in item for item in result.blockers)
    assert any("DESTINATARIO_FONTE_PEC_REQUIRED" in item for item in result.blockers)


def test_unep_richiede_canale_ufficio_destinatario_e_documenti():
    result = validate_unep_notification_request({
        "operazione": "notifica_unep",
        "tipo_richiesta_unep": "notifica_civile_pagamento",
        "tipo_notifica_unep": "telematica",
        "destinatario_nome": "Controparte",
        "atto_notificare": "atto.pdf",
        "atto_sha256": "a" * 64,
    })

    assert result.ok is False
    assert any("UFFICIO_UNEP_REQUIRED" in item for item in result.blockers)
    assert any("DESTINATARIO_PEC_REQUIRED" in item for item in result.blockers)
    assert any("DESTINATARIO_FONTE_PEC_REQUIRED" in item for item in result.blockers)
    assert any("richiesta o relata" in item for item in result.blockers)


def test_unep_telematica_con_precetto_e_pagamento_produce_piano():
    result = validate_unep_notification_request({
        "operazione": "notifica_unep",
        "tipo_richiesta_unep": "notifica_civile_pagamento",
        "tipo_notifica_unep": "telematica",
        "ufficio_unep": "UNEP - Corte d'Appello - Milano",
        "ufficio_unep_codice": "1514600637",
        "ufficio_unep_pec": "unep.ca.milano@civile.ptel.giustiziacert.it",
        "atto_notificare": "atto.pdf",
        "atto_sha256": "a" * 64,
        "richiesta_o_relata": "richiesta_unep.pdf",
        "richiesta_sha256": "b" * 64,
        "destinatario_nome": "Controparte",
        "destinatario_cf": "RSSMRA80A01H501U",
        "destinatario_pec": "controparte@example.pec.it",
        "fonte_pec_destinatario": "RegInde",
        "precetto_gia_notificato": True,
        "data_notifica_precetto": "2026-07-01",
        "spese_unep_dovute": True,
        "ricevuta_pagamento": "pagamento.pdf",
        "ricevuta_pagamento_sha256": "c" * 64,
    })

    assert result.ok is True
    assert result.template_id == "workflow_unep_notifica"
    assert result.output_plan["unepRequest"]["tipo"] == "telematica"
    assert result.output_plan["unepRequest"]["schema"] == "Atti_UNEP::AttoCivileAPagamento"
    assert result.output_plan["unepRequest"]["ufficioCodice"] == "1514600637"
    assert any(item["id"] == "recapito_destinatario" and item["status"] == "superato" for item in result.output_plan["normativeChecks"])


def test_unep_copre_tutti_i_rami_ministeriali_e_blocca_ufficio_non_coerente():
    expected_schemas = {
        "Atti_UNEP::AttoCivileAPagamento",
        "Atti_UNEP::AttoPenaleAPagamento",
        "Atti_UNEP::AttoCivileDebito",
        "Atti_UNEP::AttoPenaleDebito",
        "Atti_UNEP::AttoEsenteLavoro",
        "Atti_UNEP::PagamentoRichiestaNotifica",
        "Atti_UNEP::RichiestaPignoramentoMobiliare",
        "Atti_UNEP::RichiestaPignoramentoMobiliareADebito",
        "Atti_UNEP::RichiestaPignoramentoMobiliareMateriaLavoro",
        "Atti_UNEP::RichiestaPignoramentoImmobiliare",
        "Atti_UNEP::RichiestaPignoramentoImmobiliareADebito",
        "Atti_UNEP::RichiestaPignoramentoImmobiliareMateriaLavoro",
        "Atti_UNEP::RichiestaPignoramentoPressoTerzi",
        "Atti_UNEP::RichiestaPignoramentoPressoTerziADebito",
        "Atti_UNEP::RichiestaPignoramentoPressoTerziMateriaLavoro",
        "Atti_UNEP::PagamentoRichiestaPignoramento",
        "Atti_UNEP::RichiestaRicercaBeni",
        "Atti_UNEP::RichiestaRestituzioneSomme",
    }
    assert {item["schema"] for item in UNEP_REQUEST_TYPES.values()} == expected_schemas

    result = validate_unep_notification_request({
        "operazione": "notifica_unep",
        "tipo_richiesta_unep": "notifica_civile_pagamento",
        "tipo_notifica_unep": "mani",
        "ufficio_unep": "UNEP - Corte d'Appello - Milano",
        "ufficio_unep_codice": "1514600637",
        "ufficio_unep_pec": "unep.errata@example.pec.it",
        "destinatario_nome": "Controparte",
        "destinatario_indirizzo": "Via Roma 1",
        "destinatario_comune": "Milano",
        "atto_notificare": "atto.pdf",
        "atto_sha256": "a" * 64,
        "richiesta_o_relata": "richiesta_unep.pdf",
        "richiesta_sha256": "b" * 64,
    })

    assert result.ok is False
    assert any("UFFICIO_UNEP_REQUIRED" in item for item in result.blockers)


def test_notifica_non_pec_blocca_campi_tavola_mancanti():
    result = validate_non_pec_notification_tracking({
        "operazione": "notifica_non_pec",
        "tipo_notifica_non_pec": "raccomandata",
        "destinatario_nome": "Controparte",
    })

    assert result.ok is False
    assert any("DATA_NOTIFICA_REQUIRED" in item for item in result.blockers)
    assert any("NOTIFICA_ID_REQUIRED" in item for item in result.blockers)
    assert any("ATTO_NOTIFICATO_REQUIRED" in item for item in result.blockers)
    assert any("RACCOMANDATA_NUMERO_REQUIRED" in item for item in result.blockers)


def test_notifica_non_pec_raccomandata_traccia_data_tipo_ricevuta_id():
    result = validate_non_pec_notification_tracking({
        "operazione": "notifica_non_pec",
        "tipo_notifica_non_pec": "raccomandata",
        "data_notifica": "2026-07-03",
        "notifica_id": "N-2026-15",
        "destinatario_nome": "Controparte",
        "atto_notificato": "diffida.pdf",
        "numero_raccomandata": "1234567890",
        "data_spedizione": "2026-07-01",
        "data_ricevuta_raccomandata": "2026-07-03",
        "prova_file": "avviso_ricevimento.pdf",
        "prova_sha256": "d" * 64,
    })

    assert result.ok is True
    assert result.output_plan["historicalFields"] == {
        "DataNotifica": "2026-07-03",
        "TipoNotifica": "Raccomandata",
        "DataRicevutaRaccomandata": "2026-07-03",
        "NotificaID": "N-2026-15",
    }
    assert result.output_plan["workflowSteps"][0]["id"] == "nonpec_tipo"


def test_prova_deposito_blocca_hash_non_sha256_e_dati_atto_mancanti():
    result = validate_deposit_notification_proof({
        "atto_notificato": "ricorso.pdf",
        "atto_sha256": "abc",
        "relata_firmata": "relata_notifica.pdf.p7m",
        "relata_sha256": "b" * 64,
        "pec_inviata": "pec_inviata.eml",
        "pec_inviata_sha256": "c" * 64,
        "destinatario_nome": "Controparte",
        "destinatario_cf": "01234567890",
        "destinatario_pec": "controparte@example.pec.it",
        "fonte_pec_destinatario": "RegistroImprese",
        "rac_file": "accettazione.eml",
        "rac_sha256": "d" * 64,
        "rdac_file": "consegna.eml",
        "rdac_sha256": "e" * 64,
        "ricevuta_completa": True,
    })

    assert result.ok is False
    assert any("HASH_SHA256_INVALID" in item and "Atto notificato" in item for item in result.blockers)
    assert any("DATI_ATTO_RICEVUTE_REQUIRED" in item for item in result.blockers)


def test_api_react_notifiche_legali_espone_workflow_separati(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}

    payload_response = client.get("/api/v1/ui/notifiche-legali", headers=headers)
    invalid_response = client.post(
        "/api/v1/ui/notifiche-legali/notifica",
        json={"ruolo_destinatario": "cliente", "oggetto_pec": LEGAL_NOTIFICATION_SUBJECT},
        headers=headers,
    )
    valid_response = client.post(
        "/api/v1/ui/notifiche-legali/notifica",
        json=_legal_payload(),
        headers=headers,
    )
    attestation_response = client.post(
        "/api/v1/ui/notifiche-legali/attestazione-conformita",
        json=_legal_payload(),
        headers=headers,
    )
    send_payload = _legal_payload()
    send_payload["operazione"] = "invio_pec_l53"
    send_response = client.post(
        "/api/v1/ui/notifiche-legali/notifica",
        json=send_payload,
        headers=headers,
    )
    unverified_send_payload = _legal_payload()
    unverified_send_payload["operazione"] = "invio_pec_l53"
    unverified_send_payload["verifica_pec_mittente"] = {}
    unverified_send_payload["verifiche_pec_destinatari"] = []
    unverified_send_response = client.post(
        "/api/v1/ui/notifiche-legali/notifica",
        json=unverified_send_payload,
        headers=headers,
    )
    client_response = client.post(
        "/api/v1/ui/notifiche-legali/comunicazione-cliente",
        json={"operazione": "comunicazione_cliente_non_notifica", "cliente_nome": "Cliente", "provvedimento_descrizione": "Provvedimento depositato"},
        headers=headers,
    )
    unep_response = client.post(
        "/api/v1/ui/notifiche-legali/unep",
        json={
            "operazione": "notifica_unep",
            "tipo_richiesta_unep": "notifica_civile_pagamento",
            "tipo_notifica_unep": "telematica",
            "ufficio_unep": "UNEP - Corte d'Appello - Milano",
            "ufficio_unep_codice": "1514600637",
            "ufficio_unep_pec": "unep.ca.milano@civile.ptel.giustiziacert.it",
            "atto_notificare": "atto.pdf",
            "atto_sha256": "a" * 64,
            "richiesta_o_relata": "richiesta_unep.pdf",
            "richiesta_sha256": "b" * 64,
            "destinatario_nome": "Controparte",
            "destinatario_pec": "controparte@example.pec.it",
            "fonte_pec_destinatario": "RegInde",
        },
        headers=headers,
    )
    non_pec_response = client.post(
        "/api/v1/ui/notifiche-legali/non-pec",
        json={
            "operazione": "notifica_non_pec",
            "tipo_notifica_non_pec": "raccomandata",
            "data_notifica": "2026-07-03",
            "notifica_id": "N-2026-15",
            "destinatario_nome": "Controparte",
            "atto_notificato": "diffida.pdf",
            "numero_raccomandata": "1234567890",
            "data_spedizione": "2026-07-01",
            "data_ricevuta_raccomandata": "2026-07-03",
            "prova_file": "avviso_ricevimento.pdf",
            "prova_sha256": "d" * 64,
        },
        headers=headers,
    )
    non_pec_blocked_response = client.post(
        "/api/v1/ui/notifiche-legali/non-pec",
        json={"operazione": "notifica_non_pec", "tipo_notifica_non_pec": "raccomandata"},
        headers=headers,
    )

    payload = payload_response.get_json()
    invalid_payload = invalid_response.get_json()
    valid_payload = valid_response.get_json()
    send_result = send_response.get_json()
    unverified_send_result = unverified_send_response.get_json()
    client_payload = client_response.get_json()
    unep_payload = unep_response.get_json()
    non_pec_payload = non_pec_response.get_json()
    non_pec_blocked_payload = non_pec_blocked_response.get_json()

    assert payload_response.status_code == 200
    assert payload["mandatorySubject"] == LEGAL_NOTIFICATION_SUBJECT
    assert payload["contracts"]["clientCommunicationWithoutRelata"] is True
    assert payload["modelliRelata"][0]["previewText"]
    assert payload["automazioneGuidata"]["notifica"][0]["source"].startswith("L. 53/1994")
    assert all(item["id"] != "prova" for item in payload["automazioneGuidata"]["notifica"])
    assert not any(item["title"] == "Pacchetto prova e deposito" for item in payload["automazioneGuidata"]["notifica"])
    assert "deposito" not in payload["automazioneGuidata"]
    assert "depositProofWithOriginalReceipts" not in payload["contracts"]
    assert payload["automazioneGuidata"]["unep"][0]["id"] == "unep_canale"
    assert payload["automazioneGuidata"]["nonPec"][0]["id"] == "nonpec_tipo"
    assert any(item["id"] == "eml_ufficio" for item in payload["automazioneGuidata"]["allegati"])
    assert any(item["value"] == "telematica" for item in payload["tipiNotificaUnep"])
    assert any(
        item["value"] == "ricerca_beni"
        and item["schema"] == "Atti_UNEP::RichiestaRicercaBeni"
        for item in payload["tipiRichiestaUnep"]
    )
    pst_snapshot = json.loads((Path(__file__).parents[1] / "pct" / "data" / "uffici_pst_pubblici.json").read_text(encoding="utf-8"))
    expected_unep_codes = {
        item["codice_ufficio"]
        for item in pst_snapshot["uffici"]["civili"]
        if "UNEP" in item["descrizione"].upper()
    }
    actual_unep_codes = {item["codice"] for item in payload["ufficiUnep"]}
    assert len(expected_unep_codes) == 141
    assert actual_unep_codes == expected_unep_codes
    assert all(item["pec"] for item in payload["ufficiUnep"])
    assert any(
        item["codice"] == "02411602235"
        and item["pec"] == "unep.tribunale.vicenza@civile.ptel.giustiziacert.it"
        for item in payload["ufficiUnep"]
    )
    assert any(item["value"] == "raccomandata" for item in payload["tipiNotificaNonPec"])
    assert payload["azioni"]["unep"] == "/api/v1/ui/notifiche-legali/unep"
    assert payload["azioni"]["nonPec"] == "/api/v1/ui/notifiche-legali/non-pec"
    assert payload["azioni"]["attestazioneConformita"] == "/api/v1/ui/notifiche-legali/attestazione-conformita"
    assert payload["azioni"]["invioPecLocale"] == "/api/v1/ui/notifiche-legali/invio-pec-locale"
    assert payload["azioni"]["confermaInvioPecLocale"] == "/api/v1/ui/notifiche-legali/invio-pec-locale/conferma"
    assert "provaDeposito" not in payload["azioni"]
    assert "depositoChecklist" not in payload["azioni"]
    assert any(field["token"] == "{{ documenti_righe }}" for field in payload["campiDisponibili"])
    assert invalid_response.status_code == 400
    assert invalid_payload["ok"] is False
    assert valid_response.status_code == 200
    assert valid_payload["ok"] is True
    assert "RELATA DI NOTIFICA" in valid_payload["relataText"]
    assert valid_payload["outputPlan"]["workflowSteps"]
    assert valid_payload["outputPlan"]["auditTrail"]["documentsCount"] == 1
    assert "Relata di notifica.pdf" in valid_payload["outputPlan"]["files"]
    assert "Attestazione di conformità.pdf" in valid_payload["outputPlan"]["files"]
    assert "Ricevuta completa richiesta" not in valid_payload["checklistText"]
    assert "Avvocato ha verificato e autorizzato" not in valid_payload["checklistText"]
    assert "RAC" not in "\n".join(valid_payload["nextActions"])
    assert "RdAC" not in "\n".join(valid_payload["nextActions"])
    assert attestation_response.status_code == 200
    assert attestation_response.mimetype == "application/pdf"
    assert attestation_response.data.startswith(b"%PDF")
    assert send_response.status_code == 200
    assert send_result["ok"] is True
    assert send_result["message"] == "Piano PEC preparato dal PC locale per la notifica corrente."
    assert len(send_result["outputPlan"]["timingPlan"]["plannedAt"].split(":")) == 3
    send_delivery = send_result["outputPlan"]["deliveryPlan"]
    assert send_delivery["localSendOnly"] is True
    assert send_delivery["presidioPecAutomation"]["localSendOnly"] is True
    assert send_delivery["presidioPecAutomation"]["enabled"] is False
    assert send_delivery["presidioPecAutomation"]["phase"] == "post_invio_reale"
    assert send_delivery["mode"] == "pec_l53_controllata"
    assert send_delivery["sendPhase"] == "invio_finale"
    send_output_json = json.dumps(send_result["outputPlan"], ensure_ascii=False)
    assert "expectedReceiptSubjects" not in send_delivery
    assert "postSendEvidenceRequired" not in send_delivery
    assert send_delivery["messagesCount"] == 1
    assert send_delivery["singleMessageToAllRecipients"] is True
    assert send_delivery["separatePecRequired"] is False
    assert send_delivery["postSendDocumentArchive"]
    assert "In attesa della RAC effettiva" not in send_output_json
    assert "In attesa della RdAC effettiva" not in send_output_json
    assert unverified_send_response.status_code == 200
    assert unverified_send_result["ok"] is True
    assert not any("PEC_DESTINATARIO_VERIFICA_REQUIRED" in item for item in unverified_send_result["blockers"])
    assert client_response.status_code == 200
    assert client_payload["ok"] is True
    assert client_payload["relataText"] == ""
    assert unep_response.status_code == 200
    assert unep_payload["ok"] is True
    assert unep_payload["outputPlan"]["unepRequest"]["tipo"] == "telematica"
    assert non_pec_response.status_code == 200
    assert non_pec_payload["ok"] is True
    assert non_pec_payload["outputPlan"]["historicalFields"]["NotificaID"] == "N-2026-15"
    assert non_pec_blocked_response.status_code == 400
    assert non_pec_blocked_payload["ok"] is False
    assert any("NOTIFICA_ID_REQUIRED" in item for item in non_pec_blocked_payload["blockers"])


def test_api_react_notifiche_legali_relata_firmata_usa_pdf_generato_nella_stessa_sessione(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    with app.app_context():
        fascicolo = get_fascicoli().nuovo(
            "Pratica firma relata Studio Telematico",
            TipoFascicolo.CIVILE,
            nome_cliente="Cliente Notifica",
        )

    payload = _legal_payload()
    payload.update({
        "fascicolo_id": fascicolo.id,
        "practice_id": fascicolo.id,
        "practiceId": fascicolo.id,
    })
    source_response = client.post(
        "/api/v1/ui/notifiche-legali/relata-pdf",
        json=payload,
        headers=headers,
    )
    assert source_response.status_code == 200
    source_pdf = source_response.data
    assert hashlib.sha256(source_pdf).hexdigest() == source_response.headers["X-IUSENTRA-Document-SHA256"]

    monkeypatch.setattr(
        "pct.firme_cades.inspect_signed_document_bytes",
        lambda *, source_name, data: SimpleNamespace(
            status=SimpleNamespace(signature_verified=True),
            payload_bytes=source_pdf,
        ),
    )
    monkeypatch.setattr(
        "pct.firma.analizza_firma_documento",
        lambda data, name: [{"firmatario": "Mario Rossi", "scaduto": False}],
    )
    monkeypatch.setattr(
        react_api,
        "_extract_pdf_text_for_relata",
        lambda data: "testo estratto diverso ma PDF sorgente identico",
    )

    signed_response = client.post(
        "/api/v1/ui/notifiche-legali/relata-firmata",
        data={
            "payload": json.dumps(payload),
            "file": (io.BytesIO(b"fake-cades-signed-data"), "relata_notifica.pdf.p7m"),
        },
        headers=headers,
        content_type="multipart/form-data",
    )
    signed_payload = signed_response.get_json()

    assert signed_response.status_code == 200
    assert signed_payload["ok"] is True
    assert signed_payload["sourceSha256"] == hashlib.sha256(source_pdf).hexdigest()


def test_api_react_notifiche_legali_invio_locale_usa_allegati_reali_message_id_e_presidio(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    atto_bytes = b"%PDF-1.4\n%IUSENTRA atto da notificare\n"
    relata_bytes = b"IUSENTRA RELATA FIRMATA CADES DI TEST"
    with app.app_context():
        fascicoli = get_fascicoli()
        fascicolo = fascicoli.nuovo(
            "Pratica notifica locale Studio Telematico",
            TipoFascicolo.CIVILE,
            nome_cliente="Cliente Notifica",
        )
        atto = fascicoli.aggiungi_documento(
            fascicolo.id,
            "ricorso.pdf",
            TipoDocumento.ATTO_GIUDIZIARIO,
            atto_bytes,
            note="Atto notificato",
        )
        relata = fascicoli.aggiungi_documento(
            fascicolo.id,
            "Relata di notifica.pdf.p7m",
            TipoDocumento.NOTIFICA,
            relata_bytes,
            note="Relata firmata digitalmente",
            firmato=True,
        )

    payload = _legal_payload()
    stale_payload_hash = "0" * 64
    payload.update({
        "fascicolo_id": fascicolo.id,
        "practiceId": fascicolo.id,
        "pratica_codice": str(getattr(fascicolo, "numero", "")),
        "operazione": "invio_pec_l53",
        "conferma_invio_pec": True,
        "invio_finale": True,
        "data_ora_invio_pec": "2026-05-12T14:26:00+02:00",
        "nome_file": "ricorso.pdf",
        "descrizione_documento": "Ricorso notificato",
        "origine_documento": "nativo_digitale",
        "hash_sha256": stale_payload_hash,
        "documenti": [
            {
                "nome_file": "ricorso.pdf",
                "descrizione": "Ricorso notificato",
                "origine": "nativo_digitale",
                "hash_sha256": stale_payload_hash,
                "document_id": atto.id,
                "documentId": atto.id,
            }
        ],
        "relata_firmata": True,
        "relata_firmata_file": "Relata di notifica.pdf.p7m",
        "relata_firmata_sha256": stale_payload_hash,
        "relata_firmata_document_id": relata.id,
    })

    prepare_response = client.post(
        "/api/v1/ui/notifiche-legali/invio-pec-locale",
        json=payload,
        headers=headers,
    )
    assert prepare_response.status_code == 200
    prepared = prepare_response.get_json()
    assert prepared["ok"] is True
    assert not any("Impronta diversa" in str(item) for item in prepared.get("blockers", []))
    assert prepared["requiresLocalPec"] is True
    assert prepared["notificationId"]
    messages = prepared["localPecMessages"]
    assert len(messages) == 1
    message = messages[0]
    local_payload = message["payload"]
    assert message["endpoint"] == "http://127.0.0.1:27272/pec/send"
    assert local_payload["cc"] == []
    assert local_payload["bcc"] == []
    assert "password" not in local_payload
    assert local_payload["to"] == prepared["outputPlan"]["deliveryPlan"]["studioTelematicoTo"]
    assert "controparte@example.pec.it" in local_payload["to"]
    assert "codice fiscale:" in local_payload["to"]
    assert "pubblico elenco:" in local_payload["to"]
    assert local_payload["subject"].startswith("Notificazione ai sensi della legge n. 53/1994")
    assert "[Notifica_ID:" in local_payload["subject"]
    assert "Riferimento da citare nella risposta:" in local_payload["body"]
    assert "Pratica:" in local_payload["body"]
    attachments = local_payload["attachments"]
    assert [item["filename"] for item in attachments] == ["ricorso.pdf", "Relata di notifica.pdf.p7m"]
    assert [item["studioTelematicoArchiveRole"] for item in attachments] == ["originale_notificato", "relata_notifica"]
    assert base64.b64decode(attachments[0]["content_base64"]) == atto_bytes
    assert base64.b64decode(attachments[1]["content_base64"]) == relata_bytes
    assert attachments[0]["sha256"] == atto.hash_sha256
    assert attachments[1]["sha256"] == relata.hash_sha256

    missing_message_id_response = client.post(
        "/api/v1/ui/notifiche-legali/invio-pec-locale/conferma",
        json={
            "payload": payload,
            "notificationId": prepared["notificationId"],
            "results": [{"localMessageId": message["id"], "sentAt": "2026-05-12T14:26:11+02:00"}],
        },
        headers=headers,
    )
    assert missing_message_id_response.status_code == 400
    assert "Message-ID mancante" in missing_message_id_response.get_json()["message"]

    confirm_response = client.post(
        "/api/v1/ui/notifiche-legali/invio-pec-locale/conferma",
        json={
            "payload": payload,
            "notificationId": prepared["notificationId"],
            "results": [{
                "localMessageId": message["id"],
                "messageId": "<iusentra-test-message@pec.local>",
                "sentAt": "2026-05-12T14:26:11+02:00",
            }],
        },
        headers=headers,
    )
    assert confirm_response.status_code == 200
    confirmed = confirm_response.get_json()
    assert confirmed["ok"] is True
    assert confirmed["status"] == "SENT_WAITING_RAC"
    assert confirmed["presidioId"]
    assert "1 destinatario" in confirmed["message"]
    assert confirmed["sent"][0]["messageId"] == "<iusentra-test-message@pec.local>"
    assert confirmed["outputPlan"]["deliveryPlan"]["confirmedMessageIds"] == ["<iusentra-test-message@pec.local>"]


def test_api_react_notifiche_legali_invio_locale_recupera_relata_firmata_corrente(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(
        "web.blueprints.api_v1_react._local_rome_datetime_seconds",
        lambda: "2026-05-12T14:26:00",
    )
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    atto_bytes = b"%PDF-1.4\n%IUSENTRA atto da notificare\n"
    relata_bytes = b"IUSENTRA RELATA FIRMATA GIA SALVATA CADES DI TEST"
    with app.app_context():
        fascicoli = get_fascicoli()
        fascicolo = fascicoli.nuovo(
            "Pratica notifica con relata gia firmata",
            TipoFascicolo.CIVILE,
            nome_cliente="Cliente Notifica",
        )
        atto = fascicoli.aggiungi_documento(
            fascicolo.id,
            "ricorso.pdf",
            TipoDocumento.ATTO_GIUDIZIARIO,
            atto_bytes,
            note="Atto notificato",
        )

    payload = _legal_payload()
    payload.update({
        "fascicolo_id": fascicolo.id,
        "practiceId": fascicolo.id,
        "pratica_codice": str(getattr(fascicolo, "numero", "")),
        "operazione": "invio_pec_l53",
        "conferma_invio_pec": True,
        "invio_finale": True,
        "data_ora_invio_pec": "2026-05-12T14:26:00",
        "nome_file": "ricorso.pdf",
        "descrizione_documento": "Ricorso notificato",
        "origine_documento": "nativo_digitale",
        "hash_sha256": atto.hash_sha256,
        "documenti": [
            {
                "nome_file": "ricorso.pdf",
                "descrizione": "Ricorso notificato",
                "origine": "nativo_digitale",
                "hash_sha256": atto.hash_sha256,
                "document_id": atto.id,
                "documentId": atto.id,
            }
        ],
        "relata_firmata": False,
    })
    source_sha256 = hashlib.sha256(generate_relata_pdf_bytes(dict(payload))).hexdigest()
    with app.app_context():
        relata = get_fascicoli().aggiungi_documento(
            fascicolo.id,
            "Relata di notifica.pdf.p7m",
            TipoDocumento.NOTIFICA,
            relata_bytes,
            note="Relata firmata digitalmente",
            tags=["relata-notifica", "firma-verificata", f"relata-source-sha256:{source_sha256}"],
            firmato=True,
        )

    prepare_response = client.post(
        "/api/v1/ui/notifiche-legali/invio-pec-locale",
        json=payload,
        headers=headers,
    )

    assert prepare_response.status_code == 200
    prepared = prepare_response.get_json()
    assert prepared["ok"] is True
    blockers = " ".join(str(item) for item in prepared.get("blockers", []))
    assert "Relata firmata mancante" not in blockers
    attachments = prepared["localPecMessages"][0]["payload"]["attachments"]
    assert [item["filename"] for item in attachments] == ["ricorso.pdf", "Relata di notifica.pdf.p7m"]
    assert base64.b64decode(attachments[-1]["content_base64"]) == relata_bytes
    assert attachments[-1]["sha256"] == relata.hash_sha256


def test_api_consultazione_pubblico_elenco_salva_la_prova_nel_fascicolo(tmp_path: Path):
    app = _app(tmp_path)
    headers = {"X-API-Key": "react-test-key"}
    with app.app_context():
        fascicolo = get_fascicoli().nuovo(
            "Pratica verifica PEC",
            TipoFascicolo.CIVILE,
            nome_cliente="Cliente di prova",
        )

    response = app.test_client().post(
        "/api/v1/ui/notifiche-legali/verifica-pec-consultata",
        json={
            "fascicolo_id": fascicolo.id,
            "source": "ini_pec",
            "pec": "controparte@example.pec.it",
            "codice_fiscale": "01234567890",
            "soggetto": "Controparte S.p.A.",
            "consulted_at": datetime.now(ZoneInfo("Europe/Rome")).isoformat(),
        },
        headers=headers,
    )

    assert response.status_code == 200
    body = response.get_json()
    assert body["verified"] is True
    assert body["saved_in_practice"] is True
    with app.app_context():
        saved = get_fascicoli().get(fascicolo.id)
        rows = saved.source_snapshot["legal_notification_public_register_evidence"]
    assert len(rows) == 1
    assert rows[0]["source"] == "ini_pec"
    assert rows[0]["evidence_sha256"] == body["evidence_sha256"]


def test_api_notifiche_legali_destinatario_manuale_salva_soggetto_e_ricarica(tmp_path: Path):
    app = _app(tmp_path)
    headers = {"X-API-Key": "react-test-key"}
    with app.app_context():
        fascicolo = get_fascicoli().nuovo(
            "Pratica destinatario manuale",
            TipoFascicolo.CIVILE,
            nome_cliente="Cliente Manuale",
        )
        get_clienti().nuovo(
            TipoCliente.PERSONA_FISICA,
            nome="Destinatario",
            cognome="Manuale",
            codice_fiscale="DMNFNL26L29H224Z",
        )

    response = app.test_client().post(
        "/api/v1/ui/notifiche-legali/destinatari-manuali",
        json={
            "practiceId": fascicolo.id,
            "nome": "Destinatario Manuale",
            "codiceFiscalePiva": "DMNFNL26L29H224Z",
            "pec": "destinatario.manuale.test@pec.it",
            "ruolo": "controparte",
            "fontePecSuggerita": "ini_pec",
            "parteRappresentata": "Parte manuale",
        },
        headers=headers,
    )

    assert response.status_code == 201
    body = response.get_json()
    assert body["ok"] is True
    assert body["created"] is True
    assert body["linkedToPractice"] is True
    assert body["recipient"]["pec"] == "destinatario.manuale.test@pec.it"
    assert body["recipient"]["ruoloPratica"] == "Inserito manualmente"

    payload_response = app.test_client().get("/api/v1/ui/notifiche-legali", headers=headers)
    soggetti_response = app.test_client().get("/api/v1/ui/soggetti", headers=headers)
    payload = payload_response.get_json()
    soggetti_payload = soggetti_response.get_json()

    assert any(
        item["pec"] == "destinatario.manuale.test@pec.it"
        and item["id"] == body["subjectId"]
        for item in payload["precompilazione"]["destinatari"]
    )
    assert any(
        item["pec"] == "destinatario.manuale.test@pec.it"
        and item["id"] == body["subjectId"]
        for item in soggetti_payload["items"]
    )
    with app.app_context():
        saved = get_soggetti().get(body["subjectId"])
        parts = get_soggetti().parti_fascicolo(fascicolo.id)
    assert saved is not None
    assert "notifiche-legali-manuale" in saved.tag
    assert saved.recapiti.pec == "destinatario.manuale.test@pec.it"
    assert any(soggetto.id == body["subjectId"] for _, soggetto in parts)

    update_response = app.test_client().post(
        "/api/v1/ui/notifiche-legali/destinatari-manuali",
        json={
            "nome": "Destinatario Manuale Aggiornato",
            "codiceFiscalePiva": "DMNFNL26L29H224Z",
            "pec": "destinatario.manuale.test@pec.it",
            "ruolo": "controparte",
            "fontePecSuggerita": "ini_pec",
        },
        headers=headers,
    )
    assert update_response.status_code == 200
    assert update_response.get_json()["subjectId"] == body["subjectId"]
    with app.app_context():
        assert len(get_soggetti().tutti()) == 1


def test_api_notifiche_legali_destinatario_manuale_non_fonde_pec_diverse_stesso_ente(tmp_path: Path):
    app = _app(tmp_path)
    headers = {"X-API-Key": "react-test-key"}
    with app.app_context():
        fascicolo = get_fascicoli().nuovo(
            "Romeo Maria c. MIM",
            TipoFascicolo.CIVILE,
            nome_cliente="Romeo Maria",
        )

    client = app.test_client()
    first_response = client.post(
        "/api/v1/ui/notifiche-legali/destinatari-manuali",
        json={
            "practiceId": fascicolo.id,
            "nome": "MIM - USR Reggio Calabria",
            "codiceFiscalePiva": "80007410808",
            "pec": "usprc.contenzioso@postacert.istruzione.it",
            "ruolo": "pa",
            "fontePecSuggerita": "registro_ppaa",
            "parteRappresentata": "Ministero dell'Istruzione e del Merito",
        },
        headers=headers,
    )
    second_response = client.post(
        "/api/v1/ui/notifiche-legali/destinatari-manuali",
        json={
            "practiceId": fascicolo.id,
            "nome": "USR Calabria - Ambito Territoriale di Reggio Calabria - Ufficio VI",
            "codiceFiscalePiva": "80007410808",
            "pec": "usprc@postacert.istruzione.it",
            "ruolo": "pa",
            "fontePecSuggerita": "registro_ppaa",
            "parteRappresentata": "Ministero dell'Istruzione e del Merito",
        },
        headers=headers,
    )

    assert first_response.status_code == 201
    assert second_response.status_code == 201
    first = first_response.get_json()
    second = second_response.get_json()
    assert first["subjectId"] != second["subjectId"]
    assert second["recipient"]["pec"] == "usprc@postacert.istruzione.it"
    assert second["recipient"]["ruolo"] == "pa"
    assert second["recipient"]["fontePecSuggerita"] == "registro_ppaa"
    assert second["linkedToPractice"] is True

    payload = client.get("/api/v1/ui/notifiche-legali", headers=headers).get_json()
    saved_addresses = {
        item["pec"]
        for item in payload["precompilazione"]["destinatari"]
        if item.get("codiceFiscalePiva") == "80007410808"
    }
    assert "usprc.contenzioso@postacert.istruzione.it" in saved_addresses
    assert "usprc@postacert.istruzione.it" in saved_addresses
    with app.app_context():
        parts = get_soggetti().parti_fascicolo(fascicolo.id)
    linked_addresses = {
        soggetto.recapiti.pec
        for _, soggetto in parts
        if getattr(getattr(soggetto, "recapiti", None), "pec", "")
    }
    assert "usprc.contenzioso@postacert.istruzione.it" in linked_addresses
    assert "usprc@postacert.istruzione.it" in linked_addresses


def test_api_notifiche_legali_destinatario_manuale_sostituisce_fonte_pubblico_elenco_obsoleta(tmp_path: Path):
    app = _app(tmp_path)
    headers = {"X-API-Key": "react-test-key"}
    with app.app_context():
        fascicolo = get_fascicoli().nuovo(
            "Romeo Maria c. MIM",
            TipoFascicolo.CIVILE,
            nome_cliente="Romeo Maria",
            controparte="MINISTERO ISTRUZIONE E DEL MERITO",
        )
        existing = get_soggetti().crea(
            TipoSoggetto.PUBBLICA_AMMINISTRAZIONE,
            nome_completo="MINISTERO ISTRUZIONE E DEL MERITO",
            codice_fiscale="80185250588",
            recapiti=Recapiti(pec="usprc@postacert.istruzione.it"),
            tag=["pubblico-elenco:reginde", "notifiche-legali"],
        )
        get_soggetti().aggiungi_parte(
            fascicolo.id,
            existing.id,
            RuoloSoggetto.CONTROPARTE,
            "Ministero dell'Istruzione e del Merito",
        )

    client = app.test_client()
    response = client.post(
        "/api/v1/ui/notifiche-legali/destinatari-manuali",
        json={
            "practiceId": fascicolo.id,
            "nome": "USR Calabria - Ambito Territoriale di Reggio Calabria - Ufficio VI",
            "codiceFiscalePiva": "80007410808",
            "pec": "usprc@postacert.istruzione.it",
            "ruolo": "pa",
            "fontePecSuggerita": "registro_ppaa",
            "parteRappresentata": "Ministero dell'Istruzione e del Merito",
        },
        headers=headers,
    )

    assert response.status_code == 200
    body = response.get_json()
    assert body["subjectId"] == existing.id
    assert body["recipient"]["fontePecSuggerita"] == "registro_ppaa"

    payload = client.get("/api/v1/ui/notifiche-legali", headers=headers).get_json()
    matched = [
        item
        for item in payload["precompilazione"]["destinatari"]
        if item.get("id") == existing.id
    ]
    assert matched
    assert matched[0]["fontePecSuggerita"] == "registro_ppaa"
    with app.app_context():
        saved = get_soggetti().get(existing.id)
    assert saved is not None
    assert "pubblico-elenco:registro_ppaa" in saved.tag
    assert "pubblico-elenco:reginde" not in saved.tag


def test_api_react_notifiche_legali_salva_e_usa_modello_relata_personalizzato(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    body = "\n".join([
        "RELAZIONE DI NOTIFICAZIONE PERSONALIZZATA",
        "Avv. {{ avvocato.full_name }} notifica per {{ cliente.nome_denominazione }}.",
        "Destinatario {{ destinatario.nome_denominazione }} presso {{ destinatario.pec }}.",
        "{{ documenti_righe }}",
        "{{ blocco_procedimento }}",
        "{{ notifica.luogo }}, {{ notifica.data }}",
    ])

    save_response = client.post(
        "/api/v1/ui/notifiche-legali/modelli-relata",
        json={"label": "Relata prova studio", "description": "Uso interno studio", "body": body, "requiresProceeding": True},
        headers=headers,
    )
    saved = save_response.get_json()
    catalog = client.get("/api/v1/ui/notifiche-legali", headers=headers).get_json()
    payload = _legal_payload()
    payload["template_id"] = saved["template"]["value"]
    preview_response = client.post("/api/v1/ui/notifiche-legali/notifica", json=payload, headers=headers)
    preview = preview_response.get_json()

    assert save_response.status_code == 200
    assert saved["ok"] is True
    assert saved["template"]["custom"] is True
    assert any(item["value"] == saved["template"]["value"] and item["custom"] for item in catalog["modelliRelata"])
    assert preview_response.status_code == 200
    assert preview["ok"] is True
    assert "RELAZIONE DI NOTIFICAZIONE PERSONALIZZATA" in preview["relataText"]
    assert "Cliente S.r.l." in preview["relataText"]


def test_api_react_notifiche_legali_anteprima_relata_e_token_sicuri(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    payload = _legal_payload()
    payload["destinatario_pec"] = ""
    preview_response = client.post("/api/v1/ui/notifiche-legali/anteprima-relata", json=payload, headers=headers)
    preview = preview_response.get_json()
    dangerous = _legal_payload()
    dangerous["template_id"] = "relata_personalizzata_pericolosa"
    dangerous["template_personalizzato"] = {
        "id": "relata_personalizzata_pericolosa",
        "label": "Pericolosa",
        "custom_body": "Accesso {{ cycler.__init__.__globals__ }}",
    }
    dangerous_response = client.post("/api/v1/ui/notifiche-legali/anteprima-relata", json=dangerous, headers=headers)
    dangerous_payload = dangerous_response.get_json()

    assert preview_response.status_code == 200
    assert preview["ok"] is True
    assert "[dato mancante: PEC destinatario]" in preview["previewText"]
    assert dangerous_response.status_code == 400
    assert dangerous_payload["ok"] is False
    assert dangerous_payload["blockers"]


def test_api_react_notifiche_legali_salva_bozza_relata_e_non_modello(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    draft_response = client.post(
        "/api/v1/ui/notifiche-legali/bozze-relata",
        json={"practiceId": "fascicolo-1", "templateId": "relata_pec_base_l53", "relataText": "Bozza relata modificata per questa notifica."},
        headers=headers,
    )
    empty_response = client.post(
        "/api/v1/ui/notifiche-legali/bozze-relata",
        json={"templateId": "relata_pec_base_l53", "relataText": ""},
        headers=headers,
    )
    catalog = client.get("/api/v1/ui/notifiche-legali", headers=headers).get_json()
    draft_payload = draft_response.get_json()

    assert draft_response.status_code == 200
    assert draft_payload["ok"] is True
    assert draft_payload["draftId"]
    assert empty_response.status_code == 400
    assert not any("Bozza relata modificata" in item["previewText"] for item in catalog["modelliRelata"])
    assert (tmp_path / "notifiche" / "bozze_relata.json").exists()


def test_api_react_notifiche_legali_salva_attestazione_pdf_nel_fascicolo(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    with app.app_context():
        fascicolo = get_fascicoli().nuovo(
            "Pratica attestazione notifica",
            TipoFascicolo.CIVILE,
            nome_cliente="Cliente attestazione",
        )
    payload = _legal_payload()
    payload.update({
        "practice_id": fascicolo.id,
        "fascicolo_id": fascicolo.id,
        "documenti": [
            {
                "nome_file": "verbale_udienza.pdf",
                "descrizione": "Verbale di udienza",
                "origine": "copia_fascicolo_informatico",
                "hash_sha256": "c" * 64,
            }
        ],
        "attestazione_override_text": (
            "ATTESTAZIONE DI CONFORMITÀ\n\n"
            "Testo attestazione modificato e salvato dall'avvocato.\n\n"
            "Mario Rossi\n"
            "Firmato digitalmente"
        ),
    })

    response = client.post(
        "/api/v1/ui/notifiche-legali/attestazione-conformita-fascicolo",
        json=payload,
        headers=headers,
    )
    body = response.get_json()

    assert response.status_code == 200
    assert body["ok"] is True
    assert body["documentId"]
    assert body["fileName"].startswith("Attestazione_di_conformita")
    with app.app_context():
        saved = get_fascicoli().get(fascicolo.id)
        documents = list(getattr(saved, "documenti", []) or [])
    assert any("attestazione-conformita" in list(getattr(document, "tags", []) or []) for document in documents)


def test_api_react_notifiche_legali_salva_bozza_attestazione_e_non_modello(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    draft_response = client.post(
        "/api/v1/ui/notifiche-legali/bozze-attestazione",
        json={
            "practiceId": "fascicolo-1",
            "templateId": "relata_pec_base_l53",
            "attestationText": "Bozza attestazione modificata per questa notifica.",
        },
        headers=headers,
    )
    empty_response = client.post(
        "/api/v1/ui/notifiche-legali/bozze-attestazione",
        json={"templateId": "relata_pec_base_l53", "attestationText": ""},
        headers=headers,
    )
    catalog = client.get("/api/v1/ui/notifiche-legali", headers=headers).get_json()
    draft_payload = draft_response.get_json()

    assert draft_response.status_code == 200
    assert draft_payload["ok"] is True
    assert draft_payload["draftId"]
    assert empty_response.status_code == 400
    assert not any("Bozza attestazione modificata" in item["previewText"] for item in catalog["modelliRelata"])
    assert (tmp_path / "notifiche" / "bozze_attestazione.json").exists()


def test_bozza_relata_override_non_puo_eliminare_i_dati_obbligatori():
    payload = _legal_payload()
    payload["relata_override_text"] = "TESTO MANUALE DELLA RELATA"
    result = validate_legal_notification(payload)

    assert result.ok is False
    assert any("RELAZIONE_CONTENUTO_OBBLIGATORIO_REQUIRED" in item for item in result.blockers)


def test_notifica_richiede_data_e_ora_relata_esplicite():
    payload = _legal_payload()
    payload["data_relata"] = ""
    payload["ora_relata"] = ""

    result = validate_legal_notification(payload)

    assert result.ok is True
    assert result.blockers == []
    assert "Indica la data della relata." in result.warnings
    assert "Indica l'ora italiana della relata." in result.warnings


def test_api_react_notifiche_legali_robustezza_json_e_limiti(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    no_json = client.post("/api/v1/ui/notifiche-legali/notifica", data="non json", headers=headers)
    malformed = client.post(
        "/api/v1/ui/notifiche-legali/notifica",
        data="{",
        content_type="application/json",
        headers=headers,
    )
    label_empty = client.post(
        "/api/v1/ui/notifiche-legali/modelli-relata",
        json={"label": "", "body": "RELAZIONE\n{{ avvocato.full_name }}\n" * 10},
        headers=headers,
    )
    too_long = client.post(
        "/api/v1/ui/notifiche-legali/modelli-relata",
        json={"label": "Relata", "body": "x" * 25000},
        headers=headers,
    )
    forbidden = client.post(
        "/api/v1/ui/notifiche-legali/modelli-relata",
        json={"label": "Relata", "body": ("RELAZIONE\n{{ token_non_permesso }}\n" * 10)},
        headers=headers,
    )

    assert no_json.status_code == 400
    assert no_json.get_json()["ok"] is False
    assert malformed.status_code == 400
    assert malformed.get_json()["ok"] is False
    assert label_empty.status_code == 400
    assert too_long.status_code == 400
    assert forbidden.status_code == 400
    assert forbidden.get_json()["blockers"]


def test_api_react_notifiche_legali_modelli_cliente_separati(tmp_path: Path):
    app = _app(tmp_path)
    client = app.test_client()
    headers = {"X-API-Key": "react-test-key"}
    payload = client.get("/api/v1/ui/notifiche-legali", headers=headers).get_json()
    communication_response = client.post(
        "/api/v1/ui/notifiche-legali/comunicazione-cliente",
        json={
            "operazione": "comunicazione_cliente_non_notifica",
            "template_id": "invio_provvedimento",
            "cliente_nome": "Cliente",
            "ufficio_giudiziario": "Tribunale di Roma",
            "numero_rg": "1234",
            "anno_rg": "2026",
            "provvedimento_descrizione": "Ordinanza depositata",
        },
        headers=headers,
    )
    blocked_response = client.post(
        "/api/v1/ui/notifiche-legali/comunicazione-cliente",
        json={"template_id": "relata_pec_base_l53", "cliente_nome": "Cliente", "provvedimento_descrizione": "Atto"},
        headers=headers,
    )
    communication = communication_response.get_json()

    assert payload["clientCommunicationTemplateVersion"] != payload["templateCatalogVersion"]
    assert "2026.05.12" not in payload["clientCommunicationTemplateVersion"]
    assert payload["modelliComunicazioneCliente"]
    assert {item["value"] for item in payload["modelliComunicazioneCliente"]}.isdisjoint({item["value"] for item in payload["modelliRelata"]})
    assert communication_response.status_code == 200
    assert communication["ok"] is True
    assert communication["relataText"] == ""
    assert communication["templateId"] == "invio_provvedimento"
    assert blocked_response.status_code == 400
    assert blocked_response.get_json()["ok"] is False


def test_payload_react_notifiche_legali_precompila_da_dati_iusentra():
    cliente = SimpleNamespace(
        id="cliente-1",
        nome_completo="Cliente S.r.l.",
        identificativo_fiscale="01234567890",
        recapiti=SimpleNamespace(pec="cliente@example.pec.it"),
    )
    documento = SimpleNamespace(
        id="doc-1",
        nome_originale="20260510185021337.PDF",
        nome_portale="",
        nome="Ordinanza udienza 10 maggio 2026.pdf",
        percorso="20260510185021337.PDF",
        tipo_atto_portale="ordinanza emessa dal Tribunale di Roma",
        classificazione_portale="",
        note="",
        fonte_documento="PORTALE_TELEMATICO",
        servizio_portale="PST",
        hash_sha256="abc123",
        data_documento="2026-05-10",
        data_deposito_portale="",
        id_documento_portale="pst-doc-1",
        tags=[],
    )
    fascicolo = SimpleNamespace(
        id="fascicolo-1",
        numero="2026/001",
        titolo="Cliente S.r.l. / Alfa S.p.A.",
        id_cliente="cliente-1",
        nome_cliente="Cliente S.r.l.",
        controparte="Alfa S.p.A.",
        cf_controparte="09876543210",
        tribunale="Tribunale di Roma",
        sezione="III Civile",
        numero_rg="1234",
        anno_rg=2026,
        giudice="Dott. Verdi",
        tipo_procedimento="civile ordinario",
        documenti=[documento],
        depositi_pct=[],
    )
    soggetto = SimpleNamespace(
        id="soggetto-1",
        tipo=SimpleNamespace(value="PERSONA_GIURIDICA"),
        nome_completo="Alfa S.p.A.",
        ragione_sociale="Alfa S.p.A.",
        identificativo="09876543210",
        recapiti=SimpleNamespace(pec="alfa@example.pec.it"),
        qualifica="",
    )
    parte = SimpleNamespace(ruolo=SimpleNamespace(value="CONTROPARTE"), note="")

    clienti_repo = SimpleNamespace(tutti=lambda: [cliente], get=lambda _id: cliente)
    fascicoli_repo = SimpleNamespace(
        tutti=lambda archiviati=False: [fascicolo],
        get=lambda _id: fascicolo,
    )
    soggetti_repo = SimpleNamespace(
            tutti=lambda: [soggetto],
            parti_fascicolo=lambda id_fascicolo: [(parte, soggetto)],
    )
    payload = build_react_notifiche_legali_payload(
        get_clienti=lambda: clienti_repo,
        get_fascicoli=lambda: fascicoli_repo,
        get_soggetti=lambda: soggetti_repo,
    )
    practice_payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_payload(
        "fascicolo-1",
        get_clienti=lambda: clienti_repo,
        get_fascicoli=lambda: fascicoli_repo,
        get_soggetti=lambda: soggetti_repo,
    )

    pratica = practice_payload["pratica"]
    destinatario = pratica["destinatari"][0]
    documento_payload = pratica["documenti"][0]

    assert pratica["assistitoNome"] == "Cliente S.r.l."
    assert pratica["procedimento"]["ufficio"] == "Tribunale di Roma"
    assert pratica["procedimento"]["numeroRg"] == "1234"
    assert pratica["procedimento"]["annoRg"] == "2026"
    assert destinatario["pec"] == "alfa@example.pec.it"
    assert destinatario["fontePecSuggerita"] == "ini_pec"
    assert destinatario["parteRappresentata"] == "Alfa S.p.A."
    assert documento_payload["nomeFile"] == "Ordinanza udienza 10 maggio 2026.pdf"
    assert documento_payload["label"].startswith("Ordinanza udienza 10 maggio 2026.pdf")
    assert "20260510185021337.PDF" not in documento_payload["label"]
    assert documento_payload["casoNotificaSuggerito"] == "provvedimento_giudice"
    assert documento_payload["modelloRelataSuggerito"] == "relata_provvedimento_giudice"
    assert documento_payload["provvedimentoTipo"] == "Ordinanza"
    assert documento_payload["criterioTipoDocumento"] == "metadati portale/fascicolo"
    assert documento_payload["origine"] == "copia_fascicolo_informatico"
    assert documento_payload["riferimentoPortale"] == "pst-doc-1"
    assert documento_payload["servizioPortale"] == "PST"
    assert documento_payload["documentoUfficio"] is False
    assert documento_payload["acquisitoDaPortale"] is True
    assert documento_payload["notificaRichiesta"] is False
    assert documento_payload["necessitaAttestazione"] is True
    assert pratica["portaleAcquisizioneHref"].startswith("/portali/pst/acquisizione?")
    assert pratica["documentoUfficioMonitor"]["stato"] == "da_verificare"
    assert payload["precompilazione"]["pratiche"] == []
    assert payload["precompilazione"]["indicePratiche"][0]["id"] == "fascicolo-1"
    assert payload["azioni"]["firmaDigitale"] == "/guida/firma-digitale"


def test_payload_react_documento_legge_testo_prima_del_nome_file_fuorviante():
    documento = SimpleNamespace(
        id="doc-testo",
        nome_originale="decreto_ingiuntivo_nome_errato.pdf",
        nome_portale="",
        nome="decreto_ingiuntivo_nome_errato.pdf",
        percorso="decreto_ingiuntivo_nome_errato.pdf",
        tipo_atto_portale="",
        classificazione_portale="",
        note="",
        fonte_documento="PORTALE_TELEMATICO",
        servizio_portale="PST",
        hash_sha256="d" * 64,
        data_documento="2026-05-18",
        data_deposito_portale="",
        id_documento_portale="pst-doc-testo",
        tags=[],
        extracted_text=(
            "TRIBUNALE DI PALMI\n"
            "DECRETO DI FISSAZIONE UDIENZA\n"
            "Il Giudice fissa l'udienza del procedimento e assegna termine per la notifica."
        ),
    )
    fascicolo = SimpleNamespace(
        id="fascicolo-testo",
        numero="2026/010",
        titolo="Cliente / Ministero",
        id_cliente="",
        nome_cliente="Cliente",
        controparte="Ministero dell'Istruzione e del Merito",
        cf_controparte="",
        tribunale="Tribunale di Palmi",
        sezione="Lavoro",
        numero_rg="1477",
        anno_rg=2026,
        giudice="",
        tipo_procedimento="lavoro",
        documenti=[documento],
        depositi_pct=[],
    )

    payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_payload(
        "fascicolo-testo",
        get_clienti=lambda: SimpleNamespace(tutti=lambda: [], get=lambda _id: None),
        get_fascicoli=lambda: SimpleNamespace(get=lambda _id: fascicolo),
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: [], parti_fascicolo=lambda id_fascicolo: []),
    )

    documento_payload = payload["pratica"]["documenti"][0]
    assert documento_payload["nomeFile"] == "decreto_ingiuntivo_nome_errato.pdf"
    assert documento_payload["casoNotificaSuggerito"] == "provvedimento_giudice"
    assert documento_payload["modelloRelataSuggerito"] == "relata_provvedimento_giudice"
    assert documento_payload["provvedimentoTipo"] == "Decreto fissazione udienza"
    assert documento_payload["criterioTipoDocumento"] == "testo documento letto"
    assert documento_payload["testoDocumentoDisponibile"] is True


def test_payload_react_non_classifica_notifica_solo_dal_nome_file():
    documento = SimpleNamespace(
        id="doc-nome",
        nome_originale="decreto_ingiuntivo_nome_errato.pdf",
        nome_portale="",
        nome="decreto_ingiuntivo_nome_errato.pdf",
        percorso="decreto_ingiuntivo_nome_errato.pdf",
        tipo_atto_portale="",
        classificazione_portale="",
        note="",
        fonte_documento="PORTALE_TELEMATICO",
        servizio_portale="PST",
        hash_sha256="e" * 64,
        data_documento="",
        data_deposito_portale="",
        id_documento_portale="pst-doc-nome",
        tags=[],
        extracted_text="",
    )
    fascicolo = SimpleNamespace(
        id="fascicolo-nome",
        numero="2026/011",
        titolo="Cliente / Controparte",
        id_cliente="",
        nome_cliente="Cliente",
        controparte="Controparte",
        cf_controparte="",
        tribunale="Tribunale di Palmi",
        sezione="Lavoro",
        numero_rg="1477",
        anno_rg=2026,
        giudice="",
        tipo_procedimento="lavoro",
        documenti=[documento],
        depositi_pct=[],
    )

    payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_payload(
        "fascicolo-nome",
        get_clienti=lambda: SimpleNamespace(tutti=lambda: [], get=lambda _id: None),
        get_fascicoli=lambda: SimpleNamespace(get=lambda _id: fascicolo),
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: [], parti_fascicolo=lambda id_fascicolo: []),
    )

    documento_payload = payload["pratica"]["documenti"][0]
    assert documento_payload["nomeFile"] == "decreto_ingiuntivo_nome_errato.pdf"
    assert documento_payload["casoNotificaSuggerito"] == ""
    assert documento_payload["modelloRelataSuggerito"] == ""
    assert documento_payload["provvedimentoTipo"] == ""
    assert documento_payload["criterioTipoDocumento"] == ""


def test_destinatario_avvocatura_usa_reginde_e_parte_rappresentata_del_fascicolo():
    soggetto = SimpleNamespace(
        id="avvocatura-venezia",
        tipo=SimpleNamespace(value="PERSONA_GIURIDICA"),
        nome_completo="Avvocatura Distrettuale dello Stato di Venezia",
        ragione_sociale="",
        identificativo="94026160278",
        recapiti=SimpleNamespace(pec="ads.ve@mailcert.avvocaturastato.it"),
        qualifica="",
    )
    fascicolo = SimpleNamespace(controparte="Ministero dell'Istruzione e del Merito")

    recipient = react_notifiche_legali_bridge._recipient_from_subject(
        soggetto,
        ruolo="CONTROPARTE",
        note="Aggiunta durante l'apertura del fascicolo.",
        fascicolo=fascicolo,
    )

    assert recipient["ruolo"] == "difensore"
    assert recipient["fontePecSuggerita"] == "reginde"
    assert recipient["parteRappresentata"] == "Ministero dell'Istruzione e del Merito"
    assert recipient["pec"] == "ads.ve@mailcert.avvocaturastato.it"


def test_destinatario_avvocatura_storico_non_rappresenta_se_stesso():
    soggetto = SimpleNamespace(
        id="avvocatura-venezia",
        tipo=SimpleNamespace(value="PERSONA_GIURIDICA"),
        nome_completo="Avvocatura Distrettuale dello Stato di Venezia",
        ragione_sociale="",
        identificativo="94026160278",
        recapiti=SimpleNamespace(pec="ads.ve@mailcert.avvocaturastato.it"),
        qualifica="",
    )
    fascicolo = SimpleNamespace(
        titolo="2026/332 - Marchetti c. MIM",
        controparte="Avvocatura Distrettuale di Stato di Venezia",
    )

    recipient = react_notifiche_legali_bridge._recipient_from_subject(
        soggetto,
        ruolo="CONTROPARTE",
        fascicolo=fascicolo,
    )

    assert recipient["ruolo"] == "difensore"
    assert recipient["ruoloPratica"] == "difensore"
    assert recipient["fontePecSuggerita"] == "reginde"
    assert recipient["parteRappresentata"] == "Ministero dell'Istruzione e del Merito"


def test_destinatario_avvocatura_generico_non_inventa_la_parte_rappresentata():
    soggetto = SimpleNamespace(
        id="avvocatura-venezia",
        tipo=SimpleNamespace(value="PERSONA_GIURIDICA"),
        nome_completo="Avvocatura Distrettuale dello Stato di Venezia",
        ragione_sociale="",
        identificativo="94026160278",
        recapiti=SimpleNamespace(pec="ads.ve@mailcert.avvocaturastato.it"),
        qualifica="",
    )

    recipient = react_notifiche_legali_bridge._recipient_from_subject(soggetto)

    assert recipient["ruolo"] == "difensore"
    assert recipient["fontePecSuggerita"] == "reginde"
    assert recipient["parteRappresentata"] == ""


def test_matrice_destinatari_automatici_copre_tutte_le_categorie_e_registri_ammessi():
    cases = (
        ("Soggetto di prova", "PERSONA_FISICA", "CONTROPARTE", "", "persona@domiciliodigitale.test", "controparte", "inad"),
        ("Società di prova", "PERSONA_GIURIDICA", "CONTROPARTE", "", "societa@pec.impresa.test", "impresa", "ini_pec"),
        ("Ente privato", "ENTE", "CONTROPARTE", "", "ente@pec.test", "impresa", "ini_pec"),
        ("Professionista", "PROFESSIONISTA", "CONTROPARTE", "commercialista", "studio@pec.test", "professionista", "ini_pec"),
        ("Avvocato", "PERSONA_FISICA", "DIFENSORE", "avvocato", "avvocato@pec.test", "difensore", "reginde"),
        ("Ente pubblico", "PUBBLICA_AMMINISTRAZIONE", "CONTROPARTE", "", "protocollo@ente.gov.it", "pa", "registro_ppaa"),
        ("Ministero", "PERSONA_GIURIDICA", "CONTROPARTE", "", "notifiche@pec.istruzione.it", "pa", "registro_ppaa"),
        ("MIM - USP Milano", "PERSONA_GIURIDICA", "CONTROPARTE", "", "uspmi@postacert.istruzione.it", "pa", "registro_ppaa"),
        ("USR Catanzaro", "PERSONA_GIURIDICA", "CONTROPARTE", "", "uspcz.contenzioso@postacert.istruzione.it", "pa", "registro_ppaa"),
        ("Città Metropolitana di Reggio Calabria", "PERSONA_GIURIDICA", "CONTROPARTE", "", "protocollo@pec.cittametropolitana.rc.it", "pa", "registro_ppaa"),
        ("Agenzia delle Entrate-Riscossione", "PERSONA_GIURIDICA", "CONTROPARTE", "", "protocollo@pec.agenziariscossione.gov.it", "pa", "registro_ppaa"),
        ("Terzo", "PERSONA_FISICA", "TERZO", "", "terzo@domiciliodigitale.test", "terzo", "inad"),
    )

    for nome, tipo, ruolo, qualifica, pec, expected_role, expected_register in cases:
        soggetto = SimpleNamespace(
            tipo=SimpleNamespace(value=tipo),
            nome_completo=nome,
            ragione_sociale="",
            qualifica=qualifica,
        )
        inferred_role = react_notifiche_legali_bridge._infer_recipient_role(soggetto, ruolo, pec)
        inferred_register = react_notifiche_legali_bridge._infer_public_register(soggetto, ruolo, pec)

        assert inferred_role == expected_role
        assert inferred_register == expected_register
        assert inferred_register in RECIPIENT_NOTIFICATION_DIRECTIVES[inferred_role]["allowed_registers"]


def test_suggerimenti_notifica_escludono_email_ordinaria_spacciata_per_pec():
    ordinary = SimpleNamespace(
        id="ordinary",
        tipo=SimpleNamespace(value="PERSONA_FISICA"),
        nome_completo="Valeria",
        ragione_sociale="",
        identificativo="",
        recapiti=SimpleNamespace(pec="valeria@gmail.com"),
        qualifica="",
    )
    certified = SimpleNamespace(
        id="certified",
        tipo=SimpleNamespace(value="PERSONA_GIURIDICA"),
        nome_completo="Alfa S.p.A.",
        ragione_sociale="Alfa S.p.A.",
        identificativo="01234567890",
        recapiti=SimpleNamespace(pec="alfa@pec.impresa.it"),
        qualifica="",
    )

    payload = build_react_notifiche_legali_payload(
        get_clienti=lambda: SimpleNamespace(tutti=lambda: []),
        get_fascicoli=lambda: SimpleNamespace(tutti=lambda archiviati=False: []),
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: [ordinary, certified]),
    )

    suggestions = payload["precompilazione"]["destinatari"]
    assert [item["id"] for item in suggestions] == ["certified"]
    assert suggestions[0]["fontePecSuggerita"] == "ini_pec"


def test_suggerimenti_notifica_non_applicano_limiti_arbitrari_alla_rubrica():
    subjects = [
        SimpleNamespace(
            id=f"subject-{index}",
            tipo=SimpleNamespace(value="PERSONA_GIURIDICA"),
            nome_completo=f"Soggetto {index}",
            ragione_sociale=f"Soggetto {index}",
            identificativo=f"{index:011d}",
            recapiti=SimpleNamespace(pec=f"soggetto{index}@pec.example.it"),
            qualifica="",
            tag=["pubblico-elenco:RegistroImprese"],
        )
        for index in range(275)
    ]

    payload = build_react_notifiche_legali_payload(
        get_clienti=lambda: SimpleNamespace(tutti=lambda: []),
        get_fascicoli=lambda: SimpleNamespace(tutti=lambda archiviati=False: []),
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: subjects),
    )

    suggestions = payload["precompilazione"]["destinatari"]
    assert len(suggestions) == 275
    assert all(item["fontePecSuggerita"] == "registro_imprese" for item in suggestions)


def test_destinatari_omonimi_con_pec_diverse_restano_distinti():
    recipients = [
        react_notifiche_legali_bridge._recipient_from_plain(
            recipient_id="uno",
            name="Ministero dell'Istruzione e del Merito",
            pec="ufficio.uno@postacert.istruzione.it",
            role="pa",
            source="registro_ppaa",
        ),
        react_notifiche_legali_bridge._recipient_from_plain(
            recipient_id="due",
            name="Ministero dell'Istruzione e del Merito",
            pec="ufficio.due@postacert.istruzione.it",
            role="pa",
            source="registro_ppaa",
        ),
    ]

    merged = react_notifiche_legali_bridge._merge_notification_recipients(recipients)

    assert len(merged) == 2
    assert {item["pec"] for item in merged} == {
        "ufficio.uno@postacert.istruzione.it",
        "ufficio.due@postacert.istruzione.it",
    }


def test_destinatario_avvocatura_plain_usa_ruolo_e_registro_difensore():
    recipient = react_notifiche_legali_bridge._recipient_from_plain(
        recipient_id="avvocatura",
        name="Avvocatura Distrettuale dello Stato di Venezia",
    )

    assert recipient["ruolo"] == "difensore"
    assert recipient["fontePecSuggerita"] == "reginde"


def test_payload_react_notifiche_legali_deriva_parti_rg_destinatari_e_nomi_import_pratiche():
    assert "QuickOrganizer" not in react_notifiche_legali_bridge._display_text("QuickOrganizer")
    assert "Studio Telematico" not in react_notifiche_legali_bridge._display_text("Studio Telematico")
    sanitized_result = react_notifiche_legali_bridge.sanitize_react_notifiche_legali_payload(
        {"blockers": ["QuickOrganizer DatiAtto.xml TAVOLA Studio Telematico"]}
    )
    serialized_result = json.dumps(sanitized_result, ensure_ascii=False)
    assert "QuickOrganizer" not in serialized_result
    assert "DatiAtto.xml" not in serialized_result
    assert "TAVOLA" not in serialized_result
    assert "Studio Telematico" not in serialized_result

    documento = SimpleNamespace(
        id="doc-quick-1",
        nome_originale="20260510185021337.PDF",
        nome_portale="",
        nome="20260510185021337.PDF",
        percorso="20260510185021337.PDF",
        tipo_atto_portale="",
        classificazione_portale="Gestionale precedente",
        note=(
            "Import pratiche. Ricorso Lisciotto.pdf - Depositante: Studio. "
            "PEC: ads.rc@mailcert.avvocaturastato.it; dgosv@postcert.istruzione.it"
        ),
        fonte_documento="PORTALE_TELEMATICO",
        servizio_portale="PST QuickOrganizer DatiAtto.xml TAVOLA Studio Telematico",
        hash_sha256="f" * 64,
        data_documento="2026-05-10",
        data_deposito_portale="",
        id_documento_portale="",
        tags=[],
    )
    fascicolo = SimpleNamespace(
        id="fascicolo-quick",
        numero="2026/332",
        titolo="RG 2048/2025 - Lorenzetto II c. MIM",
        id_cliente="",
        nome_cliente="",
        controparte="",
        cf_controparte="",
        tribunale="Tribunale di Milano",
        sezione="",
        numero_rg="",
        anno_rg=0,
        giudice="",
        tipo_procedimento="",
        documenti=[documento],
        depositi_pct=[],
        note="",
        oggetto="",
    )

    practice_payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_payload(
        "fascicolo-quick",
        get_clienti=lambda: SimpleNamespace(tutti=lambda: [], get=lambda _id: None),
        get_fascicoli=lambda: SimpleNamespace(get=lambda _id: fascicolo),
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: [], parti_fascicolo=lambda id_fascicolo: []),
    )

    pratica = practice_payload["pratica"]
    recipient_names = {item["nome"] for item in pratica["destinatari"]}
    recipient_pecs = {item["pec"] for item in pratica["destinatari"] if item["pec"]}
    documento_payload = pratica["documenti"][0]

    assert pratica["assistitoNome"] == "Lorenzetto II"
    assert pratica["controparte"] == "MIM"
    assert pratica["procedimento"]["numeroRg"] == "2048"
    assert pratica["procedimento"]["annoRg"] == "2025"
    assert "MIM" not in recipient_names
    assert "ads.rc@mailcert.avvocaturastato.it" in recipient_pecs
    assert "dgosv@postcert.istruzione.it" in recipient_pecs
    assert all(item["pec"] for item in pratica["destinatari"])
    assert documento_payload["nomeFile"] == "Ricorso Lisciotto.pdf"
    assert documento_payload["label"] == "Ricorso Lisciotto.pdf"
    assert "QuickOrganizer" not in documento_payload["label"]
    serialized_payload = json.dumps(practice_payload, ensure_ascii=False)
    assert "QuickOrganizer" not in serialized_payload
    assert "DatiAtto.xml" not in serialized_payload
    assert "TAVOLA" not in serialized_payload
    assert "Studio Telematico" not in serialized_payload


def test_payload_react_notifiche_legali_deriva_rg_da_numero_fascicolo():
    fascicolo = SimpleNamespace(
        id="fascicolo-rg",
        numero="RG 466/2023",
        titolo="Alessi Robertino c. Zurich",
        id_cliente="",
        nome_cliente="Alessi Robertino",
        controparte="Zurich",
        cf_controparte="",
        tribunale="Giudice di Pace di Palmi",
        sezione="",
        numero_rg="",
        anno_rg=2026,
        giudice="",
        tipo_procedimento="",
        documenti=[],
        depositi_pct=[],
        note="",
        oggetto="",
    )

    payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_payload(
        "fascicolo-rg",
        get_clienti=lambda: SimpleNamespace(tutti=lambda: [], get=lambda _id: None),
        get_fascicoli=lambda: SimpleNamespace(get=lambda _id: fascicolo),
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: [], parti_fascicolo=lambda id_fascicolo: []),
    )

    procedimento = payload["pratica"]["procedimento"]

    assert procedimento["numeroRg"] == "466"
    assert procedimento["annoRg"] == "2023"


def test_payload_documenti_pratica_idrata_nome_timestamp_da_contenuto(monkeypatch):
    documento = SimpleNamespace(
        id="doc-ocr",
        nome="20260510185021337.PDF",
        nome_originale="20260510185021337.PDF",
        nome_portale="",
        percorso="fascicolo-1\\20260510185021337.PDF",
        tipo_atto_portale="",
        classificazione_portale="Gestionale precedente",
        note="Import pratiche. ",
        fonte_documento="IMPORT_ESTERNO",
        servizio_portale="",
        hash_sha256="",
        data_documento="",
        data_deposito_portale="",
        id_documento_portale="",
        tags=["import-pratiche"],
    )
    fascicolo = SimpleNamespace(id="fascicolo-1", documenti=[documento])

    monkeypatch.setattr(
        react_notifiche_legali_bridge,
        "_legacy_import_content_label",
        lambda _documento: "Contratto individuale di lavoro a tempo determinato",
    )

    payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_documents_payload(
        "fascicolo-1",
        get_fascicoli=lambda: SimpleNamespace(get=lambda id_fascicolo: fascicolo if id_fascicolo == "fascicolo-1" else None),
    )

    documento_payload = payload["documenti"][0]

    assert payload["ok"] is True
    assert documento_payload["label"] == "Contratto individuale di lavoro a tempo determinato"
    assert documento_payload["nomeFile"] == "20260510185021337.PDF"
    assert documento_payload["nomeOriginale"] == "20260510185021337.PDF"


def test_payload_documenti_pratica_rispetta_selezione_esplicita_oltre_primo_blocco():
    def documento(index: int) -> SimpleNamespace:
        return SimpleNamespace(
            id=f"doc-{index}",
            nome=f"Documento_{index}.pdf",
            nome_originale=f"Documento_{index}.pdf",
            nome_portale="",
            percorso="",
            tipo_atto_portale="",
            classificazione_portale="",
            note="",
            fonte_documento="PST",
            servizio_portale="PST",
            hash_sha256="",
            data_documento="",
            data_deposito_portale="",
            id_documento_portale=f"PORT-{index}",
            tags=[],
        )

    fascicolo = SimpleNamespace(id="fascicolo-1", documenti=[documento(index) for index in range(1, 56)])
    repo = SimpleNamespace(get=lambda id_fascicolo: fascicolo if id_fascicolo == "fascicolo-1" else None)

    default_payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_documents_payload(
        "fascicolo-1",
        get_fascicoli=lambda: repo,
    )
    selected_payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_documents_payload(
        "fascicolo-1",
        selected_document_ids=["doc-55", "PORT-45"],
        get_fascicoli=lambda: repo,
    )

    assert len(default_payload["documenti"]) == 40
    assert [item["id"] for item in selected_payload["documenti"]] == ["doc-45", "doc-55"]


def test_payload_pratica_notifiche_risolve_alias_fascicolo_da_link_diretto():
    documento = SimpleNamespace(
        id="doc-reddito",
        nome="Autocertificazione reddito.PDF",
        nome_originale="Autocertificazione reddito.PDF",
        nome_portale="",
        percorso="",
        tipo_atto_portale="",
        classificazione_portale="",
        note="",
        fonte_documento="PST",
        servizio_portale="PST",
        hash_sha256="",
        data_documento="2026-07-18",
        data_deposito_portale="2026-07-18",
        id_documento_portale="PORT-REDDITO",
        tags=[],
    )
    fascicolo = SimpleNamespace(
        id="fascicolo-interno",
        id_pratica="78D6022C",
        numero="RG 1428/2026",
        titolo="Romeo Maria c. MIM",
        id_cliente="",
        nome_cliente="Romeo Maria",
        controparte="Ministero dell'Istruzione e del Merito",
        cf_controparte="",
        tribunale="TRIBUNALE DI PALMI",
        sezione="",
        numero_rg="1428",
        anno_rg="2026",
        giudice="",
        tipo_procedimento="",
        oggetto="",
        note="",
        documenti=[documento],
    )
    repo = SimpleNamespace(
        get=lambda id_fascicolo: None,
        tutti=lambda archiviati=False: [fascicolo],
    )

    practice_payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_payload(
        "78D6022C",
        get_clienti=lambda: SimpleNamespace(tutti=lambda: [], get=lambda _id: None),
        get_fascicoli=lambda: repo,
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: [], parti_fascicolo=lambda _id: []),
    )
    documents_payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_documents_payload(
        "78D6022C",
        selected_document_ids=["PORT-REDDITO"],
        get_fascicoli=lambda: repo,
    )

    assert practice_payload["ok"] is True
    assert practice_payload["pratica"]["id"] == "fascicolo-interno"
    assert practice_payload["pratica"]["documenti"][0]["nomeFile"] == "Autocertificazione reddito.PDF"
    assert documents_payload["ok"] is True
    assert [item["id"] for item in documents_payload["documenti"]] == ["doc-reddito"]


def test_deriva_titolo_documento_da_testo_ocr():
    text = """
    ISTITUTO COMPRENSIVO IC 2 E 4 DI VICENZA
    Oggetto: contratto individuale di lavoro a tempo determinato stipulato tra il
    Dirigente scolastico e il sig. Rossi Mario
    """

    assert (
        react_notifiche_legali_bridge._derive_document_title_from_text(text)
        == "Contratto individuale di lavoro a tempo determinato"
    )


def test_deriva_titoli_documenti_carta_docente_da_testo_ocr():
    richiesta = """
    Oggetto: Richiesta pagamento annualità “CARTA DEL DOCENTE” Grosso Angelo Eugenio, C.F.
    """
    ricorso = """
    STUDIO LEGALE
    Ricorso per il recupero della cd. carta del docente
    Tribunale di Milano
    """

    assert react_notifiche_legali_bridge._derive_document_title_from_text(richiesta) == "Richiesta pagamento Carta del docente"
    assert react_notifiche_legali_bridge._derive_document_title_from_text(ricorso) == "Ricorso per il recupero della Carta del docente"


def test_deriva_titoli_provvedimenti_carta_docente_da_ocr_reale():
    ricorso_lavoro = """
    STUDIO LEGALE MONTAGNESE Tribunale Civile di Vicenza SEZIONE LAVORO
    Ricorso ex Art. 414 C.p.c. contro MINISTERO DELL'ISTRUZIONE E DEL MERITO.
    OGGETTO: Diritto insegnanti precari e personale educativo ad usufruire del beneficio previsto dall'art. 1 della Legge n. 107/2015.
    """
    cassazione = """
    REPUBBLICA ITALIANA IN NOME DEL POPOLO ITALIANO LA CORTE SUPREMA DI CASSAZIONE SEZIONE LAVORO.
    Oggetto: Composta dagli Ill.mi Sigg.ri Magistrati. CARTA DOCENTE. Ha pronunciato la seguente SENTENZA.
    """
    sentenza_lavoro = """
    REPUBBLICA ITALIANA TRIBUNALE ORDINARIO di VICENZA SETTORE LAVORO.
    Ha pronunciato la seguente SENTENZA. Oggetto: Altre ipotesi. docente a tempo determinato.
    """

    assert react_notifiche_legali_bridge._derive_document_title_from_text(ricorso_lavoro) == "Ricorso lavoro Carta del docente"
    assert react_notifiche_legali_bridge._derive_document_title_from_text(cassazione) == "Sentenza Cassazione Carta del docente"
    assert (
        react_notifiche_legali_bridge._derive_document_title_from_text(sentenza_lavoro)
        == "Sentenza lavoro personale docente a tempo determinato"
    )


def test_matrice_casi_notifica_legge_documento_e_non_nome_file():
    cases = {
        "OPPOSIZIONE A DECRETO INGIUNTIVO n. 18/2026": ("opposizione_decreto_ingiuntivo", "relata_opposizione_decreto_ingiuntivo"),
        "ATTO DI PRECETTO su titolo esecutivo": ("titolo_esecutivo_precetto", "relata_titolo_esecutivo_precetto"),
        "PIGNORAMENTO PRESSO TERZI": ("pignoramento_presso_terzi", "relata_pignoramento_presso_terzi"),
        "ATTO DI APPELLO avverso sentenza": ("appello_impugnazione", "relata_appello_impugnazione"),
        "RICORSO IN RIASSUNZIONE": ("riassunzione", "relata_riassunzione"),
        "CHIAMATA IN CAUSA DEL TERZO": ("chiamata_terzo", "relata_chiamata_terzo"),
        "INTEGRAZIONE DEL CONTRADDITTORIO": ("integrazione_contraddittorio", "relata_integrazione_contraddittorio"),
        "RICORSO PER SEPARAZIONE CON FIGLIO MINORE": ("famiglia_persone_minori", "relata_famiglia_persone_minori"),
        "PROVVEDIMENTO URGENTE CAUTELARE": ("provvedimento_urgente", "relata_provvedimento_urgente"),
        "ACCORDO TRANSATTIVO E DIFFIDA": ("accordo_transazione_stragiudiziale", "relata_accordo_transazione_stragiudiziale"),
    }

    for text, (case_id, template_id) in cases.items():
        kind = react_notifiche_legali_bridge._kind_from_evidence(text, content=True)
        suggestion = react_notifiche_legali_bridge._notification_suggestion_from_kind(kind, "testo documento letto")
        assert suggestion["casoNotificaSuggerito"] == case_id
        assert suggestion["modelloRelataSuggerito"] == template_id
        assert suggestion["criterioTipoDocumento"] == "testo documento letto"


def test_ui_notifiche_legali_firma_relata_direttamente_nel_flusso_operativo():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")

    assert "signRelata" in page
    assert "signCurrentRelataWithLocalSigner" in page
    assert "options: { refreshControl?: boolean } = {}" in page
    assert "if (options.refreshControl !== false)" in page
    assert 'className="iu-legal-signature-pin"' in page
    assert 'aria-label="PIN del dispositivo di firma"' in page
    assert "Il PIN resta su questo PC e viene cancellato dopo la firma della relata." in page
    assert "Firma relata" in page
    assert "Carica relata firmata" not in page
    assert "relata_firmata_sha256" in page
    assert "setNotifica((current) => ({ ...current, ...notificaOverrides, relata_firmata: true }))" in page
    control_key_block = page[page.index("function notificationControlPayloadKey"):page.index("function relataLocalSignerBaseUrl")]
    assert "delete controlled.relata_firmata_file" in control_key_block
    assert "delete controlled.relata_firmata_sha256" in control_key_block
    assert "delete controlled.relata_firmata_document_id" in control_key_block
    assert "delete controlled.relataFirmataDocumentId" in control_key_block
    assert "delete controlled.signedRelataDocumentId" in control_key_block


def test_ui_notifiche_legali_invia_pec_firma_relata_e_allega_prima_della_password():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")
    send_block = page[page.index("const sendNotification = async"):page.index("const verifyNotificationPec")]

    assert "updateLocalPecProgress('firma'" in send_block
    assert "signCurrentRelataWithLocalSigner(signatureOverrides, { refreshControl: false })" in send_block
    assert "updateLocalPecProgress('allegati'" in send_block
    assert "relata_firmata_file: signedRelataAttachment.fileName" in send_block
    assert "relata_firmata_sha256: signedRelataAttachment.sha256" in send_block
    assert "relata_firmata_document_id: signedRelataAttachment.documentId" in send_block
    assert "requestNotificationPecPassword(messages[0])" in send_block
    assert "ensureAutomaticPecVerification()" not in send_block
    assert "updateLocalPecProgress('preparazione'" not in send_block
    assert "updateLocalPecProgress('signer'" not in send_block
    assert "signedRelata?.fileName" not in send_block
    assert send_block.index("signCurrentRelataWithLocalSigner(signatureOverrides, { refreshControl: false })") < send_block.index("postLegalWorkflow(data.azioni.invioPecLocale, payload)")
    assert send_block.index("postLegalWorkflow(data.azioni.invioPecLocale, payload)") < send_block.index("requestNotificationPecPassword(messages[0])")


def test_ui_notifiche_legali_non_contiene_flusso_deposito():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")
    data_client = Path("frontend/src/notificheLegaliData.ts").read_text(encoding="utf-8")
    templates = Path("pct/data/notifiche_legali_templates.json").read_text(encoding="utf-8")
    bridge = Path("web/services/react_notifiche_legali_bridge.py").read_text(encoding="utf-8")
    blueprint = Path("web/blueprints/api_v1_react.py").read_text(encoding="utf-8")

    assert "type TabKey = 'notifica' | 'unep' | 'nonpec' | 'cliente'" in page
    assert "tab === 'deposito'" not in page
    assert "setTab('deposito')" not in page
    assert "run('deposito')" not in page
    assert "provaDeposito" not in page
    assert "Deposito prova notifica" not in page
    assert "Prova della notifica" not in page
    assert "Pacchetto prova" not in page
    assert "setDeposito" not in page
    assert "selectedDeposit" not in page
    assert "provvedimentoDataDeposito" not in page
    assert "provvedimento_data_deposito" not in page
    assert "dataDeposito" not in page
    assert "depositoId" not in page
    assert "provvedimentoDataDeposito" not in data_client
    assert "provvedimento_data_deposito" not in data_client
    assert "dataDeposito" not in data_client
    assert "depositoId" not in data_client
    assert "provaDeposito" not in data_client
    assert "depositoChecklist" not in data_client
    assert "depositProofWithOriginalReceipts" not in data_client
    assert "workflow_deposito_area_web_pst" not in templates
    assert "Deposito area web PST" not in templates
    assert "percorso di deposito" not in templates
    assert "prova del deposito" not in templates
    assert '"provaDeposito"' not in bridge
    assert '"depositoChecklist"' not in bridge
    assert '"depositProofWithOriginalReceipts"' not in bridge
    assert 'automation_payload.pop("deposito", None)' not in bridge
    assert '@api_v1_react.post("/notifiche-legali/prova-deposito")' not in blueprint


def test_ui_notifiche_legali_non_preseleziona_vecchi_destinatari_manuali():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")
    apply_practice_block = page[page.index("const applyPractice ="):page.index("setModelFields((current)", page.index("const applyPractice ="))]

    assert "function isManualNotificationRecipient" in page
    assert "const destinatariAutomatici = practice.destinatari.filter((item) => Boolean(item.pec) && !isManualNotificationRecipient(item))" in page
    assert "setSelectedRecipientIds(primoDestinatario?.id ? [primoDestinatario.id] : [])" in page
    assert "destinatario_pec: primoDestinatario?.pec || ''" in apply_practice_block
    assert "applyRecipient(savedRecipient)" in page
    assert "setSelectedRecipientIds((current) => current.includes(recipient.id) ? current : [...current, recipient.id])" in page
    assert "destinatariCompleti.map((item) => item.id).filter(Boolean)" not in page
    assert "destinatario_pec: primoDestinatario?.pec || current.destinatario_pec" not in apply_practice_block


def test_ui_notifiche_legali_ogni_controllo_porta_esito_in_vista():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")

    assert "const scrollResultIntoView" in page
    assert "scrollIntoView({ behavior: 'smooth', block: 'start' })" in page
    result_index = page.index("setResult(response)")
    scroll_index = page.index("scrollResultIntoView()", result_index)
    working_index = page.index("setWorking(false)", result_index)
    assert result_index < scroll_index < working_index


def test_ui_notifiche_legali_mostra_data_catalogo_in_formato_italiano():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")

    assert "function templateVersionDate" in page
    assert "function templateVersionLabel" in page
    assert "aggiornato il ${date}" in page
    assert "Aggiornato il ${templateVersionDate(data.templateCatalogVersion)}" in page
    assert "Versione ${data.templateCatalogVersion}" not in page
    assert "` - ${result.templateVersion}`" not in page


def test_ui_notifiche_legali_mostra_audit_in_ora_italiana():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")
    audit_block = page[page.index("function auditDateTimeText"):page.index("function deliveryPlan")]
    result_block = page[page.index("<span>Audit</span>"):page.index("{result.body ?")]

    assert "formatDateTimeIt(raw, raw)" in audit_block
    assert "auditDateTimeText(auditTrail(result.outputPlan)?.generatedAt)" in result_block
    assert "auditText(auditTrail(result.outputPlan)?.generatedAt)" not in result_block


def test_ui_notifiche_legali_rende_automatici_i_controlli_non_decisionali():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")

    assert "Verifica automatica delle PEC" not in page
    assert "Relata separata predisposta automaticamente" in page
    assert "Attestazione di conformità prodotta quando richiesta" in page
    assert "Modifica attestazione di conformità" in page
    assert "Salva attestazione per questa notifica" in page
    assert "Precisazione facoltativa dell'avvocato" not in page
    assert "Ricevuta completa prevista automaticamente" not in page
    assert "Approvazione finale dell'avvocato prima dell'invio" not in page
    assert "Conferma avvocato registrata" not in page
    assert "Lo stesso PIN verifica le PEC" not in page
    assert "Avvocato abilitato alla notifica in proprio" not in page
    assert "Data e ora verifica PEC" not in page
    assert "checked={notifica.ricevuta_completa}" not in page
    assert "checked={notifica.relata_documento_separato}" not in page


def test_ui_notifiche_legali_ha_tre_ingressi_documenti_senza_wizard_inutile():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")

    assert "documentEntrySourceFromUrl" in page
    assert "requestedDocumentEntrySource === 'presidio'" in page
    assert "const notificationDocumentEntryMode = requestedDocumentEntrySource === 'presidio'" in page
    assert "const sourceIsPresidio = requestedDocumentEntrySource === 'presidio'" in page
    assert "isPresidioNotificationDocument" in page
    assert "documenti importati`} dal presidio notifiche." in page
    assert "<em>Fascicolo</em>" in page
    assert "<em>Presidio</em>" in page
    assert "<em>Manuale</em>" in page
    assert "Scegli i documenti del fascicolo: solo quelli spuntati entrano nella relata." in page
    assert "Tutti notificabili" not in page
    assert "Usali solo quando il documento non è già nel fascicolo." in page
    assert "openDocumentPreview(documento)" in page


def test_ui_notifiche_legali_attestazione_solo_da_documenti_inclusi():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")
    block = page[
        page.index("const notificationNeedsAttestazione"):
        page.index("const guidedAutomationSteps")
    ]

    assert "selectedOrigin?.needsAttestazione" not in block
    assert "selectedNotificationDocuments.some((item) => item.necessitaAttestazione)" in block
    assert "manualNotificationDocuments.some((item) => originNeedsAttestazione(item.origine))" in block


def test_ui_notifiche_legali_pec_manuale_e_rimozione_documenti_relata():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")

    assert "const [manualRecipientSuggestions, setManualRecipientSuggestions]" in page
    assert "const [manualRecipientSaving, setManualRecipientSaving]" in page
    assert "const [manualRecipientDraft, setManualRecipientDraft]" in page
    assert "saveLegalManualRecipient(data.azioni.salvaDestinatarioManuale" in page
    assert "Salvataggio destinatario manuale nello studio in corso" in page
    assert "Destinatario PEC manuale aggiunto alla notifica" not in page
    assert 'className="iu-legal-manual-recipient iu-legal-field--wide"' in page
    assert "Inserimento manuale destinatario" in page
    assert "Aggiungi destinatario manuale" in page
    assert "changeManualRecipientDraft('pec'" in page
    assert "ruoloPratica: 'Inserito manualmente'" in page
    assert "item.ruoloPratica === 'Inserito manualmente' ? item.id" in page
    assert "Aggiungi PEC manuale alla notifica" in page
    assert "Dopo l’invio: ricevute attese dal presidio PEC" not in page
    assert "Presidio notifiche collegato" not in page
    assert "Archivio automatico nel fascicolo" not in page
    assert "const removeFinalRelataRow = (row: FinalRelataRow)" in page
    assert "removableKind: 'fascicolo' as const" in page
    assert "removableKind: 'manuale' as const" in page
    assert "onClick={() => removeFinalRelataRow(row)}" in page
    assert "setSelectedNotificationDocumentIds([]); clearNotificationDocumentFields()" in page
    assert "delete controlled.approvazione_avvocato" in page
    data_client = Path("frontend/src/notificheLegaliData.ts").read_text(encoding="utf-8")
    assert "salvaDestinatarioManuale" in data_client
    assert "/api/v1/ui/notifiche-legali/destinatari-manuali" in data_client


def test_ui_notifiche_legali_verifica_i_dati_visibili_del_destinatario_attivo():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")

    assert "const isActive = recipient.id === selectedRecipientId" in page
    assert "isActive ? notifica.fonte_pec_destinatario || recipient.fontePecSuggerita" in page
    assert "isActive ? notifica.destinatario_pec || recipient.pec" in page


def test_ui_notifiche_legali_allinea_caso_modello_e_mostra_il_blocco_locale_reale():
    page = Path("frontend/src/components/NotificheLegaliPage.tsx").read_text(encoding="utf-8")

    assert "item.templateId === practice.modelloSuggerito" in page
    assert "caso_notifica: suggestedCase?.value || current.caso_notifica" in page
    assert "syncNotificaFromNotificationDocuments(rows)" in page
    assert "documento.casoNotificaSuggerito" in page
    assert "Tipo letto da {documento.criterioTipoDocumento}" in page
    assert "pecVerificationMessage(raw)" in page
    assert "Il certificato di autenticazione del dispositivo non è disponibile." in page


def test_payload_react_notifiche_legali_segnala_pec_ufficio_da_collegare(monkeypatch):
    fascicolo = SimpleNamespace(
        id="fascicolo-portale",
        numero="2026/002",
        titolo="Cliente / Beta",
        id_cliente="",
        nome_cliente="Cliente",
        controparte="Beta S.p.A.",
        cf_controparte="",
        tribunale="Tribunale di Roma",
        sezione="",
        numero_rg="5678",
        anno_rg=2026,
        giudice="",
        tipo_procedimento="civile ordinario",
        documenti=[],
        depositi_pct=[
            SimpleNamespace(
                id="dep-relata",
                id_deposito_esterno="PST-REL-2026-0002",
                fonte="PST",
                servizio_portale="ConsultazioneFascicolo",
                tipo_atto="Ordinanza da notificare",
                data_deposito="2026-05-23",
                mittente="Tribunale di Roma",
                documenti_portale=[
                    {
                        "id_documento": "pst-doc-relata",
                        "nome": "ordinanza_da_notificare.pdf",
                        "tipo": "Ordinanza da notificare",
                        "data_deposito": "2026-05-23",
                        "disponibile": True,
                    }
                ],
            )
        ],
    )
    pec = SimpleNamespace(
        id="pec-cancelleria-5678",
        mittente="cancelleria.tribunale.roma@giustiziacert.it",
        oggetto="Tribunale di Roma R.G. 5678/2026 - ordinanza da notificare",
        data="2026-05-23T10:15:00",
        corpo_testo="Si comunica il provvedimento ordinanza_da_notificare.pdf da notificare nel procedimento R.G. 5678/2026.",
        allegati=[{"nome": "ordinanza_da_notificare.pdf", "sha256": "b" * 64}],
        message_id="<pec-cancelleria-5678@giustizia>",
    )
    monkeypatch.setattr("web.services.react_notifiche_legali_bridge._office_pec_messages", lambda: [pec])

    initial_payload = build_react_notifiche_legali_payload(
        get_clienti=lambda: SimpleNamespace(tutti=lambda: []),
        get_fascicoli=lambda: SimpleNamespace(tutti=lambda archiviati=False: [fascicolo]),
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: [], parti_fascicolo=lambda id_fascicolo: []),
    )
    payload = react_notifiche_legali_bridge.build_react_notifiche_legali_practice_payload(
        "fascicolo-portale",
        get_clienti=lambda: SimpleNamespace(tutti=lambda: [], get=lambda _id: None),
        get_fascicoli=lambda: SimpleNamespace(get=lambda _id: fascicolo),
        get_soggetti=lambda: SimpleNamespace(tutti=lambda: [], parti_fascicolo=lambda id_fascicolo: []),
    )

    pratica = payload["pratica"]

    assert initial_payload["contracts"]["officeDocumentPecEvidence"] is True
    assert pratica["documentoUfficioMonitor"]["stato"] == "da_acquisire"
    assert pratica["documentoUfficioMonitor"]["documentiDaAcquisire"] == 1
    assert pratica["documentoUfficioMonitor"]["documentiRilasciati"][0]["nome"] == "ordinanza_da_notificare.pdf"
    assert pratica["documentoUfficioMonitor"]["documentiRilasciati"][0]["pecHref"] == "/email/messaggio/pec-cancelleria-5678"
    assert "single_document=1" in pratica["documentoUfficioMonitor"]["documentiRilasciati"][0]["acquisitionHref"]
    assert "documento=ordinanza_da_notificare.pdf" in pratica["documentoUfficioMonitor"]["documentiRilasciati"][0]["acquisitionHref"]
    assert "non_duplicare_documenti=1" in pratica["documentoUfficioMonitor"]["documentiRilasciati"][0]["acquisitionHref"]
    assert "id_fasc=fascicolo-portale" in pratica["portaleAcquisizioneHref"]
    assert "numero=5678" in pratica["portaleAcquisizioneHref"]
