"""Estrazione testo best-effort per documenti leggibili da Lex."""

from __future__ import annotations

import json
import os
import re
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from email import policy
from email.message import EmailMessage, Message
from email.parser import BytesParser
from email.utils import parsedate_to_datetime
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from shutil import which
from typing import Any
from xml.etree import ElementTree as ET

from .models import DocumentAIPageText
from .pdf_quality import repair_pdf_cid_placeholders, score_extracted_text_quality
from pct.formatting import format_datetime_it


_TEXT_EXTENSIONS = {"txt", "xml", "json", "csv"}
_HTML_EXTENSIONS = {"html", "htm"}
_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "tif", "tiff", "bmp", "gif"}


def _format_email_datetime_visible(value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    try:
        return format_datetime_it(parsedate_to_datetime(raw), include_timezone=True)
    except (TypeError, ValueError, IndexError, OverflowError):
        return format_datetime_it(raw, include_timezone=True) or raw


@dataclass(slots=True)
class ExtractionResult:
    ok: bool
    text: str
    pages: list[DocumentAIPageText]
    extraction_engine: str
    warnings: list[str] = field(default_factory=list)
    error_code: str = ""
    error_message: str = ""


@dataclass(slots=True)
class DocumentAITextExtractionResult:
    text: str
    pages: list[DocumentAIPageText]
    extraction_engine: str
    page_count: int | None
    warnings: list[str] = field(default_factory=list)
    error: str | None = None


def extract_document_text(file_path: str | Path, file_type: str) -> DocumentAITextExtractionResult:
    try:
        content = Path(file_path).read_bytes()
    except OSError as exc:
        return DocumentAITextExtractionResult(
            text="",
            pages=[],
            extraction_engine="read_failed",
            page_count=None,
            warnings=["File non leggibile per l'estrazione testo."],
            error=str(exc),
        )
    result = extract_text_from_document(content, Path(file_path).name, file_type)
    return DocumentAITextExtractionResult(
        text=result.text,
        pages=result.pages,
        extraction_engine=result.extraction_engine,
        page_count=len(result.pages) if result.pages else None,
        warnings=list(result.warnings),
        error=None if result.ok else result.error_message or result.error_code or "Estrazione non completata.",
    )


def extract_text_from_document(content: bytes, filename: str, file_type: str) -> ExtractionResult:
    if _is_cades_signature_name(filename):
        payload, payload_name, unwrap_warnings = _unwrap_p7m_payload(content, filename)
        payload_type = _file_type_from_payload(payload, payload_name, fallback=file_type)
        result = extract_text_from_document(payload, payload_name, payload_type)
        if not result.ok:
            fallback = _extract_binary_best_effort(
                payload,
                payload_type or "cades",
                base_warnings=list(result.warnings),
            )
            if fallback.ok:
                result = fallback
        result.warnings[:0] = unwrap_warnings
        if result.extraction_engine and not result.extraction_engine.startswith("cades:"):
            result.extraction_engine = f"cades:{result.extraction_engine}"
        return result

    ext = str(file_type or "").lower().lstrip(".")
    if ext == "pdf":
        if not _looks_like_pdf(content):
            return _extract_mislabeled_pdf(content, filename)
        unlimited = _extract_with_unlimited_ocr_for_index(content, filename, ext)
        if unlimited is not None:
            return unlimited
        return _extract_pdf(content)
    if ext == "docx":
        return _extract_docx(content)
    if ext == "doc":
        return _extract_doc(content)
    if ext in _TEXT_EXTENSIONS:
        return _extract_textual_document(content, ext)
    if ext in _HTML_EXTENSIONS:
        return _extract_html_document(content)
    if ext == "rtf":
        return _extract_rtf_document(content)
    if ext == "odt":
        return _extract_odt(content)
    if ext in {"xlsx", "xls"}:
        return _extract_spreadsheet(content, ext)
    if ext in _IMAGE_EXTENSIONS:
        unlimited = _extract_with_unlimited_ocr_for_index(content, filename, ext)
        if unlimited is not None:
            return unlimited
        return _extract_image(content, ext)
    if ext == "eml":
        return _extract_eml(content)
    if ext == "msg":
        return _extract_msg(content)
    if ext == "zip":
        return _extract_zip(content)
    return _extract_binary_best_effort(content, ext or "bin")


def _unwrap_p7m_payload(content: bytes, filename: str) -> tuple[bytes, str, list[str]]:
    inner_name = str(filename or "").strip()
    if _is_cades_signature_name(inner_name):
        inner_name = inner_name[:-4]
    if _looks_like_pdf(content) and inner_name.lower().endswith(".pdf"):
        return content, inner_name, ["Documento con estensione .p7m: PDF firmato letto come PDF interno per l'indice Lex."]
    try:
        from pct.firme_cades import inspect_signed_document_bytes

        signed = inspect_signed_document_bytes(source_name=filename, source_path="", data=content)
        if signed.status.payload_available and signed.payload_bytes:
            payload_name = signed.status.payload_name or inner_name
            return signed.payload_bytes, payload_name, [signed.status.message]
    except Exception:
        pass
    try:
        from pct.firme_cades import (
            extract_signed_payload,
            payload_extension_from_mime,
            payload_mime_from_bytes,
            payload_name_from_source,
        )

        payload = extract_signed_payload(content)
        if payload:
            mime_type = payload_mime_from_bytes(payload, inner_name)
            payload_name = payload_name_from_source(filename, payload_extension_from_mime(mime_type, inner_name))
            return payload, payload_name, ["Contenuto del documento firmato estratto per l'indice Lex."]
    except Exception:
        pass
    return content, inner_name, ["Documento firmato letto tentando il formato del file interno dichiarato."]


def _is_cades_signature_name(filename: Any) -> bool:
    return str(filename or "").strip().lower().endswith((".p7m", ".pm7"))


def _file_type_from_payload(content: bytes, filename: str, *, fallback: str) -> str:
    lower_name = str(filename or "").strip().lower()
    for extension in (
        "pdf",
        "docx",
        "doc",
        "txt",
        "xml",
        "json",
        "csv",
        "html",
        "htm",
        "rtf",
        "odt",
        "xlsx",
        "xls",
        "png",
        "jpg",
        "jpeg",
        "tif",
        "tiff",
        "bmp",
        "gif",
        "eml",
        "msg",
        "zip",
        "p7m",
        "pm7",
    ):
        if lower_name.endswith(f".{extension}"):
            return extension
    if _looks_like_pdf(content):
        return "pdf"
    if _looks_like_image(content):
        return _image_type_from_magic(content)
    if bytes(content[:256]).lstrip().startswith((b"<?xml", b"<")):
        return "xml"
    if content.startswith(b"PK\x03\x04"):
        return "zip"
    return str(fallback or "bin").lower().lstrip(".") or "bin"


def _looks_like_pdf(content: bytes) -> bool:
    return bytes(content[:16] if content else b"").lstrip().startswith(b"%PDF")


def _extract_mislabeled_pdf(content: bytes, filename: str) -> ExtractionResult:
    """Non invia payload non-PDF ai parser PDF; recupera ZIP/XML/testo se reali."""

    warning = "File indicato come PDF ma il contenuto non è un PDF valido: parser PDF saltato."
    if bytes(content[:8] if content else b"").startswith(b"PK\x03\x04"):
        result = _extract_zip(content)
        result.warnings.insert(0, warning)
        result.extraction_engine = f"pdf-mismatch:{result.extraction_engine}"
        return result
    stripped = bytes(content[:512] if content else b"").lstrip()
    if stripped.startswith((b"<?xml", b"<")):
        result = _extract_textual_document(content, "xml")
        result.warnings.insert(0, warning)
        result.extraction_engine = f"pdf-mismatch:{result.extraction_engine}"
        return result
    result = _extract_binary_best_effort(content, "pdf-mismatch", base_warnings=[warning])
    if result.ok:
        return result
    return ExtractionResult(
        ok=False,
        text="",
        pages=[],
        extraction_engine="pdf_magic_mismatch",
        warnings=_unique_warnings(result.warnings),
        error_code="pdf_magic_mismatch",
        error_message=f"{Path(filename or 'documento.pdf').name}: il contenuto non è un PDF leggibile.",
    )


def _looks_like_image(content: bytes) -> bool:
    return _image_type_from_magic(content) != ""


def _extract_with_unlimited_ocr_for_index(content: bytes, filename: str, ext: str) -> ExtractionResult | None:
    """Usa legal_ocr/Unlimited-OCR come sorgente integrale per l'indice fascicolo."""

    try:
        from concurrent.futures import ThreadPoolExecutor, as_completed

        from legal_ocr.preprocessing import rasterize_document
        from legal_ocr.unlimited.client import UnlimitedOcrClient
        from legal_ocr.unlimited.config import UnlimitedOcrSettings
        from legal_ocr.unlimited_ocr import split_native_and_ocr_pages
    except Exception:
        return None

    settings = UnlimitedOcrSettings.from_env()
    readiness = settings.readiness()
    if not readiness.get("ok"):
        return None

    try:
        with tempfile.TemporaryDirectory(prefix="iusentra-document-ai-unlimited-") as tmpdir:
            pages = rasterize_document(content, filename or f"documento.{ext}", Path(tmpdir))
            if len(pages) > settings.max_pages:
                return None
            native_pages, ocr_pages, page_warnings = split_native_and_ocr_pages(
                pages,
                max_pages=settings.max_pages,
                max_image_bytes=settings.max_image_bytes,
            )
            text_by_page: dict[int, list[str]] = {}

            def _add_page_text(page, text: str) -> None:
                cleaned = _clean_text(text)
                if not cleaned:
                    return
                page_number = _source_page_number(page)
                bucket = text_by_page.setdefault(page_number, [])
                if cleaned not in bucket:
                    bucket.append(cleaned)

            for page in native_pages:
                _add_page_text(page, page.text_hint)
            errors: list[str] = []
            if ocr_pages:
                client = UnlimitedOcrClient(settings)

                def _read_page(page):
                    result = client.generate_for_pages([page])
                    return page, _clean_text(str(result.get("text") or "")), str(result.get("error") or "")

                with ThreadPoolExecutor(max_workers=settings.concurrency) as executor:
                    future_map = {executor.submit(_read_page, page): page for page in ocr_pages}
                    for future in as_completed(future_map):
                        page, text, error = future.result()
                        if text:
                            _add_page_text(page, text)
                        else:
                            errors.append(f"Pagina {_source_page_number(page)}: Unlimited-OCR non ha restituito testo ({error or 'risposta vuota'}).")
            source_pages = sorted({_source_page_number(page) for page in pages})
            page_texts = [DocumentAIPageText(page_number=page_number, text="\n\n".join(text_by_page.get(page_number, []))) for page_number in source_pages]
    except Exception:
        return None

    if errors:
        return None
    missing_pages = [page.page_number for page in page_texts if not str(page.text or "").strip()]
    if missing_pages:
        return None
    full_text = "\n\n".join(page.text for page in page_texts if str(page.text or "").strip())
    if not full_text.strip():
        return None
    warnings = [
        "Indice fascicolo alimentato da OCR integrale legal_ocr/Unlimited-OCR: testo completo e mappa pagine, nessun chunk OCR.",
        *list(readiness.get("warnings") or []),
        *page_warnings,
    ]
    return ExtractionResult(
        ok=True,
        text=full_text,
        pages=page_texts,
        extraction_engine="legal-ocr:unlimited-ocr:document-index",
        warnings=_unique_warnings(warnings),
    )


def _source_page_number(page: Any) -> int:
    page_number = int(getattr(page, "page", 1) or 1)
    preprocessing = list(getattr(page, "preprocessing", []) or [])
    if page_number >= 10 and "page-split" in preprocessing:
        return max(1, page_number // 10)
    return page_number


def _image_type_from_magic(content: bytes) -> str:
    sample = bytes(content[:16] if content else b"")
    if sample.startswith(b"\x89PNG\r\n\x1a\n"):
        return "png"
    if sample.startswith(b"\xff\xd8\xff"):
        return "jpg"
    if sample.startswith((b"II*\x00", b"MM\x00*")):
        return "tif"
    if sample.startswith(b"BM"):
        return "bmp"
    if sample.startswith((b"GIF87a", b"GIF89a")):
        return "gif"
    return ""


def _extract_txt(content: bytes) -> ExtractionResult:
    return _extract_textual_document(content, "txt")


def _extract_textual_document(content: bytes, kind: str) -> ExtractionResult:
    text, encoding, warnings = _decode_text_bytes(content)
    normalized_kind = str(kind or "txt").lower().lstrip(".")
    engine = f"{normalized_kind}:{encoding}"
    if normalized_kind == "json" and text.strip():
        try:
            parsed = json.loads(text)
            text = json.dumps(parsed, ensure_ascii=False, indent=2)
            engine = f"json:{encoding}"
        except Exception:
            warnings.append("JSON letto come testo perché non è stato possibile normalizzarlo.")
    elif normalized_kind == "xml":
        engine = f"xml:{encoding}"
    elif normalized_kind == "csv":
        engine = f"csv:{encoding}"
    if not text.strip():
        warnings.append("Il file non contiene contenuto testuale leggibile.")
    return ExtractionResult(
        ok=True,
        text=text,
        pages=[],
        extraction_engine=engine,
        warnings=warnings,
    )


def _extract_html_document(content: bytes) -> ExtractionResult:
    text, encoding, warnings = _decode_text_bytes(content)
    readable = _html_to_text(text)
    if not readable.strip():
        warnings.append("Il documento HTML non contiene testo leggibile.")
    return ExtractionResult(
        ok=True,
        text=readable or text,
        pages=[],
        extraction_engine=f"html:{encoding}",
        warnings=warnings,
    )


def _extract_rtf_document(content: bytes) -> ExtractionResult:
    text = _extract_rtf_text(content)
    warnings: list[str] = []
    if not text.strip():
        decoded, _encoding, decode_warnings = _decode_text_bytes(content)
        warnings.extend(decode_warnings)
        text = _clean_text(re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", decoded).replace("{", " ").replace("}", " "))
    if not text.strip():
        warnings.append("Il documento RTF non contiene testo leggibile.")
    return ExtractionResult(
        ok=True,
        text=text,
        pages=[],
        extraction_engine="rtf.best-effort",
        warnings=warnings,
    )


def _extract_odt(content: bytes) -> ExtractionResult:
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            raw = archive.read("content.xml")
    except Exception as exc:
        return _extract_binary_best_effort(
            content,
            "odt",
            base_warnings=[f"ODT non aperto come archivio OpenDocument ({exc})."],
        )
    xml_text, encoding, warnings = _decode_text_bytes(raw)
    text = _xml_plain_text(xml_text)
    if not text.strip():
        warnings.append("Il documento ODT non contiene testo leggibile.")
    return ExtractionResult(
        ok=True,
        text=text,
        pages=[],
        extraction_engine=f"odt:{encoding}",
        warnings=warnings,
    )


def _extract_spreadsheet(content: bytes, ext: str) -> ExtractionResult:
    if str(ext or "").lower() == "xlsx" or bytes(content[:4]).startswith(b"PK\x03\x04"):
        result = _extract_xlsx_zip(content)
        if result is not None:
            return result
    return _extract_binary_best_effort(
        content,
        str(ext or "xls"),
        base_warnings=["Foglio di calcolo letto con recupero testuale compatibile."],
    )


def _extract_xlsx_zip(content: bytes) -> ExtractionResult | None:
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            shared_strings = _xlsx_shared_strings(archive)
            sheet_names = sorted(
                name for name in archive.namelist() if name.startswith("xl/worksheets/") and name.endswith(".xml")
            )
            rows: list[str] = []
            for sheet_index, sheet_name in enumerate(sheet_names[:30], start=1):
                rows.extend(_xlsx_sheet_rows(archive, sheet_name, shared_strings, sheet_index=sheet_index))
    except Exception:
        return None
    text = "\n".join(row for row in rows if row.strip())
    warnings = [] if text.strip() else ["Il foglio di calcolo non contiene celle testuali leggibili."]
    return ExtractionResult(
        ok=True,
        text=text,
        pages=[],
        extraction_engine="xlsx.zip-xml",
        warnings=warnings,
    )


def _xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    except Exception:
        return []
    values: list[str] = []
    for item in root.iter():
        if _xml_local_name(item.tag) != "si":
            continue
        texts = [node.text or "" for node in item.iter() if _xml_local_name(node.tag) == "t" and node.text]
        values.append("".join(texts))
    return values


def _xlsx_sheet_rows(
    archive: zipfile.ZipFile,
    sheet_name: str,
    shared_strings: list[str],
    *,
    sheet_index: int,
) -> list[str]:
    try:
        root = ET.fromstring(archive.read(sheet_name))
    except Exception:
        return []
    rows: list[str] = [f"Foglio {sheet_index}"]
    for row in root.iter():
        if _xml_local_name(row.tag) != "row":
            continue
        cells: list[str] = []
        for cell in row:
            if _xml_local_name(cell.tag) != "c":
                continue
            cell_type = str(cell.attrib.get("t") or "")
            value = ""
            if cell_type == "inlineStr":
                value = " ".join(
                    node.text or "" for node in cell.iter() if _xml_local_name(node.tag) == "t" and node.text
                ).strip()
            else:
                raw = next((node.text or "" for node in cell if _xml_local_name(node.tag) == "v"), "")
                if cell_type == "s":
                    try:
                        value = shared_strings[int(raw)]
                    except Exception:
                        value = raw
                else:
                    value = raw
            if str(value).strip():
                cells.append(str(value).strip())
        if cells:
            rows.append(" | ".join(cells))
    return rows


def _extract_image(content: bytes, ext: str) -> ExtractionResult:
    warnings: list[str] = []
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except Exception as exc:
        return _extract_binary_best_effort(
            content,
            str(ext or "image"),
            base_warnings=[f"OCR immagine non disponibile ({exc})."],
        )
    try:
        image = Image.open(BytesIO(content))
        tess_config = _configure_tesseract_runtime(pytesseract)
        lang = _resolve_tesseract_language(
            pytesseract,
            str(os.environ.get("IUSENTRA_DOCUMENT_AI_OCR_LANG") or "ita").strip() or "ita",
            tess_config,
        )
        text, ocr_warnings = _read_tesseract_text_best(
            pytesseract,
            image,
            lang=lang,
            base_config=tess_config,
            page_number=1,
        )
        warnings.extend(ocr_warnings)
    except Exception as exc:
        return _extract_binary_best_effort(
            content,
            str(ext or "image"),
            base_warnings=[f"OCR immagine non completato ({exc})."],
        )
    if not text:
        warnings.append("OCR immagine eseguito ma senza testo leggibile.")
    return ExtractionResult(
        ok=bool(text),
        text=text,
        pages=[DocumentAIPageText(page_number=1, text=text)] if text else [],
        extraction_engine="image.ocr",
        warnings=warnings,
        error_code="" if text else "image_ocr_empty",
        error_message="" if text else "L'immagine non contiene testo recuperabile con OCR.",
    )


def _extract_msg(content: bytes) -> ExtractionResult:
    try:
        import extract_msg  # type: ignore
    except Exception:
        return _extract_binary_best_effort(
            content,
            "msg",
            base_warnings=["Messaggio Outlook letto con recupero testuale compatibile perché il parser MSG non è disponibile."],
        )
    try:
        with tempfile.TemporaryDirectory(prefix="iusentra-msg-") as tmpdir:
            msg_path = Path(tmpdir) / "messaggio.msg"
            msg_path.write_bytes(content)
            msg = extract_msg.Message(str(msg_path))
            chunks = [
                f"Oggetto: {str(getattr(msg, 'subject', '') or '').strip()}",
                f"Mittente: {str(getattr(msg, 'sender', '') or '').strip()}",
                f"Data: {_format_email_datetime_visible(getattr(msg, 'date', ''))}",
                str(getattr(msg, "body", "") or "").strip(),
            ]
            text = "\n".join(chunk for chunk in chunks if chunk.strip())
    except Exception as exc:
        return _extract_binary_best_effort(
            content,
            "msg",
            base_warnings=[f"Parser MSG non completato ({exc})."],
        )
    return ExtractionResult(
        ok=bool(text.strip()),
        text=text,
        pages=[],
        extraction_engine="extract_msg",
        warnings=[] if text.strip() else ["Il messaggio Outlook non contiene testo leggibile."],
    )


def _extract_zip(content: bytes) -> ExtractionResult:
    warnings: list[str] = []
    chunks: list[str] = []
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            entries = [entry for entry in archive.infolist() if not entry.is_dir()]
            chunks.append("File contenuti nell'archivio: " + ", ".join(entry.filename for entry in entries[:80]))
            for entry in entries[:20]:
                if entry.file_size > 5 * 1024 * 1024:
                    warnings.append(f"{entry.filename}: file interno troppo grande per estrazione immediata.")
                    continue
                try:
                    payload = archive.read(entry)
                except Exception as exc:
                    warnings.append(f"{entry.filename}: lettura non completata ({exc}).")
                    continue
                inner_type = _file_type_from_payload(payload, entry.filename, fallback=Path(entry.filename).suffix.lower().lstrip("."))
                if inner_type == "zip":
                    warnings.append(f"{entry.filename}: archivio interno non espanso per evitare ricorsione.")
                    continue
                result = extract_text_from_document(payload, Path(entry.filename).name, inner_type)
                warnings.extend(f"{entry.filename}: {warning}" for warning in result.warnings)
                if result.text.strip():
                    chunks.append(f"[{entry.filename}]\n{result.text.strip()[:4000]}")
    except Exception as exc:
        return _extract_binary_best_effort(content, "zip", base_warnings=[f"Archivio ZIP non aperto ({exc})."])
    text = "\n\n".join(chunk for chunk in chunks if chunk.strip())
    if not text.strip():
        warnings.append("Archivio ZIP senza contenuti testuali recuperabili.")
    return ExtractionResult(
        ok=bool(text.strip()),
        text=text,
        pages=[],
        extraction_engine="zip.recursive-best-effort",
        warnings=_unique_warnings(warnings),
        error_code="" if text.strip() else "zip_no_text",
        error_message="" if text.strip() else "Archivio senza testo recuperabile.",
    )


def _decode_text_bytes(content: bytes) -> tuple[str, str, list[str]]:
    warnings: list[str] = []
    for encoding in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            text = content.decode(encoding)
            if "\ufffd" in text:
                continue
            if encoding not in {"utf-8-sig", "utf-16"}:
                warnings.append("Testo decodificato con codifica di compatibilità.")
            return _clean_text(text), encoding, warnings
        except UnicodeDecodeError:
            continue
    warnings.append("Testo decodificato sostituendo caratteri non leggibili.")
    return _clean_text(content.decode("utf-8", errors="replace")), "utf-8-replace", warnings


def _clean_text(value: str) -> str:
    return str(value or "").replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n").strip()


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []

    def handle_data(self, data: str) -> None:
        clean = " ".join(str(data or "").split())
        if clean:
            self._chunks.append(clean)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"p", "br", "div", "li", "tr", "h1", "h2", "h3"}:
            self._chunks.append("\n")

    def text(self) -> str:
        return _clean_text(" ".join(self._chunks))


def _html_to_text(value: str) -> str:
    parser = _HTMLTextExtractor()
    try:
        parser.feed(value)
    except Exception:
        return _clean_text(value)
    return parser.text()


def _extract_eml(content: bytes) -> ExtractionResult:
    try:
        message = BytesParser(policy=policy.default).parsebytes(content)
    except Exception as exc:
        return ExtractionResult(
            ok=False,
            text="",
            pages=[],
            extraction_engine="eml_failed",
            warnings=["L'email non è stata letta."],
            error_code="eml_extraction_failed",
            error_message=str(exc),
        )

    warnings: list[str] = []
    chunks: list[str] = []
    headers = _email_header_lines(message)
    if headers:
        chunks.append("\n".join(headers))

    plain_parts: list[str] = []
    html_parts: list[str] = []
    attachment_chunks: list[str] = []

    for part in message.walk():
        if part.is_multipart():
            continue
        filename = part.get_filename()
        content_type = str(part.get_content_type() or "").lower()
        disposition = str(part.get_content_disposition() or "").lower()
        payload = _email_part_bytes(part)
        if filename or disposition == "attachment":
            attachment_text, attachment_warnings = _extract_eml_attachment(filename or "allegato", payload)
            warnings.extend(attachment_warnings)
            if attachment_text:
                attachment_chunks.append(attachment_text)
            continue
        if content_type == "text/plain":
            text = _email_part_text(part, payload)
            if text.strip():
                plain_parts.append(text)
        elif content_type == "text/html":
            text = _html_to_text(_email_part_text(part, payload))
            if text.strip():
                html_parts.append(text)

    body_parts = plain_parts or html_parts
    if body_parts:
        chunks.append("Corpo email:\n" + "\n\n".join(body_parts))
    if attachment_chunks:
        chunks.append("Allegati leggibili:\n" + "\n\n".join(attachment_chunks))
    if not body_parts and not attachment_chunks:
        warnings.append("Email senza corpo testuale o allegati leggibili.")
    return ExtractionResult(
        ok=True,
        text="\n\n".join(chunk for chunk in chunks if chunk.strip()),
        pages=[],
        extraction_engine="email.message",
        warnings=warnings,
    )


def _email_header_lines(message: Message | EmailMessage) -> list[str]:
    labels = [
        ("Subject", "Oggetto"),
        ("From", "Mittente"),
        ("To", "Destinatari"),
        ("Cc", "Cc"),
        ("Date", "Data"),
    ]
    lines: list[str] = []
    for key, label in labels:
        value = str(message.get(key, "") or "").strip()
        if value:
            lines.append(f"{label}: {value}")
    return lines


def _email_part_bytes(part: Message | EmailMessage) -> bytes:
    payload = part.get_payload(decode=True)
    if isinstance(payload, bytes):
        return payload
    raw = part.get_payload()
    if isinstance(raw, str):
        charset = part.get_content_charset() or "utf-8"
        try:
            return raw.encode(charset, errors="replace")
        except LookupError:
            return raw.encode("utf-8", errors="replace")
    if isinstance(raw, list):
        nested: list[bytes] = []
        for item in raw:
            if isinstance(item, Message):
                try:
                    nested.append(item.as_bytes(policy=policy.default))
                except Exception:
                    nested.append(bytes(str(item), "utf-8", errors="replace"))
        return b"\n".join(nested)
    return b""


def _email_part_text(part: Message | EmailMessage, payload: bytes) -> str:
    charset = part.get_content_charset() or "utf-8"
    try:
        return _clean_text(payload.decode(charset))
    except Exception:
        text, _, _warnings = _decode_text_bytes(payload)
        return text


def _extract_eml_attachment(filename: str, payload: bytes) -> tuple[str, list[str]]:
    clean_name = Path(str(filename or "allegato")).name
    if not payload:
        return "", [f"{clean_name}: allegato vuoto o non leggibile."]
    file_type = _file_type_from_payload(payload, clean_name, fallback=Path(clean_name).suffix.lower().lstrip("."))
    result = extract_text_from_document(payload, clean_name, file_type)
    warnings = [f"{clean_name}: {warning}" for warning in result.warnings]
    if not result.ok:
        warnings.append(f"{clean_name}: allegato non letto per Lex.")
        return "", warnings
    if not result.text.strip():
        warnings.append(f"{clean_name}: allegato senza testo estraibile.")
        return "", warnings
    return f"[Allegato: {clean_name}]\n{result.text}", warnings


def _extract_pdf(content: bytes) -> ExtractionResult:
    warnings: list[str] = []
    try:
        import pdfplumber  # type: ignore

        pages: list[DocumentAIPageText] = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages.append(DocumentAIPageText(page_number=index, text=text))
        pages, full_text, repair_warnings = _repair_pdf_text(pages)
        warnings.extend(repair_warnings)
        quality = score_extracted_text_quality(full_text)
        if not full_text.strip() or quality.cid_placeholders or quality.score < 0.55:
            if full_text.strip() and quality.cid_placeholders:
                warnings.append("Testo PDF nativo non affidabile: rilevati segnaposto CID residui.")
            elif full_text.strip():
                warnings.append("Testo PDF nativo di qualità insufficiente: tentativo OCR.")
            warnings.append("Il PDF non fornisce testo nativo affidabile: OCR richiesto per l'indice Lex.")
            ocr_result = _extract_scanned_pdf_with_ocr(content, base_warnings=warnings, base_engine="pdfplumber")
            if ocr_result is not None:
                return ocr_result
        return ExtractionResult(
            ok=True,
            text=full_text,
            pages=pages,
            extraction_engine="pdfplumber",
            warnings=warnings,
        )
    except Exception as first_exc:
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(BytesIO(content))
            pages = [
                DocumentAIPageText(page_number=index, text=page.extract_text() or "")
                for index, page in enumerate(reader.pages, start=1)
            ]
            pages, full_text, repair_warnings = _repair_pdf_text(pages)
            warnings.append("Estrattore PDF primario non disponibile; usato parser alternativo.")
            warnings.extend(repair_warnings)
            quality = score_extracted_text_quality(full_text)
            if not full_text.strip() or quality.cid_placeholders or quality.score < 0.55:
                if full_text.strip() and quality.cid_placeholders:
                    warnings.append("Testo PDF nativo non affidabile: rilevati segnaposto CID residui.")
                elif full_text.strip():
                    warnings.append("Testo PDF nativo di qualità insufficiente: tentativo OCR.")
                warnings.append("Il PDF non fornisce testo nativo affidabile: OCR richiesto per l'indice Lex.")
                ocr_result = _extract_scanned_pdf_with_ocr(content, base_warnings=warnings, base_engine="pypdf")
                if ocr_result is not None:
                    return ocr_result
            return ExtractionResult(
                ok=True,
                text=full_text,
                pages=pages,
                extraction_engine="pypdf",
                warnings=warnings,
            )
        except Exception as exc:
            return ExtractionResult(
                ok=False,
                text="",
                pages=[],
                extraction_engine="pdf_failed",
                warnings=["Estrazione PDF non completata."],
                error_code="pdf_extraction_failed",
                error_message=str(exc or first_exc),
            )


def _extract_scanned_pdf_with_ocr(
    content: bytes,
    *,
    base_warnings: list[str],
    base_engine: str,
) -> ExtractionResult | None:
    try:
        import pypdfium2 as pdfium  # type: ignore
        import pytesseract  # type: ignore
    except Exception as exc:
        base_warnings.append(f"OCR PDF non disponibile: {exc}")
        return None

    tess_config = _configure_tesseract_runtime(pytesseract)
    tesseract_probe = getattr(pytesseract, "get_tesseract_version", None)
    if callable(tesseract_probe):
        try:
            tesseract_probe()
        except Exception as exc:
            base_warnings.append(f"OCR PDF non disponibile: {exc}")
            return None

    lang = _resolve_tesseract_language(
        pytesseract,
        str(os.environ.get("IUSENTRA_DOCUMENT_AI_OCR_LANG") or "ita").strip() or "ita",
        tess_config,
    )
    max_pages = _int_from_env("IUSENTRA_DOCUMENT_AI_OCR_MAX_PAGES", default=0, minimum=0, maximum=1000)
    scale = _float_from_env("IUSENTRA_DOCUMENT_AI_OCR_SCALE", default=3.0, minimum=2.0, maximum=6.0)
    pdf = None
    pages: list[DocumentAIPageText] = []
    warnings = list(base_warnings)
    try:
        pdf = pdfium.PdfDocument(content)
        page_count = len(pdf)
        limit = min(page_count, max_pages) if max_pages else page_count
        if max_pages and page_count > max_pages:
            warnings.append(
                f"OCR limitato a {max_pages} pagine su {page_count}; aumentare "
                "IUSENTRA_DOCUMENT_AI_OCR_MAX_PAGES per indicizzare il documento completo."
            )
        for page_index in range(limit):
            page = None
            bitmap = None
            try:
                page = pdf[page_index]
                bitmap = page.render(scale=scale)
                image = bitmap.to_pil()
                text, ocr_warnings = _read_tesseract_text_best(
                    pytesseract,
                    image,
                    lang=lang,
                    base_config=tess_config,
                    page_number=page_index + 1,
                )
                warnings.extend(ocr_warnings)
            except Exception as exc:
                warnings.append(f"Pagina {page_index + 1}: OCR non completato ({exc}).")
                text = ""
            finally:
                for obj in (bitmap, page):
                    close = getattr(obj, "close", None)
                    if callable(close):
                        try:
                            close()
                        except Exception:
                            pass
            pages.append(DocumentAIPageText(page_number=page_index + 1, text=text))
    except Exception as exc:
        base_warnings.append(f"OCR PDF non completato: {exc}")
        return None
    finally:
        close = getattr(pdf, "close", None)
        if callable(close):
            try:
                close()
            except Exception:
                pass

    full_text = "\n\n".join(page.text for page in pages if page.text)
    if not full_text.strip():
        base_warnings.extend(warning for warning in warnings if warning not in base_warnings)
        base_warnings.append("OCR PDF eseguito ma senza testo leggibile.")
        return None
    warnings.append("PDF senza testo selezionabile: OCR applicato per l'indice Lex.")
    return ExtractionResult(
        ok=True,
        text=full_text,
        pages=pages,
        extraction_engine=f"{base_engine}+ocr",
        warnings=_unique_warnings(warnings),
    )


def _configure_tesseract_runtime(pytesseract: Any) -> str:
    command = _resolve_tesseract_command()
    pytesseract_module = getattr(pytesseract, "pytesseract", None)
    if command and pytesseract_module is not None and hasattr(pytesseract_module, "tesseract_cmd"):
        pytesseract_module.tesseract_cmd = command
    tessdata_dir = _resolve_tessdata_dir(command)
    if tessdata_dir:
        os.environ["TESSDATA_PREFIX"] = tessdata_dir
    return ""


def _resolve_tesseract_command() -> str:
    configured = str(os.environ.get("IUSENTRA_TESSERACT_CMD") or os.environ.get("TESSERACT_CMD") or "").strip()
    if configured and Path(configured).is_file():
        return configured
    discovered = which("tesseract")
    if discovered:
        return discovered
    if os.name == "nt":
        for root in (
            os.environ.get("ProgramFiles"),
            os.environ.get("ProgramFiles(x86)"),
            r"C:\Program Files",
            r"C:\Program Files (x86)",
        ):
            if not root:
                continue
            candidate = Path(root) / "Tesseract-OCR" / "tesseract.exe"
            if candidate.is_file():
                return str(candidate)
    return ""


def _resolve_tessdata_dir(command: str) -> str:
    candidates: list[Path] = []
    for name in ("IUSENTRA_TESSDATA_PREFIX", "TESSDATA_PREFIX"):
        configured = str(os.environ.get(name) or "").strip().strip('"')
        if configured:
            candidates.append(Path(configured))
    local_app_data = str(os.environ.get("LOCALAPPDATA") or "").strip()
    if local_app_data:
        candidates.append(Path(local_app_data) / "IUSENTRA" / "tessdata")
    if command:
        candidates.append(Path(command).resolve().parent / "tessdata")
    for candidate in candidates:
        if candidate.is_dir() and any(candidate.glob("*.traineddata")):
            return str(candidate)
    return ""


def _resolve_tesseract_language(pytesseract: Any, preferred: str, config: str) -> str:
    requested = [part for part in str(preferred or "ita").split("+") if part]
    get_languages = getattr(pytesseract, "get_languages", None)
    if not callable(get_languages):
        return "+".join(requested) or "ita"
    try:
        available = set(get_languages(config=config))
    except TypeError:
        try:
            available = set(get_languages())
        except Exception:
            return "+".join(requested) or "ita"
    except Exception:
        return "+".join(requested) or "ita"
    selected = [lang for lang in requested if lang in available]
    if selected:
        return "+".join(selected)
    if "ita" in available:
        return "ita"
    if "eng" in available:
        return "eng"
    return "+".join(requested) or "ita"


def _read_tesseract_text_best(
    pytesseract: Any,
    image: Any,
    *,
    lang: str,
    base_config: str,
    page_number: int,
) -> tuple[str, list[str]]:
    warnings: list[str] = []
    prepared = _preprocess_ocr_image(image)
    image_to_data = getattr(pytesseract, "image_to_data", None)
    if not callable(image_to_data):
        return _read_tesseract_text_fallback(pytesseract, prepared, lang=lang, config=base_config, page_number=page_number)
    output = getattr(pytesseract, "Output", None)
    if output is None:
        try:
            from pytesseract import Output as output  # type: ignore
        except Exception:
            return _read_tesseract_text_fallback(
                pytesseract,
                prepared,
                lang=lang,
                config=base_config,
                page_number=page_number,
            )
    output_dict = getattr(output, "DICT", "dict")
    candidates: list[tuple[float, str]] = []
    for config_name, config in _document_ai_tesseract_configs(base_config):
        try:
            data = image_to_data(prepared, lang=lang, config=config, output_type=output_dict)
        except TypeError:
            try:
                data = image_to_data(prepared, lang=lang, output_type=output_dict)
            except Exception as exc:
                warnings.append(f"Pagina {page_number}: OCR {config_name} non completato ({exc}).")
                continue
        except Exception as exc:
            warnings.append(f"Pagina {page_number}: OCR {config_name} non completato ({exc}).")
            continue
        text, avg_confidence = _text_from_tesseract_data(data)
        candidates.append((_score_document_ai_ocr_text(text, avg_confidence), text))
    if candidates:
        best_text = max(candidates, key=lambda item: item[0])[1].strip()
        if len(best_text) >= 40:
            return best_text, warnings
    for variant_name, variant in _low_contrast_ocr_variants(prepared):
        for config_name, config in _document_ai_tesseract_configs(base_config):
            try:
                data = image_to_data(variant, lang=lang, config=config, output_type=output_dict)
            except TypeError:
                try:
                    data = image_to_data(variant, lang=lang, output_type=output_dict)
                except Exception as exc:
                    warnings.append(
                        f"Pagina {page_number}: OCR {variant_name}/{config_name} non completato ({exc})."
                    )
                    continue
            except Exception as exc:
                warnings.append(f"Pagina {page_number}: OCR {variant_name}/{config_name} non completato ({exc}).")
                continue
            text, avg_confidence = _text_from_tesseract_data(data)
            candidates.append((_score_document_ai_ocr_text(text, avg_confidence), text))
    if candidates:
        return max(candidates, key=lambda item: item[0])[1].strip(), warnings
    text, fallback_warnings = _read_tesseract_text_fallback(
        pytesseract,
        prepared,
        lang=lang,
        config=base_config,
        page_number=page_number,
    )
    warnings.extend(fallback_warnings)
    return text, warnings


def _read_tesseract_text_fallback(
    pytesseract: Any,
    image: Any,
    *,
    lang: str,
    config: str,
    page_number: int,
) -> tuple[str, list[str]]:
    image_to_string = getattr(pytesseract, "image_to_string", None)
    if not callable(image_to_string):
        return "", [f"Pagina {page_number}: runtime Tesseract senza image_to_string."]
    try:
        return str(image_to_string(image, lang=lang, config=config) or "").strip(), []
    except TypeError:
        try:
            return str(image_to_string(image, lang=lang) or "").strip(), []
        except Exception as exc:
            return "", [f"Pagina {page_number}: OCR fallback non completato ({exc})."]
    except Exception as exc:
        return "", [f"Pagina {page_number}: OCR fallback non completato ({exc})."]


def _preprocess_ocr_image(image: Any) -> Any:
    try:
        from PIL import Image, ImageChops, ImageEnhance, ImageOps  # type: ignore

        gray = ImageOps.grayscale(image)
        white = Image.new("L", gray.size, 255)
        difference = ImageChops.difference(gray, white)
        content_mask = difference.point(lambda pixel: 255 if pixel > 10 else 0)
        content_box = content_mask.getbbox()
        if content_box:
            left, top, right, bottom = content_box
            content_width = right - left
            content_height = bottom - top
            removes_margin = content_width < gray.width * 0.9 or content_height < gray.height * 0.9
            preserves_page = content_width >= gray.width * 0.35 and content_height >= gray.height * 0.35
            if removes_margin and preserves_page:
                padding = max(8, int(min(gray.size) * 0.015))
                gray = gray.crop(
                    (
                        max(0, left - padding),
                        max(0, top - padding),
                        min(gray.width, right + padding),
                        min(gray.height, bottom + padding),
                    )
                )
        gray = ImageOps.autocontrast(gray, cutoff=1)
        gray = ImageEnhance.Sharpness(gray).enhance(1.8)
        return ImageEnhance.Contrast(gray).enhance(1.08)
    except Exception:
        return image


def _low_contrast_ocr_variants(image: Any) -> list[tuple[str, Any]]:
    try:
        from PIL import ImageOps  # type: ignore

        gray = ImageOps.grayscale(image)
        histogram = gray.histogram()
        total = sum(histogram)
        if total <= 0:
            return []
        weighted_total = sum(index * count for index, count in enumerate(histogram))
        background_weight = 0
        background_sum = 0
        best_threshold = 0
        best_variance = -1.0
        for threshold, count in enumerate(histogram):
            background_weight += count
            if background_weight <= 0:
                continue
            foreground_weight = total - background_weight
            if foreground_weight <= 0:
                break
            background_sum += threshold * count
            background_mean = background_sum / background_weight
            foreground_mean = (weighted_total - background_sum) / foreground_weight
            variance = background_weight * foreground_weight * (background_mean - foreground_mean) ** 2
            if variance > best_variance:
                best_variance = variance
                best_threshold = threshold
        thresholds = {
            max(48, min(210, round(best_threshold * 0.45))),
            max(48, min(210, round(best_threshold * 0.70))),
        }
        return [
            (
                f"contrasto-{threshold}",
                gray.point(lambda pixel, limit=threshold: 0 if pixel < limit else 255, mode="1"),
            )
            for threshold in sorted(thresholds)
        ]
    except Exception:
        return []


def _document_ai_tesseract_configs(base_config: str) -> list[tuple[str, str]]:
    prefix = (base_config.strip() + " ") if str(base_config or "").strip() else ""
    return [
        ("psm6", prefix + "--oem 1 --psm 6 -c preserve_interword_spaces=1"),
        ("psm4", prefix + "--oem 1 --psm 4 -c preserve_interword_spaces=1"),
        ("psm3", prefix + "--oem 1 --psm 3 -c preserve_interword_spaces=1"),
        ("psm11", prefix + "--oem 1 --psm 11 -c preserve_interword_spaces=1"),
    ]


def _text_from_tesseract_data(data: Any) -> tuple[str, float]:
    if not isinstance(data, dict):
        return "", 0.0
    words: list[str] = []
    confidences: list[float] = []
    raw_texts = list(data.get("text") or [])
    raw_confidences = list(data.get("conf") or [])
    for index, raw in enumerate(raw_texts):
        token = str(raw or "").strip()
        if not token:
            continue
        words.append(token)
        try:
            confidence = float(raw_confidences[index]) / 100.0
        except (IndexError, TypeError, ValueError):
            confidence = 0.0
        if confidence >= 0:
            confidences.append(max(0.0, min(1.0, confidence)))
    average = sum(confidences) / len(confidences) if confidences else 0.0
    return " ".join(words), average


def _score_document_ai_ocr_text(text: str, avg_confidence: float) -> float:
    normalized = str(text or "")
    score = min(len(normalized), 2500) / 120.0 + max(0.0, min(1.0, avg_confidence)) * 25.0
    patterns = [
        (r"\btribunale\s+di\s+[a-zàèéìòù' ]+", 8),
        (r"\b(proc\.?\s*n\.?|r\.?\s*g\.?|rgac)\s*[\w./-]+", 14),
        (r"\b[A-Z]{6}\d{2}[A-Z]\d{2}[A-Z]\d{3}[A-Z]\b", 12),
        (r"\b[\w.+-]+@[\w.-]+\.[a-z]{2,}\b", 14),
        (r"\b\d{1,2}[./]\d{1,2}[./]\d{2,4}\b", 8),
        (r"\b(?:euro|€)\s*[\.,]?\s*\d", 8),
        (r"\b(?:art\.?|dpr|c\.p\.c\.|c\.c\.)\b", 8),
    ]
    for pattern, weight in patterns:
        score += len(re.findall(pattern, normalized, flags=re.IGNORECASE)) * weight
    score -= len(re.findall(r"[|~{}_\[\]]", normalized)) * 0.75
    return score


def _extract_docx(content: bytes) -> ExtractionResult:
    try:
        from docx import Document  # type: ignore

        document = Document(BytesIO(content))
        chunks: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text:
                chunks.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))
        text = "\n".join(chunks)
        warnings = [] if text.strip() else ["Il DOCX non contiene testo estraibile."]
        return ExtractionResult(ok=True, text=text, pages=[], extraction_engine="python-docx", warnings=warnings)
    except Exception as first_exc:
        try:
            import mammoth  # type: ignore

            result: Any = mammoth.extract_raw_text(BytesIO(content))
            text = str(getattr(result, "value", "") or "")
            warnings = [str(message) for message in getattr(result, "messages", [])]
            return ExtractionResult(ok=True, text=text, pages=[], extraction_engine="mammoth", warnings=warnings)
        except Exception as exc:
            return ExtractionResult(
                ok=False,
                text="",
                pages=[],
                extraction_engine="docx_failed",
                warnings=["Estrazione DOCX non completata."],
                error_code="docx_extraction_failed",
                error_message=str(exc or first_exc),
            )


def _extract_doc(content: bytes) -> ExtractionResult:
    if bytes(content[:4]).startswith(b"PK\x03\x04"):
        result = _extract_docx(content)
        if result.extraction_engine and not result.extraction_engine.startswith("doc-compat:"):
            result.extraction_engine = f"doc-compat:{result.extraction_engine}"
        return result

    rtf_text = _extract_rtf_text(content)
    if rtf_text.strip():
        return ExtractionResult(
            ok=True,
            text=rtf_text,
            pages=[],
            extraction_engine="doc.rtf",
            warnings=["File .doc letto come documento RTF compatibile."],
        )

    external_result = _extract_doc_with_external_tool(content)
    if external_result is not None:
        return external_result

    fallback_text = _extract_legacy_doc_binary_text(content)
    if fallback_text.strip():
        return ExtractionResult(
            ok=True,
            text=fallback_text,
            pages=[],
            extraction_engine="doc.binary-text",
            warnings=[
                "DOC legacy letto con estrazione testuale compatibile; verifica l'impaginazione originale se il documento va depositato.",
            ],
        )

    return ExtractionResult(
        ok=False,
        text="",
        pages=[],
        extraction_engine="doc_legacy_unavailable",
        warnings=["Estrazione DOC non disponibile nel runtime corrente."],
        error_code="doc_legacy_extraction_unavailable",
        error_message="Formato DOC legacy ammesso in upload; estrazione testo non disponibile senza conversione locale.",
    )


def _extract_rtf_text(content: bytes) -> str:
    if not bytes(content[:12]).lstrip().startswith(b"{\\rtf"):
        return ""
    raw, _encoding, _warnings = _decode_text_bytes(content)
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
    raw = raw.replace("{", " ").replace("}", " ").replace("\\", " ")
    return _clean_text(re.sub(r"\s+", " ", raw))


def _extract_doc_with_external_tool(content: bytes) -> ExtractionResult | None:
    antiword = which("antiword")
    if antiword:
        text, _error = _run_doc_text_command([antiword, "{input}"], content)
        if text.strip():
            return ExtractionResult(
                ok=True,
                text=text,
                pages=[],
                extraction_engine="antiword",
                warnings=[],
            )

    office = which("soffice") or which("libreoffice")
    if office:
        text, error = _run_libreoffice_doc_conversion(office, content)
        if text.strip():
            return ExtractionResult(
                ok=True,
                text=text,
                pages=[],
                extraction_engine="libreoffice",
                warnings=[],
            )
    return None


def _run_doc_text_command(command: list[str], content: bytes) -> tuple[str, str]:
    timeout = _int_from_env("IUSENTRA_DOCUMENT_AI_DOC_TIMEOUT_SECONDS", default=20, minimum=3, maximum=120)
    with tempfile.TemporaryDirectory(prefix="iusentra-doc-") as tmpdir:
        input_path = Path(tmpdir) / "documento.doc"
        input_path.write_bytes(content)
        args = [str(input_path) if part == "{input}" else part for part in command]
        try:
            completed = subprocess.run(
                args,
                cwd=tmpdir,
                check=False,
                capture_output=True,
                timeout=timeout,
            )
        except Exception as exc:
            return "", str(exc)
        if completed.stdout:
            text, _encoding, _warnings = _decode_text_bytes(completed.stdout)
            return text, ""
        stderr = completed.stderr.decode("utf-8", errors="replace").strip()
        return "", stderr


def _run_libreoffice_doc_conversion(office: str, content: bytes) -> tuple[str, str]:
    timeout = _int_from_env("IUSENTRA_DOCUMENT_AI_DOC_TIMEOUT_SECONDS", default=20, minimum=3, maximum=120)
    with tempfile.TemporaryDirectory(prefix="iusentra-doc-") as tmpdir:
        input_path = Path(tmpdir) / "documento.doc"
        input_path.write_bytes(content)
        try:
            completed = subprocess.run(
                [
                    office,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    tmpdir,
                    str(input_path),
                ],
                cwd=tmpdir,
                check=False,
                capture_output=True,
                timeout=timeout,
            )
        except Exception as exc:
            return "", str(exc)
        for candidate in Path(tmpdir).glob("*.txt"):
            try:
                text, _encoding, _warnings = _decode_text_bytes(candidate.read_bytes())
                return text, ""
            except OSError:
                continue
        error = (completed.stderr or completed.stdout).decode("utf-8", errors="replace").strip()
        return "", error


def _extract_legacy_doc_binary_text(content: bytes) -> str:
    candidates: list[str] = []
    for encoding in ("utf-8", "utf-16le", "cp1252", "latin-1"):
        try:
            decoded = content.decode(encoding, errors="ignore")
        except LookupError:
            continue
        cleaned = _clean_text(decoded)
        printable = sum(1 for char in cleaned if char.isprintable() or char.isspace())
        ratio = printable / max(1, len(cleaned))
        if ratio >= 0.65:
            candidates.append(cleaned)
    byte_runs = re.findall(rb"[A-Za-z0-9\xc0-\xff][A-Za-z0-9\xc0-\xff \t\r\n.,;:!?/()'\-]{5,}", content)
    if byte_runs:
        chunks = []
        for run in byte_runs[:800]:
            text = run.decode("cp1252", errors="ignore").strip()
            if text:
                chunks.append(text)
        candidates.append(" ".join(chunks))
    candidates = [
        value
        for value in (re.sub(r"\s+", " ", candidate).strip() for candidate in candidates)
        if _legacy_doc_text_is_meaningful(value)
    ]
    if not candidates:
        return ""
    return max(candidates, key=len)


def _legacy_doc_text_is_meaningful(value: str) -> bool:
    clean = str(value or "").strip()
    if len(clean) < 24:
        return False
    words = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]{3,}", clean)
    return len(words) >= 3


def _extract_binary_best_effort(
    content: bytes,
    label: str,
    *,
    base_warnings: list[str] | None = None,
) -> ExtractionResult:
    warnings = list(base_warnings or [])
    text = _extract_legacy_doc_binary_text(content)
    if not text.strip():
        decoded, encoding, decode_warnings = _decode_text_bytes(content)
        warnings.extend(decode_warnings)
        printable = sum(1 for char in decoded if char.isprintable() or char.isspace())
        ratio = printable / max(1, len(decoded))
        if ratio >= 0.80 and _legacy_doc_text_is_meaningful(decoded):
            text = decoded
            warnings.append(f"Formato {label} letto come testo {encoding}.")
    if text.strip():
        warnings.append(
            "Documento letto con recupero testuale best-effort; verifica manualmente se serve fedeltà di impaginazione."
        )
        return ExtractionResult(
            ok=True,
            text=text,
            pages=[],
            extraction_engine=f"{label}.binary-best-effort",
            warnings=_unique_warnings(warnings),
        )
    warnings.append("Documento senza testo recuperabile automaticamente nel runtime corrente.")
    return ExtractionResult(
        ok=False,
        text="",
        pages=[],
        extraction_engine=f"{label}.binary-best-effort",
        warnings=_unique_warnings(warnings),
        error_code="no_recoverable_text",
        error_message="Documento senza testo recuperabile automaticamente.",
    )


def _xml_local_name(tag: Any) -> str:
    value = str(tag or "")
    return value.rsplit("}", 1)[-1] if "}" in value else value


def _xml_plain_text(xml_text: str) -> str:
    raw = str(xml_text or "")
    try:
        root = ET.fromstring(raw.encode("utf-8"))
        parts = [node.text or "" for node in root.iter() if str(node.text or "").strip()]
        clean = _clean_text("\n".join(parts))
        if clean:
            return clean
    except Exception:
        pass
    return _clean_text(re.sub(r"<[^>]+>", " ", raw))


def _repair_pdf_text(pages: list[DocumentAIPageText]) -> tuple[list[DocumentAIPageText], str, list[str]]:
    repaired_pages: list[DocumentAIPageText] = []
    warnings: list[str] = []
    for page in pages:
        repaired, page_warnings = repair_pdf_cid_placeholders(page.text)
        repaired_pages.append(DocumentAIPageText(page_number=page.page_number, text=repaired))
        warnings.extend(page_warnings)
    full_text = "\n\n".join(page.text for page in repaired_pages if page.text)
    original_text = "\n\n".join(page.text for page in pages if page.text)
    original_score = score_extracted_text_quality(original_text)
    repaired_score = score_extracted_text_quality(full_text)
    if repaired_score.score >= original_score.score and warnings:
        return repaired_pages, full_text, _unique_warnings(warnings)
    return pages, original_text, []


def _unique_warnings(values: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for value in values:
        clean = str(value or "").strip()
        if not clean or clean in seen:
            continue
        seen.add(clean)
        out.append(clean)
    return out


def _int_from_env(name: str, *, default: int, minimum: int, maximum: int) -> int:
    try:
        value = int(str(os.environ.get(name, default)).strip())
    except (TypeError, ValueError):
        return default
    return max(minimum, min(maximum, value))


def _float_from_env(name: str, *, default: float, minimum: float, maximum: float) -> float:
    try:
        value = float(str(os.environ.get(name, default)).strip())
    except (TypeError, ValueError):
        return default
    return max(minimum, min(maximum, value))
