"""Export governabile dei contenuti generati da Lex."""

from __future__ import annotations

import io
import re
import unicodedata
from datetime import datetime
from typing import Iterable

from docx import Document
from docx.shared import Pt

_MONTHS_IT = (
    "",
    "gennaio",
    "febbraio",
    "marzo",
    "aprile",
    "maggio",
    "giugno",
    "luglio",
    "agosto",
    "settembre",
    "ottobre",
    "novembre",
    "dicembre",
)

_HEADING_RE = re.compile(r"^\s{0,3}(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_MARKDOWN_DECORATIONS_RE = re.compile(r"[*_`>#]+")
_FILENAME_SAFE_RE = re.compile(r"[^a-z0-9]+")


def format_dataora_italiana(value: datetime | None = None) -> str:
    current = value or datetime.now()
    month = _MONTHS_IT[current.month]
    return f"{current.day:02d} {month} {current.year} {current:%H:%M}"


def sanitize_export_title(title: str) -> str:
    cleaned = " ".join(str(title or "").strip().split())
    return cleaned or "Documento generato da Lex"


def infer_export_title(*, title: str = "", question: str = "", answer: str = "") -> str:
    explicit = sanitize_export_title(title)
    if explicit != "Documento generato da Lex":
        return explicit

    for raw_line in str(answer or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = _HEADING_RE.match(line)
        if heading:
            return sanitize_export_title(heading.group(2))
        bullet = _BULLET_RE.match(line) or _NUMBERED_RE.match(line)
        if bullet:
            line = bullet.group(1).strip()
        plain = _MARKDOWN_DECORATIONS_RE.sub("", line).strip(" -:;,")
        if plain:
            return sanitize_export_title(plain[:96])

    if question:
        return sanitize_export_title(question[:96])
    return explicit


def build_export_filename(title: str, extension: str) -> str:
    normalized = unicodedata.normalize("NFKD", sanitize_export_title(title)).encode("ascii", "ignore").decode("ascii")
    slug = _FILENAME_SAFE_RE.sub("-", normalized.lower()).strip("-")
    slug = slug or "documento-generato-da-lex"
    return f"{slug}.{extension.lstrip('.')}"


def _apply_default_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    heading_style = document.styles["Heading 1"]
    heading_style.font.name = "Calibri"
    heading_style.font.size = Pt(16)


def _clean_line(value: str) -> str:
    cleaned = _MARKDOWN_DECORATIONS_RE.sub("", str(value or ""))
    return " ".join(cleaned.split())


def _iter_markdown_lines(answer: str) -> Iterable[str]:
    for line in str(answer or "").replace("\r\n", "\n").split("\n"):
        yield line.rstrip()


def build_docx_bytes(
    *,
    title: str,
    question: str,
    answer: str,
    citations: list[str] | None = None,
    context_label: str = "",
) -> bytes:
    document = Document()
    _apply_default_styles(document)

    document.add_heading(sanitize_export_title(title), level=1)
    meta_rows = [f"Generato da Lex il {format_dataora_italiana()}."]
    if context_label:
        meta_rows.append(f"Contesto operativo: {sanitize_export_title(context_label)}.")
    document.add_paragraph(" ".join(meta_rows))

    if question:
        document.add_paragraph("Richiesta iniziale", style="Heading 2")
        document.add_paragraph(_clean_line(question))

    document.add_paragraph("Contenuto generato", style="Heading 2")

    for raw_line in _iter_markdown_lines(answer):
        line = raw_line.strip()
        if not line:
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)) + 1, 3)
            document.add_paragraph(_clean_line(heading.group(2)), style=f"Heading {level}")
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            document.add_paragraph(_clean_line(bullet.group(1)), style="List Bullet")
            continue

        numbered = _NUMBERED_RE.match(line)
        if numbered:
            document.add_paragraph(_clean_line(numbered.group(1)), style="List Number")
            continue

        document.add_paragraph(_clean_line(line))

    final_citations = [str(item or "").strip() for item in citations or [] if str(item or "").strip()]
    if final_citations:
        document.add_paragraph("Fonti richiamate", style="Heading 2")
        for citation in final_citations:
            document.add_paragraph(citation, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
