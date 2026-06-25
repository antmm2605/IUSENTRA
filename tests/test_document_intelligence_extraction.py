from email.message import EmailMessage
from io import BytesIO
from pathlib import Path
import sys
from types import SimpleNamespace
import zipfile

import pytest

from pct.document_intelligence.extraction import extract_document_text, extract_text_from_document


def test_document_ai_extraction_docx_semplice(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document

    document = Document()
    document.add_paragraph("Clausola contrattuale principale")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Parte"
    table.rows[0].cells[1].text = "Valore"
    target = tmp_path / "atto.docx"
    document.save(target)

    result = extract_document_text(target, "docx")

    assert result.error is None
    assert result.extraction_engine in {"python-docx", "mammoth"}
    assert "Clausola contrattuale principale" in result.text
    assert "Parte" in result.text


def test_document_ai_extraction_doc_legacy_testuale(tmp_path: Path):
    target = tmp_path / "memoria.doc"
    target.write_bytes(b"\xd0\xcf\x11\xe0\x00\x00Memoria difensiva cliente Rossi contro Bianchi")

    result = extract_document_text(target, "doc")

    assert result.error is None
    assert result.extraction_engine in {"antiword", "libreoffice", "doc.binary-text", "doc.rtf"}
    assert "Memoria difensiva" in result.text


def test_document_ai_extraction_pdf_semplice(tmp_path: Path):
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas

    target = tmp_path / "atto.pdf"
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 720, "Testo PDF fascicolo")
    pdf.showPage()
    pdf.save()
    target.write_bytes(buffer.getvalue())

    result = extract_document_text(target, "pdf")

    assert result.error is None
    assert result.extraction_engine in {"pdfplumber", "pypdf"}
    assert result.page_count == 1
    assert "Testo PDF fascicolo" in result.text


def test_document_ai_pdf_mislabeled_pcten_non_usa_parser_pdf():
    result = extract_text_from_document(b"PCTEN" + (b"\x00" * 200), "Atto.pdf", "pdf")

    assert result.ok is False
    assert result.text == ""
    assert result.extraction_engine == "pdf_magic_mismatch"
    assert any("parser PDF saltato" in warning for warning in result.warnings)


def test_document_ai_pdf_mislabeled_zip_resta_leggibile():
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("note.txt", "Testo interno allo ZIP per Lex AI")

    result = extract_text_from_document(buffer.getvalue(), "allegato.pdf", "pdf")

    assert result.ok is True
    assert result.extraction_engine.startswith("pdf-mismatch:zip")
    assert "note.txt" in result.text
    assert "Testo interno allo ZIP" in result.text


def test_document_ai_pdf_mislabeled_xml_resta_leggibile():
    payload = b"<?xml version='1.0'?><root><titolo>Comunicazione PEC</titolo><testo>Presidio XML</testo></root>"

    result = extract_text_from_document(payload, "daticert.pdf", "pdf")

    assert result.ok is True
    assert result.extraction_engine.startswith("pdf-mismatch:")
    assert "Comunicazione PEC" in result.text
    assert "Presidio XML" in result.text


def test_document_ai_unlimited_ocr_alimenta_indice_integrale_immagine(monkeypatch):
    pytest.importorskip("PIL")
    from PIL import Image

    buffer = BytesIO()
    Image.new("RGB", (220, 120), "white").save(buffer, format="PNG")
    monkeypatch.setenv("IUSENTRA_UNLIMITED_OCR_ENABLED", "1")
    monkeypatch.setenv("IUSENTRA_UNLIMITED_OCR_ENDPOINT", "http://127.0.0.1:10000")
    monkeypatch.setenv("IUSENTRA_UNLIMITED_OCR_MAX_RETRIES", "1")

    def fake_post(url, payload, timeout, api_key):
        assert url.endswith("/v1/chat/completions")
        assert payload["messages"][0]["content"]
        return {"choices": [{"message": {"content": "TRIBUNALE DI MILANO\nR.G. n. 12345/2026"}}]}

    monkeypatch.setattr("legal_ocr.unlimited.client.post_json_request", fake_post)

    result = extract_text_from_document(buffer.getvalue(), "scansione.png", "png")

    assert result.ok is True
    assert result.extraction_engine == "legal-ocr:unlimited-ocr:document-index"
    assert "R.G. n. 12345/2026" in result.text
    assert result.pages[0].page_number == 1
    assert any("nessun chunk OCR" in warning for warning in result.warnings)


def test_document_ai_extraction_pdf_p7m_leggibile_come_pdf(tmp_path: Path):
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas

    target = tmp_path / "atto.pdf.p7m"
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 720, "Testo PDF firmato leggibile")
    pdf.showPage()
    pdf.save()
    target.write_bytes(buffer.getvalue())

    result = extract_document_text(target, "pdf")

    assert result.error is None
    assert result.extraction_engine in {"cades:pdfplumber", "cades:pypdf"}
    assert result.page_count == 1
    assert "Testo PDF firmato leggibile" in result.text
    assert any("PDF interno" in warning for warning in result.warnings)


def test_document_ai_extraction_pdf_scansionato_usa_ocr_quando_il_testo_manca(monkeypatch):
    class FakePdfPlumberPage:
        def extract_text(self):
            return ""

    class FakePdfPlumberDocument:
        pages = [FakePdfPlumberPage()]

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

    class FakeBitmap:
        def to_pil(self):
            return object()

        def close(self):
            return None

    class FakePdfiumPage:
        def render(self, *, scale: float):
            assert scale >= 1.0
            return FakeBitmap()

        def close(self):
            return None

    class FakePdfiumDocument:
        def __init__(self, content: bytes):
            assert content.startswith(b"%PDF")

        def __len__(self):
            return 1

        def __getitem__(self, index: int):
            assert index == 0
            return FakePdfiumPage()

        def close(self):
            return None

    monkeypatch.setitem(
        sys.modules,
        "pdfplumber",
        SimpleNamespace(open=lambda _stream: FakePdfPlumberDocument()),
    )
    monkeypatch.setitem(sys.modules, "pypdfium2", SimpleNamespace(PdfDocument=FakePdfiumDocument))
    monkeypatch.setitem(
        sys.modules,
        "pytesseract",
        SimpleNamespace(
            image_to_string=lambda _image, lang: "Nota Ufficio Spoglio V Sez. penale R.G. 9966/2026"
        ),
    )

    result = extract_text_from_document(b"%PDF-1.7\n% scansione", "ordinanza.pdf", "pdf")

    assert result.ok is True
    assert result.extraction_engine == "pdfplumber+ocr"
    assert result.pages[0].page_number == 1
    assert "R.G. 9966/2026" in result.text
    assert any("OCR applicato" in warning for warning in result.warnings)


def test_document_ai_extraction_pdf_con_cid_residui_usa_ocr(monkeypatch):
    class FakePdfPlumberPage:
        def extract_text(self):
            return "(cid:0)(cid:1)(cid:2)(cid:3)(cid:4)(cid:5)"

    class FakePdfPlumberDocument:
        pages = [FakePdfPlumberPage()]

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

    class FakeBitmap:
        def to_pil(self):
            return object()

        def close(self):
            return None

    class FakePdfiumPage:
        def render(self, *, scale: float):
            assert scale >= 1.0
            return FakeBitmap()

        def close(self):
            return None

    class FakePdfiumDocument:
        def __init__(self, content: bytes):
            assert content.startswith(b"%PDF")

        def __len__(self):
            return 1

        def __getitem__(self, index: int):
            assert index == 0
            return FakePdfiumPage()

        def close(self):
            return None

    monkeypatch.setitem(
        sys.modules,
        "pdfplumber",
        SimpleNamespace(open=lambda _stream: FakePdfPlumberDocument()),
    )
    monkeypatch.setitem(sys.modules, "pypdfium2", SimpleNamespace(PdfDocument=FakePdfiumDocument))
    monkeypatch.setitem(
        sys.modules,
        "pytesseract",
        SimpleNamespace(image_to_string=lambda _image, lang: "Contratto scolastico Betti Alice"),
    )

    result = extract_text_from_document(b"%PDF-1.7\n% cid", "contratto.pdf", "pdf")

    assert result.ok is True
    assert result.extraction_engine == "pdfplumber+ocr"
    assert "(cid:" not in result.text
    assert "Contratto scolastico Betti Alice" in result.text
    assert any("CID" in warning for warning in result.warnings)


def test_document_ai_extraction_p7m_esterno_usa_payload_firmato(tmp_path: Path, monkeypatch):
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 720, "Testo estratto dal payload firmato")
    pdf.showPage()
    pdf.save()
    pdf_payload = buffer.getvalue()

    class _Status:
        payload_available = True
        payload_name = "atto.pdf"
        message = "Contenuto firmato estratto."

    class _Signed:
        status = _Status()
        payload_bytes = pdf_payload

    monkeypatch.setattr(
        "pct.firme_cades.inspect_signed_document_bytes",
        lambda **_kwargs: _Signed(),
    )
    target = tmp_path / "atto.pdf.p7m"
    target.write_bytes(b"PKCS7 signed envelope")

    result = extract_document_text(target, "pdf")

    assert result.error is None
    assert result.extraction_engine in {"cades:pdfplumber", "cades:pypdf"}
    assert "Testo estratto dal payload firmato" in result.text
    assert "Contenuto firmato estratto." in result.warnings


def test_document_ai_extraction_formato_non_supportato_controllata(tmp_path: Path):
    target = tmp_path / "programma.exe"
    target.write_bytes(b"contenuto")

    result = extract_document_text(target, "exe")

    assert result.error is not None
    assert result.text == ""
    assert result.extraction_engine == "exe.binary-best-effort"


def test_document_ai_extraction_txt_utf8(tmp_path: Path):
    target = tmp_path / "note-fascicolo.txt"
    target.write_text("Promemoria fascicolo\nCredito € 1.200", encoding="utf-8")

    result = extract_document_text(target, "txt")

    assert result.error is None
    assert result.extraction_engine.startswith("txt:")
    assert "Promemoria fascicolo" in result.text
    assert "Credito € 1.200" in result.text


def test_document_ai_extraction_eml_con_corpo_e_allegato_txt(tmp_path: Path):
    message = EmailMessage()
    message["Subject"] = "Diffida inviata"
    message["From"] = "avvocato@example.it"
    message["To"] = "cliente@example.it"
    message["Date"] = "Mon, 18 May 2026 10:00:00 +0200"
    message.set_content("Corpo email con termine per adempiere.")
    message.add_attachment(
        "Testo allegato per Lex".encode("utf-8"),
        maintype="text",
        subtype="plain",
        filename="allegato.txt",
    )
    target = tmp_path / "messaggio.eml"
    target.write_bytes(message.as_bytes())

    result = extract_document_text(target, "eml")

    assert result.error is None
    assert result.extraction_engine == "email.message"
    assert "Oggetto: Diffida inviata" in result.text
    assert "Corpo email con termine per adempiere." in result.text
    assert "[Allegato: allegato.txt]" in result.text
    assert "Testo allegato per Lex" in result.text


def test_document_ai_extraction_doc_legacy_controllata(tmp_path: Path):
    target = tmp_path / "atto.doc"
    target.write_bytes(b"DOC legacy")

    result = extract_document_text(target, "doc")

    assert result.error is not None
    assert result.text == ""
    assert any("DOC" in warning for warning in result.warnings)
