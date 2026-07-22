import base64
from html import unescape
from io import BytesIO
import re
import sys
from types import SimpleNamespace
import zipfile

from asn1crypto import algos, cms
from docx import Document
from PIL import Image
import pytest

import web.services.signed_attachment_preview_images as image_preview
import web.services.signed_attachment_preview_text as text_preview
import web.services.signed_attachment_preview_word as word_preview
from web.bootstrap.fascicoli_document_helpers import preview_unavailable_html
from web.services.signed_attachment_preview import build_attachment_preview_payload


def _zip_bytes(entries: list[tuple[str, bytes]]) -> bytes:
    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, data in entries:
            archive.writestr(name, data)
    return output.getvalue()


def _signed_payload(payload: bytes | None) -> bytes:
    signed = cms.SignedData(
        {
            "version": "v1",
            "digest_algorithms": [algos.DigestAlgorithm({"algorithm": "sha256"})],
            "encap_content_info": {
                "content_type": "data",
                **({"content": payload} if payload is not None else {}),
            },
            "signer_infos": [],
        }
    )
    return cms.ContentInfo({"content_type": "signed_data", "content": signed}).dump()


def test_zip_preview_rifiuta_percorso_non_sicuro():
    payload = build_attachment_preview_payload(
        nome_file="decreto.pdf.zip",
        data=_zip_bytes([("../decreto.pdf", b"%PDF-1.4\n")]),
        mime_salvato="application/zip",
    )

    assert payload.unavailable_reason == "L'archivio ZIP contiene un percorso non sicuro."


def test_zip_preview_rifiuta_compressione_anomala():
    payload = build_attachment_preview_payload(
        nome_file="decreto.pdf.zip",
        data=_zip_bytes([("decreto.pdf", b"0" * (1024 * 1024))]),
        mime_salvato="application/zip",
    )

    assert "compressione anomala" in payload.unavailable_reason


def test_zip_preview_rifiuta_directory_centrale_con_conteggio_incoerente():
    source = bytearray(_zip_bytes([("decreto.pdf", b"%PDF-1.4\n")]))
    eocd_offset = source.rfind(b"PK\x05\x06")
    source[eocd_offset + 10 : eocd_offset + 12] = (0).to_bytes(2, "little")

    payload = build_attachment_preview_payload(
        nome_file="decreto.pdf.zip",
        data=bytes(source),
        mime_salvato="application/zip",
    )

    assert "struttura valida" in payload.unavailable_reason


def test_preview_formati_testuali_professionali_restano_nel_lettore_interno():
    samples = (
        (
            "Comunicazione.xml",
            b'<?xml version="1.0" encoding="UTF-8"?><Comunicazione><Oggetto>Sentenza</Oggetto></Comunicazione>',
            "application/xml",
            ("Documento XML", "Sentenza"),
        ),
        (
            "messaggio.eml",
            b"Subject: Comunicazione di cancelleria\r\nFrom: ufficio@example.test\r\n\r\nCorpo PEC leggibile",
            "message/rfc822",
            ("Email PEC / EML", "Corpo PEC leggibile"),
        ),
        (
            "nota.txt",
            "Attività per l'avvocato".encode("utf-8"),
            "text/plain",
            ("Documento di testo", "Attività per l'avvocato"),
        ),
    )

    for name, data, mimetype, expected in samples:
        payload = build_attachment_preview_payload(
            nome_file=name,
            data=data,
            mime_salvato=mimetype,
        )

        assert payload.unavailable_reason == ""
        assert payload.mimetype == "text/html; charset=utf-8"
        html = unescape(payload.data.decode("utf-8"))
        assert all(value in html for value in expected)


def test_preview_docx_mostra_il_testo_nel_lettore_interno():
    document = Document()
    document.add_heading("Memoria difensiva", level=1)
    document.add_paragraph("Contenuto professionale leggibile senza uscire da IUSENTRA.")
    stream = BytesIO()
    document.save(stream)

    payload = build_attachment_preview_payload(
        nome_file="memoria.docx",
        data=stream.getvalue(),
        mime_salvato="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert payload.unavailable_reason == ""
    assert payload.mimetype == "text/html; charset=utf-8"
    html = payload.data.decode("utf-8")
    assert "Documento Word" in html
    assert "Memoria difensiva" in html
    assert "Contenuto professionale leggibile" in html


def test_preview_docx_rimuove_markup_eseguibile_da_conversione_mammoth(monkeypatch):
    malicious_html = """
        <h1 id="atto">Memoria difensiva</h1>
        <script>parent.document.body.dataset.compromesso = '1'</script>
        <iframe srcdoc="<script>alert(1)</script>"></iframe>
        <form action="/logout"><button type="submit">Esci</button></form>
        <img src="x" onerror="parent.alert(1)" alt="immagine non fidata">
        <a href="javascript:parent.alert(1)" onclick="parent.alert(2)">collegamento pericoloso</a>
        <a href="https://www.normattiva.it/" target="_top">fonte ammessa</a>
        <p style="background:url(javascript:alert(1))">Contenuto leggibile</p>
    """
    fake_mammoth = SimpleNamespace(
        convert_to_html=lambda _stream: SimpleNamespace(value=malicious_html, messages=[]),
    )
    monkeypatch.setitem(sys.modules, "mammoth", fake_mammoth)
    source = _zip_bytes(
        [
            ("[Content_Types].xml", b"<Types/>"),
            ("word/document.xml", b"<w:document/>"),
        ]
    )

    payload = build_attachment_preview_payload(
        nome_file="memoria.docx",
        data=source,
        mime_salvato="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    html = payload.data.decode("utf-8")
    assert payload.unavailable_reason == ""
    assert "Memoria difensiva" in html
    assert "Contenuto leggibile" in html
    assert "fonte ammessa" in html
    assert 'href="https://www.normattiva.it/"' in html
    assert 'rel="nofollow noopener noreferrer"' in html
    assert "<script" not in html.casefold()
    assert "<iframe" not in html.casefold()
    assert "<form" not in html.casefold()
    assert "<button" not in html.casefold()
    assert "javascript:" not in html.casefold()
    assert "onerror=" not in html.casefold()
    assert "onclick=" not in html.casefold()
    assert "style=" not in html.casefold()


def test_preview_docx_blocca_ooxml_con_compressione_anomala_prima_di_mammoth(monkeypatch):
    conversion_called = False

    def _unexpected_conversion(_stream):
        nonlocal conversion_called
        conversion_called = True
        raise AssertionError("Mammoth non deve aprire un OOXML oltre budget")

    monkeypatch.setitem(
        sys.modules,
        "mammoth",
        SimpleNamespace(convert_to_html=_unexpected_conversion),
    )
    monkeypatch.setattr(word_preview, "MAX_OOXML_COMPRESSION_RATIO", 2)
    source = _zip_bytes(
        [
            ("[Content_Types].xml", b"<Types/>"),
            ("word/document.xml", b"A" * 4_096),
        ]
    )

    payload = build_attachment_preview_payload(
        nome_file="memoria-anomala.docx",
        data=source,
        mime_salvato="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert conversion_called is False
    assert payload.data == source
    assert "compressione anomala" in payload.unavailable_reason
    assert "Scarica l'originale" in payload.unavailable_reason


def test_preview_eml_applica_budget_alla_sorgente_e_al_corpo(monkeypatch):
    oversized = b"Subject: Messaggio\r\n\r\n" + (b"A" * 64)
    monkeypatch.setattr(text_preview, "MAX_MIME_SOURCE_BYTES", 32)

    rejected = build_attachment_preview_payload(
        nome_file="messaggio.eml",
        data=oversized,
        mime_salvato="message/rfc822",
    )

    assert rejected.data == oversized
    assert "messaggio MIME supera il limite" in rejected.unavailable_reason

    monkeypatch.setattr(text_preview, "MAX_MIME_SOURCE_BYTES", 1_024)
    monkeypatch.setattr(text_preview, "MAX_MIME_BODY_BYTES", 16)
    bounded = build_attachment_preview_payload(
        nome_file="messaggio.eml",
        data=oversized,
        mime_salvato="message/rfc822",
    )

    html = bounded.data.decode("utf-8")
    assert bounded.unavailable_reason == ""
    assert "A" * 16 in html
    assert "A" * 17 not in html
    assert "anteprima è stata abbreviata" in html


def test_preview_eml_blocca_troppe_boundary_prima_del_parser(monkeypatch):
    monkeypatch.setattr(text_preview, "MAX_MIME_PARTS", 2)
    source = (
        b"Content-Type: multipart/mixed; boundary=x\r\n\r\n"
        + b"\r\n--x\r\nContent-Type: text/plain\r\n\r\na" * 5
        + b"\r\n--x--\r\n"
    )

    payload = build_attachment_preview_payload(
        nome_file="messaggio.eml",
        data=source,
        mime_salvato="message/rfc822",
    )

    assert payload.data == source
    assert "troppe parti" in payload.unavailable_reason


def test_preview_riconosce_smime_dal_mime_anche_senza_estensione_p7m():
    xml = b'<?xml version="1.0"?><Comunicazione><Oggetto>Udienza</Oggetto></Comunicazione>'

    payload = build_attachment_preview_payload(
        nome_file="smime-contenuto.bin",
        data=_signed_payload(xml),
        mime_salvato="application/pkcs7-mime",
    )

    assert payload.unavailable_reason == ""
    assert payload.extracted_from_signature is True
    assert payload.mimetype == "text/html; charset=utf-8"
    assert "Udienza" in payload.data.decode("utf-8")


def test_preview_firma_smime_detached_senza_originale_resta_scaricabile():
    signature = _signed_payload(None)

    payload = build_attachment_preview_payload(
        nome_file="smime.p7s",
        data=signature,
        mime_salvato="application/pkcs7-signature",
    )

    assert payload.data == signature
    assert payload.download_name == "smime.p7s"
    assert "contenuto interno leggibile" in payload.unavailable_reason


def test_preview_non_disponibile_escapa_nome_e_blocca_url_non_interna():
    html, status, headers = preview_unavailable_html(
        '<img src=x onerror="parent.alert(1)">.docx',
        "javascript:parent.alert(2)",
    )

    assert status == 200
    assert headers["Content-Type"] == "text/html; charset=utf-8"
    assert '<img src=x onerror="parent.alert(1)">' not in html
    assert "&lt;img src=x onerror=" in html
    assert 'href="#"' in html
    assert 'href="javascript:' not in html.casefold()


def test_preview_immagini_jpeg_png_e_tiff_usa_il_lettore_interno():
    for name, image_format, mimetype in (
        ("scansione.jpg", "JPEG", "image/jpeg"),
        ("scansione.png", "PNG", "image/png"),
        ("scansione.tiff", "TIFF", "image/tiff"),
    ):
        stream = BytesIO()
        Image.new("RGB", (18, 12), color=(255, 255, 255)).save(stream, format=image_format)

        payload = build_attachment_preview_payload(
            nome_file=name,
            data=stream.getvalue(),
            mime_salvato=mimetype,
        )

        assert payload.unavailable_reason == ""
        assert payload.mimetype == "text/html; charset=utf-8"
        html = payload.data.decode("utf-8")
        assert name in html
        assert "data:image/" in html


def test_preview_tiff_rifiuta_una_pagina_con_pixel_da_decompression_bomb(monkeypatch):
    monkeypatch.setattr(image_preview, "MAX_TIFF_FRAME_PIXELS", 4_000)
    stream = BytesIO()
    Image.new("1", (80, 60), color=1).save(stream, format="TIFF", compression="group4")
    source = stream.getvalue()

    payload = build_attachment_preview_payload(
        nome_file="scansione-enorme.tiff",
        data=source,
        mime_salvato="image/tiff",
    )

    assert payload.data == source
    assert payload.mimetype == "image/tiff"
    assert "limite di sicurezza per la decompressione" in payload.unavailable_reason
    assert "Scarica l'originale" in payload.unavailable_reason


def test_preview_tiff_multiframe_rispetta_budget_pixel_complessivo(monkeypatch):
    monkeypatch.setattr(image_preview, "MAX_TIFF_TOTAL_PIXELS", 500)
    frames = [Image.new("RGB", (18, 12), color=(index * 20, 40, 80)) for index in range(3)]
    stream = BytesIO()
    frames[0].save(stream, format="TIFF", save_all=True, append_images=frames[1:])

    payload = build_attachment_preview_payload(
        nome_file="fascicolo-multipagina.tiff",
        data=stream.getvalue(),
        mime_salvato="image/tiff",
    )

    assert payload.unavailable_reason == ""
    assert payload.mimetype == "text/html; charset=utf-8"
    html = payload.data.decode("utf-8")
    assert "Pagina 1" in html
    assert "Pagina 2" in html
    assert "Pagina 3" not in html
    assert "Anteprima limitata a 2 pagine" in html


def test_preview_tiff_riduce_la_pagina_e_rispetta_il_budget_output():
    stream = BytesIO()
    Image.new("RGB", (2_400, 600), color=(245, 245, 245)).save(stream, format="TIFF")

    payload = build_attachment_preview_payload(
        nome_file="scansione-orizzontale.tiff",
        data=stream.getvalue(),
        mime_salvato="image/tiff",
    )

    assert payload.unavailable_reason == ""
    html = payload.data.decode("utf-8")
    encoded = re.search(r'data:image/jpeg;base64,([^\"]+)', html)
    assert encoded is not None
    rendered = Image.open(BytesIO(base64.b64decode(encoded.group(1))))
    assert max(rendered.size) <= image_preview.MAX_TIFF_THUMBNAIL_DIMENSION
    assert len(payload.data) <= image_preview.MAX_TIFF_PREVIEW_OUTPUT_BYTES


def test_preview_tiff_blocca_l_html_quando_supera_il_budget_output(monkeypatch):
    monkeypatch.setattr(
        image_preview,
        "MAX_TIFF_PREVIEW_OUTPUT_BYTES",
        image_preview.TIFF_PREVIEW_SHELL_RESERVE_BYTES + 128,
    )
    stream = BytesIO()
    Image.new("RGB", (32, 24), color=(30, 90, 160)).save(stream, format="TIFF")
    source = stream.getvalue()

    payload = build_attachment_preview_payload(
        nome_file="scansione-budget.tiff",
        data=source,
        mime_salvato="image/tiff",
    )

    assert payload.data == source
    assert payload.mimetype == "image/tiff"
    assert "limite di memoria previsto dal lettore" in payload.unavailable_reason
    assert "Scarica l'originale" in payload.unavailable_reason


@pytest.mark.parametrize(
    ("image_format", "extension", "mimetype"),
    (
        ("JPEG", "jpg", "image/jpeg"),
        ("PNG", "png", "image/png"),
        ("GIF", "gif", "image/gif"),
    ),
)
def test_preview_raster_compressi_rifiutano_pixel_oltre_il_limite(
    monkeypatch,
    image_format,
    extension,
    mimetype,
):
    monkeypatch.setattr(image_preview, "MAX_RASTER_FRAME_PIXELS", 4_000)
    stream = BytesIO()
    Image.new("RGB", (80, 60), color=(245, 245, 245)).save(stream, format=image_format)
    source = stream.getvalue()

    payload = build_attachment_preview_payload(
        nome_file=f"scansione-enorme.{extension}",
        data=source,
        mime_salvato=mimetype,
    )

    assert payload.data == source
    assert payload.mimetype == mimetype
    assert "limite di sicurezza per la decompressione" in payload.unavailable_reason
    assert "Scarica l'originale" in payload.unavailable_reason


@pytest.mark.parametrize(
    ("image_format", "extension", "mimetype"),
    (
        ("JPEG", "jpg", "image/jpeg"),
        ("PNG", "png", "image/png"),
    ),
)
def test_preview_jpeg_png_grandi_producono_thumbnail_governata(
    image_format,
    extension,
    mimetype,
):
    stream = BytesIO()
    Image.new("RGB", (2_000, 500), color=(236, 240, 246)).save(stream, format=image_format)

    payload = build_attachment_preview_payload(
        nome_file=f"scansione-orizzontale.{extension}",
        data=stream.getvalue(),
        mime_salvato=mimetype,
    )

    assert payload.unavailable_reason == ""
    html = payload.data.decode("utf-8")
    encoded = re.search(r'data:image/jpeg;base64,([^\"]+)', html)
    assert encoded is not None
    rendered = Image.open(BytesIO(base64.b64decode(encoded.group(1))))
    assert max(rendered.size) <= image_preview.MAX_RASTER_THUMBNAIL_DIMENSION
    assert len(payload.data) <= image_preview.MAX_RASTER_PREVIEW_OUTPUT_BYTES


def test_preview_gif_mostra_solo_il_primo_fotogramma_governato():
    first = Image.new("RGB", (40, 30), color=(220, 20, 20))
    second = Image.new("RGB", (40, 30), color=(20, 20, 220))
    stream = BytesIO()
    first.save(stream, format="GIF", save_all=True, append_images=[second], duration=100, loop=0)

    payload = build_attachment_preview_payload(
        nome_file="sequenza.gif",
        data=stream.getvalue(),
        mime_salvato="image/gif",
    )

    assert payload.unavailable_reason == ""
    html = payload.data.decode("utf-8")
    assert "primo fotogramma" in html
    assert "data:image/gif" not in html
    encoded = re.search(r'data:image/jpeg;base64,([^\"]+)', html)
    assert encoded is not None
    rendered = Image.open(BytesIO(base64.b64decode(encoded.group(1)))).convert("RGB")
    red, _green, blue = rendered.getpixel((rendered.width // 2, rendered.height // 2))
    assert red > blue


def test_preview_raster_blocca_l_html_quando_supera_il_budget_output(monkeypatch):
    monkeypatch.setattr(
        image_preview,
        "MAX_RASTER_PREVIEW_OUTPUT_BYTES",
        1_000,
    )
    stream = BytesIO()
    Image.new("RGB", (32, 24), color=(30, 90, 160)).save(stream, format="PNG")
    source = stream.getvalue()

    payload = build_attachment_preview_payload(
        nome_file="scansione-budget.png",
        data=source,
        mime_salvato="image/png",
    )

    assert payload.data == source
    assert payload.mimetype == "image/png"
    assert "limite di memoria previsto dal lettore" in payload.unavailable_reason
    assert "Scarica l'originale" in payload.unavailable_reason
