"""
pct/editor.py — Conversione documenti per l'editor web.

Funzionalità:
  - docx_to_html()  : .docx → HTML con mammoth (fedele ai formati Word)
  - pdf_to_html()   : .pdf → HTML con pdfplumber (testo + struttura)
  - html_to_docx()  : HTML → .docx con python-docx + lxml
  - html_to_pdf()   : HTML → PDF con reportlab
  - txt_to_html()   : .txt → HTML semplice

Tutte le funzioni lavorano su bytes già decifrati.
Le dipendenze (mammoth, python-docx, reportlab) sono opzionali:
se non presenti, si solleva ImportError con messaggio chiaro.

Nota sul supporto PDF:
  Il PDF è un formato di presentazione, non di editing. La conversione
  PDF → HTML preserva testo e struttura di base (titoli, paragrafi,
  grassetto) ma non layout complessi, immagini o tabelle grafiche.
  Per PDF scansionati viene usato Tesseract OCR (lingua italiana).
"""
from __future__ import annotations

import base64
import io
import re
from pathlib import Path
from typing import Optional

# Estensioni supportate dall'editor
ESTENSIONI_EDITABILI = {".docx", ".txt", ".html", ".htm", ".pdf"}
_CID_TOKEN_RE = re.compile(r"\(cid:\s*\d+\)", re.IGNORECASE)


def estensione_editabile(nome_file: str) -> bool:
    return Path(nome_file).suffix.lower() in ESTENSIONI_EDITABILI


# ─────────────────────────────────────────────── docx → HTML

def docx_to_html(data: bytes) -> tuple[str, list[str]]:
    """
    Converte un file .docx in HTML tramite mammoth.

    Returns:
        (html, avvisi) — html pronto per TipTap, lista di avvisi di conversione
    """
    try:
        import mammoth
    except ImportError:
        return (
            "<p><em>Libreria mammoth non disponibile. "
            "Esegui: pip install mammoth</em></p>",
            ["mammoth non installato"]
        )
    try:
        result = mammoth.convert_to_html(io.BytesIO(data))
        html = result.value or "<p></p>"
        avvisi = [str(m) for m in result.messages]
        return html, avvisi
    except Exception as e:
        return f"<p><em>Errore conversione .docx: {e}</em></p>", [str(e)]


# ─────────────────────────────────────────────── pdf → HTML

# ─────────────────────────────────────────────── pdf → HTML

def pdf_to_html(data: bytes) -> tuple[str, list[str], bool, int]:
    """
    Converte un file PDF in HTML per l'editing.

    Strategia:
      1. Prova estrazione testo nativo con pdfplumber (PDF digitali)
      2. Se una pagina è vuota (PDF scansionato), usa Tesseract OCR
      3. Usa la dimensione del font per rilevare titoli (H1/H2/H3)
      4. Preserva grassetto/corsivo dai font names
      5. Estrae tabelle con formattazione HTML
      6. Raggruppa il testo in paragrafi per riga

    Returns:
        (html, avvisi, is_scanned, n_pagine)
        - html       : contenuto HTML pronto per TipTap
        - avvisi     : lista di messaggi informativi
        - is_scanned : True se almeno una pagina è stata processata via OCR
        - n_pagine   : numero totale di pagine nel PDF
    """
    try:
        import pdfplumber
    except ImportError:
        return (
            "<p><em>pdfplumber non disponibile.</em></p>",
            ["pdfplumber non installato"],
            False,
            0
        )

    avvisi: list[str] = []
    html_parti: list[str] = []
    is_scanned = False
    n_pagine = 0

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            n_pagine = len(pdf.pages)
            motivi_layout = _motivi_layout_pdf_non_fedele(pdf.pages)
            if motivi_layout:
                motivo = "; ".join(motivi_layout[:3])
                if len(motivi_layout) > 3:
                    motivo += f"; altre {len(motivi_layout) - 3} pagine con layout complesso"
                avvisi.append(
                    "Il PDF contiene elementi grafici o impaginazione non lineare "
                    f"({motivo}). L'editor usa l'anteprima originale per preservare "
                    "la resa del documento."
                )
                html_parti.append(_html_pdf_non_modificabile(1, "layout PDF complesso", visuale=True))
                return "\n".join(html_parti), avvisi, is_scanned, n_pagine
            
            for i, pagina in enumerate(pdf.pages):
                # ── Estrae tabelle con formattazione ──────────────
                tabelle = pagina.extract_tables()
                if tabelle:
                    for tabella in tabelle:
                        html_tabella = _estrai_tabella_html(tabella)
                        if html_tabella:
                            html_parti.append(html_tabella)

                # ── Raccoglie caratteri con font info ──────────────
                chars = pagina.chars
                testo_plain = pagina.extract_text() or ""

                if not testo_plain.strip() or not _testo_pdf_affidabile(testo_plain):
                    motivo = (
                        "font CID senza mappa Unicode"
                        if testo_plain.strip()
                        else "pagina priva di testo nativo"
                    )
                    testo_fallback = ""
                    if testo_plain.strip():
                        testo_alternativo = _estrai_testo_pymupdf(data, i)
                        if _testo_pdf_affidabile(testo_alternativo):
                            testo_fallback = testo_alternativo
                            avvisi.append(
                                f"Pagina {i+1}: testo nativo non affidabile ({motivo}), "
                                "estratto con motore PDF alternativo."
                            )
                    if not testo_fallback:
                        testo_ocr = _ocr_pagina(data, i, pagina)
                        if _testo_pdf_affidabile(testo_ocr):
                            testo_fallback = testo_ocr
                            is_scanned = True
                            avvisi.append(
                                f"Pagina {i+1}: testo nativo non affidabile ({motivo}), estratto via OCR."
                            )
                    if testo_fallback:
                        for blocco in _splitta_paragrafi(testo_fallback):
                            html_parti.append(f"<p>{_escape_html(blocco)}</p>")
                    else:
                        html_parti.append(_html_pdf_non_modificabile(i + 1, motivo))
                        avvisi.append(
                            f"Pagina {i+1}: testo PDF non leggibile automaticamente. "
                            "Apri l'anteprima originale o importa una versione DOCX/testo prima di modificare."
                        )
                    if n_pagine > 1 and i < n_pagine - 1:
                        html_parti.append('<hr class="page-break">')
                    continue

                # ── Estrae righe con dimensione font media ─────────
                righe = _estrai_righe_con_font(chars, pagina.extract_text_lines() or [])

                if not righe:
                    # Fallback: testo piano
                    for blocco in _splitta_paragrafi(testo_plain):
                        html_parti.append(f"<p>{_escape_html(blocco)}</p>")
                else:
                    html_parti.extend(_righe_to_html(righe))

                # Separatore di pagina visivo (tranne dopo l'ultima)
                if n_pagine > 1 and i < n_pagine - 1:
                    html_parti.append('<hr class="page-break">')

    except Exception as e:
        avvisi.append(f"Errore estrazione PDF: {e}")
        html_parti.append(f"<p><em>Errore: {_escape_html(str(e))}</em></p>")

    html = "\n".join(html_parti) if html_parti else "<p></p>"
    return html, avvisi, is_scanned, n_pagine


def _estrai_tabella_html(tabella: list) -> str:
    """
    Converte una tabella pdfplumber in HTML con formattazione.
    """
    if not tabella:
        return ""
    
    html = '<table class="pdf-table" style="border-collapse:collapse;width:100%;margin:1rem 0;">\n'
    
    for ri, riga in enumerate(tabella):
        html += '  <tr>\n'
        for cella in riga:
            if cella is not None and str(cella).strip():
                tag = 'th' if ri == 0 else 'td'
                contenuto = _escape_html(str(cella).strip())
                html += f'    <{tag} style="border:1px solid #999;padding:6px 8px;">{contenuto}</{tag}>\n'
        html += '  </tr>\n'
    
    html += '</table>\n'
    return html


def _testo_pdf_affidabile(testo: str) -> bool:
    """Riconosce estrazioni PDF inutilizzabili, in particolare token CID."""
    pulito = (testo or "").strip()
    if not pulito:
        return False
    cid_count = len(_CID_TOKEN_RE.findall(pulito))
    if cid_count >= 3:
        return False
    senza_cid = _CID_TOKEN_RE.sub("", pulito)
    lettere = len(re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9]", senza_cid))
    if lettere < 8 and len(pulito) > 40:
        return False
    return True


def _motivi_layout_pdf_non_fedele(pagine: list) -> list[str]:
    """Identifica pagine PDF che non possono essere rese fedelmente come HTML editabile."""
    motivi: list[str] = []
    for indice, pagina in enumerate(pagine, start=1):
        motivo = _motivo_layout_pagina_non_fedele(pagina)
        if motivo:
            motivi.append(f"pagina {indice}: {motivo}")
    return motivi


def _motivo_layout_pagina_non_fedele(pagina) -> str:
    """Rileva immagini, timbri, testo ruotato e disegni che richiedono preview nativa."""
    chars = list(getattr(pagina, "chars", []) or [])
    testo_chars = [c for c in chars if str(c.get("text", "")).strip()]
    non_upright = sum(1 for c in testo_chars if not c.get("upright", True))
    immagini = len(getattr(pagina, "images", []) or [])
    rects = len(getattr(pagina, "rects", []) or [])
    curves = len(getattr(pagina, "curves", []) or [])
    lines = len(getattr(pagina, "lines", []) or [])
    elementi_vettoriali = rects + curves + lines

    ragioni: list[str] = []
    if immagini:
        ragioni.append("immagini/stemmi")
    if non_upright >= 12 or (testo_chars and non_upright / max(1, len(testo_chars)) >= 0.03):
        ragioni.append("testo ruotato o laterale")
    if elementi_vettoriali >= 12:
        ragioni.append("riquadri, timbri o segni grafici")
    return ", ".join(ragioni)


def _estrai_testo_pymupdf(data: bytes, page_index: int) -> str:
    """Secondo motore di estrazione testo, utile su alcuni PDF con font embedded."""
    try:
        import fitz
    except ImportError:
        return ""
    doc = None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        if page_index >= len(doc):
            return ""
        return (doc[page_index].get_text("text") or "").strip()
    except Exception:
        return ""
    finally:
        if doc is not None:
            try:
                doc.close()
            except Exception:
                pass


def _ocr_pagina(data: bytes, page_index: int, pagina=None) -> str:
    """OCR su pagina PDF via pytesseract, con renderer PDF robusto."""
    try:
        import pytesseract
    except ImportError:
        return ""

    if pagina is not None:
        try:
            img = pagina.to_image(resolution=220).original
            testo = pytesseract.image_to_string(img, lang="ita").strip()
            if testo:
                return testo
        except Exception:
            pass

    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(data)
        page = pdf[page_index]
        bitmap = page.render(scale=2.3)
        try:
            img = bitmap.to_pil()
            return pytesseract.image_to_string(img, lang="ita").strip()
        finally:
            for obj in (bitmap, page, pdf):
                close = getattr(obj, "close", None)
                if callable(close):
                    try:
                        close()
                    except Exception:
                        pass
    except Exception:
        return ""


def _html_pdf_non_modificabile(numero_pagina: int, motivo: str, *, visuale: bool = False) -> str:
    if visuale:
        testo = (
            f"Pagina {numero_pagina}: il PDF contiene {motivo}. Per non alterare "
            "intestazioni, stemmi, timbri, testo verticale o spaziature, l'editor "
            "mostra l'anteprima originale e blocca il salvataggio inline. Per "
            "modificare il contenuto importa una versione DOCX/testo verificata."
        )
    else:
        testo = (
            f"Pagina {numero_pagina}: il testo del PDF non e' modificabile automaticamente "
            f"perche' l'estrazione ha restituito {motivo}. Apri l'anteprima originale, "
            "oppure importa una versione DOCX/testo verificata prima di salvare modifiche."
        )
    return (
        f'<section data-editor-disabled="true" data-editor-disabled-reason="{_escape_html(motivo)}">'
        "<p><strong>Testo PDF non modificabile automaticamente.</strong></p>"
        f"<p>{_escape_html(testo)}</p>"
        "</section>"
    )


def _estrai_righe_con_font(chars: list, lines_raw: list) -> list[dict]:
    """
    Costruisce lista di righe con testo, dimensione media font e flag bold/italic.
    Ogni dict: {testo, size, bold, italic}
    """
    if not chars:
        return []

    # Raggruppa chars per y (riga)
    riga_map: dict[float, list] = {}
    for c in chars:
        y = round(c.get("top", 0), 1)
        riga_map.setdefault(y, []).append(c)

    righe = []
    for y in sorted(riga_map.keys()):
        gruppo = sorted(riga_map[y], key=lambda c: c.get("x0", 0))
        testo = "".join(c.get("text", "") for c in gruppo).strip()
        if not testo:
            continue
        sizes = [c.get("size", 10) for c in gruppo if c.get("size")]
        avg_size = sum(sizes) / len(sizes) if sizes else 10
        fonts = [c.get("fontname", "").lower() for c in gruppo]
        is_bold = any("bold" in f or "black" in f for f in fonts)
        is_italic = any("italic" in f or "oblique" in f for f in fonts)
        righe.append({
            "testo": testo,
            "size": avg_size,
            "bold": is_bold,
            "italic": is_italic
        })

    return righe


def _righe_to_html(righe: list[dict]) -> list[str]:
    """
    Converte righe con font info in tag HTML.
    Usa dimensione relativa per classificare titoli:
      size > 1.4x media → H1
      size > 1.2x media → H2
      size > 1.05x media → H3
      altrimenti → <p>
    Preserva grassetto e corsivo.
    """
    if not righe:
        return []

    sizes = [r["size"] for r in righe]
    media = sum(sizes) / len(sizes)

    html: list[str] = []
    paragrafo: list[str] = []

    def _flush_paragrafo():
        if paragrafo:
            html.append("<p>" + " ".join(_escape_html(t) for t in paragrafo) + "</p>")
            paragrafo.clear()

    for r in righe:
        testo = r["testo"]
        size = r["size"]
        bold = r.get("bold", False)
        italic = r.get("italic", False)

        # Formatta testo con grassetto/corsivo
        testo_format = _escape_html(testo)
        if bold and italic:
            testo_format = f"<strong><em>{testo_format}</em></strong>"
        elif bold:
            testo_format = f"<strong>{testo_format}</strong>"
        elif italic:
            testo_format = f"<em>{testo_format}</em>"

        if size >= media * 1.4:
            _flush_paragrafo()
            html.append(f"<h1>{testo_format}</h1>")
        elif size >= media * 1.2:
            _flush_paragrafo()
            html.append(f"<h2>{testo_format}</h2>")
        elif size >= media * 1.05:
            _flush_paragrafo()
            html.append(f"<h3>{testo_format}</h3>")
        elif bold and len(testo) < 120:
            _flush_paragrafo()
            html.append(f"<h4>{testo_format}</h4>")
        elif testo == "":
            _flush_paragrafo()
        else:
            paragrafo.append(testo_format)

    _flush_paragrafo()
    return html


def _splitta_paragrafi(testo: str) -> list[str]:
    """Divide testo plain in paragrafi su righe vuote."""
    blocchi, buf = [], []
    for riga in testo.splitlines():
        if riga.strip():
            buf.append(riga.strip())
        elif buf:
            blocchi.append(" ".join(buf))
            buf = []
    if buf:
        blocchi.append(" ".join(buf))
    return blocchi or [testo.strip()]


def _escape_html(s: str) -> str:
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;")
             .replace('"', "&quot;"))


# ─────────────────────────────────────────────── txt → HTML

def txt_to_html(data: bytes) -> tuple[str, list[str]]:
    """Converte testo plain in HTML con paragrafi."""
    try:
        testo = data.decode("utf-8", errors="replace")
    except Exception:
        testo = ""
    righe = testo.splitlines()
    blocchi: list[str] = []
    buf: list[str] = []
    for riga in righe:
        if riga.strip():
            buf.append(riga)
        else:
            if buf:
                blocchi.append("<p>" + " ".join(buf) + "</p>")
                buf = []
    if buf:
        blocchi.append("<p>" + " ".join(buf) + "</p>")
    html = "\n".join(blocchi) if blocchi else "<p></p>"
    return html, []


# ─────────────────────────────────────────────── bytes → HTML (dispatcher)

def documento_to_html(data: bytes, nome_file: str) -> tuple[str, list[str], dict]:
    """
    Dispatcher: converte il documento in HTML in base all'estensione.

    Returns:
        (html, avvisi, meta)
        meta contiene: {tipo_originale, is_scanned, n_pagine, n_caratteri, ...}
    """
    ext = Path(nome_file).suffix.lower()

    if ext == ".docx":
        html, avvisi = docx_to_html(data)
        return html, avvisi, {
            "tipo_originale": "docx",
            "is_scanned": False,
            "n_pagine": 1,
            "n_caratteri": len(html)
        }

    if ext == ".pdf":
        html, avvisi, is_scanned, n_pagine = pdf_to_html(data)
        editor_disabled = 'data-editor-disabled="true"' in html
        disabled_reason = ""
        if editor_disabled:
            match = re.search(r'data-editor-disabled-reason="([^"]*)"', html)
            disabled_reason = match.group(1) if match else ""
        return html, avvisi, {
            "tipo_originale": "pdf",
            "is_scanned": is_scanned,
            "n_pagine": n_pagine,
            "n_caratteri": len(html),
            "layout_preservato": not is_scanned and not editor_disabled,
            "testo_affidabile": not editor_disabled and "(cid:" not in html.lower(),
            "editor_disabled": editor_disabled,
            "editor_disabled_reason": disabled_reason,
            "anteprima_originale_obbligatoria": editor_disabled,
        }

    if ext in (".txt",):
        html, avvisi = txt_to_html(data)
        return html, avvisi, {
            "tipo_originale": "txt",
            "is_scanned": False,
            "n_caratteri": len(html)
        }

    if ext in (".html", ".htm"):
        html = data.decode("utf-8", errors="replace")
        return html, [], {
            "tipo_originale": "html",
            "is_scanned": False,
            "n_caratteri": len(html)
        }

    return (
        "<p><em>Formato non supportato per la modifica inline.</em></p>",
        [],
        {"tipo_originale": ext.lstrip("."), "is_scanned": False, "n_caratteri": 0},
    )


# ─────────────────────────────────────────────── HTML → .docx

def html_to_docx(html: str, titolo: str = "Documento") -> bytes:
    """
    Converte HTML in formato .docx tramite python-docx.

    Supporta: paragrafi, H1–H4, grassetto, corsivo, sottolineato,
              liste puntate/numerate, tabelle (struttura di base).

    Returns:
        bytes del file .docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx non installato. Esegui: pip install python-docx"
        )
    try:
        from lxml import etree as ET
        from lxml.html import fragment_fromstring, fragments_fromstring
    except ImportError:
        raise ImportError("lxml non disponibile")

    doc = Document()

    # Margini pagina (A4 con margini legali italiani)
    for sec in doc.sections:
        sec.top_margin    = Inches(1.18)   # 3 cm
        sec.bottom_margin = Inches(0.98)   # 2.5 cm
        sec.left_margin   = Inches(1.57)   # 4 cm
        sec.right_margin  = Inches(0.98)   # 2.5 cm

    # Stile default paragrafo
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(12)

    # Parse HTML
    html_clean = f"<div>{html}</div>"
    try:
        root = ET.fromstring(html_clean.encode("utf-8"),
                             parser=ET.HTMLParser(encoding="utf-8"))
    except Exception:
        p = doc.add_paragraph()
        p.add_run(_strip_tags(html))
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    body = root.find(".//body") or root.find(".//div") or root

    def _process_node(el, paragraph=None):
        tag = (el.tag or "").lower().split("}")[-1]

        if tag in ("h1", "h2", "h3", "h4"):
            level = int(tag[1])
            hdg_style = f"Heading {level}"
            p = doc.add_paragraph(style=hdg_style)
            _add_runs(p, el)
            return

        if tag in ("ul", "ol"):
            for li in el.findall("li"):
                p = doc.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
                _add_runs(p, li)
            return

        if tag == "table":
            rows = el.findall(".//tr")
            if not rows:
                return
            cols = max(len(r.findall("td") + r.findall("th")) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = "Table Grid"
            for ri, row in enumerate(rows):
                cells = row.findall("th") + row.findall("td")
                for ci, cell in enumerate(cells[:cols]):
                    tbl.rows[ri].cells[ci].text = _strip_tags(
                        ET.tostring(cell, encoding="unicode", method="text")
                    )
            return

        if tag in ("p", "div"):
            p = doc.add_paragraph()
            # allineamento
            align = (el.get("style") or "")
            if "center" in align:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "right" in align:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif "justify" in align:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_runs(p, el)
            return

        if tag in ("br",):
            doc.add_paragraph()
            return

        # Fallback: processa figli
        for child in el:
            _process_node(child)

    def _add_runs(paragraph, el):
        """Aggiunge runs al paragrafo con formattazione inline."""
        # Testo diretto nell'elemento
        if el.text and el.text.strip():
            r = paragraph.add_run(el.text)
            _apply_inline(r, el.tag.lower())

        for child in el:
            child_tag = (child.tag or "").lower().split("}")[-1]
            if child_tag in ("strong", "b"):
                r = paragraph.add_run(child.text_content() if hasattr(child, 'text_content') else "")
                r.bold = True
            elif child_tag in ("em", "i"):
                r = paragraph.add_run(child.text_content() if hasattr(child, 'text_content') else "")
                r.italic = True
            elif child_tag in ("u",):
                r = paragraph.add_run(child.text_content() if hasattr(child, 'text_content') else "")
                r.underline = True
            elif child_tag == "span":
                r = paragraph.add_run(child.text_content() if hasattr(child, 'text_content') else "")
                # Colore testo da style inline
                style = child.get("style", "")
                m = re.search(r"color:\s*#([0-9a-fA-F]{6})", style)
                if m:
                    rgb = int(m.group(1), 16)
                    r.font.color.rgb = RGBColor(
                        (rgb >> 16) & 0xFF,
                        (rgb >> 8) & 0xFF,
                        rgb & 0xFF
                    )
            else:
                txt = child.text_content() if hasattr(child, 'text_content') else (child.text or "")
                if txt:
                    paragraph.add_run(txt)
            # tail (testo dopo il tag inline)
            if child.tail and child.tail.strip():
                paragraph.add_run(child.tail)

    def _apply_inline(run, tag):
        if tag in ("strong", "b"):
            run.bold = True
        elif tag in ("em", "i"):
            run.italic = True
        elif tag == "u":
            run.underline = True

    # Processa tutti i figli del body
    for child in (body if body is not None else []):
        try:
            _process_node(child)
        except Exception:
            pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────── HTML → PDF
# Usa reportlab (già in requirements) — nessuna dipendenza di sistema aggiuntiva

def html_to_pdf(html: str, titolo: str = "Documento", layout: Optional[dict] = None) -> bytes:
    """
    Converte HTML in PDF tramite reportlab (già installato).

    Gestisce: paragrafi, H1-H4, grassetto, corsivo, liste,
    tabelle di base, linee orizzontali e immagini base64.

    Il parametro opzionale ``layout`` permette di allineare font,
    interlinea e margini al preset dell'editor atti.

    Returns:
        bytes del file PDF
    Raises:
        ImportError se reportlab non è installato
    """
    try:
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, HRFlowable, ListFlowable, ListItem, Image as RLImage,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT, TA_LEFT
    except ImportError:
        raise ImportError("reportlab non installato. Esegui: pip install reportlab")

    try:
        from lxml import etree
    except ImportError:
        raise ImportError("lxml non installato")

    try:
        from pct.template_atti import font_editor, normalizza_editor_layout
        layout_cfg = normalizza_editor_layout(layout or {})
        font_meta = font_editor(layout_cfg["font_family"])
    except Exception:
        layout_cfg = {
            "font_size_pt": 12,
            "line_height": 1.9,
            "text_align": "justify",
            "margin_top_mm": 25,
            "margin_right_mm": 22,
            "margin_bottom_mm": 25,
            "margin_left_mm": 32,
            "paragraph_spacing_pt": 8,
        }
        font_meta = {"pdf_family": "times"}

    pdf_family = (font_meta.get("pdf_family") or "times").lower()
    font_bundle = {
        "times": {
            "normal": "Times-Roman",
            "bold": "Times-Bold",
            "italic": "Times-Italic",
            "bold_italic": "Times-BoldItalic",
        },
        "helvetica": {
            "normal": "Helvetica",
            "bold": "Helvetica-Bold",
            "italic": "Helvetica-Oblique",
            "bold_italic": "Helvetica-BoldOblique",
        },
        "courier": {
            "normal": "Courier",
            "bold": "Courier-Bold",
            "italic": "Courier-Oblique",
            "bold_italic": "Courier-BoldOblique",
        },
    }.get(pdf_family, {
        "normal": "Times-Roman",
        "bold": "Times-Bold",
        "italic": "Times-Italic",
        "bold_italic": "Times-BoldItalic",
    })

    font_size = layout_cfg["font_size_pt"]
    leading = round(font_size * layout_cfg["line_height"], 1)
    paragraph_spacing = layout_cfg["paragraph_spacing_pt"]
    alignment = {
        "justify": TA_JUSTIFY,
        "left": TA_LEFT,
        "center": TA_CENTER,
        "right": TA_RIGHT,
    }.get(layout_cfg["text_align"], TA_JUSTIFY)

    # ── Stili ────────────────────────────────────────────────
    base = getSampleStyleSheet()
    st_normal = ParagraphStyle(
        "LegalNormal", parent=base["Normal"],
        fontName=font_bundle["normal"], fontSize=font_size, leading=leading,
        spaceAfter=paragraph_spacing, alignment=alignment,
    )
    st_h1 = ParagraphStyle(
        "LegalH1", parent=base["Heading1"],
        fontName=font_bundle["bold"], fontSize=round(font_size * 1.45, 1), leading=round(leading * 1.15, 1),
        spaceBefore=12, spaceAfter=6,
        alignment=TA_CENTER,
    )
    st_h2 = ParagraphStyle(
        "LegalH2", parent=base["Heading2"],
        fontName=font_bundle["bold"], fontSize=round(font_size * 1.2, 1), leading=round(leading * 1.08, 1),
        spaceBefore=10, spaceAfter=4,
    )
    st_h3 = ParagraphStyle(
        "LegalH3", parent=base["Heading3"],
        fontName=font_bundle["bold"], fontSize=round(font_size * 1.08, 1), leading=round(leading * 1.02, 1),
        spaceBefore=8, spaceAfter=4,
    )
    st_h4 = ParagraphStyle(
        "LegalH4", parent=base["Normal"],
        fontName=font_bundle["bold"], fontSize=round(font_size * 1.02, 1), leading=round(leading, 1),
        spaceBefore=6, spaceAfter=4,
        underlineWidth=0.5,
    )
    st_li = ParagraphStyle(
        "LegalLI", parent=st_normal,
        leftIndent=18, spaceAfter=3,
    )
    st_quote = ParagraphStyle(
        "LegalQuote", parent=st_normal,
        leftIndent=36, rightIndent=18,
        fontName=font_bundle["italic"], textColor=colors.HexColor("#555555"),
    )
    st_caption = ParagraphStyle(
        "LegalCaption", parent=base["Normal"],
        fontName=font_bundle["italic"], fontSize=max(font_size - 2, 8), leading=max(leading - 2, 10),
        alignment=TA_CENTER, textColor=colors.HexColor("#64748b"), spaceAfter=6,
    )

    # Mappa tag → stile
    HEADING_STYLES = {"h1": st_h1, "h2": st_h2, "h3": st_h3, "h4": st_h4}

    # ── Parse HTML ───────────────────────────────────────────
    html_clean = f"<div>{html}</div>"
    try:
        root = etree.fromstring(
            html_clean.encode("utf-8"),
            parser=etree.HTMLParser(encoding="utf-8"),
        )
        body = root.find(".//body")
        if body is None:
            body = root
    except Exception:
        body = None

    # ── Conversione nodo → testo reportlab rich text ─────────
    def _node_to_rich(el) -> str:
        """Converte un elemento HTML in markup reportlab (para XML)."""
        tag = (el.tag or "").lower().split("}")[-1]
        text = (el.text or "")
        parts = [text]
        for child in el:
            child_tag = (child.tag or "").lower().split("}")[-1]
            inner = _node_to_rich(child)
            if child_tag in ("strong", "b"):
                parts.append(f"<b>{inner}</b>")
            elif child_tag in ("em", "i"):
                parts.append(f"<i>{inner}</i>")
            elif child_tag in ("u",):
                parts.append(f"<u>{inner}</u>")
            else:
                parts.append(inner)
            if child.tail:
                parts.append(child.tail)
        return "".join(parts)

    def _image_flowable_from_src(src: str):
        if not src:
            return None
        match = re.match(r"^data:image/[^;]+;base64,(.+)$", src, flags=re.I | re.S)
        if not match:
            return None
        try:
            image_bytes = base64.b64decode(match.group(1))
        except Exception:
            return None
        try:
            img = RLImage(io.BytesIO(image_bytes))
            usable_width = A4[0] - ((layout_cfg["margin_left_mm"] + layout_cfg["margin_right_mm"]) / 10.0 * cm)
            if img.drawWidth > usable_width:
                ratio = usable_width / float(img.drawWidth)
                img.drawWidth = usable_width
                img.drawHeight = img.drawHeight * ratio
            return img
        except Exception:
            return None

    # ── Costruisce flowables ──────────────────────────────────
    story = []

    def _process(el):
        tag = (el.tag or "").lower().split("}")[-1]

        if tag in HEADING_STYLES:
            rich = _node_to_rich(el)
            story.append(Paragraph(rich, HEADING_STYLES[tag]))
            return

        if tag in ("p", "div"):
            child_tags = [(child.tag or "").lower().split("}")[-1] for child in el]
            rich = _node_to_rich(el)
            if any(child_tag in ("img", "figure") for child_tag in child_tags) and not rich.strip():
                for child in el:
                    _process(child)
                return
            if rich.strip():
                story.append(Paragraph(rich, st_normal))
            else:
                story.append(Spacer(1, 0.3 * cm))
            return

        if tag in ("ul", "ol"):
            items = []
            for li in el.findall("li"):
                rich = _node_to_rich(li)
                items.append(ListItem(Paragraph(rich, st_li), bulletType="bullet" if tag == "ul" else "1"))
            if items:
                story.append(ListFlowable(items, bulletType="bullet" if tag == "ul" else "1",
                                          leftIndent=18, bulletFontSize=10))
            return

        if tag == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#999999"), spaceAfter=6))
            return

        if tag == "blockquote":
            rich = _node_to_rich(el)
            story.append(Paragraph(rich, st_quote))
            return

        if tag == "figure":
            for child in el:
                _process(child)
            story.append(Spacer(1, 0.15 * cm))
            return

        if tag == "figcaption":
            rich = _node_to_rich(el)
            if rich.strip():
                story.append(Paragraph(rich, st_caption))
            return

        if tag == "img":
            image = _image_flowable_from_src(el.get("src", ""))
            if image:
                story.append(image)
                story.append(Spacer(1, 0.2 * cm))
            return

        if tag == "table":
            rows_data = []
            for tr in el.findall(".//tr"):
                row = []
                for cell in tr.findall("th") + tr.findall("td"):
                    testo_cella = etree.tostring(cell, encoding="unicode", method="text").strip()
                    row.append(Paragraph(testo_cella, st_li))
                if row:
                    rows_data.append(row)
            if rows_data:
                col_n = max(len(r) for r in rows_data)
                for r in rows_data:
                    while len(r) < col_n:
                        r.append(Paragraph("", st_li))
                tbl = Table(rows_data, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("GRID",      (0, 0), (-1, -1), 0.5, colors.HexColor("#999999")),
                    ("BACKGROUND",(0, 0), (-1, 0),  colors.HexColor("#f0f0f0")),
                    ("FONTNAME",  (0, 0), (-1, 0),  font_bundle["bold"]),
                    ("FONTSIZE",  (0, 0), (-1, -1), max(font_size - 1, 9)),
                    ("TOPPADDING",(0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 0.3 * cm))
            return

        # Fallback: processa figli
        for child in el:
            try:
                _process(child)
            except Exception:
                pass

    if body is not None:
        for child in body:
            try:
                _process(child)
            except Exception:
                pass
    else:
        story.append(Paragraph(_strip_tags(html), st_normal))

    if not story:
        story.append(Paragraph(titolo, st_normal))

    # ── Genera PDF ───────────────────────────────────────────
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=(layout_cfg["margin_right_mm"] / 10.0) * cm,
        leftMargin=(layout_cfg["margin_left_mm"] / 10.0) * cm,
        topMargin=(layout_cfg["margin_top_mm"] / 10.0) * cm,
        bottomMargin=(layout_cfg["margin_bottom_mm"] / 10.0) * cm,
        title=titolo,
        author="IUSENTRA",
    )
    doc.build(story)
    return buf.getvalue()


# ─────────────────────────────────────────────── utility

def _strip_tags(html: str) -> str:
    """Rimuove tutti i tag HTML restituendo solo il testo."""
    return re.sub(r"<[^>]+>", "", html or "")

