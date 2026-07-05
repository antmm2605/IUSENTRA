"""
web/blueprints/template_atti.py - Generatore atti legali da template.

URL base: /template-atti/
"""
from __future__ import annotations

import io
import json
import os
import re
import zipfile
from datetime import date
from html import escape
from html.parser import HTMLParser
from types import SimpleNamespace
from typing import Any
from xml.etree import ElementTree

from flask import (Blueprint, abort, current_app, flash, g, jsonify,
                   redirect, render_template, request, send_file, url_for)
from werkzeug.utils import secure_filename

from web.helpers import get_clienti, get_fascicoli, get_soggetti, get_utenti
from web.blueprints.react_shell import render_react_shell_response
from web.services.app_v2_routing import is_safe_internal_path

template_atti = Blueprint("template_atti", __name__, url_prefix="/template-atti")


_UI_LABEL_REPLACEMENTS: tuple[tuple[str, str], ...] = (
    (r"\bdemo\b", "di prova"),
    (r"\bsample\b", "di prova"),
    (r"\brepository\b", "archivio"),
    (r"\bbackend\b", "servizio applicativo"),
    (r"\bfrontend\b", "interfaccia"),
    (r"\blegacy\b", "percorso precedente"),
    (r"\bpayload\b", "dati"),
    (r"\bruntime\b", "ambiente"),
    (r"\bjson_api\b", "servizio dati"),
    (r"\bprovider\b", "canale"),
    (r"\bwebhook\b", "notifica automatica"),
)


@template_atti.app_template_filter("template_atti_ui_label")
def template_atti_ui_label(value) -> str:
    """Rende sicure per la UI etichette provenienti da dati storici dello studio."""
    text = str(value or "")
    for pattern, replacement in _UI_LABEL_REPLACEMENTS:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return text


def _cfg_path(key: str, default: str = "") -> str:
    paths = getattr(g, "data_paths", {}) or {}
    if key in paths:
        return str(paths[key] or default)
    if getattr(g, "tenant_context_missing", False):
        raise RuntimeError(
            "Contesto studio non disponibile per la richiesta corrente. "
            "Accesso ai dati bloccato per evitare letture cross-studio."
        )
    return str(current_app.config.get(key, default) or default)


def _get_gt():
    from pct.template_atti import GestioneTemplateAtti
    return GestioneTemplateAtti(
        db_path=_cfg_path("TEMPLATE_ATTI_DB", "./template_atti/templates.json")
    )


def _get_gp():
    from pct.template_atti import GestionePreferenzeTemplateAtti, percorso_preferenze_editor

    prefs_path = _cfg_path("TEMPLATE_ATTI_PREFS_DB")
    if not prefs_path:
        prefs_path = percorso_preferenze_editor(
            _cfg_path("TEMPLATE_ATTI_DB", "./template_atti/templates.json")
        )
    return GestionePreferenzeTemplateAtti(prefs_path=prefs_path)


def _get_assistente_redazionale():
    from pct.assistente_redazionale import AssistenteRedazionale

    return AssistenteRedazionale(
        audit_db_path=_cfg_path("REDACTION_ASSISTANT_DB", "./intelligence/assistente_redazionale.json"),
        office_cache_path=current_app.config.get("UFFICI_GIUDIZIARI_DB", "") or current_app.config.get("REGINDE_DB", ""),
        pst_wsdl_catalog_zip_path=current_app.config.get("PST_WSDL_CATALOG_ZIP", ""),
        pst_official_cache_path=current_app.config.get("PST_OFFICIAL_CACHE", ""),
        pst_catalog_endpoint=current_app.config.get("PST_SOAP_CATALOGO_UG_ENDPOINT", ""),
        pst_timeout=float(current_app.config.get("PST_SOAP_TIMEOUT", 8.0) or 8.0),
    )


def _get_config_studio_manager():
    core_runtime = current_app.extensions.get("core_runtime", {}) or {}
    loader = core_runtime.get("get_config_studio")
    if callable(loader):
        return loader()
    from pct.config_studio import GestioneConfigStudio

    return GestioneConfigStudio(config_path=_cfg_path("CONFIG_STUDIO_DB", "./config/studio.json"))


def _studio_config_for_prefill() -> dict:
    """Espone al compilatore i dati studio reali, inclusi quelli salvati in Impostazioni."""
    config = dict(current_app.config)
    try:
        manager = _get_config_studio_manager()
        config_studio = getattr(manager, "config", None)
        studio = getattr(config_studio, "studio", None)
        pec = getattr(config_studio, "pec", None)
    except Exception:
        studio = None
        pec = None

    if studio is not None:
        studio_payload = {
            "nome": str(getattr(studio, "nome", "") or "").strip(),
            "avvocato": str(getattr(studio, "avvocato", "") or "").strip(),
            "indirizzo": str(getattr(studio, "indirizzo", "") or "").strip(),
            "cf": str(getattr(studio, "cf", "") or "").strip(),
            "piva": str(getattr(studio, "piva", "") or "").strip(),
            "codice_fiscale_avvocato": str(getattr(studio, "codice_fiscale_avvocato", "") or "").strip(),
        }
        config["studio"] = studio_payload
        if studio_payload["nome"]:
            config["STUDIO_NOME"] = studio_payload["nome"]
        if studio_payload["avvocato"]:
            config["STUDIO_AVVOCATO"] = studio_payload["avvocato"]
            config["PCT_STUDIO_AVVOCATO"] = studio_payload["avvocato"]
        if studio_payload["indirizzo"]:
            config["STUDIO_INDIRIZZO"] = studio_payload["indirizzo"]
        if studio_payload["cf"]:
            config["STUDIO_CF"] = studio_payload["cf"]
        if studio_payload["piva"]:
            config["STUDIO_PIVA"] = studio_payload["piva"]
    if pec is not None:
        pec_payload = {"indirizzo": str(getattr(pec, "indirizzo", "") or "").strip()}
        config["pec"] = pec_payload
        if pec_payload["indirizzo"]:
            config["PCT_STUDIO_PEC"] = pec_payload["indirizzo"]
            config.setdefault("SMTP_FROM", pec_payload["indirizzo"])
    return config


def _get_studio_timbro():
    from pct.studio_timbro import build_studio_timbro

    try:
        config_manager = _get_config_studio_manager()
        config_studio = getattr(config_manager, "config", None)
    except Exception:
        config_studio = None
    return build_studio_timbro(
        db_path=_cfg_path("STUDIO_TIMBRO_DB", "./config/studio_timbro.db"),
        config_studio=config_studio,
        app_config=current_app.config,
    )


@template_atti.context_processor
def _inject_editor_preferences():
    from pct.template_atti import DEFAULT_EDITOR_LAYOUT, catalogo_font_editor

    return {
        "editor_preferences": _get_gp().carica(),
        "editor_default_preferences": DEFAULT_EDITOR_LAYOUT,
        "editor_font_choices": catalogo_font_editor(),
    }


def _richiedi_login(f):
    from functools import wraps
    @wraps(f)
    def w(*a, **kw):
        if not g.get("utente_corrente"):
            return redirect(url_for("login"))
        return f(*a, **kw)
    return w


def _variabili_base(config):
    timbro = _get_studio_timbro()
    return {
        "data_oggi":         date.today().strftime("%d/%m/%Y"),
        "studio_nome":       config.get("STUDIO_NOME", "IUSENTRA"),
        "studio_indirizzo":  config.get("STUDIO_INDIRIZZO", ""),
        "studio_iban":       config.get("STUDIO_IBAN", ""),
        "avvocato_nome":     config.get("STUDIO_AVVOCATO", config.get("STUDIO_NOME", "Avvocato")),
        "studio_timbro":     timbro,
        "studio_timbro_payload": timbro.to_payload(),
        "studio_timbro_lines": timbro.to_lines(),
        "studio_timbro_text": timbro.to_text(),
        "studio_timbro_html": timbro.to_html(),
    }


def _namespace_from_mapping(mapping: dict | None = None):
    return SimpleNamespace(**(mapping or {}))


def _soggetto_template_namespace(soggetto, *, parte=None):
    if not soggetto:
        return None
    ruolo = getattr(parte, "ruolo", None)
    note = getattr(parte, "note", "") if parte else ""
    return SimpleNamespace(
        id=getattr(soggetto, "id", ""),
        nome_completo=getattr(soggetto, "nome_completo", ""),
        identificativo=getattr(soggetto, "identificativo", ""),
        codice_fiscale=getattr(soggetto, "codice_fiscale", ""),
        partita_iva=getattr(soggetto, "partita_iva", ""),
        qualifica=getattr(soggetto, "qualifica", ""),
        email=getattr(getattr(soggetto, "recapiti", None), "email", ""),
        pec=getattr(getattr(soggetto, "recapiti", None), "pec", ""),
        telefono=(
            getattr(getattr(soggetto, "recapiti", None), "telefono", "")
            or getattr(getattr(soggetto, "recapiti", None), "cellulare", "")
        ),
        indirizzo=str(getattr(soggetto, "indirizzo", "") or ""),
        ruolo=getattr(ruolo, "value", ""),
        ruolo_label=getattr(ruolo, "label", ""),
        note=note,
    )


def _build_parti_template_context(id_fascicolo: str):
    empty = _namespace_from_mapping(
        {
            "elenco": [],
            "assistiti": [],
            "controparti": [],
            "difensori_controparte": [],
            "altri": [],
            "assistito_principale": None,
            "controparte_principale": None,
            "difensore_controparte_principale": None,
        }
    )
    if not id_fascicolo:
        return [], empty

    pairs = get_soggetti().parti_fascicolo(id_fascicolo)
    elenco = [_soggetto_template_namespace(soggetto, parte=parte) for parte, soggetto in pairs]
    assistiti = [item for item in elenco if item and item.ruolo == "ASSISTITO"]
    controparti = [item for item in elenco if item and item.ruolo in {"CONTROPARTE", "CREDITORE", "DEBITORE"}]
    difensori_controparte = [item for item in elenco if item and item.ruolo == "DIFENSORE_CONTROPARTE"]
    altri = [
        item
        for item in elenco
        if item and item.ruolo not in {"ASSISTITO", "CONTROPARTE", "CREDITORE", "DEBITORE", "DIFENSORE_CONTROPARTE"}
    ]
    return elenco, _namespace_from_mapping(
        {
            "elenco": elenco,
            "assistiti": assistiti,
            "controparti": controparti,
            "difensori_controparte": difensori_controparte,
            "altri": altri,
            "assistito_principale": assistiti[0] if assistiti else None,
            "controparte_principale": controparti[0] if controparti else None,
            "difensore_controparte_principale": difensori_controparte[0] if difensori_controparte else None,
        }
    )


def _valore_form(value):
    if isinstance(value, list):
        return "\n".join([str(item) for item in value if str(item).strip()])
    return value or ""


def _split_csv(value: str) -> list[str]:
    return [item.strip() for item in str(value or "").split(",") if item.strip()]


def _build_correction_context() -> dict:
    intent = (request.args.get("intent", "") or "").strip()
    focus_field = (request.args.get("focus_field", "") or "").strip()
    highlight_fields = _split_csv(request.args.get("highlight_fields", ""))
    title = (request.args.get("correction_title", "") or "").strip()
    help_text = (request.args.get("correction_help", "") or "").strip()
    if not any([intent, focus_field, highlight_fields, title, help_text]):
        return {}
    return {
        "active": True,
        "intent": intent,
        "focus_field": focus_field,
        "highlight_fields": highlight_fields,
        "title": title or "Correzione guidata",
        "help": help_text or "Completa i campi evidenziati per superare il controllo di conformita'.",
    }


def _fallback_template_fields() -> list[dict]:
    return [
        {"name": "destinatario_nome", "label": "Destinatario", "type": "text", "placeholder": "Nome e cognome / ragione sociale", "section": "Campi rapidi", "rows": 1},
        {"name": "destinatario_indirizzo", "label": "Indirizzo / PEC", "type": "text", "placeholder": "Via, CAP, citta o PEC", "section": "Campi rapidi", "rows": 1},
        {"name": "oggetto_diffida", "label": "Oggetto / descrizione", "type": "textarea", "placeholder": "Descrivi l'oggetto della richiesta o del documento", "section": "Campi rapidi", "rows": 4},
        {"name": "importo_dovuto", "label": "Importo", "type": "text", "placeholder": "€ 0,00", "section": "Campi rapidi", "rows": 1},
        {"name": "titolo_credito", "label": "Titolo / causale", "type": "text", "placeholder": "es. canone, compenso, fattura", "section": "Campi rapidi", "rows": 1},
        {"name": "termine_giorni", "label": "Termine (giorni)", "type": "number", "placeholder": "15", "section": "Campi rapidi", "rows": 1},
        {"name": "tribunale_competente", "label": "Autorita / tribunale competente", "type": "text", "placeholder": "es. Tribunale di Milano", "section": "Campi rapidi", "rows": 1},
        {"name": "durata_anni", "label": "Durata anni", "type": "number", "placeholder": "5", "section": "Campi rapidi", "rows": 1},
    ]


def _group_template_fields(fields: list[dict] | None) -> list[dict]:
    grouped: dict[str, list[dict]] = {}
    ordered_sections: list[str] = []
    for field in fields or []:
        section = field.get("section") or "Contenuto"
        if section not in grouped:
            grouped[section] = []
            ordered_sections.append(section)
        grouped[section].append(field)
    return [{"title": title, "fields": grouped[title]} for title in ordered_sections]


def _prefill_template_fields(template, *, cliente=None, fascicolo=None, parti=None) -> dict[str, str]:
    assistito = getattr(parti, "assistito_principale", None) if parti else None
    controparte = getattr(parti, "controparte_principale", None) if parti else None
    values = {
        "autorita_competente": getattr(fascicolo, "tribunale", "") or "",
        "procedimento_riferimento": getattr(fascicolo, "numero_rg", "") or "",
        "oggetto_atto": getattr(fascicolo, "titolo", "") or "",
        "controparte_nome": getattr(controparte, "nome_completo", "") or "",
        "assistito_nome": getattr(assistito, "nome_completo", "") or getattr(cliente, "nome_completo", "") or "",
        "destinatario_nome": getattr(controparte, "nome_completo", "") or "",
        "domicilio_eletto": current_app.config.get("STUDIO_INDIRIZZO", "") or "",
        "termine_giorni": "15",
        "durata_anni": "5",
        "tribunale_competente": getattr(fascicolo, "tribunale", "") or "",
    }
    try:
        from pct.template_atti_prefill import resolve_template_prefill

        field_specs = list(getattr(template, "campi_guidati", []) or [])
        legacy_campi = [
            str(field.get("label") or field.get("name") or "")
            for field in field_specs
            if isinstance(field, dict)
        ]
        resolution = resolve_template_prefill(
            model_code=getattr(template, "link_compilatore_code", "") or getattr(template, "codice", ""),
            legacy_campi=legacy_campi,
            fascicolo=fascicolo,
            cliente=cliente,
            utente=g.get("utente_corrente"),
            config=_studio_config_for_prefill(),
            studio_timbro=_get_studio_timbro(),
            parti=parti,
            legacy_payload=values,
        )
        for key, result in resolution.get("fields", {}).items():
            mapped_key = f"dato_{key}"
            value = result.get("value")
            if value and mapped_key not in values:
                values[mapped_key] = _valore_form(value)
            if value and key not in values:
                values[key] = _valore_form(value)
        values["_prefill_resolution"] = resolution
    except Exception:
        pass
    return values


def _contesto_compilatore(model_code: str, *, payload: dict, selected_cliente=None,
                          selected_fascicolo=None, errors: dict | None = None,
                          assistant_analysis: dict | None = None,
                          correction_context: dict | None = None):
    from pct.compilatore_atti import (
        opzioni_campo,
        wizard_schema_modello,
    )
    schema = wizard_schema_modello(model_code)
    model = schema["model"]
    clienti = get_clienti().tutti()
    fascicoli = get_fascicoli().tutti()
    utenti = get_utenti().tutti(solo_attivi=True)
    base_fields = schema["base_fields"]
    extra_fields = schema["extra_fields"]
    field_options = {}
    for field in base_fields + extra_fields:
        if field["type"] == "select":
            field_options[field["name"]] = opzioni_campo(
                field["name"],
                fascicoli=fascicoli,
                utenti=utenti,
                model=model,
            )
    form_values = {key: _valore_form(value) for key, value in payload.items()}
    prefill_resolution = payload.get("_prefill_resolution") if isinstance(payload.get("_prefill_resolution"), dict) else {}
    try:
        timbro = _get_studio_timbro()
        studio_timbro_preview = {
            "lines": timbro.to_lines(),
            "html": timbro.to_html(),
            "text": timbro.to_text(),
        }
    except Exception:
        studio_timbro_preview = {"lines": [], "html": "", "text": ""}
    return {
        "model": model,
        "clienti": clienti,
        "fascicoli": fascicoli,
        "utenti": utenti,
        "base_fields": base_fields,
        "extra_fields": extra_fields,
        "field_options": field_options,
        "form_values": form_values,
        "payload": payload,
        "errors": errors or {},
        "selected_cliente": selected_cliente,
        "selected_fascicolo": selected_fascicolo,
        "guidance": schema["guidance"],
        "suggested_attachments": schema["suggested_attachments"],
        "suggested_clauses": schema["suggested_clauses"],
        "validation_rules": schema["validation_rules"],
        "sections": schema["sections"],
        "renderer_name": schema["renderer"],
        "prefill_map": schema["prefill_map"],
        "prefill_resolution": prefill_resolution,
        "studio_timbro_preview": studio_timbro_preview,
        "assistant_analysis": assistant_analysis or {},
        "correction_context": correction_context or {},
    }


def _resolve_compiler_context(model_code: str):
    from pct.compilatore_atti import (
        merge_payload_with_form,
        prefill_payload,
    )

    if request.method == "POST":
        id_cliente = (
            request.form.get("id_cliente", "").strip()
            or request.args.get("id_cliente", "").strip()
        )
        id_fascicolo = (
            request.form.get("id_fascicolo", "").strip()
            or request.form.get("case_id", "").strip()
            or request.args.get("id_fascicolo", "").strip()
        )
    else:
        id_cliente = (
            request.args.get("id_cliente", "").strip()
            or request.form.get("id_cliente", "").strip()
        )
        id_fascicolo = (
            request.args.get("id_fascicolo", "").strip()
            or request.form.get("id_fascicolo", "").strip()
            or request.form.get("case_id", "").strip()
        )

    clienti_repo = get_clienti()
    fascicoli_repo = get_fascicoli()
    selected_cliente = clienti_repo.get(id_cliente) if id_cliente else None
    selected_fascicolo = fascicoli_repo.get(id_fascicolo) if id_fascicolo else None
    if selected_fascicolo and not selected_cliente and getattr(selected_fascicolo, "id_cliente", ""):
        selected_cliente = clienti_repo.get(selected_fascicolo.id_cliente)
        id_cliente = getattr(selected_cliente, "id", "") if selected_cliente else id_cliente
    _, parti_prefill = _build_parti_template_context(id_fascicolo)

    initial_payload = prefill_payload(
        model_code,
        fascicolo=selected_fascicolo,
        cliente=selected_cliente,
        utente=g.get("utente_corrente"),
        config=_studio_config_for_prefill(),
        studio_timbro=_get_studio_timbro(),
        parti=parti_prefill,
    )
    payload = initial_payload
    if request.method == "POST":
        form_data = request.form.to_dict(flat=True)
        form_data["case_id"] = form_data.get("case_id", "").strip() or id_fascicolo
        if id_cliente and not form_data.get("client_or_sender"):
            form_data["client_or_sender"] = getattr(selected_cliente, "nome_completo", "")
        payload = merge_payload_with_form(
            model_code,
            initial_payload=initial_payload,
            form_data=form_data,
        )
    return {
        "id_cliente": id_cliente,
        "id_fascicolo": id_fascicolo,
        "selected_cliente": selected_cliente,
        "selected_fascicolo": selected_fascicolo,
        "initial_payload": initial_payload,
        "payload": payload,
        "correction_context": _build_correction_context(),
    }


def _build_assistant_analysis(model_code: str, *, payload: dict, selected_cliente=None, selected_fascicolo=None):
    result = _get_assistente_redazionale().analyze(
        model_code,
        payload,
        fascicolo=selected_fascicolo,
        cliente=selected_cliente,
        utente=g.get("utente_corrente"),
    ).to_dict()
    try:
        result["contextual_compliance"] = _build_contextual_compliance(
            model_code,
            payload=payload,
            selected_cliente=selected_cliente,
            selected_fascicolo=selected_fascicolo,
        ).to_dict()
    except Exception:
        current_app.logger.debug("Compliance contestuale non disponibile per %s", model_code, exc_info=True)
    return result


def _extract_editor_placeholders(text: str) -> list[str]:
    source = str(text or "")[:100_000]
    matches: set[str] = set()
    index = 0
    while index < len(source):
        if source.startswith("{{", index):
            end = source.find("}}", index + 2, index + 160)
            if end != -1:
                token = source[index:end + 2]
                if token[2:-2].strip():
                    matches.add(token)
                index = end + 2
                continue
        if source[index] == "[":
            end = source.find("]", index + 1, index + 90)
            if end != -1:
                candidate = source[index + 1:end]
                if candidate and all(ch == "_" or ch.isdigit() or "A" <= ch <= "Z" or "À" <= ch <= "Ù" for ch in candidate):
                    matches.add(source[index:end + 1])
                index = end + 1
                continue
        index += 1
    return sorted(matches)


def _editor_lex_enrichment(analysis: dict[str, Any], *, action: str, mode: str, text: str, html: str, instructions: str) -> dict[str, Any]:
    clean_text = " ".join(str(text or "").split())
    placeholder_matches = _extract_editor_placeholders(text)
    pii_matches = sorted(set(re.findall(r"\b[\w.+-]+@[\w.-]+\.\w+\b|\b\d{3}\s?\d{3}\s?\d{4}\b|\b[A-Z]{6}\d{2}[A-Z]\d{2}[A-Z]\d{3}[A-Z]\b", text or "", re.I)))
    has_legal_basis = bool(re.search(r"\bart\.|\barticolo\b|codice|d\.lgs\.|legge|regolamento|giurisprudenza|cassazione|consiglio di stato", clean_text, re.I))
    lowered = f"{action} {mode}".casefold()
    issues = list(analysis.get("issues") or [])
    next_steps = list(analysis.get("next_steps") or [])
    focus = "revisione redazionale"
    suggestion = "Rileggere il periodo selezionato, mantenendo fatti, domande e riferimenti del fascicolo separati."
    if "placeholder" in lowered:
        focus = "controllo placeholder"
        suggestion = "Sostituire o confermare ogni segnaposto residuo prima della versione finale."
        if placeholder_matches:
            issues.insert(0, {
                "level": "WARN",
                "title": "Segnaposto ancora presenti",
                "detail": ", ".join(placeholder_matches[:8]),
                "suggested_action": suggestion,
            })
    elif "privacy" in lowered:
        focus = "controllo privacy e dati personali"
        suggestion = "Verificare pertinenza, minimizzazione e base giuridica dei dati personali citati nell'atto."
        if pii_matches:
            issues.insert(0, {
                "level": "WARN",
                "title": "Dati personali individuati nel testo",
                "detail": ", ".join(pii_matches[:6]),
                "suggested_action": suggestion,
            })
    elif "normativ" in lowered or "legal_basis" in lowered:
        focus = "controllo fondamento normativo"
        suggestion = "Esplicitare norme, fonte, collegamento al fatto e conseguenza processuale senza automatismi non verificati."
        if not has_legal_basis:
            issues.insert(0, {
                "level": "WARN",
                "title": "Base normativa non evidente",
                "detail": "Nel testo corrente non risultano riferimenti normativi o giurisprudenziali riconoscibili.",
                "suggested_action": suggestion,
            })
    elif "formale" in lowered:
        focus = "formalizzazione del tono"
        suggestion = "Rendere la formulazione più istituzionale, evitando espressioni colloquiali e conclusioni non dimostrate."
    elif "chiaro" in lowered or "cliente" in lowered:
        focus = "chiarezza per il cliente"
        suggestion = "Separare fatto, diritto e prossima azione con periodi brevi e senza abbreviazioni non spiegate."
    elif "incisivo" in lowered:
        focus = "maggiore incisività"
        suggestion = "Portare in apertura il punto decisivo, poi documenti, norma e richiesta conclusiva."
    elif "clausola" in lowered:
        focus = "generazione clausola"
        suggestion = "Generare una clausola coerente con materia, parti, oggetto e limiti del fascicolo."
    elif "final" in lowered:
        focus = "controllo finale"
        suggestion = "Bloccare la versione finale finché restano segnaposti, dati mancanti o riferimenti privacy/normativi non verificati."
    if suggestion not in next_steps:
        next_steps.insert(0, suggestion)
    summary_parts = [
        f"Lex ha eseguito {focus} sul documento corrente.",
        f"Testo analizzato: {len(clean_text)} caratteri.",
    ]
    if placeholder_matches:
        summary_parts.append(f"Segnaposti residui: {len(placeholder_matches)}.")
    if pii_matches:
        summary_parts.append(f"Dati personali individuati: {len(pii_matches)}.")
    if instructions:
        summary_parts.append(f"Istruzioni dell'avvocato considerate: {instructions}.")
    analysis["summary"] = " ".join(summary_parts)
    analysis["issues"] = issues
    analysis["next_steps"] = next_steps[:8]
    analysis["editor_context"] = {
        "action": action,
        "mode": mode,
        "has_html": bool(html),
        "characters": len(clean_text),
        "placeholders": placeholder_matches[:20],
        "pii_candidates": pii_matches[:20],
        "legal_basis_detected": has_legal_basis,
    }
    return analysis


def _build_contextual_compliance(model_code: str, *, payload: dict, selected_cliente=None, selected_fascicolo=None):
    from pct.template_normative_compliance import analyze_template_compliance

    documents = []
    try:
        documents = list(getattr(selected_fascicolo, "documenti", []) or [])
    except Exception:
        documents = []
    return analyze_template_compliance(
        model_code,
        payload,
        fascicolo=selected_fascicolo,
        cliente=selected_cliente,
        documents=documents,
        studio_timbro_payload=payload.get("_studio_timbro") if isinstance(payload, dict) else None,
    )


def _audit_template_event(action: str, resource_type: str = "", resource_id: str = "", details: str = "") -> None:
    audit = (current_app.extensions.get("core_runtime", {}) or {}).get("audit")
    if not callable(audit):
        return
    try:
        audit(action, resource_type, resource_id, dettagli=details)
    except TypeError:
        try:
            audit(action, resource_type, resource_id, details)
        except Exception:
            current_app.logger.debug("Audit template atti non registrato per %s", action, exc_info=True)
    except Exception:
        current_app.logger.debug("Audit template atti non registrato per %s", action, exc_info=True)


def _safe_editor_filename(title: str, model_code: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", str(title or "").lower()).strip("._-")
    if not stem:
        stem = re.sub(r"[^A-Za-z0-9._-]+", "_", str(model_code or "atto").lower()).strip("._-") or "atto"
    return f"{stem[:70]}_{date.today().strftime('%Y%m%d')}.html"


def _sanitize_editor_html(value: str) -> str:
    try:
        from pct.editor_ai.editor_renderer import sanitize_editor_html

        return sanitize_editor_html(value)
    except Exception:
        return value or "<p></p>"


def _importa_compilazione_editor_professionale(
    *,
    model_code: str,
    model: dict,
    payload: dict,
    testo_generato: str,
    selected_fascicolo: Any,
    compliance_result: Any = None,
    requested_draft: str = "final_draft",
    confirmed_warning: bool = False,
    editor_html_builder: Any = None,
) -> dict[str, str] | None:
    id_fascicolo = str(getattr(selected_fascicolo, "id", "") or payload.get("case_id") or "").strip()
    if not id_fascicolo:
        return None

    from pct.template_atti_lex_service import create_editor_draft
    from web.services.document_crypto import encrypt_doc

    utente = g.get("utente_corrente")
    created = create_editor_draft(
        model_code,
        payload,
        id_fascicolo,
        getattr(utente, "username", "") if utente else "",
        rendered_text=testo_generato,
        model=model,
        fascicoli_repo=get_fascicoli(),
        encrypt_func=encrypt_doc,
        editor_html_builder=editor_html_builder or (lambda text: _sanitize_editor_html(_to_editor_html(text))),
        audit_callback=_audit_template_event,
        compliance_result=compliance_result,
        requested_draft=requested_draft,
        confirmed_warning=confirmed_warning,
        tenant_id=str(getattr(g, "tenant_id", "") or getattr(g, "studio_id", "") or ""),
        cliente_id=str(payload.get("id_cliente") or ""),
    )
    if not created.ok:
        raise RuntimeError("; ".join(created.errors or [created.message or "Importazione editor non riuscita."]))
    verify_repo = get_fascicoli()
    verify_fascicolo = verify_repo.get(id_fascicolo)
    verify_document = next(
        (
            item
            for item in (getattr(verify_fascicolo, "documenti", []) or [])
            if str(getattr(item, "id", "") or "").strip() == created.document_id
        ),
        None,
    ) if verify_fascicolo else None
    if verify_document is None:
        raise RuntimeError("Documento creato ma non leggibile dal fascicolo collegato.")
    try:
        from audit.integrations import emit_act_generated

        percorso = verify_repo.percorso_documento(id_fascicolo, created.document_id)
        emit_act_generated(
            fascicolo_id=id_fascicolo,
            atto_type=str(model.get("category") or model.get("name") or model_code or "ATTO"),
            template_id=str(model_code or ""),
            template_version=str(model.get("version") or model.get("updated") or "corrente"),
            editor_version="iusentra-editor",
            output_format="html",
            file_path=percorso,
            storage_ref=f"dms://fascicoli/{id_fascicolo}/documenti/{created.document_id}",
            idempotency_key=f"ACT_GENERATED:{id_fascicolo}:{created.document_id}",
        )
    except Exception:
        if current_app.config.get("AUDIT_ENABLED") or str(os.getenv("AUDIT_ENABLED", "")).lower() in {"1", "true", "yes", "on"}:
            raise
        current_app.logger.debug("Audit probatorio ACT_GENERATED non attivo per %s", created.document_id, exc_info=True)
    editor_url = url_for("editor_documento", id_fasc=id_fascicolo, id_doc=created.document_id)
    signature_url = url_for("firma_documento", id_fasc=id_fascicolo, id_doc=created.document_id)
    return {
        "document_id": created.document_id,
        "filename": created.filename,
        "open_url": editor_url,
        "editor_url": editor_url,
        "signature_url": signature_url,
        "created_as": created.created_as,
        "compliance_state": created.compliance_state,
    }


def _request_payload() -> dict:
    payload = request.get_json(silent=True)
    return payload if isinstance(payload, dict) else {}


def _safe_identifier(value: Any, *, max_length: int = 120) -> str:
    raw = str(value or "").strip()
    allowed = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789_.:-")
    if not raw or len(raw) > max_length or any(char not in allowed for char in raw):
        return ""
    return raw


def _canonical_route_template_code(value: Any) -> str:
    raw = str(value or "").strip()
    if not re.fullmatch(r"[A-Za-z0-9_.:-]{1,120}", raw):
        return ""
    requested = raw.upper()
    try:
        from pct.template_atti_unified_catalog import iter_all_templates

        for item in iter_all_templates():
            candidates = (
                item.get("codice"),
                item.get("original_id"),
                item.get("slug"),
                item.get("link_compilatore_code"),
                item.get("fallback_compilatore_code"),
            )
            for candidate in candidates:
                candidate_code = _safe_identifier(candidate)
                if candidate_code and candidate_code.upper() == requested:
                    return candidate_code
    except Exception:
        return ""
    return ""


def _safe_pdf_download_name(title: Any) -> str:
    base = secure_filename(str(title or "").strip().lower().replace(" ", "_"))
    if not base:
        base = "atto"
    if not base.endswith(".pdf"):
        base = f"{base}.pdf"
    return base


def _safe_word_download_name(title: Any) -> str:
    base = secure_filename(str(title or "").strip().lower().replace(" ", "_"))
    if not base:
        base = "atto"
    if not base.endswith(".docx"):
        base = f"{base}.docx"
    return base


def _json_safe_payload(value: Any) -> Any:
    if isinstance(value, str):
        return escape(value, quote=True)
    if isinstance(value, list):
        return [_json_safe_payload(item) for item in value]
    if isinstance(value, dict):
        return {str(key): _json_safe_payload(item) for key, item in value.items()}
    return value


def _value_present(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, (list, tuple, set, dict)):
        return bool(value)
    return str(value).strip() != ""


def _prefill_values(prefill: Any) -> dict[str, Any]:
    if not isinstance(prefill, dict):
        return {}
    values = prefill.get("values")
    return dict(values) if isinstance(values, dict) else {}


def _prefill_fields(prefill: Any) -> dict[str, Any]:
    if not isinstance(prefill, dict):
        return {}
    fields = prefill.get("fields")
    return dict(fields) if isinstance(fields, dict) else {}


def _preserved_server_field_names(user_values: dict[str, Any], prefill: Any) -> list[str]:
    allowed = sorted(set(_prefill_values(prefill)) | set(_prefill_fields(prefill)))
    return [field for field in allowed if _value_present(user_values.get(field))]


def _resolve_template_context_payload(payload: dict | None = None) -> dict:
    payload = payload or {}
    id_cliente = (
        str(payload.get("id_cliente") or payload.get("cliente_id") or "").strip()
        or request.args.get("id_cliente", "").strip()
    )
    id_fascicolo = (
        str(payload.get("id_fascicolo") or payload.get("case_id") or payload.get("pratica_id") or "").strip()
        or request.args.get("id_fascicolo", "").strip()
        or request.args.get("case_id", "").strip()
    )
    clienti_repo = get_clienti()
    fascicoli_repo = get_fascicoli()
    selected_cliente = clienti_repo.get(id_cliente) if id_cliente else None
    selected_fascicolo = fascicoli_repo.get(id_fascicolo) if id_fascicolo else None
    if selected_fascicolo and not selected_cliente and getattr(selected_fascicolo, "id_cliente", ""):
        selected_cliente = clienti_repo.get(selected_fascicolo.id_cliente)
        id_cliente = getattr(selected_cliente, "id", "") if selected_cliente else id_cliente
    _, parti_prefill = _build_parti_template_context(id_fascicolo)
    return {
        "id_cliente": id_cliente,
        "id_fascicolo": id_fascicolo,
        "selected_cliente": selected_cliente,
        "selected_fascicolo": selected_fascicolo,
        "parti": parti_prefill,
    }


def _resolve_template_prefill_response(codice: str, payload: dict | None = None) -> tuple[dict, int]:
    from pct.template_atti_prefill import resolve_template_prefill
    from pct.template_atti_unified_catalog import get_unified_template_item

    payload = payload or {}
    safe_code = _safe_identifier(codice)
    if not safe_code:
        return {"ok": False, "errore": "Template non trovato."}, 404
    item = get_unified_template_item(safe_code)
    if not item:
        return {"ok": False, "errore": "Template non trovato."}, 404
    context = _resolve_template_context_payload(payload)
    values = payload.get("values") if isinstance(payload.get("values"), dict) else {}
    dati = payload.get("dati") if isinstance(payload.get("dati"), dict) else {}
    legacy_payload = {**dati, **values}
    resolution = resolve_template_prefill(
        model_code=item.get("link_compilatore_code") or item.get("codice") or safe_code,
        bindings=item.get("prefill_bindings") if isinstance(item.get("prefill_bindings"), dict) else None,
        required_fields=list(item.get("required_prefill_fields") or item.get("prefill_required_fields") or []),
        optional_fields=list(item.get("optional_prefill_fields") or []),
        legacy_campi=list(item.get("dati_obbligatori") or []),
        fascicolo=context["selected_fascicolo"],
        cliente=context["selected_cliente"],
        utente=g.get("utente_corrente"),
        config=_studio_config_for_prefill(),
        studio_timbro=_get_studio_timbro(),
        parti=context["parti"],
        legacy_payload=legacy_payload,
    )
    return {
        "ok": True,
        "codice": _safe_identifier(item.get("codice") or safe_code) or "template",
        "titolo": item.get("titolo"),
        "context": {
            "id_cliente": getattr(context["selected_cliente"], "id", "") if context.get("selected_cliente") else "",
            "id_fascicolo": getattr(context["selected_fascicolo"], "id", "") if context.get("selected_fascicolo") else "",
        },
        "prefill": resolution,
    }, 200


# ================================================================ LISTA

@template_atti.route("/", methods=["GET"])
@_richiedi_login
def lista():
    gt = _get_gt()
    templates = gt.tutti()
    from pct.template_atti import CATEGORIE
    from pct.compilatore_atti import modelli_per_area
    builtin_templates = [t for t in templates if t.builtin]
    custom_templates = [t for t in templates if not t.builtin]
    area_options = sorted({t.area for t in builtin_templates if t.area})
    branch_options = sorted({t.branca for t in builtin_templates if t.branca})
    channel_options = sorted({t.canale_telematico for t in builtin_templates if t.canale_telematico})
    profile_options = sorted({t.profilo_deposito for t in builtin_templates if t.profilo_deposito})
    collections: dict[str, list] = {}
    for tmpl in builtin_templates:
        key = tmpl.collezione or tmpl.categoria or "Altro"
        collections.setdefault(key, []).append(tmpl)
    workspace_sections = [
        {
            "title": title,
            "count": len(items),
            "templates": sorted(items, key=lambda item: (item.ordine, item.titolo.lower())),
        }
        for title, items in collections.items()
    ]
    workspace_sections.sort(key=lambda section: section["templates"][0].ordine if section["templates"] else 9999)
    featured_titles = {
        "Atto di citazione",
        "Ricorso per decreto ingiuntivo",
        "Atto di precetto",
        "Ricorso per separazione giudiziale",
        "Ricorso lavoro",
        "Nomina del difensore di fiducia",
    }
    featured_templates = [t for t in builtin_templates if t.titolo in featured_titles][:6]
    stats = {
        "totale": len(builtin_templates),
        "con_compilatore": sum(1 for t in builtin_templates if t.link_compilatore_code),
        "canali": len(channel_options),
        "collezioni": len(workspace_sections),
        "custom": len(custom_templates),
    }
    return render_template(
        "template_atti/lista.html",
        templates=templates,
        builtin_templates=builtin_templates,
        custom_templates=custom_templates,
        workspace_sections=workspace_sections,
        area_options=area_options,
        branch_options=branch_options,
        channel_options=channel_options,
        profile_options=profile_options,
        featured_templates=featured_templates,
        stats=stats,
        categorie=CATEGORIE,
        modelli_compilatore=modelli_per_area(),
    )


@template_atti.route("/scheda/<id_template>", methods=["GET"])
@_richiedi_login
def scheda(id_template: str):
    gt = _get_gt()
    template = gt.get(id_template)
    if not template:
        abort(404)
    related = []
    for candidate in gt.tutti():
        if candidate.id == template.id:
            continue
        if template.branca and candidate.branca == template.branca:
            related.append(candidate)
        elif template.area and candidate.area == template.area:
            related.append(candidate)
        if len(related) >= 8:
            break
    compiler_model = None
    if template.link_compilatore_code:
        from pct.compilatore_atti import get_modello

        compiler_model = get_modello(template.link_compilatore_code)
    field_sections = _group_template_fields(template.campi_guidati or _fallback_template_fields())
    return render_template(
        "template_atti/scheda.html",
        t=template,
        related_templates=related,
        compiler_model=compiler_model,
        field_sections=field_sections,
    )


# ================================================================ CATALOGO ATTI

@template_atti.route("/catalogo", methods=["GET"])
@_richiedi_login
def catalogo():
    from pct.compilatore_atti import MODELS, AREA_LABELS, AREA_ORDINE, get_essential_docs
    from pct.template_catalog_service import build_template_catalog_page_context

    # Raggruppa i modelli del compilatore per area con metadati
    area_groups: list[dict] = []
    for area_key in AREA_ORDINE:
        modelli = []
        for m in MODELS:
            if m["area"] == area_key:
                modelli.append({**m, "essential_docs": get_essential_docs(m["code"])})
        if modelli:
            area_groups.append({
                "area_key": area_key,
                "area_label": AREA_LABELS.get(area_key, area_key.title()),
                "modelli": modelli,
            })

    # Catalogo piatto per JS (ricerca client-side)
    catalogo_flat = []
    for grp in area_groups:
        for m in grp["modelli"]:
            catalogo_flat.append({
                "code":  m["code"],
                "name":  m["name"],
                "area":  grp["area_key"],
                "area_label": grp["area_label"],
                "url":   url_for("template_atti.compila", model_code=m["code"]),
            })

    # Tipologie del wizard preventivi con mapping verso atti (per la sezione "da pratica")
    try:
        from pct.motore_preventivo import catalogo_wizard as _cw
        wizard_cat = _cw()
    except Exception:
        wizard_cat = {}
    from pct.compilatore_atti import PRATICA_TO_MODELS, MODEL_INDEX
    wizard_tipologie = []
    for area_name, items in wizard_cat.items():
        for tip in items:
            pid = tip.get("id", "")
            codes = PRATICA_TO_MODELS.get(pid, [])
            modelli_atti = [{"code": c, "name": MODEL_INDEX[c]["name"]} for c in codes if c in MODEL_INDEX]
            if modelli_atti:
                wizard_tipologie.append({
                    "id": pid,
                    "label": tip.get("label", pid),
                    "area": area_name,
                    "modelli_atti": modelli_atti,
                })

    catalog_context = build_template_catalog_page_context()

    return render_template(
        "template_atti/catalogo.html",
        area_groups=area_groups,
        catalogo_flat=catalogo_flat,
        wizard_tipologie=wizard_tipologie,
        studio_timbro_preview={
            "lines": _get_studio_timbro().to_lines(),
            "text": _get_studio_timbro().to_text(),
        },
        **catalog_context,
        oggi=date.today(),
    )


@template_atti.route("/catalogo/data", methods=["GET"])
@_richiedi_login
def catalogo_data():
    from pct.template_catalog_service import build_template_catalog_page_context
    from pct.template_atti_inventory import build_template_inventory, template_inventory_stats

    context = build_template_catalog_page_context()
    records = build_template_inventory()
    source_records = build_template_inventory(include_source_duplicates=True)
    inventory = template_inventory_stats(records, source_records=source_records)
    return jsonify(
        {
            "ok": True,
            "suite": context["suite_summary"],
            "inventory": inventory,
            "templates": context["template_suite"],
            "filters": context["template_filters"],
        }
    )


@template_atti.route("/catalogo/filters", methods=["GET"])
@_richiedi_login
def catalogo_filters():
    from pct.template_catalog_service import build_template_catalog_page_context

    context = build_template_catalog_page_context()
    return jsonify({"ok": True, "filters": context["template_filters"], "quick_filters": context["quick_filters"]})


@template_atti.route("/inventory/data", methods=["GET"])
@_richiedi_login
def template_inventory_data():
    from pct.template_atti_inventory import (
        build_template_inventory,
        detect_duplicate_templates,
        detect_missing_cartabia_profile,
        detect_missing_prefill_bindings,
        detect_missing_timbro_support,
        detect_templates_not_reachable_from_ui,
        template_inventory_stats,
    )

    records = build_template_inventory()
    source_records = build_template_inventory(include_source_duplicates=True)
    stats = template_inventory_stats(records, source_records=source_records)
    return jsonify(
        {
            "ok": True,
            "stats": stats,
            "records": records[:250],
            "samples": {
                "senza_cartabia": detect_missing_cartabia_profile(records)[:50],
                "senza_prefill": detect_missing_prefill_bindings(records)[:50],
                "senza_timbro": detect_missing_timbro_support(records)[:50],
                "duplicati": detect_duplicate_templates(source_records)[:25],
                "non_raggiungibili_ui": detect_templates_not_reachable_from_ui(records)[:50],
            },
        }
    )


@template_atti.route("/<codice>/prefill", methods=["GET"])
@_richiedi_login
def template_prefill_data(codice: str):
    safe_code = _canonical_route_template_code(codice)
    if not safe_code:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    payload, status = _resolve_template_prefill_response(safe_code)
    return jsonify(_json_safe_payload(payload)), status


@template_atti.route("/<codice>/prefill/resolve", methods=["POST"])
@_richiedi_login
def template_prefill_resolve(codice: str):
    safe_code = _canonical_route_template_code(codice)
    if not safe_code:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    payload, status = _resolve_template_prefill_response(safe_code, {})
    return jsonify(_json_safe_payload(payload)), status


@template_atti.route("/<codice>/prefill/merge", methods=["POST"])
@_richiedi_login
def template_prefill_merge(codice: str):
    from pct.template_atti_prefill import merge_prefill_values

    safe_code = _canonical_route_template_code(codice)
    if not safe_code:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    request_payload = _request_payload()
    resolved_payload, status = _resolve_template_prefill_response(safe_code, {})
    if status != 200:
        return jsonify(_json_safe_payload(resolved_payload)), status
    user_values = request_payload.get("user_values") if isinstance(request_payload.get("user_values"), dict) else {}
    if not user_values:
        user_values = request_payload.get("values") if isinstance(request_payload.get("values"), dict) else {}
    prefill = resolved_payload.get("prefill") or {}
    merge = merge_prefill_values({}, _prefill_values(prefill), _prefill_fields(prefill))
    merge["preserved_user_inputs"] = _preserved_server_field_names(user_values, prefill)
    return jsonify(_json_safe_payload({"ok": True, "codice": resolved_payload.get("codice"), "merge": merge, "prefill": prefill}))


@template_atti.route("/<codice>/cartabia-compliance", methods=["GET"])
@_richiedi_login
def template_cartabia_compliance(codice: str):
    from pct.template_cartabia_rules import verifica_cartabia_template
    from pct.template_atti_unified_catalog import get_unified_template_item

    safe_code = _canonical_route_template_code(codice)
    if not safe_code:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    item = get_unified_template_item(safe_code) if safe_code else None
    if not item:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    return jsonify(_json_safe_payload({"ok": True, "codice": _safe_identifier(item.get("codice") or safe_code) or "template", "cartabia": verifica_cartabia_template(item)}))


@template_atti.route("/<codice>/verifica-cartabia", methods=["POST"])
@_richiedi_login
def template_verifica_cartabia(codice: str):
    from pct.template_cartabia_rules import verifica_cartabia_template
    from pct.template_atti_unified_catalog import get_unified_template_item

    safe_code = _canonical_route_template_code(codice)
    if not safe_code:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    item = get_unified_template_item(safe_code) if safe_code else None
    if not item:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    prefill_payload, _ = _resolve_template_prefill_response(safe_code, {})
    prefill = prefill_payload.get("prefill") if prefill_payload.get("ok") else {}
    result = verifica_cartabia_template(item, prefill_resolution=prefill, payload=_prefill_values(prefill), strict_data_check=True)
    return jsonify(_json_safe_payload({"ok": result.get("ok", False), "codice": _safe_identifier(item.get("codice") or safe_code) or "template", "cartabia": result}))


@template_atti.route("/<codice>/verifica-completa", methods=["POST"])
@_richiedi_login
def template_verifica_completa(codice: str):
    from pct.template_atti_prefill import merge_prefill_values
    from pct.template_catalog_service import verifica_deposito_template as _verifica_deposito
    from pct.template_cartabia_rules import verifica_cartabia_template
    from pct.template_atti_unified_catalog import get_unified_template_item

    safe_code = _canonical_route_template_code(codice)
    if not safe_code:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    request_payload = _request_payload()
    item = get_unified_template_item(safe_code) if safe_code else None
    if not item:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    prefill_payload, status = _resolve_template_prefill_response(safe_code, {})
    if status != 200:
        return jsonify(_json_safe_payload(prefill_payload)), status
    prefill = prefill_payload.get("prefill") or {}
    user_values = request_payload.get("values") if isinstance(request_payload.get("values"), dict) else {}
    merge = merge_prefill_values({}, _prefill_values(prefill), _prefill_fields(prefill))
    merge["preserved_user_inputs"] = _preserved_server_field_names(user_values, prefill)
    cartabia = verifica_cartabia_template(item, prefill_resolution=prefill, payload=merge.get("values"), strict_data_check=True)
    deposito = _verifica_deposito(
        item.get("codice") or safe_code,
        {
            "dati": merge.get("values") or {},
            "allegati": {},
        },
    )
    blocking = bool(cartabia.get("controlli_bloccanti") or deposito.get("dati_mancanti") or deposito.get("allegati_mancanti"))
    return jsonify(_json_safe_payload({
        "ok": not blocking,
        "codice": _safe_identifier(item.get("codice") or safe_code) or "template",
        "prefill": prefill,
        "merge": merge,
        "cartabia": cartabia,
        "deposito": deposito,
        "stato": "verificato" if not blocking else "richiede_revisione",
    }))


@template_atti.route("/<codice>/compliance", methods=["GET"])
@_richiedi_login
def template_compliance(codice: str):
    from pct.template_catalog_service import get_template_catalog_item

    item = get_template_catalog_item(codice)
    if not item:
        return jsonify({"ok": False, "errore": "Template non trovato."}), 404
    try:
        from pct.template_normative_compliance import analyze_template_compliance

        contextual = analyze_template_compliance(
            item.get("link_compilatore_code") or item.get("codice") or codice,
            {},
            template_name=item.get("titolo") or codice,
            studio_timbro_payload=_get_studio_timbro().to_payload(),
            strict_payload=False,
        ).to_dict()
    except Exception:
        contextual = {}
    return jsonify(
        {
            "ok": True,
            "codice": item["codice"],
            "titolo": item["titolo"],
            "canale_deposito": item["canale_deposito"],
            "portale_deposito": item["portale_deposito"],
            "ruleset_version": item["compliance_summary"]["ruleset_version"],
            "cartabia_ruleset_version": item.get("cartabia_ruleset_version"),
            "cartabia": item.get("cartabia_verifica"),
            "contextual_compliance": contextual,
            "stato_conformita": item.get("stato_conformita"),
            "controlli": item["controlli_conformita_dettaglio"],
        }
    )


@template_atti.route("/<codice>/verifica-deposito", methods=["POST"])
@_richiedi_login
def verifica_deposito_template(codice: str):
    from pct.template_catalog_service import verifica_deposito_template as _verifica

    payload = request.get_json(silent=True) or {}
    result = _verifica(codice, payload)
    return jsonify(result), 200 if result.get("ok") or result.get("codice") else 404


@template_atti.route("/api/modelli-per-pratica/<id_pratica>", methods=["GET"])
@_richiedi_login
def api_modelli_per_pratica(id_pratica: str):
    from pct.compilatore_atti import get_modelli_per_pratica
    modelli = get_modelli_per_pratica(id_pratica)
    return jsonify([{"code": m["code"], "name": m["name"], "area": m["area"]} for m in modelli])


# ================================================================ NUOVO TEMPLATE

@template_atti.route("/nuovo", methods=["GET", "POST"])
@_richiedi_login
def nuovo():
    from pct.template_atti import CATEGORIE

    if request.method == "POST":
        gt = _get_gt()
        titolo   = request.form.get("titolo", "").strip()
        categoria = request.form.get("categoria", "Altro")
        corpo    = request.form.get("corpo", "").strip()
        note     = request.form.get("note", "").strip()
        if not titolo or not corpo:
            flash("Titolo e corpo obbligatori.", "danger")
            return render_template("template_atti/form.html", categorie=CATEGORIE,
                                   t=None, form=request.form)
        t = gt.crea(titolo=titolo, categoria=categoria, corpo=corpo, note=note)
        flash(f"Template '{t.titolo}' creato.", "success")
        return redirect(url_for("template_atti.lista"))
    return render_template("template_atti/form.html", categorie=CATEGORIE, t=None, form={})


# ================================================================ MODIFICA

@template_atti.route("/<id_template>/modifica", methods=["GET", "POST"])
@_richiedi_login
def modifica(id_template: str):
    from pct.template_atti import CATEGORIE
    gt = _get_gt()
    t = gt.get(id_template)
    if not t:
        abort(404)
    if t.builtin:
        flash("I modelli integrati non possono essere modificati. Clonalo prima.", "warning")
        return redirect(url_for("template_atti.lista"))
    if request.method == "POST":
        try:
            gt.aggiorna(id_template,
                        titolo=request.form.get("titolo", t.titolo),
                        categoria=request.form.get("categoria", t.categoria),
                        corpo=request.form.get("corpo", t.corpo),
                        note=request.form.get("note", t.note))
            flash("Template aggiornato.", "success")
        except ValueError as e:
            current_app.logger.warning("Template atti non aggiornato: %s", e)
            flash("Template non aggiornato: controlla i dati inseriti.", "danger")
        return redirect(url_for("template_atti.lista"))
    return render_template("template_atti/form.html", categorie=CATEGORIE, t=t, form=t.to_dict())


# ================================================================ CLONA

@template_atti.route("/<id_template>/clona", methods=["POST"])
@_richiedi_login
def clona(id_template: str):
    gt = _get_gt()
    t = gt.get(id_template)
    if not t:
        abort(404)
    nuovo = gt.crea(
        titolo=f"[Copia] {t.titolo}",
        categoria=t.categoria,
        corpo=t.corpo,
        note=t.note,
    )
    flash(f"Template clonato: '{nuovo.titolo}'.", "success")
    return redirect(url_for("template_atti.modifica", id_template=nuovo.id))


# ================================================================ ELIMINA

@template_atti.route("/<id_template>/elimina", methods=["POST"])
@_richiedi_login
def elimina(id_template: str):
    gt = _get_gt()
    try:
        gt.elimina(id_template)
        flash("Template eliminato.", "success")
    except ValueError as e:
        current_app.logger.warning("Template atti non eliminato: %s", e)
        flash("Template non eliminato: elemento non disponibile.", "warning")
    return redirect(url_for("template_atti.lista"))


# ================================================================ USA TEMPLATE (form variabili)

@template_atti.route("/<id_template>/usa", methods=["GET", "POST"])
@_richiedi_login
def usa(id_template: str):
    gt = _get_gt()
    t = gt.get(id_template)
    if not t:
        abort(404)
    clienti   = get_clienti().tutti()
    fascicoli_tutti = get_fascicoli().tutti()
    selected_cliente = None
    selected_fascicolo = None
    selected_cliente_id = request.args.get("id_cliente", "").strip()
    selected_fascicolo_id = request.args.get("id_fascicolo", "").strip()
    if selected_cliente_id:
        selected_cliente = get_clienti().get(selected_cliente_id)
    if selected_fascicolo_id:
        selected_fascicolo = get_fascicoli().get(selected_fascicolo_id)
    if selected_fascicolo and not selected_cliente and getattr(selected_fascicolo, "id_cliente", ""):
        selected_cliente = get_clienti().get(selected_fascicolo.id_cliente)
        selected_cliente_id = getattr(selected_cliente, "id", selected_cliente_id) if selected_cliente else selected_cliente_id
    _, parti_prefill = _build_parti_template_context(selected_fascicolo_id)
    field_specs = t.campi_guidati or _fallback_template_fields()
    field_sections = _group_template_fields(field_specs)
    form_values = _prefill_template_fields(
        t,
        cliente=selected_cliente,
        fascicolo=selected_fascicolo,
        parti=parti_prefill,
    )

    if request.method == "POST":
        f = request.form
        id_cliente   = f.get("id_cliente", "").strip()
        id_fascicolo = f.get("id_fascicolo", "").strip()

        cliente  = get_clienti().get(id_cliente) if id_cliente else None
        fascicolo = get_fascicoli().get(id_fascicolo) if id_fascicolo else None

        variabili = _variabili_base(current_app.config)
        soggetti_fascicolo, parti = _build_parti_template_context(id_fascicolo)
        variabili["cliente"]             = cliente
        variabili["fascicolo"]           = fascicolo
        variabili["soggetti"]            = soggetti_fascicolo
        variabili["parti"]               = parti
        submitted_fields = {}
        for field in field_specs:
            value = f.get(field["name"], "")
            variabili[field["name"]] = value
            submitted_fields[field["name"]] = value
        for legacy_name, default in {
            "destinatario_nome": "",
            "destinatario_indirizzo": "",
            "oggetto_diffida": "",
            "importo_dovuto": "",
            "titolo_credito": "",
            "termine_giorni": "15",
            "tribunale_competente": "",
            "durata_anni": "5",
        }.items():
            variabili.setdefault(legacy_name, f.get(legacy_name, default))
            submitted_fields.setdefault(legacy_name, variabili[legacy_name])

        try:
            testo_generato = gt.renderizza(id_template, variabili)
        except Exception as e:
            flash(f"Errore nella generazione: {e}", "danger")
            testo_generato = ""

        return render_template(
            "template_atti/anteprima.html",
            t=t,
            testo_generato=testo_generato,
            editor_html=_to_editor_html(testo_generato),
            id_template=id_template,
            variabili_json=_variabili_safe(variabili),
            cliente=cliente,
            fascicolo=fascicolo,
            submitted_fields=submitted_fields,
        )

    fascicoli = []
    if selected_cliente_id:
        fascicoli = [f for f in fascicoli_tutti if f.id_cliente == selected_cliente_id]

    return render_template(
        "template_atti/usa.html",
        t=t,
        clienti=clienti,
        fascicoli=fascicoli,
        field_sections=field_sections,
        form_values=form_values,
        selected_cliente_id=selected_cliente_id,
        selected_fascicolo_id=selected_fascicolo_id,
    )


@template_atti.route("/editor", methods=["GET"])
@template_atti.route("/editor-libero", methods=["GET"])
@_richiedi_login
def editor_libero():
    return render_react_shell_response(request.path.lstrip("/"))


@template_atti.route("/compila/<model_code>", methods=["GET", "POST"])
@_richiedi_login
def compila(model_code: str):
    from pct.compilatore_atti import (
        get_modello,
        validate_payload,
        render_compiled_act,
    )
    model = get_modello(model_code)
    if not model:
        abort(404)
    resolved = _resolve_compiler_context(model_code)
    selected_cliente = resolved["selected_cliente"]
    selected_fascicolo = resolved["selected_fascicolo"]
    payload = resolved["payload"]
    correction_context = resolved["correction_context"]
    assistant_analysis = _build_assistant_analysis(
        model_code,
        payload=payload,
        selected_cliente=selected_cliente,
        selected_fascicolo=selected_fascicolo,
    )
    guide_mode = bool(
        str(request.form.get("guida_pratica") or request.args.get("guida_pratica") or "").strip()
        or str(request.args.get("origine") or request.form.get("origine") or "").strip().lower() == "guida_pratica"
    )

    if request.method == "POST":
        errors = validate_payload(model_code, payload)
        blockers = [issue for issue in assistant_analysis.get("issues", []) if issue.get("level") == "BLOCK"]
        compliance_result = _build_contextual_compliance(
            model_code,
            payload=payload,
            selected_cliente=selected_cliente,
            selected_fascicolo=selected_fascicolo,
        )
        if errors:
            flash("Completa i campi obbligatori evidenziati prima di generare l'atto.", "warning")
            if request.form.get("_react_return"):
                return jsonify({
                    "ok": False,
                    "message": "Completa i campi obbligatori evidenziati prima di generare l'atto.",
                    "errors": errors,
                    "compliance": compliance_result.to_dict(),
                }), 400
            ctx = _contesto_compilatore(
                model_code,
                payload=payload,
                selected_cliente=selected_cliente,
                selected_fascicolo=selected_fascicolo,
                errors=errors,
                assistant_analysis=assistant_analysis,
                correction_context=correction_context,
            )
            return render_template(
                "template_atti/compilatore.html",
                **ctx,
            )
        if blockers and not guide_mode:
            first = blockers[0]
            flash(
                f"Bozza bloccata: {first.get('title')}. {first.get('suggested_action', '')}".strip(),
                "warning",
            )
            if request.form.get("_react_return"):
                return jsonify({
                    "ok": False,
                    "message": f"Bozza bloccata: {first.get('title')}. {first.get('suggested_action', '')}".strip(),
                    "compliance": compliance_result.to_dict(),
                }), 400
            ctx = _contesto_compilatore(
                model_code,
                payload=payload,
                selected_cliente=selected_cliente,
                selected_fascicolo=selected_fascicolo,
                assistant_analysis=assistant_analysis,
                correction_context=correction_context,
            )
            return render_template("template_atti/compilatore.html", **ctx)

        requested_draft = request.form.get("requested_draft", "final_draft").strip() or "final_draft"
        confirmed_warning = request.form.get("confirmed_warning") in {"1", "true", "on", "si"}
        if guide_mode:
            requested_draft = "working_draft"
            confirmed_warning = True
        if compliance_result.overall_state == "block" and not guide_mode:
            message = "La bozza finale è bloccata dai controlli applicabili."
            flash(message, "warning")
            if request.form.get("_react_return"):
                return jsonify({
                    "ok": False,
                    "message": message,
                    "created_as": "blocked",
                    "compliance_state": "block",
                    "compliance": compliance_result.to_dict(),
                }), 400
            ctx = _contesto_compilatore(
                model_code,
                payload=payload,
                selected_cliente=selected_cliente,
                selected_fascicolo=selected_fascicolo,
                assistant_analysis={**assistant_analysis, "contextual_compliance": compliance_result.to_dict()},
                correction_context=correction_context,
            )
            return render_template("template_atti/compilatore.html", **ctx)
        if compliance_result.overall_state == "warning" and not confirmed_warning:
            message = "I controlli richiedono conferma: posso creare solo una bozza di lavoro da revisionare."
            flash(message, "warning")
            if request.form.get("_react_return"):
                return jsonify({
                    "ok": False,
                    "message": message,
                    "requires_confirmation": True,
                    "created_as": "warning_requires_confirmation",
                    "compliance_state": "warning",
                    "compliance": compliance_result.to_dict(),
                }), 409
            ctx = _contesto_compilatore(
                model_code,
                payload=payload,
                selected_cliente=selected_cliente,
                selected_fascicolo=selected_fascicolo,
                assistant_analysis={**assistant_analysis, "contextual_compliance": compliance_result.to_dict()},
                correction_context=correction_context,
            )
            return render_template("template_atti/compilatore.html", **ctx)
        if compliance_result.overall_state == "warning":
            requested_draft = "working_draft"

        testo_generato = render_compiled_act(model_code, payload)
        if selected_fascicolo:
            try:
                editor_import = _importa_compilazione_editor_professionale(
                    model_code=model_code,
                    model=model,
                    payload=payload,
                    testo_generato=testo_generato,
                    selected_fascicolo=selected_fascicolo,
                    compliance_result=compliance_result,
                    requested_draft=requested_draft,
                    confirmed_warning=confirmed_warning,
                )
                if editor_import:
                    document_id = str(editor_import["document_id"])
                    id_fascicolo = str(getattr(selected_fascicolo, "id", "") or "").strip()
                    editor_url = url_for("editor_documento", id_fasc=id_fascicolo, id_doc=document_id)
                    signature_url = url_for("firma_documento", id_fasc=id_fascicolo, id_doc=document_id)
                    if not is_safe_internal_path(editor_url):
                        raise RuntimeError("Percorso editor generato non valido.")
                    flash("Bozza importata nell'editor professionale. Puoi impaginarla e salvarla nel fascicolo.", "success")
                    if request.form.get("_react_return"):
                        return jsonify({
                            "ok": True,
                            "message": "Bozza creata nell'editor professionale.",
                            "document_id": document_id,
                            "editor_url": editor_url,
                            "signature_url": signature_url,
                            "redirect": editor_url,
                            "created_as": editor_import.get("created_as") or requested_draft,
                            "compliance_state": editor_import.get("compliance_state") or compliance_result.overall_state,
                            "compliance": compliance_result.to_dict(),
                        })
                    return redirect(editor_url)
            except Exception as exc:
                current_app.logger.exception("Import editor professionale non riuscito per %s: %s", model_code, exc)
                flash("Bozza creata, ma l'importazione nell'editor professionale non è riuscita. Resta disponibile l'anteprima.", "warning")
        ctx = _contesto_compilatore(
            model_code,
            payload=payload,
            selected_cliente=selected_cliente,
            selected_fascicolo=selected_fascicolo,
            assistant_analysis=assistant_analysis,
            correction_context=correction_context,
        )
        return render_template(
            "template_atti/anteprima_compilatore.html",
            model=model,
            payload=payload,
            form_values={key: _valore_form(value) for key, value in payload.items()},
            testo_generato=testo_generato,
            editor_html=_to_editor_html(testo_generato),
            selected_cliente=selected_cliente,
            selected_fascicolo=selected_fascicolo,
            guidance=ctx["guidance"],
            suggested_attachments=ctx["suggested_attachments"],
            suggested_clauses=ctx["suggested_clauses"],
            sections=ctx["sections"],
            renderer_name=ctx["renderer_name"],
            assistant_analysis=assistant_analysis,
            correction_context=correction_context,
        )

    return render_template(
        "template_atti/compilatore.html",
        **_contesto_compilatore(
            model_code,
            payload=payload,
            selected_cliente=selected_cliente,
            selected_fascicolo=selected_fascicolo,
            assistant_analysis=assistant_analysis,
            correction_context=correction_context,
        ),
    )


@template_atti.route("/compila/<model_code>/guida-pratica/salva", methods=["POST"])
@_richiedi_login
def compila_guida_pratica_salva(model_code: str):
    from pct.compilatore_atti import get_modello, render_compiled_act
    from pct.template_normative_compliance import analyze_template_compliance

    model = get_modello(model_code)
    if not model:
        abort(404)
    resolved = _resolve_compiler_context(model_code)
    selected_cliente = resolved["selected_cliente"]
    selected_fascicolo = resolved["selected_fascicolo"]
    payload = dict(resolved["payload"] or {})
    if not selected_fascicolo:
        return jsonify({
            "ok": False,
            "message": "Seleziona prima il fascicolo: il documento deve rientrare in una pratica reale.",
        }), 400

    testo_generato = request.form.get("testo_generato", "").strip()
    if not testo_generato:
        testo_generato = render_compiled_act(model_code, payload, include_timbro=False)
    payload["title"] = request.form.get("title", "").strip() or payload.get("title") or model.get("name") or model_code
    payload["case_id"] = request.form.get("case_id", "").strip() or request.form.get("id_fascicolo", "").strip() or getattr(selected_fascicolo, "id", "")
    payload["id_cliente"] = request.form.get("id_cliente", "").strip() or payload.get("id_cliente") or getattr(selected_fascicolo, "id_cliente", "")
    if request.form.get("guida_pratica"):
        payload["guida_pratica"] = request.form.get("guida_pratica", "").strip()

    try:
        compliance_result = analyze_template_compliance(
            model_code,
            payload=payload,
            fascicolo=selected_fascicolo,
            cliente=selected_cliente,
            documents=list(getattr(selected_fascicolo, "documenti", []) or []),
            studio_timbro_payload=payload.get("_studio_timbro") if isinstance(payload, dict) else None,
            strict_payload=False,
        )
        editor_import = _importa_compilazione_editor_professionale(
            model_code=model_code,
            model=model,
            payload=payload,
            testo_generato=testo_generato,
            selected_fascicolo=selected_fascicolo,
            compliance_result=compliance_result,
            requested_draft="working_draft",
            confirmed_warning=True,
        )
    except Exception as exc:
        current_app.logger.exception("Salvataggio anteprima Guida Pratica non riuscito per %s: %s", model_code, exc)
        return jsonify({
            "ok": False,
            "message": "Documento non salvato nel fascicolo. Controlla i dati essenziali e riprova.",
        }), 500

    if not editor_import:
        return jsonify({
            "ok": False,
            "message": "Documento non salvato: fascicolo non collegato.",
        }), 400
    return jsonify({
        "ok": True,
        "message": "Documento salvato nel fascicolo come bozza modificabile.",
        "document_id": editor_import["document_id"],
        "editor_url": editor_import["editor_url"],
        "signature_url": editor_import["signature_url"],
        "redirect_url": editor_import["editor_url"],
        "created_as": editor_import.get("created_as") or "working_draft",
        "compliance_state": editor_import.get("compliance_state") or "warning",
    })


@template_atti.route("/compila/<model_code>/guida-pratica/anteprima", methods=["POST"])
@_richiedi_login
def compila_guida_pratica_anteprima(model_code: str):
    from pct.compilatore_atti import get_modello, render_compiled_act

    model = get_modello(model_code)
    if not model:
        abort(404)
    try:
        resolved = _resolve_compiler_context(model_code)
        payload = dict(resolved["payload"] or {})
        testo_generato = render_compiled_act(model_code, payload, include_timbro=False)
    except Exception as exc:
        current_app.logger.exception("Anteprima Guida Pratica non generata per %s: %s", model_code, exc)
        return jsonify({
            "ok": False,
            "message": "Anteprima non aggiornata. Il testo già visibile resta modificabile.",
        }), 500
    return jsonify({
        "ok": True,
        "message": "Anteprima aggiornata dal template compilato.",
        "testo_generato": testo_generato,
        "text": testo_generato,
    })


@template_atti.route("/api/editor-layout", methods=["POST"])
@_richiedi_login
def salva_editor_layout():
    try:
        from pct.template_atti import normalizza_editor_layout

        payload = request.get_json(silent=True) or {}
        layout = payload.get("layout", payload)
        salvato = _get_gp().salva(normalizza_editor_layout(layout))
        return jsonify({"ok": True, "layout": salvato}), 200
    except Exception as e:
        current_app.logger.exception("Errore salvataggio layout editor template atti: %s", e)
        return jsonify({"ok": False, "errore": "Layout non salvato. Riprova tra poco."}), 200


@template_atti.route("/api/editor-layout/reset", methods=["POST"])
@_richiedi_login
def reset_editor_layout():
    try:
        ripristinato = _get_gp().reset()
        return jsonify({"ok": True, "layout": ripristinato}), 200
    except Exception as e:
        current_app.logger.exception("Errore reset layout editor template atti: %s", e)
        return jsonify({"ok": False, "errore": "Layout non ripristinato. Riprova tra poco."}), 200


@template_atti.route("/api/scanner/windows-scan", methods=["POST"])
@_richiedi_login
def scanner_windows_scan():
    try:
        from pct.template_atti import acquisisci_da_scanner_windows

        scan = acquisisci_da_scanner_windows()
        return jsonify({"ok": True, "scan": scan}), 200
    except Exception as e:
        current_app.logger.exception("Errore scanner desktop template atti: %s", e)
        return jsonify({"ok": False, "errore": "Acquisizione da scanner non completata."}), 200


# ================================================================ IMPORTA DOCUMENTO

def _decode_document_bytes(data: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            return data.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), "utf-8 con caratteri sostitutivi"


def _extract_docx_fonts(data: bytes) -> list[str]:
    fonts: set[str] = set()
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            for name in ("word/document.xml", "word/styles.xml"):
                if name not in archive.namelist():
                    continue
                root = ElementTree.fromstring(archive.read(name))
                for node in root.iter():
                    if node.tag.rsplit("}", 1)[-1] != "rFonts":
                        continue
                    for attr, value in node.attrib.items():
                        if attr.rsplit("}", 1)[-1] in {"ascii", "hAnsi", "cs", "eastAsia"} and str(value or "").strip():
                            fonts.add(str(value).strip())
    except Exception:
        current_app.logger.debug("Font DOCX non rilevati", exc_info=True)
    return sorted(fonts, key=str.casefold)


def _rtf_control_group_span(raw: str, control: str, *, max_scan: int = 300_000) -> tuple[int, int] | None:
    start = raw.find("{" + control)
    if start < 0:
        return None
    depth = 0
    limit = min(len(raw), start + max_scan)
    for index in range(start, limit):
        char = raw[index]
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return start, index + 1
    return None


def _remove_rtf_control_group(raw: str, control: str) -> str:
    span = _rtf_control_group_span(raw, control)
    if span is None:
        return raw
    start, end = span
    return raw[:start] + raw[end:]


def _strip_rtf_controls(value: str) -> str:
    output: list[str] = []
    index = 0
    while index < len(value):
        char = value[index]
        if char == "\\":
            index += 1
            if index >= len(value):
                break
            escaped = value[index]
            if escaped in "{}\\":
                output.append(escaped)
                index += 1
                continue
            if escaped == "'":
                index += 3
                continue
            if not escaped.isalpha():
                index += 1
                continue
            while index < len(value) and value[index].isalpha():
                index += 1
            if index < len(value) and value[index] == "-":
                index += 1
            while index < len(value) and value[index].isdigit():
                index += 1
            if index < len(value) and value[index] == " ":
                index += 1
            continue
        if char not in "{}":
            output.append(char)
        index += 1
    return "".join(output)


def _extract_rtf_fonts(raw: str) -> list[str]:
    fonts: set[str] = set()
    span = _rtf_control_group_span(raw, "\\fonttbl", max_scan=60_000)
    table = raw[span[0]: span[1]] if span else raw[:60_000]
    for entry in table.split(";"):
        if "\\f" not in entry:
            continue
        name = _strip_rtf_controls(entry).strip(" ;")
        if name:
            fonts.add(name)
    return sorted(fonts, key=str.casefold)


def _rtf_to_text(raw: str) -> str:
    text = _remove_rtf_control_group(raw, "\\fonttbl")
    output: list[str] = []
    index = 0
    while index < len(text):
        char = text[index]
        if char == "\\":
            index += 1
            if index >= len(text):
                break
            marker = text[index]
            if marker == "'":
                hex_value = text[index + 1: index + 3]
                try:
                    output.append(bytes.fromhex(hex_value).decode("cp1252"))
                except Exception:
                    pass
                index += 3
                continue
            if marker in "{}\\":
                output.append(marker)
                index += 1
                continue
            if not marker.isalpha():
                index += 1
                continue
            control_start = index
            while index < len(text) and text[index].isalpha():
                index += 1
            control = text[control_start:index]
            for known_control in ("pard", "par", "line", "tab", "u"):
                if control.startswith(known_control) and control != known_control:
                    index = control_start + len(known_control)
                    control = known_control
                    break
            negative = False
            if index < len(text) and text[index] == "-":
                negative = True
                index += 1
            number_start = index
            while index < len(text) and text[index].isdigit():
                index += 1
            number = text[number_start:index]
            if index < len(text) and text[index] == " ":
                index += 1
            if control in {"par", "pard", "line"}:
                output.append("\n")
            elif control == "tab":
                output.append("\t")
            elif control == "u" and number:
                code = int(("-" if negative else "") + number)
                if code < 0:
                    code += 65536
                output.append(chr(code))
                if index < len(text) and text[index] not in "\\{}\r\n":
                    index += 1
            continue
        if char not in "{}":
            output.append(char)
        index += 1
    cleaned = "".join(output).replace("\r\n", "\n").replace("\r", "\n")
    cleaned = "\n".join(line.rstrip(" \t") for line in cleaned.split("\n"))
    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")
    return cleaned.strip()

@template_atti.route("/api/importa-documento", methods=["POST"])
@_richiedi_login
def api_importa_documento():
    """Converte un documento DOCX, PDF, RTF o TXT in testo/HTML per l'editor template."""
    try:
        file = request.files.get("file")
        if not file or not file.filename:
            return jsonify({"errore": "Nessun file ricevuto."}), 200

        filename = file.filename
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext == "docx":
            import mammoth
            data = file.read()
            result = mammoth.convert_to_html(io.BytesIO(data))
            html = result.value or ""
            # Stripping di tag vuoti lasciati da mammoth
            html = re.sub(r"<p>\s*</p>", "", html)
            fonts = _extract_docx_fonts(data)
            return jsonify({
                "ok": True,
                "tipo": "html",
                "contenuto": html,
                "filename": filename,
                "fonts": fonts,
                "detectedFonts": fonts,
                "encoding": "DOCX UTF-8 interno",
                "note": "Documento DOCX importato: font e testo sono disponibili per creare un template personalizzato.",
            })

        if ext == "pdf":
            data = file.read()
            html, fonts, layout_preserved = _pdf_layout_html_from_bytes(data)
            return jsonify({
                "ok": True,
                "tipo": "pdf_layout" if layout_preserved else "html",
                "contenuto": html,
                "filename": filename,
                "fonts": fonts,
                "detectedFonts": fonts,
                "encoding": "PDF testo vettoriale",
                "layoutPreserved": layout_preserved,
                "note": "PDF importato in pagine e righe modificabili: font e struttura disponibili quando dichiarati dal file.",
            })

        if ext in {"txt", "text"}:
            data = file.read()
            testo, encoding = _decode_document_bytes(data)
            return jsonify({"ok": True, "tipo": "testo", "contenuto": testo, "filename": filename, "fonts": [], "encoding": encoding})

        if ext == "rtf":
            raw, encoding = _decode_document_bytes(file.read())
            fonts = _extract_rtf_fonts(raw)
            testo = _rtf_to_text(raw)
            return jsonify({"ok": True, "tipo": "testo", "contenuto": testo, "filename": filename, "fonts": fonts, "detectedFonts": fonts, "encoding": encoding})

        return jsonify({"errore": f"Formato '.{ext}' non supportato. Usa DOCX, PDF, RTF o TXT."}), 200

    except Exception as e:
        current_app.logger.exception("Errore api_importa_documento: %s", e)
        return jsonify({"errore": "Importazione documento non completata."}), 200


# ================================================================ PDF

@template_atti.route("/<id_template>/pdf", methods=["POST"])
@_richiedi_login
def pdf(id_template: str):
    gt = _get_gt()
    t = gt.get(id_template)
    if not t:
        abort(404)
    testo, html = _request_editor_content(request.form.get("testo_generato", ""))
    titolo = t.titolo
    layout = _parse_editor_layout(request.form.get("testo_generato__editor_layout"))
    buf = _genera_pdf(titolo, testo, current_app.config, layout=layout, timbro=_get_studio_timbro(), html_content=html)
    nome_file = _safe_pdf_download_name(titolo)
    return send_file(buf, mimetype="application/pdf",
                     as_attachment=False, download_name=nome_file)


@template_atti.route("/compila/<model_code>/pdf", methods=["POST"])
@_richiedi_login
def compila_pdf(model_code: str):
    from pct.compilatore_atti import get_modello, render_compiled_act
    model = get_modello(model_code)
    if not model:
        abort(404)
    testo, html = _request_editor_content()
    if not testo:
        resolved = _resolve_compiler_context(model_code)
        testo = render_compiled_act(model_code, resolved["payload"], include_timbro=False)
    titolo = request.form.get("title", "") or model["name"]
    layout = _parse_editor_layout(request.form.get("testo_generato__editor_layout"))
    buf = _genera_pdf(titolo, testo, current_app.config, layout=layout, timbro=_get_studio_timbro(), html_content=html)
    nome_file = _safe_pdf_download_name(titolo)
    return send_file(buf, mimetype="application/pdf",
                     as_attachment=False, download_name=nome_file)


@template_atti.route("/compila/<model_code>/word", methods=["POST"])
@_richiedi_login
def compila_word(model_code: str):
    from pct.compilatore_atti import get_modello, render_compiled_act

    model = get_modello(model_code)
    if not model:
        abort(404)
    testo, html = _request_editor_content()
    if not testo:
        resolved = _resolve_compiler_context(model_code)
        testo = render_compiled_act(model_code, resolved["payload"], include_timbro=False)
    titolo = request.form.get("title", "") or model["name"]
    layout = _parse_editor_layout(request.form.get("testo_generato__editor_layout")) or {}
    try:
        from pct.template_atti import font_editor

        font_meta = font_editor(str(layout.get("font_family") or ""))
    except Exception:
        font_meta = {"docx_family": "Times New Roman"}
    try:
        font_size = max(9.0, min(22.0, float(layout.get("font_size_pt") or layout.get("font_size") or layout.get("fontSize") or 12)))
    except (TypeError, ValueError):
        font_size = 12.0
    try:
        line_height = max(1.0, min(2.2, float(layout.get("line_height") or layout.get("lineHeight") or 1.45)))
    except (TypeError, ValueError):
        line_height = 1.45
    timbro = _get_studio_timbro()
    docx_font = str(font_meta.get("docx_family") or layout.get("fallback_font_family") or "Times New Roman")
    text_align = str(layout.get("text_align") or layout.get("textAlign") or "justify").lower()
    stamp_position = str(layout.get("stamp_position") or layout.get("stampPosition") or "top-center").lower()
    download_name = _safe_word_download_name(titolo)
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        document = Document()
        stamp_alignment = WD_ALIGN_PARAGRAPH.CENTER
        if stamp_position.endswith("left"):
            stamp_alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stamp_position.endswith("right"):
            stamp_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            stamp_lines = timbro.to_lines()
        except Exception:
            stamp_lines = []
        if not stamp_lines and getattr(timbro, "text", ""):
            stamp_lines = [{"text": line} for line in str(timbro.text).splitlines() if line.strip()]
        try:
            stamp_offset = int(layout.get("stamp_offset_y_mm") or 0)
        except (TypeError, ValueError):
            stamp_offset = 0
        if stamp_position.startswith("middle"):
            for _ in range(7):
                document.add_paragraph("")
        elif stamp_offset > 8:
            for _ in range(min(5, max(1, stamp_offset // 12))):
                document.add_paragraph("")
        for line in stamp_lines:
            text_value = str(line.get("text") if isinstance(line, dict) else getattr(line, "text", "")).strip()
            if not text_value:
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = stamp_alignment
            run = paragraph.add_run(text_value)
            run.bold = bool(line.get("bold")) if isinstance(line, dict) else bool(getattr(line, "bold", False))
            run.font.name = docx_font
            run.font.size = Pt(9)
        if stamp_lines:
            document.add_paragraph("")
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        body_alignment = align_map.get(text_align, WD_ALIGN_PARAGRAPH.JUSTIFY)
        if html:
            _append_html_to_docx(
                document,
                html,
                docx_font=docx_font,
                font_size=font_size,
                line_height=line_height,
                default_alignment=body_alignment,
                align_map=align_map,
            )
        else:
            for block in str(testo or "").splitlines():
                paragraph = document.add_paragraph()
                paragraph.alignment = body_alignment
                paragraph.paragraph_format.line_spacing = line_height
                run = paragraph.add_run(block)
                run.font.name = docx_font
                run.font.size = Pt(font_size)
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=download_name,
        )
    except Exception:
        current_app.logger.exception("Generazione DOCX non completata per %s", model_code)
        html = (
            "<!doctype html><html><head><meta charset=\"utf-8\">"
            f"<title>{escape(str(titolo), quote=True)}</title>"
            "</head><body>"
            f"<pre>{escape(testo)}</pre>"
            "</body></html>"
        )
        response = current_app.response_class(html.encode("utf-8"), mimetype="application/msword; charset=utf-8")
        response.headers["Content-Disposition"] = f'attachment; filename="{download_name}"'
        return response


def _rtf_escape(value: str) -> str:
    parts = []
    for char in value:
        code = ord(char)
        if char == "\\":
            parts.append("\\\\")
        elif char == "{":
            parts.append("\\{")
        elif char == "}":
            parts.append("\\}")
        elif char == "\n":
            parts.append("\\par\n")
        elif code > 127:
            parts.append(f"\\u{code}?")
        else:
            parts.append(char)
    return "".join(parts)


def _rtf_content_from_text(
    testo: str,
    layout: dict[str, Any] | None = None,
    *,
    timbro: Any = None,
    html_content: str | None = None,
) -> str:
    layout = layout or {}
    try:
        from pct.template_atti import font_editor

        font_meta = font_editor(str(layout.get("font_family") or ""))
    except Exception:
        font_meta = {"rtf_family": "Times New Roman"}
    try:
        font_size = max(9.0, min(22.0, float(layout.get("font_size") or layout.get("fontSize") or layout.get("font_size_pt") or 12)))
    except (TypeError, ValueError):
        font_size = 12.0
    rtf_font = str(font_meta.get("rtf_family") or "Times New Roman").replace(";", "")
    stamp_text = _stamp_text(timbro)
    body_rtf = _rtf_body_from_html(html_content or "") if html_content else _rtf_escape(testo)
    stamp_rtf = ""
    if stamp_text:
        stamp_rtf = f"{_rtf_align_command(layout, stamp=True)}\\fs18\n{_rtf_escape(stamp_text)}\\par\\par\n"
    body_align = _rtf_align_command(layout)
    return (
        "{\\rtf1\\ansi\\deff0"
        f"{{\\fonttbl{{\\f0 {rtf_font};}}}}"
        f"\\fs{int(round(font_size * 2))}\n"
        f"{stamp_rtf}{body_align}\n{body_rtf}"
        "}"
    )


@template_atti.route("/compila/<model_code>/rtf", methods=["POST"])
@_richiedi_login
def compila_rtf(model_code: str):
    from pct.compilatore_atti import get_modello, render_compiled_act

    model = get_modello(model_code)
    if not model:
        abort(404)
    testo, html = _request_editor_content()
    if not testo:
        resolved = _resolve_compiler_context(model_code)
        testo = render_compiled_act(model_code, resolved["payload"], include_timbro=False)
    titolo = request.form.get("title", "") or model["name"]
    layout = _parse_editor_layout(request.form.get("testo_generato__editor_layout")) or {}
    content = _rtf_content_from_text(testo, layout, timbro=_get_studio_timbro(), html_content=html)
    download_name = _safe_word_download_name(titolo).rsplit(".", 1)[0] + ".rtf"
    response = current_app.response_class(content.encode("utf-8"), mimetype="application/rtf; charset=utf-8")
    response.headers["Content-Disposition"] = f'attachment; filename="{download_name}"'
    return response


@template_atti.route("/api/compilazione-multipla-rtf", methods=["POST"])
@_richiedi_login
def api_compilazione_multipla_rtf():
    from pct.compilatore_atti import get_modello, render_compiled_act

    raw_codes = request.form.getlist("model_codes")
    seen_codes: set[str] = set()
    model_codes: list[str] = []
    for raw_code in raw_codes:
        code = str(raw_code or "").strip()
        if code and code not in seen_codes:
            model_codes.append(code)
            seen_codes.add(code)
        if len(model_codes) >= 20:
            break
    if not model_codes:
        return jsonify({"ok": False, "message": "Seleziona almeno un modello da compilare."}), 400

    layout = _parse_editor_layout(request.form.get("testo_generato__editor_layout")) or {}
    timbro = _get_studio_timbro()
    archive = io.BytesIO()
    written = 0
    used_names: dict[str, int] = {}
    with zipfile.ZipFile(archive, mode="w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for code in model_codes:
            model = get_modello(code)
            if not model:
                continue
            resolved = _resolve_compiler_context(code)
            testo = render_compiled_act(code, resolved["payload"], include_timbro=False)
            base_name = _safe_word_download_name(str(model.get("name") or code)).rsplit(".", 1)[0] or code
            file_name = f"{base_name}.rtf"
            duplicate_index = used_names.get(file_name, 0)
            used_names[file_name] = duplicate_index + 1
            if duplicate_index:
                file_name = f"{base_name}-{duplicate_index + 1}.rtf"
            zip_file.writestr(file_name, _rtf_content_from_text(testo, layout, timbro=timbro).encode("utf-8"))
            written += 1
    if not written:
        return jsonify({"ok": False, "message": "Nessun modello selezionato è disponibile per l'esportazione RTF."}), 404
    archive.seek(0)
    return send_file(
        archive,
        mimetype="application/zip",
        as_attachment=True,
        download_name="compilazione_multipla_template_atti.zip",
    )


# ================================================================ API assistente redazionale

@template_atti.route("/api/assistente-redazionale/<model_code>", methods=["POST"])
@_richiedi_login
def api_assistente_redazionale(model_code: str):
    try:
        from pct.compilatore_atti import get_modello

        if not get_modello(model_code):
            return jsonify({"ok": False, "errore": "Modello non trovato."}), 200
        resolved = _resolve_compiler_context(model_code)
        analysis = _build_assistant_analysis(
            model_code,
            payload=resolved["payload"],
            selected_cliente=resolved["selected_cliente"],
            selected_fascicolo=resolved["selected_fascicolo"],
        )
        editor_text, editor_html = _request_editor_content()
        analysis = _editor_lex_enrichment(
            analysis,
            action=str(request.form.get("lex_action") or request.form.get("azione") or "").strip(),
            mode=str(request.form.get("lex_mode") or request.form.get("modalita") or "").strip(),
            text=editor_text,
            html=editor_html,
            instructions=str(request.form.get("istruzioni_lex") or "").strip(),
        )
        return jsonify({"ok": True, "analysis": analysis}), 200
    except Exception as e:
        current_app.logger.exception("Errore api_assistente_redazionale: %s", e)
        return jsonify({"ok": False, "errore": "Analisi redazionale non completata."}), 200


# ================================================================ helpers

def _variabili_safe(v: dict) -> dict:
    """Serializza le variabili per il form nascosto (solo stringhe)."""
    safe = {}
    for k, val in v.items():
        if isinstance(val, str):
            safe[k] = val
    return safe


def _parse_editor_layout(raw: str | None) -> dict | None:
    if not raw:
        return None
    try:
        from pct.template_atti import normalizza_editor_layout

        parsed = json.loads(raw)
        if not isinstance(parsed, dict):
            return None
        return normalizza_editor_layout(parsed)
    except Exception:
        return None


def _to_editor_html(content: str) -> str:
    text = (content or "").strip()
    if not text:
        return "<p></p>"
    if _looks_like_html(text):
        return text
    return _plain_text_to_editor_html(text)


def _looks_like_html(content: str) -> bool:
    lowered = (content or "").strip().lower()
    return any(
        tag in lowered
        for tag in ("<p", "<div", "<h1", "<h2", "<h3", "<ul", "<ol", "<li", "<table", "<blockquote", "<br", "<hr", "<img", "<figure")
    )


def _plain_text_to_editor_html(content: str) -> str:
    normalized = (content or "").replace("\r", "").strip()
    if not normalized:
        return "<p></p>"

    blocks = [
        [line.strip() for line in block.split("\n") if line.strip()]
        for block in re.split(r"\n\s*\n", normalized)
    ]
    blocks = [lines for lines in blocks if lines]
    html_blocks: list[str] = []

    for index, lines in enumerate(blocks):
        if len(lines) == 1:
            line = lines[0]
            if index == 0 and _is_upper_heading(line):
                html_blocks.append(f'<div class="legal-doc-kicker">{escape(line)}</div>')
                continue
            if (index == 0 or index == 1) and len(line) <= 90 and not line.endswith(":"):
                html_blocks.append(f"<h1>{escape(line)}</h1>")
                continue
            if _is_upper_heading(line):
                html_blocks.append(f"<h2>{escape(line)}</h2>")
                continue
            if line.endswith(":") and len(line) <= 60:
                html_blocks.append(f'<p class="legal-doc-label"><strong>{escape(line)}</strong></p>')
                continue

        if all(re.match(r"^\d+\.\s+", line) for line in lines):
            items = [re.sub(r"^\d+\.\s+", "", line) for line in lines]
            html_blocks.append("<ol>" + "".join(f"<li>{escape(item)}</li>" for item in items) + "</ol>")
            continue

        if all(line.startswith("- ") for line in lines):
            items = [line[2:].strip() for line in lines]
            html_blocks.append("<ul>" + "".join(f"<li>{escape(item)}</li>" for item in items) + "</ul>")
            continue

        css_class = ' class="legal-doc-party"' if len(lines) > 1 else ""
        html_blocks.append(f"<p{css_class}>" + "<br>".join(escape(line) for line in lines) + "</p>")

    return "\n".join(html_blocks) or "<p></p>"


def _is_upper_heading(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped or len(stripped) > 90:
        return False
    letters = [char for char in stripped if char.isalpha()]
    if not letters:
        return False
    return stripped == stripped.upper()


class _EditorTextParser(HTMLParser):
    _BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "blockquote", "tr", "table", "ul", "ol"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []

    def _append_break(self) -> None:
        if self.parts and not self.parts[-1].endswith("\n"):
            self.parts.append("\n")

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag == "br":
            self._append_break()
        elif tag == "li":
            self._append_break()
            self.parts.append("- ")
        elif tag in self._BLOCK_TAGS:
            self._append_break()

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in self._BLOCK_TAGS or tag.lower() == "li":
            self._append_break()

    def handle_data(self, data: str) -> None:
        if data:
            self.parts.append(data)

    def get_text(self) -> str:
        text = "".join(self.parts).replace("\r\n", "\n").replace("\r", "\n")
        while "\n\n\n" in text:
            text = text.replace("\n\n\n", "\n\n")
        return text.strip()


def _editor_html_to_text(value: str) -> str:
    parser = _EditorTextParser()
    parser.feed(value or "")
    parser.close()
    return parser.get_text()


def _request_editor_content(default_text: str = "") -> tuple[str, str]:
    html = (request.form.get("testo_generato_html") or "").strip()
    text = (request.form.get("testo_generato") or default_text or "").strip()
    if html and _looks_like_html(html):
        return text or _editor_html_to_text(html), _sanitize_editor_html(html)
    return text, ""


def _stamp_lines(stamp: Any) -> list[dict[str, Any]]:
    if not stamp or not getattr(stamp, "enabled", False):
        return []
    try:
        lines = stamp.to_lines()
    except Exception:
        lines = []
    if lines:
        return [line for line in lines if str(line.get("text") or "").strip()]
    text = str(getattr(stamp, "text", "") or "")
    return [{"text": line.strip()} for line in text.splitlines() if line.strip()]


def _stamp_text(stamp: Any) -> str:
    return "\n".join(str(line.get("text") or "").strip() for line in _stamp_lines(stamp)).strip()


def _document_text_with_stamp(text: str, stamp: Any) -> str:
    stamp_text = _stamp_text(stamp)
    body = str(text or "").strip()
    if not stamp_text:
        return body
    if body.startswith(stamp_text.splitlines()[0]):
        return body
    return f"{stamp_text}\n\n{body}".strip()


def _append_html_to_docx(
    document: Any,
    html: str,
    *,
    docx_font: str,
    font_size: float,
    line_height: float,
    default_alignment: Any,
    align_map: dict[str, Any],
) -> None:
    try:
        from docx.shared import Pt
        from lxml import etree
    except Exception:
        for block in _editor_html_to_text(html).splitlines():
            paragraph = document.add_paragraph()
            paragraph.alignment = default_alignment
            run = paragraph.add_run(block)
            run.font.name = docx_font
            run.font.size = Pt(font_size)
        return

    try:
        root = etree.fromstring(
            f"<div>{html}</div>".encode("utf-8"),
            parser=etree.HTMLParser(encoding="utf-8"),
        )
        body = root.find(".//body") or root.find(".//div") or root
    except Exception:
        body = None

    def _alignment_from_style(style: str) -> Any:
        lowered = (style or "").lower()
        if "text-align:center" in lowered or "text-align: center" in lowered:
            return align_map.get("center", default_alignment)
        if "text-align:right" in lowered or "text-align: right" in lowered:
            return align_map.get("right", default_alignment)
        if "text-align:left" in lowered or "text-align: left" in lowered:
            return align_map.get("left", default_alignment)
        if "text-align:justify" in lowered or "text-align: justify" in lowered:
            return align_map.get("justify", default_alignment)
        return default_alignment

    def _add_run(paragraph: Any, text: str, *, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.name = docx_font
        run.font.size = Pt(font_size)

    def _walk_inline(paragraph: Any, node: Any, *, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
        tag = str(getattr(node, "tag", "") or "").lower().split("}")[-1]
        next_bold = bold or tag in {"strong", "b"}
        next_italic = italic or tag in {"em", "i"}
        next_underline = underline or tag == "u"
        if getattr(node, "text", None):
            _add_run(paragraph, str(node.text), bold=next_bold, italic=next_italic, underline=next_underline)
        for child in list(node):
            child_tag = str(getattr(child, "tag", "") or "").lower().split("}")[-1]
            if child_tag == "br":
                paragraph.add_run().add_break()
            else:
                _walk_inline(paragraph, child, bold=next_bold, italic=next_italic, underline=next_underline)
            if getattr(child, "tail", None):
                _add_run(paragraph, str(child.tail), bold=next_bold, italic=next_italic, underline=next_underline)

    def _add_paragraph(node: Any, style: str | None = None, list_style: str | None = None) -> None:
        paragraph = document.add_paragraph(style=list_style or style)
        paragraph.alignment = _alignment_from_style(str(node.get("style") or ""))
        paragraph.paragraph_format.line_spacing = line_height
        _walk_inline(paragraph, node)

    def _process(node: Any) -> None:
        tag = str(getattr(node, "tag", "") or "").lower().split("}")[-1]
        if tag in {"h1", "h2", "h3", "h4"}:
            _add_paragraph(node, style=f"Heading {tag[1]}")
            return
        if tag == "ul":
            for li in node.findall("li"):
                _add_paragraph(li, list_style="List Bullet")
            return
        if tag == "ol":
            for li in node.findall("li"):
                _add_paragraph(li, list_style="List Number")
            return
        if tag in {"p", "div", "section", "blockquote"}:
            if len(node) == 0 and not str(getattr(node, "text", "") or "").strip():
                document.add_paragraph("")
            else:
                _add_paragraph(node)
            return
        for child in list(node):
            _process(child)

    if body is None:
        return
    for child in list(body):
        _process(child)


def _rtf_align_command(layout: dict[str, Any] | None, *, stamp: bool = False) -> str:
    align = str((layout or {}).get("text_align") or "justify").lower()
    if stamp:
        position = str((layout or {}).get("stamp_position") or "top-center").lower()
        if position.endswith("left"):
            return "\\qc"
        if position.endswith("right"):
            return "\\qc"
        return "\\qc"
    return {"left": "\\ql", "center": "\\qc", "right": "\\qr", "justify": "\\qj"}.get(align, "\\qj")


class _RtfHtmlParser(HTMLParser):
    _BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "blockquote"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._list_stack: list[str] = []

    def _paragraph(self) -> None:
        if self.parts and not self.parts[-1].endswith("\\par\n"):
            self.parts.append("\\par\n")

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in self._BLOCK_TAGS:
            self._paragraph()
            if tag in {"h1", "h2", "h3", "h4"}:
                self.parts.append("\\b ")
        elif tag == "br":
            self.parts.append("\\line ")
        elif tag in {"strong", "b"}:
            self.parts.append("\\b ")
        elif tag in {"em", "i"}:
            self.parts.append("\\i ")
        elif tag == "u":
            self.parts.append("\\ul ")
        elif tag in {"ul", "ol"}:
            self._list_stack.append(tag)
        elif tag == "li":
            self._paragraph()
            self.parts.append("- ")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"strong", "b"}:
            self.parts.append("\\b0 ")
        elif tag in {"em", "i"}:
            self.parts.append("\\i0 ")
        elif tag == "u":
            self.parts.append("\\ul0 ")
        elif tag in {"h1", "h2", "h3", "h4"}:
            self.parts.append("\\b0 ")
            self._paragraph()
        elif tag in self._BLOCK_TAGS or tag == "li":
            self._paragraph()
        elif tag in {"ul", "ol"} and self._list_stack:
            self._list_stack.pop()

    def handle_data(self, data: str) -> None:
        if data:
            self.parts.append(_rtf_escape(data))

    def get_rtf(self) -> str:
        return "".join(self.parts).strip()


def _rtf_body_from_html(html: str) -> str:
    parser = _RtfHtmlParser()
    parser.feed(html or "")
    parser.close()
    return parser.get_rtf()


def _pdf_layout_html_from_bytes(data: bytes) -> tuple[str, list[str], bool]:
    import io as _io
    import pdfplumber

    pages: list[str] = []
    fonts: set[str] = set()
    used_layout = False
    with pdfplumber.open(_io.BytesIO(data)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            page_lines: list[str] = []
            try:
                words = page.extract_words(
                    keep_blank_chars=False,
                    use_text_flow=True,
                    extra_attrs=["fontname", "size"],
                )
            except TypeError:
                words = page.extract_words(keep_blank_chars=False, use_text_flow=True)
            if words:
                used_layout = True
                grouped: dict[int, list[dict[str, Any]]] = {}
                for word in words:
                    text = str(word.get("text") or "").strip()
                    if not text:
                        continue
                    font_name = str(word.get("fontname") or "").split("+")[-1].strip()
                    if font_name:
                        fonts.add(font_name)
                    top = int(round(float(word.get("top") or 0) / 3) * 3)
                    grouped.setdefault(top, []).append(word)
                for _top, line_words in sorted(grouped.items(), key=lambda item: item[0]):
                    ordered = sorted(line_words, key=lambda item: float(item.get("x0") or 0))
                    text = " ".join(str(item.get("text") or "").strip() for item in ordered if str(item.get("text") or "").strip())
                    if not text:
                        continue
                    font_name = str(ordered[0].get("fontname") or "").split("+")[-1].strip()
                    size = ordered[0].get("size")
                    style_bits = []
                    if font_name:
                        style_bits.append(f"font-family:{escape(font_name, quote=True)}")
                    try:
                        style_bits.append(f"font-size:{max(8, min(22, float(size))):.1f}pt")
                    except (TypeError, ValueError):
                        pass
                    style_attr = f' style="{";".join(style_bits)}"' if style_bits else ""
                    page_lines.append(f"<p{style_attr}>{escape(text)}</p>")
            else:
                text = page.extract_text() or ""
                page_lines.extend(f"<p>{escape(line.strip())}</p>" for line in text.splitlines() if line.strip())
            pages.append(
                '<section class="iusentra-imported-pdf-page" data-page="{}">'.format(page_index)
                + "\n".join(page_lines or ["<p></p>"])
                + "</section>"
            )
    return "\n".join(pages), sorted(fonts), used_layout


def _fallback_pdf_from_text(titolo: str, testo: str, config: dict, layout: dict | None = None, timbro=None) -> io.BytesIO:
    safe_text = testo or ""
    if timbro is not None and hasattr(timbro, "to_text"):
        timbro_text = timbro.to_text()
        if timbro_text and safe_text.strip().startswith(timbro_text.splitlines()[0]):
            safe_text = safe_text.strip()[len(timbro_text):].lstrip()
    safe_text = _editor_html_to_text(safe_text)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                         Spacer, HRFlowable)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER

        PRIMARY = colors.HexColor("#1a3a5c")
        buf = io.BytesIO()
        try:
            from pct.legal_layout_profiles import make_pdf_stamp_callback, top_margin_with_stamp

            stamp_callback = make_pdf_stamp_callback(timbro, layout=layout or {})
            top_margin = top_margin_with_stamp({"margin_top_mm": 25, **(layout or {})}, timbro)
        except Exception:
            stamp_callback = None
            top_margin = 25
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                 leftMargin=25*mm, rightMargin=25*mm,
                                 topMargin=top_margin*mm, bottomMargin=25*mm)
        styles = getSampleStyleSheet()
        style_body = ParagraphStyle("body", parent=styles["Normal"],
                                    fontSize=10, leading=15,
                                    fontName="Helvetica")
        style_title = ParagraphStyle("title", parent=styles["Normal"],
                                     fontSize=13, leading=17,
                                     fontName="Helvetica-Bold",
                                     textColor=PRIMARY,
                                     alignment=TA_CENTER)
        story = []
        story.append(Paragraph(titolo.upper(), style_title))
        story.append(Spacer(1, 6*mm))
        story.append(HRFlowable(width="100%", thickness=1.5, color=PRIMARY))
        story.append(Spacer(1, 6*mm))
        for riga in safe_text.split("\n"):
            if riga.strip():
                story.append(Paragraph(riga.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"),
                                       style_body))
            else:
                story.append(Spacer(1, 4*mm))
        studio_nome = config.get("STUDIO_NOME", "IUSENTRA")
        story.append(Spacer(1, 8*mm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#94a3b8")))
        story.append(Spacer(1, 2*mm))
        story.append(Paragraph(studio_nome, ParagraphStyle(
            "footer", parent=style_body, fontSize=8,
            textColor=colors.HexColor("#6b7280"), alignment=TA_CENTER)))
        if stamp_callback is not None:
            doc.build(story, onFirstPage=stamp_callback, onLaterPages=stamp_callback)
        else:
            doc.build(story)
        buf.seek(0)
        return buf
    except ImportError:
        buf = io.BytesIO(safe_text.encode("utf-8"))
        buf.seek(0)
        return buf


def _genera_pdf(
    titolo: str,
    testo: str,
    config: dict,
    layout: dict | None = None,
    timbro=None,
    html_content: str | None = None,
) -> io.BytesIO:
    html = _sanitize_editor_html(html_content) if html_content and _looks_like_html(html_content) else _to_editor_html(testo)
    try:
        from pct.editor import html_to_pdf

        try:
            from pct.template_normative_compliance import analyze_template_compliance

            compliance = analyze_template_compliance(
                str((layout or {}).get("model_code") or ""),
                {},
                studio_timbro_payload=timbro.to_payload() if hasattr(timbro, "to_payload") else None,
                strict_payload=False,
            )
            layout_profile_id = (compliance.layout_compliance.layout_profile_id if compliance.layout_compliance else "")
            from pct.legal_layout_profiles import get_layout_profile

            layout_profile = get_layout_profile(layout_profile_id)
        except Exception:
            layout_profile = None
        pdf_bytes = html_to_pdf(html, titolo, layout=layout, studio_timbro=timbro, layout_profile=layout_profile)
        buf = io.BytesIO(pdf_bytes)
        buf.seek(0)
        return buf
    except Exception as exc:
        current_app.logger.exception("Errore generazione PDF HTML template atti: %s", exc)
        return _fallback_pdf_from_text(titolo, html or testo, config, layout=layout, timbro=timbro)
